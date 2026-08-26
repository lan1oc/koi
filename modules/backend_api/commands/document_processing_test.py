#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from modules.backend_api.commands import document_processing
from modules.backend_api.commands.document_processing import _collect_notice_pdf_candidates, _doc_pdf_extract_run


class PdfExtractCommandTest(unittest.TestCase):
    def tearDown(self) -> None:
        with document_processing._NOTICE_TASK_LOCK:
            document_processing._NOTICE_TASKS.clear()

    def test_multiple_pdf_default_output_uses_first_source_name(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            first = Path(temp_dir) / "report-a.pdf"
            second = Path(temp_dir) / "report-b.pdf"
            first.touch()
            second.touch()
            selections = [
                {"file_path": str(first), "page_num": 1, "order": 1},
                {"file_path": str(second), "page_num": 1, "order": 2},
            ]

            with patch(
                "modules.backend_api.commands.document_processing.merge_pages_from_multiple_pdfs",
                return_value=(2, 2),
            ) as merge_pages:
                result = _doc_pdf_extract_run({
                    "pdf_files": [str(first), str(second)],
                    "page_selections": selections,
                })

            self.assertTrue(result["success"])
            self.assertEqual(Path(result["output_file"]), Path(temp_dir) / "report-a_merged_pages.pdf")
            self.assertEqual(merge_pages.call_args.args[1], Path(temp_dir) / "report-a_merged_pages.pdf")

    def test_notice_pdf_conversion_discovers_generated_docs_without_failures(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "company"
            output_dir.mkdir()
            generated = output_dir / "授权委托书-test.docx"
            generated.touch()
            logs: list[str] = []

            candidates = _collect_notice_pdf_candidates(Path(temp_dir), [], logs, scan_target=True)

            self.assertIn(generated.resolve(), candidates)
            self.assertTrue(any("1 个可转换Word文件" in line for line in logs))

    def test_notice_pdf_conversion_explicit_empty_list_does_not_scan_target(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "company"
            output_dir.mkdir()
            (output_dir / "授权委托书-test.docx").touch()
            logs: list[str] = []

            candidates = _collect_notice_pdf_candidates(Path(temp_dir), [], logs)

            self.assertEqual(candidates, [])
            self.assertTrue(any("失败列表中未找到" in line for line in logs))

    def test_notice_pdf_conversion_explicit_item_does_not_include_backup_or_siblings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            selected = root / "授权委托书-selected.docx"
            backup = root / "授权委托书-selected.final_backup.docx"
            sibling = root / "授权委托书-sibling.docx"
            selected.touch()
            backup.touch()
            sibling.touch()
            logs: list[str] = []

            candidates = _collect_notice_pdf_candidates(root, [{
                "file": str(selected),
                "backup_file": str(backup),
            }], logs)

            self.assertEqual(candidates, [selected.resolve()])

    def test_notice_pdf_conversion_relocates_item_after_classification(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            old_path = root / "宁波众翮科技有限公司" / "关于宁波众翮科技有限公司存在漏洞的通报.docx"
            moved_path = root / "南部商务区" / old_path.parent.name / old_path.name
            moved_path.parent.mkdir(parents=True)
            moved_path.touch()
            logs: list[str] = []

            candidates = _collect_notice_pdf_candidates(root, [{"output_file": str(old_path)}], logs)

            self.assertEqual(candidates, [moved_path.resolve()])
            self.assertTrue(any("已重新定位分类后的Word文件" in line for line in logs))

    def test_generated_authorization_word_is_deleted_after_pdf_conversion(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "授权委托书（执法调查类）.docx"
            source.touch()

            def convert(file_map, overwrite=True):
                for _, destination in file_map:
                    destination.touch()
                return len(file_map), 0, []

            with patch(
                "modules.backend_api.commands.document_processing.convert_with_word_com",
                side_effect=convert,
            ):
                result = document_processing._convert_generated_docs_to_pdf(work_dir, [])

            self.assertEqual(result["converted"], 1)
            self.assertFalse(source.exists())
            self.assertTrue(source.with_suffix(".pdf").exists())

    def test_manual_pdf_conversion_deletes_relocated_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            old_path = root / "宁波众翮科技有限公司" / "关于宁波众翮科技有限公司存在漏洞的通报.docx"
            moved_path = root / "南部商务区" / old_path.parent.name / old_path.name
            moved_path.parent.mkdir(parents=True)
            moved_path.touch()

            def convert(file_map, overwrite=True):
                for _, destination in file_map:
                    destination.touch()
                return len(file_map), 0, []

            with patch(
                "modules.backend_api.commands.document_processing.convert_with_word_com",
                side_effect=convert,
            ):
                result = document_processing._doc_notice_convert_failed_pdf({
                    "target_path": str(root),
                    "failed_files": [{"output_file": str(old_path)}],
                })

            self.assertTrue(result["success"])
            self.assertFalse(moved_path.exists())
            self.assertTrue(moved_path.with_suffix(".pdf").exists())
            self.assertEqual(result["deleted_files"], [str(moved_path.resolve())])

    def test_notice_task_start_reuses_active_task_for_same_target(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir)
            progress = document_processing.NoticeProgress()
            progress.set(42, "处理中")
            with document_processing._NOTICE_TASK_LOCK:
                document_processing._NOTICE_TASKS["existing-task"] = {
                    "task_id": "existing-task",
                    "running": True,
                    "done": False,
                    "target_key": document_processing._notice_target_key(target),
                    "progress": progress,
                }

            result = document_processing._doc_notice_process_start({"target_path": str(target)})

            self.assertTrue(result["success"])
            self.assertTrue(result["already_running"])
            self.assertEqual(result["task_id"], "existing-task")
            self.assertEqual(result["progress"], 42)

    def test_notice_mutating_commands_reject_active_target(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir)
            progress = document_processing.NoticeProgress()
            with document_processing._NOTICE_TASK_LOCK:
                document_processing._NOTICE_TASKS["active-task"] = {
                    "task_id": "active-task",
                    "running": True,
                    "done": False,
                    "target_key": document_processing._notice_target_key(target),
                    "progress": progress,
                }

            classify = document_processing._doc_notice_classify({"target_path": str(target)})
            convert = document_processing._doc_notice_convert_failed_pdf({
                "target_path": str(target),
                "failed_files": [],
            })

            self.assertEqual(classify["error_code"], "notice_task_active")
            self.assertEqual(convert["error_code"], "notice_task_active")

    def test_notice_status_marks_missing_task_explicitly(self) -> None:
        result = document_processing._doc_notice_process_status({"task_id": "missing-task"})

        self.assertTrue(result["done"])
        self.assertEqual(result["error_code"], "notice_task_not_found")

    def test_classification_backfills_copy_to_only_for_rewritten_notices(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            company_dir = root / "南部商务区" / "宁波众翮科技有限公司"
            company_dir.mkdir(parents=True)
            rewritten_path = company_dir / "关于宁波众翮科技有限公司存在漏洞的通报.docx"
            ordinary_path = company_dir / "普通材料.docx"

            rewritten = Document()
            rewritten.add_paragraph("抄送：")
            rewritten.core_properties.comments = "koi.notice.rewritten.v1"
            rewritten.save(rewritten_path)

            ordinary = Document()
            ordinary.add_paragraph("抄送：")
            ordinary.save(ordinary_path)

            logs: list[str] = []
            stats = document_processing._fill_rewritten_notice_copy_to(
                root,
                [["宁波众翮科技有限公司", "南部商务区"]],
                logs,
            )

            self.assertEqual(stats["updated"], 1)
            self.assertEqual(Document(rewritten_path).paragraphs[0].text, "抄送：南部商务区")
            self.assertEqual(Document(ordinary_path).paragraphs[0].text, "抄送：")

    def test_classification_does_not_overwrite_existing_copy_to(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            company_dir = root / "首南街道" / "宁波众翮科技有限公司"
            company_dir.mkdir(parents=True)
            notice_path = company_dir / "关于宁波众翮科技有限公司存在漏洞的通报.docx"
            notice = Document()
            notice.add_paragraph("抄送：原街道")
            notice.core_properties.comments = "koi.notice.rewritten.v1"
            notice.save(notice_path)

            stats = document_processing._fill_rewritten_notice_copy_to(
                root,
                [["宁波众翮科技有限公司", "首南街道"]],
                [],
            )

            self.assertEqual(stats["updated"], 0)
            self.assertEqual(Document(notice_path).paragraphs[0].text, "抄送：原街道")

    @staticmethod
    def _write_rewritten_notice(path: Path, *, marker: bool = True) -> None:
        doc = Document()
        doc.add_paragraph("鄞州区网络安全预警通报")
        doc.add_paragraph("〔2026〕第1期")
        doc.add_paragraph("关于测试有限公司所属系统存在未授权访问漏洞的通报")
        doc.add_paragraph("抄送：首南街道")
        doc.add_paragraph("验证情况")
        doc.add_paragraph("处置措施")
        if marker:
            doc.core_properties.comments = "koi.notice.rewritten.v1"
        doc.save(path)

    @staticmethod
    def _write_authorization(path: Path) -> None:
        doc = Document()
        doc.add_paragraph("授权委托书")
        doc.add_paragraph("关于测试有限公司所属系统存在未授权访问漏洞的通报")
        doc.save(path)

    @staticmethod
    def _write_pdf(path: Path, text: str = "授权委托书 测试有限公司") -> None:
        from reportlab.pdfgen import canvas

        canvas_instance = canvas.Canvas(str(path))
        canvas_instance.drawString(72, 760, text)
        canvas_instance.save()

    def _successful_notice_side_effects(self, work_dir: Path):
        def rewrite(source, **_kwargs):
            output = work_dir / Path(source).name.lstrip("0123456789")
            self._write_rewritten_notice(output)
            return {"success": True, "output_file": str(output), "backup_file": None}

        def authorization(*_args, **_kwargs):
            self._write_authorization(work_dir / "授权委托书（执法调查类）.docx")
            return True

        def convert(directory, _logs, _progress=None, skip_paths=None):
            source = Path(directory) / "授权委托书（执法调查类）.docx"
            target = source.with_suffix(".pdf")
            self._write_pdf(target)
            source.unlink()
            return {"converted": 1, "skipped": 0, "failures": [], "output_files": [str(target)]}

        return rewrite, authorization, convert

    def test_rewrite_failure_keeps_original_for_retry(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            Document().save(source)

            with patch(
                "modules.Document_Processing.Report_Rewrite.rewrite_report.rewrite_report",
                return_value={"success": False, "skip_reason": "测试失败"},
            ):
                result = document_processing._process_report_batch(
                    [source], "测试有限公司", {"rewrite": None, "authorization": None, "rectification": None, "disposal": None},
                    {"测试有限公司"}, [], work_dir=work_dir,
                )

            self.assertTrue(source.exists())
            self.assertTrue(result["failures"])
            state = document_processing._load_notice_state(work_dir, "测试有限公司", [])
            self.assertFalse(state["stages"]["rewrite"])
            self.assertFalse(state["complete"])

    def test_restart_resumes_after_rewrite_without_rewriting_or_renumbering(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            Document().save(source)
            rewrite, authorization, convert = self._successful_notice_side_effects(work_dir)

            with patch(
                "modules.Document_Processing.Report_Rewrite.rewrite_report.rewrite_report", side_effect=rewrite,
            ) as rewrite_mock, patch(
                "modules.Document_Processing.Report_Rewrite.edit_authorization.edit_authorization", return_value=False,
            ):
                first = document_processing._process_report_batch(
                    [source], "测试有限公司", {"rewrite": None, "authorization": None, "rectification": None, "disposal": None},
                    {"测试有限公司"}, [], work_dir=work_dir,
                )
            self.assertTrue(first["failures"])
            self.assertTrue(source.exists())
            self.assertEqual(rewrite_mock.call_count, 1)

            with patch(
                "modules.Document_Processing.Report_Rewrite.rewrite_report.rewrite_report",
                side_effect=AssertionError("重启续跑不应再次改写"),
            ) as rewrite_again, patch(
                "modules.Document_Processing.Report_Rewrite.edit_authorization.edit_authorization", side_effect=authorization,
            ), patch.object(document_processing, "_convert_generated_docs_to_pdf", side_effect=convert):
                second = document_processing._process_report_batch(
                    [source], "测试有限公司", {"rewrite": None, "authorization": None, "rectification": None, "disposal": None},
                    {"测试有限公司"}, [], work_dir=work_dir,
                )

            self.assertFalse(second["failures"])
            self.assertEqual(rewrite_again.call_count, 0)
            self.assertFalse(source.exists())
            state = document_processing._load_notice_state(work_dir, "测试有限公司", [])
            self.assertTrue(state["complete"])

    def test_completed_company_is_skipped_after_process_restart(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            rewritten = work_dir / "关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            self._write_rewritten_notice(rewritten)
            self._write_pdf(work_dir / "授权委托书（执法调查类）.pdf")
            state = document_processing._new_notice_state("测试有限公司")
            state["stages"] = {stage: True for stage in document_processing.NOTICE_PROCESS_STAGES}
            state["complete"] = True
            document_processing._save_notice_state(work_dir, state)

            with patch(
                "modules.Document_Processing.Report_Rewrite.edit_authorization.edit_authorization",
                side_effect=AssertionError("完成企业不应重新生成"),
            ) as authorization:
                result = document_processing._process_report_batch(
                    [], "测试有限公司", {"rewrite": None, "authorization": None, "rectification": None, "disposal": None},
                    {"测试有限公司"}, [], work_dir=work_dir,
                )

            self.assertTrue(result["skipped_complete"])
            self.assertEqual(authorization.call_count, 0)

    def test_legacy_rewritten_notice_without_marker_is_migrated(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            rewritten = work_dir / "关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            Document().save(source)
            self._write_rewritten_notice(rewritten, marker=False)

            discovered = document_processing._discover_report_files(work_dir, [])
            state = document_processing._new_notice_state("测试有限公司")
            artifact = document_processing._rewrite_artifact_for_source(work_dir, source, state)
            document_processing._mark_legacy_rewrite_artifact(rewritten)

            self.assertEqual(discovered, [source])
            self.assertEqual(artifact, rewritten.resolve())
            self.assertTrue(document_processing._is_rewritten_notice_file(rewritten))

    def test_changed_same_name_source_invalidates_saved_rewrite(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            source.write_bytes(b"old-source")
            rewritten = work_dir / "关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            self._write_rewritten_notice(rewritten)
            state = document_processing._new_notice_state("测试有限公司")
            document_processing._set_rewrite_state_item(work_dir, state, source, rewritten)
            state["input_signature"] = document_processing._source_signature([source])
            document_processing._save_notice_state(work_dir, state)
            source.write_bytes(b"new-corrected-source")

            reloaded = document_processing._load_notice_state(work_dir, "测试有限公司", [])
            artifact = document_processing._rewrite_artifact_for_source(work_dir, source, reloaded)

            self.assertIsNone(artifact)

    def test_restart_recovers_rewrite_written_before_process_was_closed(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            source.write_bytes(b"corrected-source")
            old_rewritten = work_dir / "关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            self._write_rewritten_notice(old_rewritten)
            old_rewritten.touch()

            state = document_processing._new_notice_state("测试有限公司")
            fingerprint = document_processing._file_fingerprint(source)
            state["input_signature"] = [fingerprint]
            state["rewrite_required"] = [fingerprint]
            document_processing._start_notice_stage(work_dir, state, "rewrite", source)
            self._write_rewritten_notice(old_rewritten)

            reloaded = document_processing._load_notice_state(work_dir, "测试有限公司", [])
            artifact = document_processing._rewrite_artifact_for_source(work_dir, source, reloaded)

            self.assertEqual(artifact, old_rewritten.resolve())

    def test_valid_authorization_written_before_error_is_not_regenerated(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            work_dir = Path(temp_dir)
            source = work_dir / "123关于测试有限公司所属系统存在未授权访问漏洞的通报.docx"
            Document().save(source)
            rewrite, _authorization, convert = self._successful_notice_side_effects(work_dir)

            def authorization_then_error(*_args, **_kwargs):
                self._write_authorization(work_dir / "授权委托书（执法调查类）.docx")
                raise RuntimeError("保存后的尾部错误")

            with patch(
                "modules.Document_Processing.Report_Rewrite.rewrite_report.rewrite_report", side_effect=rewrite,
            ), patch(
                "modules.Document_Processing.Report_Rewrite.edit_authorization.edit_authorization",
                side_effect=authorization_then_error,
            ), patch.object(document_processing, "_convert_generated_docs_to_pdf", side_effect=convert):
                result = document_processing._process_report_batch(
                    [source], "测试有限公司", {"rewrite": None, "authorization": None, "rectification": None, "disposal": None},
                    {"测试有限公司"}, [], work_dir=work_dir,
                )

            self.assertFalse(result["failures"])
            self.assertFalse(source.exists())
            state = document_processing._load_notice_state(work_dir, "测试有限公司", [])
            self.assertTrue(state["complete"])


if __name__ == "__main__":
    unittest.main()
