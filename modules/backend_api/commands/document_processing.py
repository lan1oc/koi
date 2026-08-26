#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import base64
import contextlib
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import traceback
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Sequence

from modules.Document_Processing.doc_pdf import (
    compute_output_path,
    convert_pdf_to_word,
    convert_with_word_com,
    list_document_files,
)
from modules.Document_Processing.pdf_extract import (
    build_compressed_output_path,
    build_default_output_path,
    compress_pdf,
    extract_pages,
    merge_pages_from_multiple_pdfs,
    parse_page_ranges,
)
from modules.Document_Processing.Report_Rewrite.notice_name_utils import (
    NOTICE_ISSUE_KEYWORDS,
    filename_has_notice_issue,
)

DOCUMENT_PROCESSING_COMMANDS = {
    "doc.convert.run",
    "doc.pdf_extract.preview",
    "doc.pdf_extract.run",
    "doc.pdf_extract.compress",
    "doc.notice.process",
    "doc.notice.process.start",
    "doc.notice.process.status",
    "doc.notice.counters.save",
    "doc.notice.classify",
    "doc.notice.convert_failed_pdf",
    "doc.open_path",
}

DEFAULT_TEMPLATE_SKIP_KEYWORDS = ["漏洞隐患处置文件模板", "app整改模板", "处置文件模板"]
SUPPORTED_ARCHIVE_SUFFIXES = {".zip", ".7z", ".rar"}
WORD_SUFFIXES = {".doc", ".docx"}
NOTICE_PDF_NAME_RULES = (
    lambda name: name.startswith("授权委托书"),
    lambda name: name.startswith("责令整改"),
    lambda name: name.startswith("关于"),
    lambda name: "通报" in name,
    lambda name: _notification_name(name),
)
NOTICE_BACKUP_MARKERS = (".clean_backup.", ".final_backup.", ".backup.", ".temp.")
NOTICE_VULNERABILITY_KEYWORDS = NOTICE_ISSUE_KEYWORDS
NOTICE_PROCESS_STATE_FILENAME = ".koi_notice_process_state.json"
NOTICE_PROCESS_STATE_VERSION = 1
NOTICE_PROCESS_STAGES = ("rewrite", "authorization", "rectification", "disposal", "pdf")


class _ProgressTee(io.TextIOBase):
    def __init__(self, original: Any, progress: "NoticeProgress"):
        self.original = original
        self.progress = progress
        self._buffer = ""

    @property
    def encoding(self) -> str:
        return getattr(self.original, "encoding", "utf-8")

    def writable(self) -> bool:
        return True

    def write(self, text: str) -> int:
        value = str(text)
        self._buffer += value
        while "\n" in self._buffer:
            line, self._buffer = self._buffer.split("\n", 1)
            if line.strip():
                self.progress.log(line.rstrip())
        if self.original is not None and self.original is not sys.stdout:
            try:
                self.original.write(value)
            except Exception:
                pass
        return len(value)

    def flush(self) -> None:
        if self._buffer.strip():
            self.progress.log(self._buffer.rstrip())
        self._buffer = ""
        if self.original is not None and self.original is not sys.stdout:
            try:
                self.original.flush()
            except Exception:
                pass


class NoticeProgress:
    def __init__(self, total: int = 0):
        self.lock = threading.RLock()
        self.logs: List[str] = []
        self.progress = 0
        self.message = "等待开始..."
        self.processed = 0
        self.total = max(0, int(total or 0))

    def snapshot(self) -> Dict[str, Any]:
        with self.lock:
            return {
                "progress": self.progress,
                "message": self.message,
                "logs": list(self.logs),
                "processed": self.processed,
                "total_reports": self.total,
            }

    def log(self, message: Any) -> None:
        text = str(message or "").strip("\r\n")
        if not text:
            return
        with self.lock:
            self.logs.append(text)
            self.message = text

    def extend(self, messages: Iterable[Any]) -> None:
        for message in messages:
            self.log(message)

    def set(self, progress: int | float | None = None, message: Any | None = None) -> None:
        with self.lock:
            if progress is not None:
                self.progress = max(0, min(100, int(round(float(progress)))))
            if message is not None:
                self.message = str(message)

    def set_total(self, total: int) -> None:
        with self.lock:
            self.total = max(0, int(total or 0))

    def set_processed(self, processed: int, total: int | None = None, message: Any | None = None) -> None:
        with self.lock:
            self.processed = max(0, int(processed or 0))
            if total is not None:
                self.total = max(0, int(total or 0))
            total_value = self.total
            if total_value > 0:
                self.progress = max(self.progress, min(98, 20 + int(self.processed / total_value * 75)))
            if message is not None:
                self.message = str(message)


class ProgressLogList(list):
    def __init__(self, progress: NoticeProgress | None = None):
        super().__init__()
        self.progress = progress

    def append(self, item: Any) -> None:
        super().append(str(item))
        if self.progress:
            self.progress.log(item)

    def extend(self, items: Iterable[Any]) -> None:
        for item in items:
            self.append(item)

    def extend_silent(self, items: Iterable[Any]) -> None:
        for item in items:
            super().append(str(item))


_NOTICE_TASKS: Dict[str, Dict[str, Any]] = {}
_NOTICE_TASK_LOCK = threading.RLock()
NOTICE_TASK_MAX_ACTIVE = 4
NOTICE_TASK_MAX_COMPLETED = 24
NOTICE_TASK_RETENTION_SECONDS = 6 * 60 * 60


def _prune_notice_tasks_locked(now: float | None = None) -> None:
    current_time = time.time() if now is None else float(now)
    completed: List[tuple[float, str]] = []

    for task_id, task in list(_NOTICE_TASKS.items()):
        if bool(task.get("running")) and not bool(task.get("done")):
            continue
        try:
            finished_at = float(task.get("finished_at") or task.get("created_at") or 0)
        except (TypeError, ValueError):
            finished_at = 0
        if finished_at <= 0 or current_time - finished_at > NOTICE_TASK_RETENTION_SECONDS:
            _NOTICE_TASKS.pop(task_id, None)
            continue
        completed.append((finished_at, task_id))

    overflow = len(completed) - NOTICE_TASK_MAX_COMPLETED
    if overflow > 0:
        for _, task_id in sorted(completed)[:overflow]:
            _NOTICE_TASKS.pop(task_id, None)


def _notice_target_key(path: Path | str) -> str:
    try:
        resolved = Path(path).expanduser().resolve()
    except Exception:
        resolved = Path(path).expanduser().absolute()
    return os.path.normcase(str(resolved))


def _active_notice_task_for_target_locked(target_path: Path | str, exclude_task_id: str = "") -> Dict[str, Any] | None:
    target_key = _notice_target_key(target_path)
    for task in _NOTICE_TASKS.values():
        if exclude_task_id and str(task.get("task_id") or "") == exclude_task_id:
            continue
        if not bool(task.get("running")) or bool(task.get("done")):
            continue
        if str(task.get("target_key") or "") == target_key:
            return task
    return None


def _notice_active_conflict(target_path: Path | str, operation: str, exclude_task_id: str = "") -> Dict[str, Any] | None:
    with _NOTICE_TASK_LOCK:
        _prune_notice_tasks_locked()
        task = _active_notice_task_for_target_locked(target_path, exclude_task_id)
        if not task:
            return None
        return {
            "success": False,
            "message": f"该目标仍有通报处理任务在运行，暂不能{operation}",
            "error_code": "notice_task_active",
            "task_id": str(task.get("task_id") or ""),
            "running": True,
            "done": False,
            "logs": [],
        }


def is_document_processing_command(command: str | None) -> bool:
    return str(command or "") in DOCUMENT_PROCESSING_COMMANDS


def handle_document_processing_command(command: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    if command == "doc.convert.run":
        return _doc_convert_run(payload)
    if command == "doc.pdf_extract.preview":
        return _doc_pdf_extract_preview(payload)
    if command == "doc.pdf_extract.run":
        return _doc_pdf_extract_run(payload)
    if command == "doc.pdf_extract.compress":
        return _doc_pdf_compress_run(payload)
    if command == "doc.notice.process":
        return _doc_notice_process(payload)
    if command == "doc.notice.process.start":
        return _doc_notice_process_start(payload)
    if command == "doc.notice.process.status":
        return _doc_notice_process_status(payload)
    if command == "doc.notice.counters.save":
        return _doc_notice_counters_save(payload)
    if command == "doc.notice.classify":
        return _doc_notice_classify(payload)
    if command == "doc.notice.convert_failed_pdf":
        return _doc_notice_convert_failed_pdf(payload)
    if command == "doc.open_path":
        return _doc_open_path(payload)

    raise ValueError(f"未知文档处理命令: {command}")


def _required_text(payload: Dict[str, Any], key: str, message: str) -> str:
    value = str(payload.get(key) or "").strip()
    if not value:
        raise ValueError(message)
    return value


def _optional_text(payload: Dict[str, Any], key: str) -> str | None:
    value = str(payload.get(key) or "").strip()
    return value or None


def _normalize_conversion_type(value: Any) -> str:
    text = str(value or "word_to_pdf").strip()
    if text in {"word_to_pdf", "Word转PDF", "word-pdf"}:
        return "word_to_pdf"
    if text in {"pdf_to_word", "PDF转Word", "pdf-word"}:
        return "pdf_to_word"
    raise ValueError(f"不支持的转换类型: {text}")


def _path_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    if not text:
        return []
    return [item.strip() for item in text.split(";") if item.strip()]


def _keyword_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in str(value).split(",") if item.strip()]


def _failure_dicts(failures: Iterable[tuple[Path, str]]) -> List[Dict[str, str]]:
    return [{"file": str(src), "name": src.name, "reason": str(reason)} for src, reason in failures]


def _captured_lines(buffer: io.StringIO) -> List[str]:
    return [line for line in buffer.getvalue().splitlines() if line.strip()]


def _store_captured_lines(logs: List[str], lines: Iterable[Any]) -> None:
    if isinstance(logs, ProgressLogList):
        logs.extend_silent(lines)
    else:
        logs.extend(str(line) for line in lines)


def _append_log(logs: List[str], message: Any, progress: NoticeProgress | None = None) -> None:
    text = str(message or "")
    logs.append(text)
    if progress:
        progress.log(text)


def _extend_logs(logs: List[str], messages: Iterable[Any], progress: NoticeProgress | None = None) -> None:
    for message in messages:
        _append_log(logs, message, progress)


def _run_with_progress_capture(func: Any, progress: NoticeProgress | None = None) -> tuple[Any, List[str]]:
    buffer = io.StringIO()
    stream: Any = _ProgressTee(buffer, progress) if progress else buffer
    with contextlib.redirect_stdout(stream):
        result = func()
    try:
        stream.flush()
    except Exception:
        pass
    return result, _captured_lines(buffer)


def _call_with_progress_capture(func: Any, progress: NoticeProgress | None = None) -> tuple[Any, List[str], BaseException | None]:
    buffer = io.StringIO()
    stream: Any = _ProgressTee(buffer, progress) if progress else buffer
    result = None
    error: BaseException | None = None
    with contextlib.redirect_stdout(stream):
        try:
            result = func()
        except BaseException as exc:
            error = exc
    try:
        stream.flush()
    except Exception:
        pass
    return result, _captured_lines(buffer), error


def _format_exception(error: BaseException | None) -> str:
    if error is None:
        return ""
    return "".join(traceback.format_exception(type(error), error, error.__traceback__))


def _parse_int(value: Any) -> int | None:
    try:
        text = str(value or "").strip()
        if not text:
            return None
        return int(text)
    except (TypeError, ValueError):
        return None


def _parse_number_ranges(value: Any) -> List[int]:
    if value is None:
        return []
    if isinstance(value, list):
        raw_parts = [str(item).strip() for item in value]
    else:
        raw_parts = re.split(r"[,，;；\s]+", str(value))

    numbers: set[int] = set()
    for part in raw_parts:
        if not part:
            continue
        if "-" in part:
            start_text, end_text = part.split("-", 1)
            start = _parse_int(start_text)
            end = _parse_int(end_text)
            if start is None or end is None:
                continue
            low, high = sorted((start, end))
            numbers.update(range(low, high + 1))
            continue
        number = _parse_int(part)
        if number is not None:
            numbers.add(number)
    return sorted(number for number in numbers if number > 0)


def _default_report_counters() -> Dict[str, Any]:
    return {
        "notification_number": 1,
        "rectification_number": 1,
        "unavailable_notification_numbers": [],
        "unavailable_rectification_numbers": [],
        "year": datetime.now().year,
        "last_updated": "",
    }


def _ensure_report_counters(config: Dict[str, Any]) -> Dict[str, Any]:
    counters = config.get("report_counters")
    if not isinstance(counters, dict):
        counters = {}
    normalized = _default_report_counters()
    normalized.update(counters)
    if not isinstance(normalized.get("unavailable_notification_numbers"), list):
        normalized["unavailable_notification_numbers"] = []
    if not isinstance(normalized.get("unavailable_rectification_numbers"), list):
        normalized["unavailable_rectification_numbers"] = []
    config["report_counters"] = normalized
    return normalized


def _merge_number_lists(existing: Any, incoming: Any) -> List[int]:
    return sorted(set(_parse_number_ranges(existing)) | set(_parse_number_ranges(incoming)))


def _clean_used_unavailable_numbers(counters: Dict[str, Any]) -> None:
    notification_next = _parse_int(counters.get("notification_number")) or 1
    rectification_next = _parse_int(counters.get("rectification_number")) or 1
    counters["unavailable_notification_numbers"] = [
        number for number in _parse_number_ranges(counters.get("unavailable_notification_numbers")) if number >= notification_next
    ]
    counters["unavailable_rectification_numbers"] = [
        number for number in _parse_number_ranges(counters.get("unavailable_rectification_numbers")) if number >= rectification_next
    ]


def _unavailable_counter_key(unavailable_type: Any) -> str:
    text = str(unavailable_type or "通报")
    if "责令" in text or "整改" in text:
        return "unavailable_rectification_numbers"
    return "unavailable_notification_numbers"


def _merge_unavailable_update(updates: Dict[str, Any], counters: Dict[str, Any], key: str, incoming: Any) -> None:
    numbers = _parse_number_ranges(incoming)
    if not numbers:
        return
    existing = updates.get(key, counters.get(key))
    updates[key] = _merge_number_lists(existing, numbers)


def _build_report_counter_updates(payload: Dict[str, Any], counters: Dict[str, Any]) -> Dict[str, Any]:
    updates: Dict[str, Any] = {}

    if "notice_number" in payload:
        notice_number = _parse_int(payload.get("notice_number"))
        if notice_number is not None and notice_number > 0:
            updates["notification_number"] = notice_number

    if "rectification_number" in payload:
        rectification_number = _parse_int(payload.get("rectification_number"))
        if rectification_number is not None and rectification_number > 0:
            updates["rectification_number"] = rectification_number

    if "unavailable_numbers" in payload:
        _merge_unavailable_update(
            updates,
            counters,
            _unavailable_counter_key(payload.get("unavailable_type")),
            payload.get("unavailable_numbers"),
        )

    if "unavailable_notification_numbers" in payload:
        _merge_unavailable_update(
            updates,
            counters,
            "unavailable_notification_numbers",
            payload.get("unavailable_notification_numbers"),
        )

    if "unavailable_rectification_numbers" in payload:
        _merge_unavailable_update(
            updates,
            counters,
            "unavailable_rectification_numbers",
            payload.get("unavailable_rectification_numbers"),
        )

    if updates:
        updates.setdefault("year", counters.get("year") or datetime.now().year)
        updates["last_updated"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    return updates


def _save_report_counter_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.config.config_manager import ConfigManager

    manager = ConfigManager()
    config = manager.load_config()
    counters = _ensure_report_counters(config)
    updates = _build_report_counter_updates(payload, counters)

    if not updates:
        return {
            "success": True,
            "updated": False,
            "message": "没有可保存的编号配置修改",
            "report_counters": counters,
        }

    if not manager.save_config({"report_counters": updates}):
        return {
            "success": False,
            "updated": False,
            "message": "编号配置保存失败",
            "report_counters": counters,
        }

    refreshed = manager.load_config()
    refreshed_counters = _ensure_report_counters(refreshed)
    return {
        "success": True,
        "updated": True,
        "message": "编号配置已保存到 report_counters",
        "report_counters": refreshed_counters,
    }


def _apply_report_counter_payload(payload: Dict[str, Any], logs: List[str]) -> None:
    counter_keys = {
        "notice_number",
        "rectification_number",
        "unavailable_numbers",
        "unavailable_notification_numbers",
        "unavailable_rectification_numbers",
    }
    if not any(key in payload for key in counter_keys):
        return

    try:
        result = _save_report_counter_payload(payload)
        logs.append(str(result.get("message") or "编号配置保存完成"))
    except Exception as exc:
        logs.append(f"编号配置保存失败: {exc}")


def _doc_notice_counters_save(payload: Dict[str, Any]) -> Dict[str, Any]:
    try:
        result = _save_report_counter_payload(payload)
        return {
            "success": bool(result.get("success")),
            "updated": bool(result.get("updated")),
            "message": str(result.get("message") or "编号配置保存完成"),
            "report_counters": result.get("report_counters") or {},
            "logs": [str(result.get("message") or "编号配置保存完成")],
        }
    except Exception as exc:
        return {
            "success": False,
            "updated": False,
            "message": f"编号配置保存失败: {exc}",
            "report_counters": {},
            "logs": [traceback.format_exc()],
        }


def _normalize_rewrite_result(raw_result: Any, report_file: Path) -> Dict[str, Any]:
    if isinstance(raw_result, dict):
        raw_result.setdefault("success", bool(raw_result.get("success")))
        raw_result.setdefault("output_file", None)
        raw_result.setdefault("backup_file", None)
        raw_result.setdefault("needs_manual_processing", False)
        raw_result.setdefault("skip_reason", None if raw_result.get("success") else "通报改写失败")
        return raw_result

    if raw_result is True:
        return {
            "success": True,
            "output_file": None,
            "backup_file": None,
            "needs_manual_processing": False,
            "skip_reason": None,
        }

    return {
        "success": False,
        "output_file": None,
        "backup_file": None,
        "needs_manual_processing": False,
        "skip_reason": "通报改写未返回详细结果，可能是模板缺少插入标记 * 或文档格式不符合要求",
    }


def _normalize_grouping_result(raw_result: Any) -> Dict[str, Any]:
    if isinstance(raw_result, dict):
        raw_result.setdefault("moved", 0)
        raw_result.setdefault("skipped_exist", 0)
        raw_result.setdefault("errors", 0)
        return raw_result
    if raw_result is True:
        return {"moved": 0, "skipped_exist": 0, "errors": 0, "log": ["分类命令已完成"]}
    return {"moved": 0, "skipped_exist": 0, "errors": 1, "log": ["分类命令未返回详细结果"]}


def _open_path_in_system(path: Path) -> tuple[bool, str | None]:
    try:
        if os.name == "nt":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
        return True, None
    except Exception as exc:
        return False, str(exc)


def _doc_convert_run(payload: Dict[str, Any]) -> Dict[str, Any]:
    conversion_type = _normalize_conversion_type(payload.get("conversion_type"))
    input_path = Path(_required_text(payload, "input_path", "请选择输入路径")).expanduser()
    output_dir = _optional_text(payload, "output_dir")
    output_root = Path(output_dir).expanduser() if output_dir else None
    recursive = bool(payload.get("recursive", True))
    overwrite = bool(payload.get("overwrite", True))
    skip_template = bool(payload.get("skip_template", True))
    skip_keywords = _keyword_list(payload.get("skip_keywords"))

    if not input_path.exists():
        return {"success": False, "message": f"输入路径不存在: {input_path}", "logs": []}
    if output_root is not None and output_root.exists() and not output_root.is_dir():
        return {"success": False, "message": f"输出路径不是目录: {output_root}", "logs": []}

    if conversion_type == "word_to_pdf":
        file_type = "word"
        expected_extensions = {".doc", ".docx"}
        file_label = "Word"
        if skip_template:
            skip_keywords = DEFAULT_TEMPLATE_SKIP_KEYWORDS + skip_keywords
    else:
        file_type = "pdf"
        expected_extensions = {".pdf"}
        file_label = "PDF"

    if input_path.is_file():
        if input_path.suffix.lower() not in expected_extensions:
            return {"success": False, "message": f"输入文件不是{file_label}文件: {input_path}", "logs": []}
        if skip_keywords and any(keyword in input_path.name for keyword in skip_keywords):
            return {"success": False, "message": f"输入文件命中跳过关键词: {input_path.name}", "logs": []}
        input_root = input_path.parent
        input_files = [input_path]
    else:
        input_root = input_path
        input_files = list_document_files(input_root, recursive=recursive, file_type=file_type, skip_keywords=skip_keywords)

    if not input_files:
        return {"success": False, "message": f"未找到可转换的{file_label}文件", "logs": []}

    file_map = [(src, compute_output_path(src, input_root, output_root, conversion_type)) for src in input_files]
    logs = [
        f"开始转换: {'Word转PDF' if conversion_type == 'word_to_pdf' else 'PDF转Word'}",
        f"找到 {len(file_map)} 个文件",
    ]
    if output_root:
        logs.append(f"输出目录: {output_root}")

    buffer = io.StringIO()
    try:
        with contextlib.redirect_stdout(buffer):
            if conversion_type == "word_to_pdf":
                converted, skipped, failures = convert_with_word_com(file_map, overwrite=overwrite)
            else:
                converted, skipped, failures = convert_pdf_to_word(file_map, overwrite=overwrite)
    except RuntimeError as exc:
        logs.extend(_captured_lines(buffer))
        return {
            "success": False,
            "message": str(exc),
            "converted": 0,
            "skipped": 0,
            "failures": [],
            "logs": logs,
            "output_files": [],
        }

    logs.extend(_captured_lines(buffer))
    for src, reason in failures:
        logs.append(f"失败: {src.name} -> {reason}")

    failure_paths = {src.resolve() for src, _ in failures}
    output_files = [
        str(dst)
        for src, dst in file_map
        if src.resolve() not in failure_paths and dst.exists()
    ]
    message = f"转换完成：成功 {converted}，跳过 {skipped}，失败 {len(failures)}"

    return {
        "success": len(failures) == 0,
        "message": message,
        "converted": converted,
        "skipped": skipped,
        "failures": _failure_dicts(failures),
        "logs": logs,
        "output_files": output_files,
        "total": len(file_map),
    }


def _render_pdf_page_thumbnails(pdf_file: Path, page_count: int, max_width: int, limit: int) -> Dict[int, str]:
    thumbnails: Dict[int, str] = {}
    if limit <= 0:
        return thumbnails
    try:
        import fitz  # type: ignore
    except Exception:
        return thumbnails

    try:
        with fitz.open(str(pdf_file)) as document:
            for page_index in range(min(page_count, limit, document.page_count)):
                page = document.load_page(page_index)
                rect = page.rect
                zoom = max_width / rect.width if rect.width else 0.2
                zoom = max(0.12, min(0.45, zoom))
                pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                encoded = base64.b64encode(pixmap.tobytes("png")).decode("ascii")
                thumbnails[page_index + 1] = f"data:image/png;base64,{encoded}"
    except Exception:
        return thumbnails
    return thumbnails


def _read_pdf_page_info(pdf_file: Path, include_thumbnails: bool = False, thumbnail_limit: int = 80) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError("未安装 pypdf，请先安装 pypdf 后再使用 PDF 页面提取") from exc

    reader = PdfReader(str(pdf_file))
    page_count = len(reader.pages)
    thumbnails = _render_pdf_page_thumbnails(pdf_file, page_count, 180, thumbnail_limit) if include_thumbnails else {}
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        width = None
        height = None
        try:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
        except Exception:
            pass
        pages.append({
            "page_number": index,
            "label": f"第 {index} 页 / 共 {page_count} 页",
            "width": width,
            "height": height,
            "thumbnail": thumbnails.get(index),
        })
    return {
        "path": str(pdf_file),
        "name": pdf_file.name,
        "page_count": page_count,
        "pages": pages,
    }


def _doc_pdf_extract_preview(payload: Dict[str, Any]) -> Dict[str, Any]:
    pdf_files = _path_list(payload.get("pdf_files") or payload.get("pdf_file"))
    include_thumbnails = bool(payload.get("include_thumbnails", False))
    thumbnail_limit = int(payload.get("thumbnail_limit") or 80)
    if not pdf_files:
        return {"success": False, "message": "请先选择PDF文件", "files": []}

    files = []
    failures = []
    for item in pdf_files:
        pdf_file = Path(item).expanduser()
        if not pdf_file.exists() or pdf_file.suffix.lower() != ".pdf":
            failures.append({"file": str(pdf_file), "reason": "文件不存在或不是PDF"})
            continue
        try:
            files.append(_read_pdf_page_info(pdf_file, include_thumbnails=include_thumbnails, thumbnail_limit=thumbnail_limit))
        except Exception as exc:
            failures.append({"file": str(pdf_file), "reason": str(exc)})

    total_pages = sum(int(file_info.get("page_count") or 0) for file_info in files)
    success = bool(files) and not failures
    message = f"预览加载完成，共 {len(files)} 个文件、{total_pages} 页" if files else "预览加载失败"
    if failures and files:
        message = f"预览部分完成，{len(failures)} 个文件加载失败"

    return {
        "success": success,
        "message": message,
        "files": files,
        "failures": failures,
        "total_pages": total_pages,
    }


def _normalize_page_selection(selection: Dict[str, Any], index: int) -> Dict[str, Any]:
    file_path = str(selection.get("file_path") or selection.get("path") or "").strip()
    page_num = int(selection.get("page_num") or selection.get("page_number") or 0)
    order = int(selection.get("order") or index + 1)
    if not file_path or page_num < 1:
        raise ValueError("页面选择数据不完整")
    return {"file_path": file_path, "page_num": page_num, "order": order}


def _doc_pdf_extract_run(payload: Dict[str, Any]) -> Dict[str, Any]:
    pdf_files = _path_list(payload.get("pdf_files") or payload.get("pdf_file"))
    page_ranges = str(payload.get("page_ranges") or "").strip()
    output_file = _optional_text(payload, "output_file")
    raw_selections = payload.get("page_selections") or []

    if not pdf_files:
        return {"success": False, "message": "请先选择PDF文件", "logs": []}
    if not isinstance(raw_selections, list):
        return {"success": False, "message": "页面选择数据格式不正确", "logs": []}

    page_selections = [_normalize_page_selection(selection, index) for index, selection in enumerate(raw_selections)]
    buffer = io.StringIO()

    if page_selections:
        first_input = Path(pdf_files[0]).expanduser()
        output_path = Path(output_file).expanduser() if output_file else first_input.parent / f"{first_input.stem}_merged_pages.pdf"
        logs = [f"开始合并 {len(page_selections)} 页"]
        try:
            with contextlib.redirect_stdout(buffer):
                merged_count, file_count = merge_pages_from_multiple_pdfs(page_selections, output_path)
        except RuntimeError as exc:
            logs.extend(_captured_lines(buffer))
            return {"success": False, "message": str(exc), "logs": logs, "output_file": str(output_path)}
        logs.extend(_captured_lines(buffer))
        message = f"已合并 {merged_count} 页（来自 {file_count} 个文件）"
        return {
            "success": True,
            "message": message,
            "output_file": str(output_path),
            "merged_count": merged_count,
            "file_count": file_count,
            "logs": logs,
        }

    if len(pdf_files) != 1:
        return {"success": False, "message": "多文件提取请先加载预览并选择页面", "logs": []}

    input_path = Path(pdf_files[0]).expanduser()
    if not input_path.exists() or input_path.suffix.lower() != ".pdf":
        return {"success": False, "message": f"PDF文件不存在或格式不正确: {input_path}", "logs": []}
    if not page_ranges:
        return {"success": False, "message": "请输入页码范围，或先加载预览选择页面", "logs": []}

    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError("未安装 pypdf，请先安装 pypdf 后再使用 PDF 页面提取") from exc

    total_pages = len(PdfReader(str(input_path)).pages)
    page_numbers = parse_page_ranges(page_ranges, total_pages)
    output_path = Path(output_file).expanduser() if output_file else build_default_output_path(input_path, page_ranges)
    logs = [f"开始提取 {input_path.name} 的 {len(page_numbers)} 页"]
    try:
        with contextlib.redirect_stdout(buffer):
            extracted, total = extract_pages(input_path, output_path, page_numbers)
    except RuntimeError as exc:
        logs.extend(_captured_lines(buffer))
        return {"success": False, "message": str(exc), "logs": logs, "output_file": str(output_path)}

    logs.extend(_captured_lines(buffer))
    return {
        "success": True,
        "message": f"已从 {input_path.name} 提取 {extracted}/{total} 页",
        "output_file": str(output_path),
        "extracted": extracted,
        "total_pages": total,
        "logs": logs,
    }


def _doc_pdf_compress_run(payload: Dict[str, Any]) -> Dict[str, Any]:
    pdf_files = _path_list(payload.get("pdf_files") or payload.get("pdf_file"))
    output_file = _optional_text(payload, "output_file")
    output_dir = _optional_text(payload, "output_dir")
    compression_mode = str(payload.get("compression_mode") or "standard").strip().lower()

    if not pdf_files:
        return {"success": False, "message": "请先选择PDF文件", "logs": [], "output_files": []}
    if compression_mode not in {"standard", "strong"}:
        return {"success": False, "message": f"不支持的压缩模式: {compression_mode}", "logs": [], "output_files": []}

    output_root = Path(output_dir).expanduser() if output_dir else None
    if output_root is not None and output_root.exists() and not output_root.is_dir():
        return {"success": False, "message": f"输出路径不是目录: {output_root}", "logs": [], "output_files": []}

    logs: List[str] = [f"开始压缩 {len(pdf_files)} 个PDF文件"]
    results: List[Dict[str, Any]] = []
    failures: List[Dict[str, str]] = []

    for index, item in enumerate(pdf_files, start=1):
        input_path = Path(item).expanduser()
        if not input_path.exists() or input_path.suffix.lower() != ".pdf":
            failures.append({"file": str(input_path), "reason": "文件不存在或不是PDF"})
            continue

        if output_file and len(pdf_files) == 1:
            output_path = Path(output_file).expanduser()
        else:
            output_path = build_compressed_output_path(input_path, output_root)

        logs.append(f"[{index}/{len(pdf_files)}] 压缩 {input_path.name}")
        try:
            result = compress_pdf(input_path, output_path, mode=compression_mode)
            results.append(result)
            saved_percent = result.get("saved_percent", 0)
            logs.append(
                f"完成: {result.get('original_size_text')} -> {result.get('compressed_size_text')} "
                f"(节省 {saved_percent}%)"
            )
            if float(saved_percent or 0) <= 0:
                logs.append("提示: 该文件已经比较紧凑，压缩后体积未明显降低")
        except Exception as exc:
            failures.append({"file": str(input_path), "reason": str(exc)})
            logs.append(f"失败: {input_path.name} -> {exc}")

    output_files = [str(result["output_file"]) for result in results if result.get("output_file")]
    success = bool(results) and not failures
    total_original = sum(int(result.get("original_size") or 0) for result in results)
    total_compressed = sum(int(result.get("compressed_size") or 0) for result in results)
    total_saved = total_original - total_compressed
    total_saved_percent = round(total_saved / total_original * 100, 2) if total_original else 0.0
    message = f"压缩完成：成功 {len(results)} 个，失败 {len(failures)} 个，整体节省 {total_saved_percent}%"
    if failures and results:
        message = f"压缩部分完成：成功 {len(results)} 个，失败 {len(failures)} 个"
    elif failures and not results:
        message = f"压缩失败：{len(failures)} 个文件未处理成功"

    return {
        "success": success,
        "message": message,
        "logs": logs,
        "output_file": output_files[0] if output_files else None,
        "output_files": output_files,
        "results": results,
        "failures": failures,
        "total_original_size": total_original,
        "total_compressed_size": total_compressed,
        "total_saved_bytes": total_saved,
        "total_saved_percent": total_saved_percent,
    }


def _project_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _template_dir() -> Path:
    try:
        from modules.utils.resource_path import get_report_template_dir

        return get_report_template_dir()
    except Exception:
        return _project_root() / "Report_Template"


def _find_template(keyword: str) -> Path | None:
    template_root = _template_dir()
    if not template_root.exists():
        return None
    candidates = sorted(
        item
        for item in template_root.iterdir()
        if item.is_file() and item.suffix.lower() in {".doc", ".docx"}
    )
    return next((item for item in candidates if keyword in item.name), None)


def _safe_chdir(target_dir: Path):
    class _Chdir:
        def __enter__(self):
            self.previous = Path.cwd()
            os.chdir(target_dir)
            return self

        def __exit__(self, exc_type, exc, tb):
            os.chdir(self.previous)

    return _Chdir()


def _notification_name(filename: str) -> bool:
    return filename_has_notice_issue(filename)


def _company_group_map(value: Any) -> Dict[str, str]:
    """Normalize the classification result into company -> township."""
    mapping: Dict[str, str] = {}
    if not isinstance(value, list):
        return mapping
    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf
    except Exception:
        gf = None
    for item in value:
        if not isinstance(item, (list, tuple)) or len(item) < 2:
            continue
        company, township = str(item[0] or '').strip(), str(item[1] or '').strip()
        if not company or not township:
            continue
        normalized = gf.normalize_company(company) if gf else company
        mapping[normalized or company] = township
    return mapping


def _doc_has_rewrite_marker(doc: Any) -> bool:
    try:
        comments = str(doc.core_properties.comments or '')
    except Exception:
        return False
    return 'koi.notice.rewritten.v1' in {item.strip() for item in comments.split(';')}


def _is_rewritten_notice_file(path: Path) -> bool:
    if not path.is_file() or path.suffix.lower() != ".docx" or path.name.startswith(("~$", ".")):
        return False
    try:
        from docx import Document

        return _doc_has_rewrite_marker(Document(str(path)))
    except Exception:
        return False


def _is_legacy_rewritten_notice_file(path: Path) -> bool:
    """Recognize strong template fingerprints from releases before metadata existed."""
    if not path.is_file() or path.suffix.lower() != ".docx" or path.name.startswith(("~$", ".")):
        return False
    text = _document_text(path)
    if not text:
        return False
    return (
        "鄞州区网络安全预警通报" in text
        and bool(re.search(r"〔\d{4}〕第\d+期", text))
        and "*" not in text
        and "验证情况" in text
        and "处置措施" in text
        and "抄送" in text
    )


def _is_any_rewritten_notice_file(path: Path) -> bool:
    return _is_rewritten_notice_file(path) or _is_legacy_rewritten_notice_file(path)


def _rewritten_notice_files(directory: Path) -> List[Path]:
    return sorted(
        (path.resolve() for path in directory.glob("*.docx") if _is_any_rewritten_notice_file(path)),
        key=lambda path: (".backup." in path.name, str(path).lower()),
    )


def _notice_state_path(directory: Path) -> Path:
    return directory / NOTICE_PROCESS_STATE_FILENAME


def _new_notice_state(company_name: str) -> Dict[str, Any]:
    return {
        "version": NOTICE_PROCESS_STATE_VERSION,
        "company_name": company_name,
        "stages": {stage: False for stage in NOTICE_PROCESS_STAGES},
        "rewrite_items": [],
        "rewrite_required": [],
        "input_signature": [],
        "complete": False,
        "updated_at": None,
        "_loaded_from_disk": False,
    }


def _load_notice_state(directory: Path, company_name: str, logs: List[str]) -> Dict[str, Any]:
    state_path = _notice_state_path(directory)
    state = _new_notice_state(company_name)
    if not state_path.exists():
        return state
    try:
        loaded = json.loads(state_path.read_text(encoding="utf-8"))
        if not isinstance(loaded, dict):
            raise ValueError("状态内容不是对象")
        state.update(loaded)
        state["version"] = NOTICE_PROCESS_STATE_VERSION
        state["company_name"] = company_name or str(state.get("company_name") or "")
        loaded_stages = loaded.get("stages") if isinstance(loaded.get("stages"), dict) else {}
        state["stages"] = {stage: bool(loaded_stages.get(stage)) for stage in NOTICE_PROCESS_STAGES}
        state["rewrite_items"] = loaded.get("rewrite_items") if isinstance(loaded.get("rewrite_items"), list) else []
        state["rewrite_required"] = loaded.get("rewrite_required") if isinstance(loaded.get("rewrite_required"), list) else []
        state["input_signature"] = loaded.get("input_signature") if isinstance(loaded.get("input_signature"), list) else []
        state["_loaded_from_disk"] = True
        return state
    except Exception as exc:
        logs.append(f"处理状态文件无效，将根据现有产物重新判断: {state_path} -> {exc}")
        return state


def _save_notice_state(directory: Path, state: Dict[str, Any]) -> None:
    state_path = _notice_state_path(directory)
    temp_path = state_path.with_name(f"{state_path.name}.{uuid.uuid4().hex}.tmp")
    payload = {key: value for key, value in state.items() if not str(key).startswith("_")}
    payload["version"] = NOTICE_PROCESS_STATE_VERSION
    payload["updated_at"] = datetime.now().astimezone().isoformat(timespec="seconds")
    try:
        temp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        os.replace(temp_path, state_path)
        state.update(payload)
    finally:
        if temp_path.exists():
            try:
                temp_path.unlink()
            except OSError:
                pass


def _file_fingerprint(path: Path) -> Dict[str, Any]:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return {"name": path.name, "size": path.stat().st_size, "sha256": digest.hexdigest()}


def _source_signature(report_files: Sequence[Path]) -> List[Dict[str, Any]]:
    signature: List[Dict[str, Any]] = []
    for path in sorted(report_files, key=lambda item: item.name.lower()):
        try:
            signature.append(_file_fingerprint(path))
        except OSError:
            signature.append({"name": path.name, "size": None, "sha256": None})
    return signature


def _state_relative_path(directory: Path, path: Path | None) -> str | None:
    if path is None:
        return None
    try:
        return str(path.resolve().relative_to(directory.resolve()))
    except ValueError:
        return str(path.resolve())


def _state_artifact_path(directory: Path, value: Any) -> Path | None:
    text = str(value or "").strip()
    if not text:
        return None
    path = Path(text)
    return (directory / path).resolve() if not path.is_absolute() else path.resolve()


def _rewrite_artifact_for_source(directory: Path, source: Path, state: Dict[str, Any]) -> Path | None:
    try:
        fingerprint = _file_fingerprint(source)
    except OSError:
        fingerprint = {"name": source.name, "size": None, "sha256": None}
    matching_entries = [
        item for item in state.get("rewrite_items", [])
        if isinstance(item, dict) and str(item.get("source", {}).get("name") or "") == source.name
    ]
    for item in matching_entries:
        if item.get("source") != fingerprint:
            continue
        artifact = _state_artifact_path(directory, item.get("artifact"))
        if artifact and _is_any_rewritten_notice_file(artifact):
            return artifact
    requires_rewrite = fingerprint in state.get("rewrite_required", [])
    if requires_rewrite and not (
        state.get("active_stage") == "rewrite"
        and state.get("active_source") == fingerprint
    ):
        return None
    if matching_entries:
        return None

    expected_name = re.sub(r"^\d+", "", source.name) or source.name
    expected = directory / expected_name
    candidates = [expected, expected.with_suffix(".backup.docx")]
    candidates.extend(path for path in _rewritten_notice_files(directory) if path.name.startswith(expected.stem))
    for artifact in dict.fromkeys(candidates):
        if not _is_any_rewritten_notice_file(artifact):
            continue
        try:
            minimum_mtime = source.stat().st_mtime
            if requires_rewrite:
                minimum_mtime = max(minimum_mtime, float(state.get("stage_started_at") or 0))
            if artifact.stat().st_mtime + 1 < minimum_mtime:
                continue
        except OSError:
            continue
        return artifact.resolve()
    return None


def _set_rewrite_state_item(directory: Path, state: Dict[str, Any], source: Path, artifact: Path) -> None:
    fingerprint = _file_fingerprint(source)
    items = [
        item for item in state.get("rewrite_items", [])
        if not (isinstance(item, dict) and str(item.get("source", {}).get("name") or "") == source.name)
    ]
    items.append({"source": fingerprint, "artifact": _state_relative_path(directory, artifact)})
    state["rewrite_items"] = items
    state["rewrite_required"] = [item for item in state.get("rewrite_required", []) if item != fingerprint]


def _document_text(path: Path) -> str | None:
    try:
        from docx import Document

        doc = Document(str(path))
        parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception:
        return None


def _valid_pdf(path: Path) -> bool:
    if not path.is_file() or path.suffix.lower() != ".pdf":
        return False
    try:
        with path.open("rb") as source:
            header = source.read(5)
        if path.stat().st_size < 8 or header != b"%PDF-":
            return False
        from pypdf import PdfReader

        return len(PdfReader(str(path)).pages) > 0
    except Exception:
        return False


def _pdf_matches_text(path: Path, required_terms: Sequence[str], *, allow_unextractable: bool = False) -> bool:
    if not _valid_pdf(path):
        return False
    if allow_unextractable:
        return True
    try:
        from pypdf import PdfReader

        text = "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages).strip()
    except Exception:
        return False
    compact_text = re.sub(r"\s+", "", text)
    if not compact_text:
        return False
    return all(re.sub(r"\s+", "", term) in compact_text for term in required_terms if term)


def _authorization_artifacts(directory: Path, company_name: str, *, trust_pdf_structure: bool = False) -> List[Path]:
    artifacts: List[Path] = []
    for path in sorted(directory.glob("授权委托书*")):
        if path.name.startswith("~$") or any(marker in path.name for marker in NOTICE_BACKUP_MARKERS):
            continue
        if path.suffix.lower() == ".pdf" and _pdf_matches_text(
            path, ("授权委托书", company_name), allow_unextractable=trust_pdf_structure
        ):
            artifacts.append(path.resolve())
        elif path.suffix.lower() == ".docx":
            text = _document_text(path)
            if text and "*" not in text and company_name in text:
                artifacts.append(path.resolve())
    return artifacts


def _rectification_artifacts(
    directory: Path,
    company_name: str,
    vuln_type: str | None,
    *,
    trust_pdf_structure: bool = False,
) -> List[Path]:
    artifacts: List[Path] = []
    for path in sorted(directory.glob("责令整改*")):
        if path.name.startswith("~$") or any(marker in path.name for marker in NOTICE_BACKUP_MARKERS):
            continue
        if path.suffix.lower() == ".pdf" and _pdf_matches_text(
            path, ("责令整改", company_name), allow_unextractable=trust_pdf_structure
        ):
            artifacts.append(path.resolve())
        elif path.suffix.lower() == ".docx" and _rectification_manual_reason(path, company_name, vuln_type) is None:
            artifacts.append(path.resolve())
    return artifacts


def _disposal_artifacts(directory: Path) -> List[Path]:
    artifacts: List[Path] = []
    for path in sorted(directory.glob("*处置*")):
        if path.name.startswith("~$"):
            continue
        if path.suffix.lower() == ".pdf" and _valid_pdf(path):
            artifacts.append(path.resolve())
            continue
        if path.suffix.lower() != ".docx":
            continue
        text = _document_text(path)
        if text and "鄞州区网信办：" in text:
            artifacts.append(path.resolve())
    return artifacts


def _artifacts_created_after(artifacts: Sequence[Path], started_at: Any) -> bool:
    try:
        started = float(started_at)
    except (TypeError, ValueError):
        return False
    for path in artifacts:
        try:
            if path.stat().st_mtime >= started - 1:
                return True
        except OSError:
            continue
    return False


def _stage_verified(
    state: Dict[str, Any],
    stage: str,
    artifacts: Sequence[Path],
    *,
    optional: bool = False,
) -> bool:
    valid = optional or bool(artifacts)
    if not valid:
        return False
    if bool((state.get("stages") or {}).get(stage)):
        return True
    if not state.get("_loaded_from_disk"):
        return True
    return (
        state.get("active_stage") == stage
        and _artifacts_created_after(artifacts, state.get("stage_started_at"))
    )


def _start_notice_stage(directory: Path, state: Dict[str, Any], stage: str, source: Path | None = None) -> None:
    state["complete"] = False
    state["active_stage"] = stage
    state["stage_started_at"] = time.time()
    state["active_source"] = _file_fingerprint(source) if source is not None else None
    state.setdefault("stages", {})[stage] = False
    _save_notice_state(directory, state)


def _finish_notice_stage(directory: Path, state: Dict[str, Any], stage: str, success: bool) -> None:
    state.setdefault("stages", {})[stage] = bool(success)
    if state.get("active_stage") == stage:
        state["active_stage"] = None
        state["stage_started_at"] = None
        state["active_source"] = None
        state["active_source"] = None
    _save_notice_state(directory, state)


def _mark_legacy_rewrite_artifact(path: Path) -> None:
    try:
        from docx import Document

        doc = Document(str(path))
        if _doc_has_rewrite_marker(doc):
            return
        comments = str(doc.core_properties.comments or "").strip()
        markers = [item.strip() for item in comments.split(";") if item.strip()]
        markers.append("koi.notice.rewritten.v1")
        doc.core_properties.comments = ";".join(dict.fromkeys(markers))
        doc.save(str(path))
    except Exception:
        return


def _generated_word_files(directory: Path) -> List[Path]:
    files: List[Path] = []
    for pattern in ("授权委托书*.docx", "责令整改*.docx"):
        for path in directory.glob(pattern):
            if path.name.startswith("~$") or any(marker in path.name for marker in NOTICE_BACKUP_MARKERS):
                continue
            files.append(path.resolve())
    return sorted(set(files), key=lambda path: str(path).lower())


def _cleanup_words_with_valid_pdf(
    directory: Path,
    company_name: str,
    logs: List[str],
    *,
    minimum_pdf_mtime: float | None = None,
) -> List[Dict[str, str]]:
    failures: List[Dict[str, str]] = []
    for word_path in _generated_word_files(directory):
        pdf_path = word_path.with_suffix(".pdf")
        required_terms = ("授权委托书", company_name) if word_path.name.startswith("授权委托书") else ("责令整改", company_name)
        if not _pdf_matches_text(pdf_path, required_terms, allow_unextractable=True):
            continue
        try:
            if minimum_pdf_mtime is not None and pdf_path.stat().st_mtime + 1 < minimum_pdf_mtime:
                continue
        except OSError:
            continue
        try:
            word_path.unlink()
            logs.append(f"检测到已有有效PDF，已补删原Word文件: {word_path.name}")
        except OSError as exc:
            failures.append({"file": str(word_path), "reason": f"PDF已存在，但删除原Word失败: {exc}"})
    return failures


def _can_recover_pdf_cleanup(state: Dict[str, Any]) -> bool:
    return bool((state.get("stages") or {}).get("pdf")) or state.get("active_stage") == "pdf"


def _pdf_stage_artifacts(
    directory: Path,
    company_name: str,
    is_soe: bool,
    *,
    trust_pdf_structure: bool = False,
    minimum_mtime: float | None = None,
) -> List[Path]:
    auth_pdfs = [
        path for path in directory.glob("授权委托书*.pdf")
        if _pdf_matches_text(path, ("授权委托书", company_name), allow_unextractable=trust_pdf_structure)
    ]
    rect_pdfs = [
        path for path in directory.glob("责令整改*.pdf")
        if _pdf_matches_text(path, ("责令整改", company_name), allow_unextractable=trust_pdf_structure)
    ]
    if minimum_mtime is not None:
        auth_pdfs = [path for path in auth_pdfs if path.stat().st_mtime + 1 >= minimum_mtime]
        rect_pdfs = [path for path in rect_pdfs if path.stat().st_mtime + 1 >= minimum_mtime]
    if not auth_pdfs or (not is_soe and not rect_pdfs):
        return []
    if _generated_word_files(directory):
        return []
    return [*(path.resolve() for path in auth_pdfs), *(path.resolve() for path in rect_pdfs)]


def _delete_completed_notice_sources(report_files: Sequence[Path], logs: List[str]) -> List[Dict[str, str]]:
    failures: List[Dict[str, str]] = []
    for source in report_files:
        if not source.name[:1].isdigit() or not source.exists():
            continue
        last_error: OSError | None = None
        for _ in range(5):
            try:
                source.unlink()
                logs.append(f"企业流程已全部完成，删除原始通报: {source.name}")
                last_error = None
                break
            except OSError as exc:
                last_error = exc
                time.sleep(0.2)
        if last_error is not None:
            failures.append({"file": str(source), "reason": f"全部步骤已完成，但删除原始通报失败: {last_error}"})
    return failures


def _replace_blank_notice_copy_to(doc: Any, township: str) -> bool:
    """Fill only an empty ``抄送：`` in a document produced by rewrite."""
    if not township:
        return False
    changed = False
    for paragraph in doc.paragraphs:
        text = (paragraph.text or '').strip()
        if not re.match(r'^抄送\s*[:：]\s*$', text):
            continue
        if paragraph.runs:
            paragraph.runs[0].text = f'抄送：{township}'
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            paragraph.add_run(f'抄送：{township}')
        changed = True
    return changed


def _fill_rewritten_notice_copy_to(target_path: Path, company_groups: Any, logs: List[str]) -> Dict[str, int]:
    """Backfill street copy-to only for documents previously rewritten."""
    company_map = _company_group_map(company_groups)
    stats = {'updated': 0, 'skipped_no_group': 0, 'errors': 0}
    if not company_map:
        return stats
    try:
        from docx import Document
        from modules.Document_Processing.Report_Rewrite import group_folders as gf
    except Exception as exc:
        logs.append(f'抄送补写模块不可用: {exc}')
        stats['errors'] += 1
        return stats

    for doc_path in target_path.rglob('*.docx'):
        if doc_path.name.startswith(('~$', '.')):
            continue
        try:
            doc = Document(str(doc_path))
            if not _doc_has_rewrite_marker(doc):
                continue
            company_name = gf.normalize_company(doc_path.name) or gf.normalize_company(doc_path.parent.name)
            township = company_map.get(company_name or '')
            if not township:
                stats['skipped_no_group'] += 1
                continue
            if _replace_blank_notice_copy_to(doc, township):
                doc.save(str(doc_path))
                stats['updated'] += 1
                logs.append(f'已补写抄送：{doc_path.name} -> {township}')
        except Exception as exc:
            stats['errors'] += 1
            logs.append(f'补写抄送失败 {doc_path.name}: {exc}')
    return stats


def _count_notification_docs(directory: Path) -> int:
    source_count = 0
    source_dirs: set[Path] = set()
    resumable_dirs: set[Path] = set()
    for docx_file in directory.rglob("*.docx"):
        if docx_file.name.startswith(("~$", ".")) or not _notification_name(docx_file.name):
            continue
        if _is_any_rewritten_notice_file(docx_file):
            resumable_dirs.add(docx_file.parent.resolve())
            continue
        source_count += 1
        source_dirs.add(docx_file.parent.resolve())
    for state_file in directory.rglob(NOTICE_PROCESS_STATE_FILENAME):
        if state_file.parent.resolve() not in source_dirs:
            resumable_dirs.add(state_file.parent.resolve())
    resumable_dirs.difference_update(source_dirs)
    return source_count + len(resumable_dirs)


def _is_supported_archive(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() in SUPPORTED_ARCHIVE_SUFFIXES


def _iter_supported_archives(directory: Path) -> List[Path]:
    return sorted(
        (item for item in directory.rglob("*") if _is_supported_archive(item)),
        key=lambda item: (len(item.relative_to(directory).parts), str(item).lower()),
    )


def _unique_child_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem} ({index}){suffix}")
        if not candidate.exists():
            return candidate
    return path.with_name(f"{stem}_{int(time.time())}{suffix}")


def _safe_archive_member_target(extract_dir: Path, member_name: str) -> Path:
    target = (extract_dir / member_name).resolve()
    root = extract_dir.resolve()
    if target == root or root not in target.parents:
        raise RuntimeError(f"压缩包包含不安全路径: {member_name}")
    return target


def _extract_zip_archive(archive_path: Path, extract_dir: Path) -> None:
    with zipfile.ZipFile(archive_path, "r") as zip_ref:
        for info in zip_ref.infolist():
            target = _safe_archive_member_target(extract_dir, info.filename)
            if info.is_dir():
                target.mkdir(parents=True, exist_ok=True)
                continue
            target = _unique_child_path(target)
            target.parent.mkdir(parents=True, exist_ok=True)
            with zip_ref.open(info) as src, open(target, "wb") as dst:
                shutil.copyfileobj(src, dst)


def _extract_7z_archive(archive_path: Path, extract_dir: Path) -> None:
    try:
        import py7zr  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少 py7zr 依赖，无法解压 7z。请重新安装依赖或重新打包应用。") from exc

    with py7zr.SevenZipFile(archive_path, mode="r") as archive:
        archive.extractall(path=extract_dir)


def _extract_rar_archive(archive_path: Path, extract_dir: Path) -> None:
    try:
        import rarfile  # type: ignore
    except ImportError as exc:
        raise RuntimeError("缺少 rarfile 依赖，无法解压 rar。请重新安装依赖或重新打包应用。") from exc

    try:
        with rarfile.RarFile(archive_path, "r") as archive:
            archive.extractall(extract_dir)
    except rarfile.RarCannotExec as exc:  # type: ignore[attr-defined]
        raise RuntimeError("当前应用没有可用的 RAR 解码器，无法直接解压 rar；请先转成 zip/7z 或把 rar 解码器随应用打包。") from exc


def _merge_extracted_tree(source_dir: Path, target_dir: Path) -> None:
    source_root = source_dir.resolve()
    target_root = target_dir.resolve()
    for child in sorted(source_dir.rglob("*"), key=lambda item: (len(item.relative_to(source_dir).parts), str(item).lower())):
        child_resolved = child.resolve()
        if child_resolved == source_root or source_root not in child_resolved.parents:
            raise RuntimeError(f"解压临时目录包含不安全路径: {child}")
        relative_path = child.relative_to(source_dir)
        target = (target_dir / relative_path).resolve()
        if target == target_root or target_root not in target.parents:
            raise RuntimeError(f"解压目标包含不安全路径: {relative_path}")
        if child.is_dir():
            target.mkdir(parents=True, exist_ok=True)
            continue
        target = _unique_child_path(target)
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(child), str(target))


def _safe_folder_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(value or "")).strip(" .")
    return cleaned or "未识别企业"


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(2, 1000):
        candidate = path.with_name(f"{path.name} ({index})")
        if not candidate.exists():
            return candidate
    return path.with_name(f"{path.name}_{int(time.time())}")


def _normalize_company_from_text(value: Any) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf

        company_name = gf.normalize_company(text)
    except Exception:
        company_name = None
    return company_name.strip() if company_name else None


def _detect_company_name_for_extract(archive_path: Path, extract_dir: Path) -> tuple[str | None, List[str]]:
    candidates: List[str] = []

    def add_candidate(value: Any) -> None:
        company_name = _normalize_company_from_text(value)
        if company_name and company_name not in candidates:
            candidates.append(company_name)

    add_candidate(archive_path.stem)
    for child in sorted(extract_dir.iterdir(), key=lambda item: item.name.lower()):
        if child.name.startswith(("~$", ".")):
            continue
        add_candidate(child.name)

    for doc_file in sorted(extract_dir.rglob("*.docx"), key=lambda item: str(item).lower()):
        if doc_file.name.startswith(("~$", ".")):
            continue
        if any(keyword in doc_file.name for keyword in ["模板", "授权委托书", "责令整改", "处置"]):
            continue
        add_candidate(doc_file.name)

    if len(candidates) == 1:
        return candidates[0], candidates
    if candidates:
        archive_company = _normalize_company_from_text(archive_path.stem)
        if archive_company in candidates:
            return archive_company, candidates
    return None, candidates


def _rename_extracted_dir_by_company(archive_path: Path, extract_dir: Path, logs: List[str]) -> Path:
    company_name, candidates = _detect_company_name_for_extract(archive_path, extract_dir)
    if not company_name:
        if candidates:
            logs.append(f"解压目录企业名不唯一，保留原目录名: {', '.join(candidates)}")
        else:
            logs.append("未识别到企业名，保留原解压目录名")
        return extract_dir

    target_dir = extract_dir.parent / _safe_folder_name(company_name)
    if target_dir.resolve() == extract_dir.resolve():
        logs.append(f"解压目录已按企业名命名: {target_dir.name}")
        return extract_dir

    target_dir = _unique_path(target_dir)
    try:
        shutil.move(str(extract_dir), str(target_dir))
        logs.append(f"解压目录已按企业名命名: {target_dir}")
        return target_dir
    except Exception as exc:
        logs.append(f"按企业名重命名解压目录失败，保留原目录: {exc}")
        return extract_dir


def _delete_archive_file(archive_path: Path, logs: List[str]) -> None:
    try:
        if archive_path.exists():
            archive_path.unlink()
            logs.append(f"已删除压缩包: {archive_path.name}")
    except Exception as exc:
        logs.append(f"删除压缩包失败 {archive_path.name}: {exc}")


def _extract_archive(archive_path: Path, logs: List[str]) -> tuple[Path | None, bool]:
    suffix = archive_path.suffix.lower()
    if suffix not in SUPPORTED_ARCHIVE_SUFFIXES:
        logs.append(f"暂不支持 {archive_path.suffix} 格式，请先手动解压")
        return None, False

    extract_dir = archive_path.parent
    logs.append(f"解压到压缩包所在目录: {extract_dir}")
    if suffix == ".zip":
        _extract_zip_archive(archive_path, extract_dir)
    elif suffix in {".7z", ".rar"}:
        temp_dir = archive_path.parent / f".koi_extract_{archive_path.stem}_{uuid.uuid4().hex[:8]}"
        temp_dir.mkdir(parents=True, exist_ok=False)
        try:
            if suffix == ".7z":
                _extract_7z_archive(archive_path, temp_dir)
            else:
                _extract_rar_archive(archive_path, temp_dir)
            _merge_extracted_tree(temp_dir, extract_dir)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
    logs.append(f"解压完成，已保留原压缩包: {archive_path.name}")
    return extract_dir, True


def _extract_archives_before_notice_processing(root_dir: Path, logs: List[str], progress: NoticeProgress | None = None) -> List[Dict[str, str]]:
    failures: List[Dict[str, str]] = []
    processed: set[Path] = set()

    while True:
        archives = [item for item in _iter_supported_archives(root_dir) if item.resolve() not in processed]
        if not archives:
            return failures

        logs.append(f"预解压压缩包: 发现 {len(archives)} 个")
        if progress:
            progress.set(message=f"预解压压缩包: 发现 {len(archives)} 个")

        for archive_path in archives:
            processed.add(archive_path.resolve())
            logs.append(f"预解压: {archive_path}")
            try:
                extract_dir, _ = _extract_archive(archive_path, logs)
                if extract_dir and progress:
                    progress.set(message=f"已解压: {archive_path.name}")
            except Exception as exc:
                failures.append({"file": str(archive_path), "reason": str(exc)})
                logs.append(f"解压失败 {archive_path.name}: {exc}")


def _latest_rectification_doc(directory: Path) -> Path | None:
    candidates = [
        path
        for suffix in WORD_SUFFIXES
        for path in directory.glob(f"责令整改*{suffix}")
    ]
    if not candidates:
        return None
    candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return candidates[0]


def _resolve_work_path(work_dir: Path, value: Any) -> Path | None:
    text = str(value or "").strip()
    if not text:
        return None
    path = Path(text).expanduser()
    if not path.is_absolute():
        path = work_dir / path
    return path.resolve()


def _manual_path_values(files: Iterable[Dict[str, Any]]) -> set[Path]:
    paths: set[Path] = set()
    for item in files:
        if not isinstance(item, dict):
            continue
        for key in ("file", "output_file", "backup_file"):
            path = _resolve_work_path(Path.cwd(), item.get(key))
            if path:
                paths.add(path)
    return paths


def _copy_template_for_manual_edit(template_file: Path | None, work_dir: Path, output_name: str, logs: List[str], reason: str) -> Path | None:
    if template_file is None:
        logs.append(f"{reason}，且未找到可保留的模板")
        return None

    source = Path(template_file)
    if not source.exists():
        logs.append(f"{reason}，模板文件不存在: {source}")
        return None

    suffix = source.suffix if source.suffix.lower() in WORD_SUFFIXES else ".docx"
    target_name = Path(output_name).with_suffix(suffix).name
    target = work_dir / target_name
    if target.exists():
        logs.append(f"{reason}，已保留现有待编辑文件: {target.name}")
        return target.resolve()

    try:
        shutil.copy2(source, target)
        logs.append(f"{reason}，已保留模板供手动编辑: {target}")
        return target.resolve()
    except Exception as exc:
        logs.append(f"保留模板失败 {source.name}: {exc}")
        return None


def _rectification_manual_reason(doc_path: Path, expected_company: str | None = None, expected_vuln: str | None = None) -> str | None:
    try:
        from docx import Document

        doc = Document(str(doc_path))
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as exc:
        return f"无法检查责令整改文档内容: {exc}"

    issues: List[str] = []
    placeholder_markers = ("【公司名】", "【漏洞类型】", "{公司名}", "{漏洞类型}", "公司名】", "漏洞类型】")
    if any(marker in full_text for marker in placeholder_markers):
        issues.append("仍包含公司名或漏洞类型占位符")

    if expected_company and expected_company not in full_text:
        issues.append("未写入企业名")

    if expected_vuln:
        vuln_needles = [expected_vuln]
        if expected_vuln.startswith("存在"):
            vuln_needles.append(expected_vuln[2:])
        else:
            vuln_needles.append(f"存在{expected_vuln}")
        if not any(needle and needle in full_text for needle in vuln_needles):
            issues.append("未写入漏洞类型")

    return "；".join(dict.fromkeys(issues)) if issues else None


def _is_notice_pdf_candidate(path: Path) -> bool:
    if path.name.startswith(("~$", ".")):
        return False
    if path.suffix.lower() not in WORD_SUFFIXES:
        return False
    if any(part == "Report_Template" for part in path.parts):
        return False
    if any(keyword in path.name for keyword in DEFAULT_TEMPLATE_SKIP_KEYWORDS) or "模板" in path.name:
        return False
    if path.name[0].isdigit():
        return False
    if any(marker in path.name for marker in NOTICE_BACKUP_MARKERS):
        return False
    return any(rule(path.name) for rule in NOTICE_PDF_NAME_RULES)


def _is_explicit_manual_pdf_candidate(path: Path) -> bool:
    if path.name.startswith(("~$", ".")):
        return False
    if path.suffix.lower() not in WORD_SUFFIXES:
        return False
    if any(part == "Report_Template" for part in path.parts):
        return False
    if any(keyword in path.name for keyword in DEFAULT_TEMPLATE_SKIP_KEYWORDS) or "模板" in path.name:
        return False
    return True


def _collect_notice_pdf_candidates(
    target_path: Path,
    failed_files: Any,
    logs: List[str],
    *,
    scan_target: bool = False,
) -> List[Path]:
    candidates: set[Path] = set()
    has_explicit_failures = isinstance(failed_files, list)

    if isinstance(failed_files, list):
        for item in failed_files:
            if isinstance(item, dict):
                values: List[Any] = [item.get("output_file") or item.get("file")]
            else:
                values = [item]
            for value in values:
                path = _resolve_work_path(target_path, value)
                if path and path.exists() and _is_explicit_manual_pdf_candidate(path):
                    candidates.add(path)
                    continue
                if path is None or not path.name:
                    continue

                # Classification moves root/company/file to
                # root/township/company/file.  Re-locate a manual item by its
                # original company directory and filename after that move.
                matches = [
                    candidate.resolve()
                    for candidate in target_path.rglob(path.name)
                    if candidate.is_file() and _is_explicit_manual_pdf_candidate(candidate)
                ]
                company_matches = [candidate for candidate in matches if candidate.parent.name == path.parent.name]
                relocated = company_matches if company_matches else matches
                if len(relocated) == 1:
                    candidates.add(relocated[0])
                    logs.append(f"已重新定位分类后的Word文件: {path} -> {relocated[0]}")
                elif len(relocated) > 1:
                    logs.append(f"分类后存在多个同名Word文件，无法确定目标: {path.name}")

    if scan_target:
        for file_path in target_path.rglob("*"):
            if file_path.is_file() and _is_notice_pdf_candidate(file_path):
                candidates.add(file_path.resolve())

    sorted_candidates = sorted(candidates, key=lambda item: str(item).lower())
    if sorted_candidates:
        source = "失败列表" if has_explicit_failures and not scan_target else "目标目录"
        logs.append(f"从{source}找到 {len(sorted_candidates)} 个可转换Word文件")
    elif has_explicit_failures and not scan_target:
        logs.append("失败列表中未找到仍存在且可转换的Word文件")
    return sorted_candidates


def _convert_generated_docs_to_pdf(
    work_dir: Path,
    logs: List[str],
    progress: NoticeProgress | None = None,
    skip_paths: Iterable[Path] | None = None,
) -> Dict[str, Any]:
    skip_resolved = {Path(path).resolve() for path in (skip_paths or [])}
    docx_files = []
    for pattern in ("授权委托书*.docx", "责令整改*.docx"):
        for file_path in work_dir.glob(pattern):
            if file_path.name.startswith("~$"):
                continue
            if any(marker in file_path.name for marker in (".clean_backup.docx", ".final_backup.docx", ".backup.docx")):
                continue
            if file_path.resolve() in skip_resolved:
                logs.append(f"跳过需手动处理的Word文件: {file_path.name}")
                continue
            docx_files.append(file_path)

    if not docx_files:
        logs.append("当前目录未找到需要转换的授权委托书或责令整改通知书")
        return {"converted": 0, "skipped": 0, "failures": [], "output_files": []}

    file_map = [(src, src.with_suffix(".pdf")) for src in docx_files]
    raw_result, captured, error = _call_with_progress_capture(
        lambda: convert_with_word_com(file_map, overwrite=True),
        progress,
    )
    if error is not None:
        raise error
    converted, skipped, failures = raw_result
    failures = list(failures)
    _store_captured_lines(logs, captured)

    failed_files = {src.resolve() for src, _ in failures}
    output_files = []
    for src, pdf_path in file_map:
        if src.resolve() in failed_files or not pdf_path.exists():
            continue
        output_files.append(str(pdf_path))
        try:
            src.unlink()
            logs.append(f"PDF转换成功，已删除原Word文件: {src.name}")
        except Exception as exc:
            failures.append((src, f"PDF已生成，但删除原Word失败: {exc}"))
            logs.append(f"PDF已生成，但删除原Word失败 {src.name}: {exc}")

    for src, reason in failures:
        logs.append(f"转换失败 {src.name}: {reason}")

    return {
        "converted": converted,
        "skipped": skipped,
        "failures": _failure_dicts(failures),
        "output_files": output_files,
    }


def _process_report_batch_legacy(
    report_files: Sequence[Path],
    company_name: str,
    template_paths: Dict[str, Path | None],
    soe_companies: set[str],
    logs: List[str],
    copy_to: str | None = None,
    progress: NoticeProgress | None = None,
    progress_base: float = 20,
    progress_span: float = 75,
    processed_before: int = 0,
    total_reports: int = 0,
) -> Dict[str, Any]:
    generated_files: List[str] = []
    manual_files: List[Dict[str, Any]] = []
    failures: List[Dict[str, str]] = []
    pdf_outputs: List[str] = []

    if not report_files:
        return {"generated_files": generated_files, "manual_files": manual_files, "failures": failures, "pdf_outputs": pdf_outputs}

    work_dir = report_files[0].parent
    logs.append("=" * 80)
    logs.append(f"处理企业: {company_name} (共 {len(report_files)} 个文档)")
    if progress:
        progress.set(progress_base, f"处理企业: {company_name}")

    def update_step(step_percent: float, message: str) -> None:
        if not progress:
            return
        total_value = total_reports or progress.total or len(report_files)
        batch_ratio = (processed_before / total_value) if total_value else 0
        batch_span = (len(report_files) / total_value * progress_span) if total_value else progress_span
        progress.set(progress_base + batch_ratio * progress_span + batch_span * step_percent / 100, message)

    with _safe_chdir(work_dir):
        try:
            from modules.Document_Processing.Report_Rewrite.edit_authorization import edit_authorization
            from modules.Document_Processing.Report_Rewrite.edit_disposal import process_disposal
            from modules.Document_Processing.Report_Rewrite.edit_rectification import edit_rectification, extract_info_from_filename
            from modules.Document_Processing.Report_Rewrite.rewrite_report import rewrite_report
        except Exception as exc:
            failures.append({"file": str(work_dir), "reason": f"导入通报处理模块失败: {exc}"})
            logs.append(f"导入通报处理模块失败: {exc}")
            return {"generated_files": generated_files, "manual_files": manual_files, "failures": failures, "pdf_outputs": pdf_outputs}

        collected_vulns: List[str] = []
        logs.append("步骤1/5: 通报改写")
        update_step(0, "步骤1/5: 通报改写")
        for report_file in report_files:
            logs.append(f"改写文档: {report_file.name}")
            _, vuln = extract_info_from_filename(str(report_file))
            if vuln:
                collected_vulns.append(vuln)

            raw_result, captured, error = _call_with_progress_capture(
                lambda: rewrite_report(
                    str(report_file),
                    template_file=str(template_paths["rewrite"]) if template_paths["rewrite"] else None,
                    start_para=1,
                    copy_to=copy_to,
                ),
                progress,
            )
            if error is None:
                result = _normalize_rewrite_result(raw_result, report_file)
            else:
                result = {
                    "success": False,
                    "output_file": None,
                    "backup_file": None,
                    "needs_manual_processing": False,
                    "skip_reason": f"执行错误: {error}",
                }
                logs.append(_format_exception(error))
            _store_captured_lines(logs, captured)

            output_path = _resolve_work_path(work_dir, result.get("output_file"))
            backup_path = _resolve_work_path(work_dir, result.get("backup_file"))
            if output_path and output_path.exists():
                generated_files.append(str(output_path))
            if backup_path and backup_path.exists():
                generated_files.append(str(backup_path))

            if not result.get("success"):
                reason = str(result.get("skip_reason") or "通报改写失败")
                fallback_path = output_path if output_path and output_path.exists() else None
                if fallback_path is None:
                    fallback_name = re.sub(r"^\d+", "", report_file.name) or report_file.name
                    fallback_path = _copy_template_for_manual_edit(
                        template_paths["rewrite"],
                        work_dir,
                        fallback_name,
                        logs,
                        f"通报改写失败: {report_file.name}",
                    )
                    if fallback_path:
                        generated_files.append(str(fallback_path))
                failures.append({"file": str(report_file), "reason": reason})
                manual_files.append({
                    "file": str(report_file.resolve()),
                    "reason": reason,
                    "backup_file": str(backup_path) if backup_path and backup_path.exists() else None,
                    "output_file": str(fallback_path) if fallback_path and fallback_path.exists() else None,
                })
                logs.append(f"通报改写失败: {report_file.name} -> {reason}")
            elif result.get("needs_manual_processing"):
                manual_output = output_path if output_path and output_path.exists() else backup_path
                manual_files.append({
                    "file": str(report_file.resolve()),
                    "reason": result.get("skip_reason") or "需要手动处理",
                    "backup_file": str(backup_path) if backup_path and backup_path.exists() else None,
                    "output_file": str(manual_output) if manual_output and manual_output.exists() else None,
                })
                logs.append(f"需要手动处理: {report_file.name} -> {result.get('skip_reason')}")
        update_step(20, "步骤1/5完成")

        time.sleep(0.5)
        target_report = report_files[0]
        override_name = f"{company_name}存在多个漏洞" if len(report_files) > 1 else None

        logs.append("步骤2/5: 生成授权委托书")
        update_step(20, "步骤2/5: 生成授权委托书")
        ok, captured, error = _call_with_progress_capture(
            lambda: edit_authorization(
                    str(target_report),
                    template_file=str(template_paths["authorization"]) if template_paths["authorization"] else None,
                    override_name=override_name,
                ),
            progress,
        )
        if error is not None:
            ok = False
            logs.append(_format_exception(error))
            failures.append({"file": str(target_report), "reason": f"授权委托书生成异常: {error}"})
        _store_captured_lines(logs, captured)
        if not ok:
            failures.append({"file": str(target_report), "reason": "授权委托书生成失败"})
        generated_files.extend(str(path.resolve()) for path in work_dir.glob("授权委托书*.docx"))
        update_step(40, "步骤2/5完成")

        if company_name in soe_companies:
            logs.append(f"检测到国企: {company_name}，跳过责令整改通知书")
        else:
            logs.append("步骤3/5: 生成责令整改通知书")
            update_step(40, "步骤3/5: 生成责令整改通知书")
            combined_vulns = None
            if len(report_files) > 1 and collected_vulns:
                combined_vulns = "、".join(sorted(set(collected_vulns)))
                if not combined_vulns.endswith(("漏洞", "风险")):
                    combined_vulns += "漏洞"
            elif collected_vulns:
                combined_vulns = collected_vulns[0]

            rectification_started_at = time.time()
            ok, captured, error = _call_with_progress_capture(
                lambda: edit_rectification(
                        str(target_report),
                        template_file=str(template_paths["rectification"]) if template_paths["rectification"] else None,
                        company_name=company_name,
                        vuln_type=combined_vulns,
                    ),
                progress,
            )
            if error is not None:
                ok = False
                logs.append(_format_exception(error))
                failures.append({"file": str(target_report), "reason": f"责令整改通知书生成异常: {error}"})
            _store_captured_lines(logs, captured)
            if not ok:
                reason = "责令整改通知书生成失败"
                failures.append({"file": str(target_report), "reason": reason})
                fallback_rect = _copy_template_for_manual_edit(
                    template_paths["rectification"],
                    work_dir,
                    "责令整改通知书.docx",
                    logs,
                    reason,
                )
                if fallback_rect:
                    generated_files.append(str(fallback_rect))
                    manual_files.append({
                        "file": str(fallback_rect),
                        "reason": "责令整改通知书生成失败，请基于该模板手动补齐后再转换PDF",
                        "output_file": str(fallback_rect),
                    })
            latest_rect = _latest_rectification_doc(work_dir)
            if latest_rect:
                generated_files.append(str(latest_rect.resolve()))
                manual_reason = _rectification_manual_reason(latest_rect, company_name, combined_vulns)
                if manual_reason and latest_rect.stat().st_mtime >= rectification_started_at - 1:
                    latest_rect_resolved = str(latest_rect.resolve())
                    if not any(item.get("output_file") == latest_rect_resolved or item.get("file") == latest_rect_resolved for item in manual_files):
                        manual_files.append({
                            "file": latest_rect_resolved,
                            "reason": f"责令整改文档需要手动校正: {manual_reason}",
                            "output_file": latest_rect_resolved,
                        })
                    logs.append(f"责令整改文档需要手动校正: {manual_reason}")
        update_step(60, "步骤3/5完成")

        logs.append("步骤4/5: 处理处置文件")
        update_step(60, "步骤4/5: 处理处置文件")
        if not list(work_dir.glob("*处置*.docx")) and not list(work_dir.glob("*处置*.pdf")):
            if template_paths["disposal"]:
                ok, captured, error = _call_with_progress_capture(
                    lambda: process_disposal(str(template_paths["disposal"]), target_directory=work_dir),
                    progress,
                )
                if error is not None:
                    ok = False
                    logs.append(_format_exception(error))
                    failures.append({"file": str(work_dir), "reason": f"处置文件处理异常: {error}"})
                _store_captured_lines(logs, captured)
                if not ok:
                    failures.append({"file": str(work_dir), "reason": "处置文件处理失败"})
            else:
                logs.append("未找到处置文件模板，跳过")
        else:
            logs.append("处置文件已存在，跳过")
        generated_files.extend(str(path.resolve()) for path in work_dir.glob("*处置*.docx"))
        update_step(80, "步骤4/5完成")

        logs.append("步骤5/5: 转换授权委托书与责令整改通知书为PDF")
        update_step(80, "步骤5/5: 转换授权委托书与责令整改通知书为PDF")
        pdf_result = _convert_generated_docs_to_pdf(work_dir, logs, progress, skip_paths=_manual_path_values(manual_files))
        pdf_outputs.extend(pdf_result["output_files"])
        failures.extend(pdf_result["failures"])
        update_step(95, "步骤5/5完成")

    return {
        "generated_files": sorted(set(generated_files)),
        "manual_files": manual_files,
        "failures": failures,
        "pdf_outputs": pdf_outputs,
    }


def _process_report_batch(
    report_files: Sequence[Path],
    company_name: str,
    template_paths: Dict[str, Path | None],
    soe_companies: set[str],
    logs: List[str],
    copy_to: str | None = None,
    progress: NoticeProgress | None = None,
    progress_base: float = 20,
    progress_span: float = 75,
    processed_before: int = 0,
    total_reports: int = 0,
    work_dir: Path | None = None,
) -> Dict[str, Any]:
    generated_files: List[str] = []
    manual_files: List[Dict[str, Any]] = []
    failures: List[Dict[str, str]] = []
    pdf_outputs: List[str] = []
    sources = [Path(path).resolve() for path in report_files if Path(path).exists()]
    if work_dir is None:
        if not sources:
            return {"generated_files": [], "manual_files": [], "failures": [], "pdf_outputs": [], "skipped_complete": False}
        work_dir = sources[0].parent
    work_dir = work_dir.resolve()
    is_soe = company_name in soe_companies
    state = _load_notice_state(work_dir, company_name, logs)
    current_signature = _source_signature(sources)
    saved_signature = state.get("input_signature") or []

    logs.append("=" * 80)
    logs.append(f"处理企业: {company_name} (原始通报 {len(sources)} 个)")
    if progress:
        progress.set(progress_base, f"处理企业: {company_name}")

    def update_step(step_percent: float, message: str) -> None:
        if not progress:
            return
        total_value = total_reports or progress.total or max(1, len(sources))
        batch_count = max(1, len(sources))
        batch_ratio = processed_before / total_value if total_value else 0
        batch_span = batch_count / total_value * progress_span if total_value else progress_span
        progress.set(progress_base + batch_ratio * progress_span + batch_span * step_percent / 100, message)

    try:
        from modules.Document_Processing.Report_Rewrite.edit_authorization import edit_authorization
        from modules.Document_Processing.Report_Rewrite.edit_disposal import process_disposal
        from modules.Document_Processing.Report_Rewrite.edit_rectification import edit_rectification, extract_info_from_filename
        from modules.Document_Processing.Report_Rewrite.rewrite_report import rewrite_report
    except Exception as exc:
        reason = f"导入通报处理模块失败: {exc}"
        logs.append(reason)
        failures.append({"file": str(work_dir), "reason": reason})
        return {"generated_files": [], "manual_files": [], "failures": failures, "pdf_outputs": [], "skipped_complete": False}

    collected_vulns: List[str] = []
    for path in [*sources, *_rewritten_notice_files(work_dir)]:
        _, vuln = extract_info_from_filename(str(path))
        if vuln:
            collected_vulns.append(vuln)
    combined_vulns = None
    if collected_vulns:
        combined_vulns = "、".join(sorted(set(collected_vulns))) if len(set(collected_vulns)) > 1 else collected_vulns[0]
        if len(set(collected_vulns)) > 1 and not combined_vulns.endswith(("漏洞", "风险")):
            combined_vulns += "漏洞"

    if sources and saved_signature and saved_signature != current_signature:
        logs.append("检测到原始通报内容与上次记录不同，将重新核验并处理新文件")
        valid_items: List[Dict[str, Any]] = []
        rewrite_required: List[Dict[str, Any]] = []
        for fingerprint in current_signature:
            matched_item = next((
                item for item in state.get("rewrite_items", [])
                if isinstance(item, dict)
                and item.get("source") == fingerprint
                and (lambda artifact: bool(artifact and _is_any_rewritten_notice_file(artifact)))(
                    _state_artifact_path(work_dir, item.get("artifact"))
                )
            ), None)
            if matched_item:
                valid_items.append(matched_item)
            else:
                rewrite_required.append(fingerprint)
        state["complete"] = False
        state["active_stage"] = None
        state["stage_started_at"] = None
        state["stages"] = {stage: False for stage in NOTICE_PROCESS_STAGES}
        state["rewrite_items"] = valid_items
        state["rewrite_required"] = rewrite_required
        state["input_signature"] = current_signature
    elif sources and not saved_signature:
        state["input_signature"] = current_signature

    rewrite_artifacts: List[Path] = []
    for source in sources:
        artifact = _rewrite_artifact_for_source(work_dir, source, state)
        if artifact:
            _mark_legacy_rewrite_artifact(artifact)
            _set_rewrite_state_item(work_dir, state, source, artifact)
            rewrite_artifacts.append(artifact)
    if not sources:
        rewrite_artifacts = _rewritten_notice_files(work_dir)
        for artifact in rewrite_artifacts:
            _mark_legacy_rewrite_artifact(artifact)

    legacy_pdf_recovery = not state.get("_loaded_from_disk") and not sources
    trust_existing_pdf = legacy_pdf_recovery or bool((state.get("stages") or {}).get("pdf")) or state.get("active_stage") == "pdf"
    recovery_pdf_minimum_mtime = (
        float(state.get("stage_started_at") or 0)
        if state.get("active_stage") == "pdf"
        else None
    )
    auth_artifacts = _authorization_artifacts(work_dir, company_name, trust_pdf_structure=trust_existing_pdf)
    rect_artifacts = _rectification_artifacts(work_dir, company_name, combined_vulns, trust_pdf_structure=trust_existing_pdf)
    disposal_artifacts = _disposal_artifacts(work_dir)
    cleanup_failures = _cleanup_words_with_valid_pdf(
        work_dir,
        company_name,
        logs,
        minimum_pdf_mtime=recovery_pdf_minimum_mtime,
    ) if (_can_recover_pdf_cleanup(state) or legacy_pdf_recovery) else []
    failures.extend(cleanup_failures)
    pdf_artifacts = _pdf_stage_artifacts(
        work_dir,
        company_name,
        is_soe,
        trust_pdf_structure=trust_existing_pdf,
        minimum_mtime=recovery_pdf_minimum_mtime,
    )

    rewrite_verified = bool(rewrite_artifacts) and (not sources or len(rewrite_artifacts) == len(sources))
    completed_artifacts_valid = (
        rewrite_verified
        and _stage_verified(state, "authorization", auth_artifacts)
        and _stage_verified(state, "rectification", rect_artifacts, optional=is_soe)
        and _stage_verified(state, "disposal", disposal_artifacts, optional=template_paths.get("disposal") is None)
        and _stage_verified(state, "pdf", pdf_artifacts)
    )
    same_inputs = not sources or not saved_signature or saved_signature == current_signature
    if bool(state.get("complete")) and same_inputs and completed_artifacts_valid and not cleanup_failures:
        cleanup_errors = _delete_completed_notice_sources(sources, logs)
        failures.extend(cleanup_errors)
        if not cleanup_errors:
            logs.append(f"已识别为完整处理过的企业，跳过重复生成和编号: {company_name}")
            generated_files.extend(str(path) for path in [*rewrite_artifacts, *disposal_artifacts])
            pdf_outputs.extend(str(path) for path in pdf_artifacts)
            return {
                "generated_files": sorted(set(generated_files)),
                "manual_files": [],
                "failures": failures,
                "pdf_outputs": sorted(set(pdf_outputs)),
                "skipped_complete": True,
            }

    if not sources and not rewrite_artifacts:
        failures.append({"file": str(work_dir), "reason": "找到处理状态或后续产物，但缺少原始通报和有效改写件，无法自动续跑"})
        return {"generated_files": [], "manual_files": [], "failures": failures, "pdf_outputs": [], "skipped_complete": False}

    with _safe_chdir(work_dir):
        logs.append("步骤1/5: 通报改写")
        update_step(0, "步骤1/5: 通报改写")
        rewrite_failed = False
        if sources:
            for source in sources:
                artifact = _rewrite_artifact_for_source(work_dir, source, state)
                if artifact:
                    _mark_legacy_rewrite_artifact(artifact)
                    _set_rewrite_state_item(work_dir, state, source, artifact)
                    rewrite_artifacts.append(artifact)
                    logs.append(f"已识别改写完成，跳过: {source.name} -> {artifact.name}")
                    continue

                _start_notice_stage(work_dir, state, "rewrite", source)
                logs.append(f"改写文档: {source.name}")
                raw_result, captured, error = _call_with_progress_capture(
                    lambda source=source: rewrite_report(
                        str(source),
                        template_file=str(template_paths["rewrite"]) if template_paths["rewrite"] else None,
                        start_para=1,
                        copy_to=copy_to,
                    ),
                    progress,
                )
                _store_captured_lines(logs, captured)
                result = _normalize_rewrite_result(raw_result, source) if error is None else {
                    "success": False,
                    "output_file": None,
                    "backup_file": None,
                    "needs_manual_processing": False,
                    "skip_reason": f"执行错误: {error}",
                }
                if error is not None:
                    logs.append(_format_exception(error))
                output_path = _resolve_work_path(work_dir, result.get("output_file"))
                backup_path = _resolve_work_path(work_dir, result.get("backup_file"))
                artifact = next((path for path in (output_path, backup_path) if path and _is_any_rewritten_notice_file(path)), None)
                if artifact is None:
                    artifact = _rewrite_artifact_for_source(work_dir, source, state)
                rewrite_output_recovered = bool(
                    artifact
                    and (
                        result.get("success")
                        or _artifacts_created_after([artifact], state.get("stage_started_at"))
                    )
                )
                if rewrite_output_recovered and artifact:
                    _set_rewrite_state_item(work_dir, state, source, artifact)
                    rewrite_artifacts.append(artifact)
                    generated_files.append(str(artifact))
                    if not result.get("success"):
                        logs.append(f"改写返回异常，但检测到本次已生成有效改写件，按完成恢复: {artifact.name}")
                    _finish_notice_stage(work_dir, state, "rewrite", True)
                    if result.get("needs_manual_processing"):
                        manual_files.append({
                            "file": str(source),
                            "reason": result.get("skip_reason") or "通报需手动确认",
                            "backup_file": str(backup_path) if backup_path and backup_path.exists() else None,
                            "output_file": str(artifact),
                        })
                else:
                    rewrite_failed = True
                    reason = str(result.get("skip_reason") or "通报改写失败")
                    failures.append({"file": str(source), "reason": reason})
                    manual_files.append({"file": str(source), "reason": reason, "output_file": str(output_path) if output_path and output_path.exists() else None})
                    logs.append(f"通报改写失败，已保留原件供下次重试: {source.name} -> {reason}")
                    _finish_notice_stage(work_dir, state, "rewrite", False)
        rewrite_artifacts = sorted(set(_rewritten_notice_files(work_dir)), key=lambda path: str(path).lower())
        rewrite_verified = bool(rewrite_artifacts) and (not sources or all(_rewrite_artifact_for_source(work_dir, source, state) for source in sources))
        state["input_signature"] = current_signature or state.get("input_signature") or []
        state["stages"]["rewrite"] = rewrite_verified
        _save_notice_state(work_dir, state)
        update_step(20, "步骤1/5完成")
        if rewrite_failed or not rewrite_verified:
            logs.append("通报尚未全部改写成功，本次不执行后续步骤")
            return {
                "generated_files": sorted(set(generated_files + [str(path) for path in rewrite_artifacts])),
                "manual_files": manual_files,
                "failures": failures,
                "pdf_outputs": [],
                "skipped_complete": False,
            }

        target_report = sources[0] if sources else rewrite_artifacts[0]
        override_name = f"{company_name}存在多个漏洞" if max(len(sources), len(rewrite_artifacts)) > 1 else None

        auth_artifacts = _authorization_artifacts(work_dir, company_name, trust_pdf_structure=_can_recover_pdf_cleanup(state))
        if _stage_verified(state, "authorization", auth_artifacts):
            logs.append("步骤2/5: 已识别授权委托书生成完成，跳过重复生成")
            state["stages"]["authorization"] = True
            _save_notice_state(work_dir, state)
        else:
            logs.append("步骤2/5: 生成授权委托书")
            update_step(20, "步骤2/5: 生成授权委托书")
            _start_notice_stage(work_dir, state, "authorization")
            ok, captured, error = _call_with_progress_capture(
                lambda: edit_authorization(
                    str(target_report),
                    template_file=str(template_paths["authorization"]) if template_paths["authorization"] else None,
                    override_name=override_name,
                ),
                progress,
            )
            _store_captured_lines(logs, captured)
            if error is not None:
                ok = False
                logs.append(_format_exception(error))
            auth_artifacts = _authorization_artifacts(work_dir, company_name, trust_pdf_structure=False)
            auth_ok = bool(auth_artifacts and (ok or _artifacts_created_after(auth_artifacts, state.get("stage_started_at"))))
            if auth_ok and not ok:
                logs.append("授权委托书返回异常，但检测到本次已生成且内容有效，按完成恢复")
            _finish_notice_stage(work_dir, state, "authorization", auth_ok)
            if not auth_ok:
                failures.append({"file": str(target_report), "reason": f"授权委托书生成失败{f': {error}' if error else ''}"})
                return {
                    "generated_files": sorted(set(generated_files + [str(path) for path in rewrite_artifacts])),
                    "manual_files": manual_files,
                    "failures": failures,
                    "pdf_outputs": [],
                    "skipped_complete": False,
                }
        generated_files.extend(str(path) for path in auth_artifacts if path.suffix.lower() == ".docx")
        update_step(40, "步骤2/5完成")

        if is_soe:
            logs.append(f"步骤3/5: 检测到国企 {company_name}，无需责令整改通知书")
            state["stages"]["rectification"] = True
            _save_notice_state(work_dir, state)
        else:
            rect_artifacts = _rectification_artifacts(work_dir, company_name, combined_vulns, trust_pdf_structure=_can_recover_pdf_cleanup(state))
            if _stage_verified(state, "rectification", rect_artifacts):
                logs.append("步骤3/5: 已识别责令整改通知书生成完成，跳过重复编号")
                state["stages"]["rectification"] = True
                _save_notice_state(work_dir, state)
            else:
                logs.append("步骤3/5: 生成责令整改通知书")
                update_step(40, "步骤3/5: 生成责令整改通知书")
                _start_notice_stage(work_dir, state, "rectification")
                ok, captured, error = _call_with_progress_capture(
                    lambda: edit_rectification(
                        str(target_report),
                        template_file=str(template_paths["rectification"]) if template_paths["rectification"] else None,
                        company_name=company_name,
                        vuln_type=combined_vulns,
                    ),
                    progress,
                )
                _store_captured_lines(logs, captured)
                if error is not None:
                    ok = False
                    logs.append(_format_exception(error))
                rect_artifacts = _rectification_artifacts(work_dir, company_name, combined_vulns, trust_pdf_structure=False)
                rect_ok = bool(rect_artifacts and (ok or _artifacts_created_after(rect_artifacts, state.get("stage_started_at"))))
                if rect_ok and not ok:
                    logs.append("责令整改返回异常，但检测到本次已生成且内容有效，按完成恢复，避免重复编号")
                _finish_notice_stage(work_dir, state, "rectification", rect_ok)
                if not rect_ok:
                    latest_rect = _latest_rectification_doc(work_dir)
                    reason = f"责令整改通知书生成或内容校验失败{f': {error}' if error else ''}"
                    failures.append({"file": str(target_report), "reason": reason})
                    if latest_rect:
                        manual_files.append({"file": str(latest_rect), "reason": reason, "output_file": str(latest_rect)})
                    return {
                        "generated_files": sorted(set(generated_files + [str(path) for path in rewrite_artifacts + rect_artifacts])),
                        "manual_files": manual_files,
                        "failures": failures,
                        "pdf_outputs": [],
                        "skipped_complete": False,
                    }
            generated_files.extend(str(path) for path in rect_artifacts if path.suffix.lower() == ".docx")
        update_step(60, "步骤3/5完成")

        disposal_artifacts = _disposal_artifacts(work_dir)
        disposal_optional = template_paths.get("disposal") is None
        if _stage_verified(state, "disposal", disposal_artifacts, optional=disposal_optional):
            logs.append("步骤4/5: 已识别处置文件完成，跳过")
            state["stages"]["disposal"] = True
            _save_notice_state(work_dir, state)
        elif disposal_optional:
            logs.append("步骤4/5: 未找到处置模板，按现有规则跳过")
            state["stages"]["disposal"] = True
            _save_notice_state(work_dir, state)
        else:
            logs.append("步骤4/5: 处理处置文件")
            update_step(60, "步骤4/5: 处理处置文件")
            _start_notice_stage(work_dir, state, "disposal")
            ok, captured, error = _call_with_progress_capture(
                lambda: process_disposal(str(template_paths["disposal"]), target_directory=work_dir),
                progress,
            )
            _store_captured_lines(logs, captured)
            if error is not None:
                ok = False
                logs.append(_format_exception(error))
            disposal_artifacts = _disposal_artifacts(work_dir)
            disposal_ok = bool(disposal_artifacts and (ok or _artifacts_created_after(disposal_artifacts, state.get("stage_started_at"))))
            if disposal_ok and not ok:
                logs.append("处置文件返回异常，但检测到本次已生成且内容有效，按完成恢复")
            _finish_notice_stage(work_dir, state, "disposal", disposal_ok)
            if not disposal_ok:
                failures.append({"file": str(work_dir), "reason": f"处置文件处理或内容校验失败{f': {error}' if error else ''}"})
                return {
                    "generated_files": sorted(set(generated_files + [str(path) for path in rewrite_artifacts + disposal_artifacts])),
                    "manual_files": manual_files,
                    "failures": failures,
                    "pdf_outputs": [],
                    "skipped_complete": False,
                }
        generated_files.extend(str(path) for path in disposal_artifacts if path.suffix.lower() == ".docx")
        update_step(80, "步骤4/5完成")

        cleanup_errors = _cleanup_words_with_valid_pdf(
            work_dir,
            company_name,
            logs,
            minimum_pdf_mtime=float(state.get("stage_started_at") or 0) or None,
        ) if (_can_recover_pdf_cleanup(state) or legacy_pdf_recovery) else []
        failures.extend(cleanup_errors)
        pdf_artifacts = _pdf_stage_artifacts(
            work_dir,
            company_name,
            is_soe,
            trust_pdf_structure=_can_recover_pdf_cleanup(state),
            minimum_mtime=(
                float(state.get("stage_started_at") or 0)
                if state.get("active_stage") == "pdf"
                else None
            ),
        )
        if not cleanup_errors and _stage_verified(state, "pdf", pdf_artifacts):
            logs.append("步骤5/5: 已识别有效PDF且原Word已清理，跳过重复转换")
            state["stages"]["pdf"] = True
            _save_notice_state(work_dir, state)
        else:
            logs.append("步骤5/5: 转换授权委托书与责令整改通知书为PDF")
            update_step(80, "步骤5/5: 转换PDF")
            _start_notice_stage(work_dir, state, "pdf")
            pdf_started_at = float(state.get("stage_started_at") or 0) or None
            pdf_result = _convert_generated_docs_to_pdf(work_dir, logs, progress, skip_paths=_manual_path_values(manual_files))
            cleanup_errors = _cleanup_words_with_valid_pdf(
                work_dir,
                company_name,
                logs,
                minimum_pdf_mtime=pdf_started_at,
            )
            pdf_artifacts = _pdf_stage_artifacts(
                work_dir,
                company_name,
                is_soe,
                trust_pdf_structure=True,
                minimum_mtime=pdf_started_at,
            )
            pdf_ok = not cleanup_errors and bool(pdf_artifacts)
            _finish_notice_stage(work_dir, state, "pdf", pdf_ok)
            if not pdf_ok:
                failures.extend(pdf_result["failures"])
                failures.extend(cleanup_errors)
                failures.append({"file": str(work_dir), "reason": "PDF转换未全部完成，将在下次从此步骤继续"})
            elif pdf_result["failures"]:
                logs.append("PDF转换曾返回异常，但已核验全部PDF有效且原Word已清理，按完成恢复")
        pdf_outputs.extend(str(path) for path in pdf_artifacts)
        update_step(95, "步骤5/5完成")

    if not failures and all(bool(state["stages"].get(stage)) for stage in NOTICE_PROCESS_STAGES):
        cleanup_errors = _delete_completed_notice_sources(sources, logs)
        failures.extend(cleanup_errors)
        state["complete"] = not cleanup_errors
        _save_notice_state(work_dir, state)
        if state["complete"]:
            logs.append(f"企业全部阶段已完成并保存断点状态: {company_name}")
    else:
        state["complete"] = False
        _save_notice_state(work_dir, state)

    generated_files.extend(str(path) for path in rewrite_artifacts)
    return {
        "generated_files": sorted(set(generated_files)),
        "manual_files": manual_files,
        "failures": failures,
        "pdf_outputs": sorted(set(pdf_outputs)),
        "skipped_complete": False,
    }


def _discover_report_files(directory: Path, logs: List[str]) -> List[Path]:
    report_files: List[Path] = []
    for item in list(directory.glob("*.docx")):
        if item.name.startswith("~$"):
            continue
        if any(keyword in item.name for keyword in ["模板", "授权委托书", "责令整改", "处置"]):
            continue
        if not _notification_name(item.name):
            continue
        if _is_any_rewritten_notice_file(item):
            continue
        if not item.name[0].isdigit():
            new_name = f"{str(int(time.time() * 1000))[-10:]}{item.name}"
            new_path = item.parent / new_name
            try:
                item.rename(new_path)
                logs.append(f"重命名原始通报: {item.name} -> {new_name}")
                item = new_path
            except Exception as exc:
                logs.append(f"重命名失败 {item.name}: {exc}")
                continue
        report_files.append(item)
    return report_files


def _directory_has_resumable_notice(directory: Path) -> bool:
    return _notice_state_path(directory).exists() or bool(_rewritten_notice_files(directory))


def _infer_resumable_company_name(directory: Path, report_files: Sequence[Path]) -> str:
    state_path = _notice_state_path(directory)
    if state_path.exists():
        try:
            state = json.loads(state_path.read_text(encoding="utf-8"))
            company_name = str(state.get("company_name") or "").strip()
            if company_name:
                return company_name
        except Exception:
            pass
    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf

        for path in [*report_files, *_rewritten_notice_files(directory)]:
            company_name = gf.normalize_company(path.name)
            if company_name:
                return company_name
        company_name = gf.normalize_company(directory.name)
        if company_name:
            return company_name
    except Exception:
        pass
    return directory.name


def _process_notice_directory(
    directory: Path,
    template_paths: Dict[str, Path | None],
    soe_companies: set[str],
    logs: List[str],
    company_groups: Dict[str, str] | None = None,
    processed_dirs: set[Path] | None = None,
    progress: NoticeProgress | None = None,
    total_reports: int = 0,
    processed_offset: int = 0,
) -> Dict[str, Any]:
    processed_dirs = processed_dirs or set()
    company_groups = company_groups or {}
    directory = directory.resolve()
    if directory in processed_dirs:
        return {"processed": 0, "generated_files": [], "manual_files": [], "failures": [], "pdf_outputs": []}
    processed_dirs.add(directory)

    processed = 0
    generated_files: List[str] = []
    manual_files: List[Dict[str, Any]] = []
    failures: List[Dict[str, str]] = []
    pdf_outputs: List[str] = []

    logs.append(f"扫描目录: {directory}")
    if progress:
        progress.set(message=f"扫描目录: {directory}")
    report_files = _discover_report_files(directory, logs)
    if report_files:
        report_groups: Dict[str, List[Path]] = {}
        for report_file in report_files:
            company_name = report_file.parent.name
            if company_name == report_file.parent.parent.name:
                try:
                    from modules.Document_Processing.Report_Rewrite import group_folders as gf

                    company_name = gf.normalize_company(report_file.name) or company_name
                except Exception:
                    pass
            report_groups.setdefault(company_name, []).append(report_file)

        for company_name, files in report_groups.items():
            result = _process_report_batch(
                files,
                company_name,
                template_paths,
                soe_companies,
                logs,
                copy_to=company_groups.get(company_name),
                progress=progress,
                processed_before=processed_offset + processed,
                total_reports=total_reports,
            )
            processed += len(files)
            if progress:
                progress.set_processed(processed_offset + processed, total_reports, f"已完成 {processed_offset + processed}/{total_reports} 个文档")
            generated_files.extend(result["generated_files"])
            manual_files.extend(result["manual_files"])
            failures.extend(result["failures"])
            pdf_outputs.extend(result["pdf_outputs"])
    elif _directory_has_resumable_notice(directory):
        company_name = _infer_resumable_company_name(directory, [])
        logs.append(f"检测到旧处理产物或断点状态，尝试续跑企业: {company_name}")
        result = _process_report_batch(
            [],
            company_name,
            template_paths,
            soe_companies,
            logs,
            copy_to=company_groups.get(company_name),
            progress=progress,
            processed_before=processed_offset + processed,
            total_reports=total_reports,
            work_dir=directory,
        )
        processed += 1
        if progress:
            progress.set_processed(processed_offset + processed, total_reports, f"已完成 {processed_offset + processed}/{total_reports} 个文档")
        generated_files.extend(result["generated_files"])
        manual_files.extend(result["manual_files"])
        failures.extend(result["failures"])
        pdf_outputs.extend(result["pdf_outputs"])
    else:
        logs.append(f"未在 {directory.name} 找到符合规则的通报文档")

    for subdir in [item for item in directory.iterdir() if item.is_dir() and not item.name.startswith(".")]:
        result = _process_notice_directory(
            subdir,
            template_paths,
            soe_companies,
            logs,
            company_groups,
            processed_dirs,
            progress=progress,
            total_reports=total_reports,
            processed_offset=processed_offset + processed,
        )
        processed += result["processed"]
        if progress:
            progress.set_processed(processed_offset + processed, total_reports, f"已完成 {processed_offset + processed}/{total_reports} 个文档")
        generated_files.extend(result["generated_files"])
        manual_files.extend(result["manual_files"])
        failures.extend(result["failures"])
        pdf_outputs.extend(result["pdf_outputs"])

    return {
        "processed": processed,
        "generated_files": sorted(set(generated_files)),
        "manual_files": manual_files,
        "failures": failures,
        "pdf_outputs": sorted(set(pdf_outputs)),
    }


def _doc_notice_classify(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_path = Path(_required_text(payload, "target_path", "请选择要分类的目录")).expanduser()
    if not target_path.exists() or not target_path.is_dir():
        return {"success": False, "message": f"目标目录不存在或不是目录: {target_path}", "logs": []}
    conflict = _notice_active_conflict(target_path, "执行一键分类", str(payload.get("_notice_task_id") or ""))
    if conflict:
        return conflict

    logs: List[str] = []
    buffer = io.StringIO()
    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf

        with contextlib.redirect_stdout(buffer):
            raw_result = gf.run_grouping(
                str(target_path),
                entries=str(payload.get("entries") or "both"),
                pattern=str(payload.get("pattern") or "exact"),
                groups_source=str(payload.get("groups_source") or "db"),
            )
        result = _normalize_grouping_result(raw_result)
    except Exception as exc:
        logs.extend(_captured_lines(buffer))
        logs.append(traceback.format_exc())
        return {"success": False, "message": f"分类失败: {exc}", "logs": logs}

    logs.extend(_captured_lines(buffer))
    logs.extend(str(line) for line in (result.get("log") or []) if str(line).strip())
    copy_stats = _fill_rewritten_notice_copy_to(
        target_path,
        result.get("company_group_list"),
        logs,
    )
    result["copy_to_updated"] = copy_stats["updated"]
    result["copy_to_skipped_no_group"] = copy_stats["skipped_no_group"]
    result["copy_to_errors"] = copy_stats["errors"]
    message = f"分类完成：移动 {result.get('moved', 0)} 个，跳过 {result.get('skipped_exist', 0)} 个，错误 {result.get('errors', 0)} 个"
    if copy_stats["updated"]:
        message += f"，补写抄送 {copy_stats['updated']} 个"
    if copy_stats["errors"]:
        message += f"，抄送补写错误 {copy_stats['errors']} 个"
    return {
        "success": result.get("errors", 0) == 0 and copy_stats["errors"] == 0,
        "message": message,
        "logs": logs,
        "result": result,
    }


def _doc_notice_process(payload: Dict[str, Any], progress: NoticeProgress | None = None) -> Dict[str, Any]:
    target_path = Path(_required_text(payload, "target_path", "请选择文件夹或压缩包")).expanduser()
    auto_group = bool(payload.get("auto_group", True))
    logs: List[str] = ProgressLogList(progress) if progress else []
    company_groups: Dict[str, str] = {}

    if not target_path.exists():
        return {"success": False, "message": f"目标路径不存在: {target_path}", "logs": []}

    if progress:
        progress.set(3, "正在准备通报处理任务...")
    _apply_report_counter_payload(payload, logs)

    template_paths = {
        "rewrite": _find_template("通报模板"),
        "authorization": _find_template("授权委托书"),
        "rectification": _find_template("责令整改"),
        "disposal": _find_template("处置"),
    }
    for key, value in template_paths.items():
        logs.append(f"模板 {key}: {value if value else '未找到'}")
    if progress:
        progress.set(8, "模板检查完成")

    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf

        soe_companies = gf.get_soe_companies()
    except Exception:
        soe_companies = set()

    if target_path.is_file():
        logs.append(f"检测到压缩包: {target_path.name}")
        if progress:
            progress.set(10, f"正在解压: {target_path.name}")
        try:
            target_path, _ = _extract_archive(target_path, logs)
            if target_path is None:
                return {"success": False, "message": "解压失败", "logs": logs}
        except Exception as exc:
            logs.append(traceback.format_exc())
            return {"success": False, "message": f"解压失败: {exc}", "logs": logs}

    if not target_path.is_dir():
        return {"success": False, "message": "只支持文件夹或ZIP压缩包", "logs": logs}

    if progress:
        progress.set(11, "正在预解压目录中的压缩包...")
    archive_failures = _extract_archives_before_notice_processing(target_path, logs, progress)

    if auto_group:
        logs.append("执行自动分类")
        if progress:
            progress.set(12, "执行自动分类")
        classify_result = _doc_notice_classify({
            "target_path": str(target_path),
            "_notice_task_id": str(payload.get("_notice_task_id") or ""),
        })
        logs.extend(classify_result.get("logs") or [])
        logs.append(classify_result.get("message") or "")
        classify_data = classify_result.get("result") or {}
        company_groups = _company_group_map(classify_data.get("company_group_list"))

    if not company_groups:
        try:
            from modules.Document_Processing.Report_Rewrite import group_folders as gf
            company_groups = _company_group_map(gf.collect_all_company_groups(str(target_path)))
        except Exception:
            company_groups = {}

    total_reports = _count_notification_docs(target_path)
    logs.append(f"共发现 {total_reports} 个通报文档")
    if progress:
        progress.set_total(total_reports)
        progress.set(20 if total_reports else 95, f"共发现 {total_reports} 个通报文档")
    result = _process_notice_directory(
        target_path,
        template_paths,
        soe_companies,
        logs,
        company_groups,
        progress=progress,
        total_reports=total_reports,
    )
    if archive_failures:
        result["failures"].extend(archive_failures)

    message = f"处理完成：处理 {result['processed']} 个文档，失败 {len(result['failures'])} 个，需手动处理 {len(result['manual_files'])} 个"
    if progress:
        progress.set(100, message)
    return {
        "success": len(result["failures"]) == 0,
        "message": message,
        "target_path": str(target_path),
        "total_reports": total_reports,
        "processed": result["processed"],
        "generated_files": result["generated_files"],
        "manual_files": result["manual_files"],
        "failures": result["failures"],
        "pdf_outputs": result["pdf_outputs"],
        "logs": logs,
    }


def _notice_task_worker(task_id: str, payload: Dict[str, Any]) -> None:
    with _NOTICE_TASK_LOCK:
        task = _NOTICE_TASKS.get(task_id)
    if not task:
        return

    progress = task["progress"]
    pythoncom = None
    try:
        import pythoncom as _pythoncom  # type: ignore

        pythoncom = _pythoncom
        pythoncom.CoInitialize()
    except Exception:
        pythoncom = None
    try:
        worker_payload = dict(payload)
        worker_payload["_notice_task_id"] = task_id
        result = _doc_notice_process(worker_payload, progress=progress)
        with _NOTICE_TASK_LOCK:
            task.update({
                "running": False,
                "done": True,
                "success": bool(result.get("success")),
                "message": result.get("message") or ("处理完成" if result.get("success") else "处理失败"),
                "result": result,
                "finished_at": time.time(),
            })
            _prune_notice_tasks_locked()
    except Exception as exc:
        progress.log(_format_exception(exc))
        progress.set(0, f"处理失败: {exc}")
        result = {
            "success": False,
            "message": f"处理失败: {exc}",
            "logs": progress.snapshot().get("logs", []),
            "failures": [{"file": str(payload.get("target_path") or ""), "reason": str(exc)}],
            "generated_files": [],
            "manual_files": [],
            "pdf_outputs": [],
        }
        with _NOTICE_TASK_LOCK:
            task.update({
                "running": False,
                "done": True,
                "success": False,
                "message": result["message"],
                "result": result,
                "error": str(exc),
                "finished_at": time.time(),
            })
            _prune_notice_tasks_locked()
    finally:
        if pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


def _doc_notice_process_start(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_path = Path(_required_text(payload, "target_path", "请选择文件夹或ZIP压缩包")).expanduser()
    if not target_path.exists():
        return {"success": False, "message": f"目标路径不存在: {target_path}", "logs": []}
    target_path = target_path.resolve()
    target_key = _notice_target_key(target_path)

    task_id = uuid.uuid4().hex
    progress = NoticeProgress()
    progress.set(1, "任务已创建，正在启动...")
    task = {
        "task_id": task_id,
        "running": True,
        "done": False,
        "success": False,
        "message": "任务已创建，正在启动...",
        "progress": progress,
        "result": None,
        "target_path": str(target_path),
        "target_key": target_key,
        "created_at": time.time(),
        "finished_at": None,
    }
    with _NOTICE_TASK_LOCK:
        _prune_notice_tasks_locked()
        existing = _active_notice_task_for_target_locked(target_path)
        if existing:
            snapshot = existing["progress"].snapshot()
            return {
                "success": True,
                "already_running": True,
                "task_id": str(existing.get("task_id") or ""),
                "running": True,
                "done": False,
                "message": "该目标已有任务正在运行，已连接到原任务",
                "progress": snapshot["progress"],
                "logs": snapshot["logs"],
                "processed": snapshot["processed"],
                "total_reports": snapshot["total_reports"],
            }
        active_count = sum(
            1
            for existing in _NOTICE_TASKS.values()
            if bool(existing.get("running")) and not bool(existing.get("done"))
        )
        if active_count >= NOTICE_TASK_MAX_ACTIVE:
            return {
                "success": False,
                "message": f"已有 {active_count} 个通报任务正在运行，请等待任务结束后重试",
                "logs": [],
            }
        _NOTICE_TASKS[task_id] = task

    worker = threading.Thread(target=_notice_task_worker, args=(task_id, dict(payload)), daemon=True)
    task["thread"] = worker
    worker.start()
    snapshot = progress.snapshot()
    return {
        "success": True,
        "task_id": task_id,
        "running": True,
        "done": False,
        "message": snapshot["message"],
        "progress": snapshot["progress"],
        "logs": snapshot["logs"],
        "processed": 0,
        "total_reports": 0,
    }


def _doc_notice_process_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    task_id = _required_text(payload, "task_id", "缺少任务ID")
    with _NOTICE_TASK_LOCK:
        _prune_notice_tasks_locked()
        task = _NOTICE_TASKS.get(task_id)
        if not task:
            return {
                "success": False,
                "task_id": task_id,
                "done": True,
                "running": False,
                "message": "任务不存在或已过期，可能是后端已经重启",
                "error_code": "notice_task_not_found",
                "logs": [],
            }
        progress = task["progress"]
        snapshot = progress.snapshot()
        result = task.get("result") or {}
        done = bool(task.get("done"))
        running = bool(task.get("running"))
        success = bool(task.get("success")) if done else True
        message = str(task.get("message") or snapshot.get("message") or "")

    response = {
        "success": success,
        "task_id": task_id,
        "running": running,
        "done": done,
        "message": message,
        "progress": 100 if done and success else snapshot["progress"],
        "logs": snapshot["logs"],
        "processed": result.get("processed", snapshot["processed"]),
        "total_reports": result.get("total_reports", snapshot["total_reports"]),
        "target_path": result.get("target_path"),
        "generated_files": result.get("generated_files", []),
        "manual_files": result.get("manual_files", []),
        "failures": result.get("failures", []),
        "pdf_outputs": result.get("pdf_outputs", []),
        "error": task.get("error"),
    }
    if done:
        response["result"] = result
    return response


def _doc_notice_convert_failed_pdf(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_path = Path(_required_text(payload, "target_path", "请选择要转换的目录")).expanduser()
    failed_files = payload.get("failed_files") if "failed_files" in payload else None
    logs: List[str] = []
    if not target_path.exists() or not target_path.is_dir():
        return {"success": False, "message": f"目标目录不存在或不是目录: {target_path}", "logs": []}
    conflict = _notice_active_conflict(target_path, "转换PDF")
    if conflict:
        return conflict
    if failed_files is not None and not isinstance(failed_files, list):
        return {"success": False, "message": "失败文件列表格式无效", "error_code": "invalid_failed_files", "logs": []}

    candidates = _collect_notice_pdf_candidates(
        target_path,
        failed_files,
        logs,
        scan_target=bool(payload.get("scan_target", failed_files is None)),
    )
    file_map = [(src, src.with_suffix(".pdf")) for src in candidates]
    if not file_map:
        return {"success": False, "message": "未找到可转换的Word文档", "logs": logs}

    buffer = io.StringIO()
    try:
        with contextlib.redirect_stdout(buffer):
            converted, skipped, failures = convert_with_word_com(file_map, overwrite=True)
    except RuntimeError as exc:
        logs.extend(_captured_lines(buffer))
        return {
            "success": False,
            "message": str(exc),
            "converted": 0,
            "skipped": 0,
            "failures": [],
            "output_files": [],
            "deleted_files": [],
            "logs": logs,
        }

    logs.extend(_captured_lines(buffer))
    logs.extend(f"转换失败 {src.name}: {reason}" for src, reason in failures)
    failures = list(failures)
    failed_paths = {src.resolve() for src, _ in failures}
    output_files: List[str] = []
    deleted_files: List[str] = []
    for src, dst in file_map:
        if src.resolve() in failed_paths or not dst.exists():
            continue
        output_files.append(str(dst.resolve()))
        try:
            src.unlink()
            deleted_files.append(str(src.resolve()))
            logs.append(f"PDF转换成功，已删除原Word文件: {src}")
        except Exception as exc:
            failures.append((src, f"PDF已生成，但删除原Word失败: {exc}"))
            logs.append(f"PDF已生成，但删除原Word失败 {src}: {exc}")

    return {
        "success": len(failures) == 0,
        "message": f"转换完成：成功 {converted}，跳过 {skipped}，失败 {len(failures)}，删除原Word {len(deleted_files)} 个",
        "converted": converted,
        "skipped": skipped,
        "failures": _failure_dicts(failures),
        "output_files": output_files,
        "deleted_files": deleted_files,
        "logs": logs,
    }


def _doc_open_path(payload: Dict[str, Any]) -> Dict[str, Any]:
    target = Path(_required_text(payload, "path", "请选择要打开的路径")).expanduser()
    if not target.exists():
        return {"success": False, "message": f"路径不存在: {target}", "path": str(target)}
    open_target = target.parent if target.is_file() else target
    opened, error = _open_path_in_system(open_target)
    return {
        "success": opened,
        "message": f"已打开: {open_target}" if opened else f"无法打开: {error}",
        "path": str(open_target),
    }
