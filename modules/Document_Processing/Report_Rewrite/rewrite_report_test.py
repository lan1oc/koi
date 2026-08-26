#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import unittest

from docx import Document

from modules.Document_Processing.Report_Rewrite.rewrite_report import (
    REWRITTEN_NOTICE_METADATA,
    _find_insert_marker,
    _mark_rewritten_notice,
    _replace_cc_placeholder,
)


class RewriteReportCopyToTest(unittest.TestCase):
    def test_copy_to_placeholder_is_not_used_as_insert_marker(self) -> None:
        doc = Document()
        doc.add_paragraph("抄送：*")
        body_marker = doc.add_paragraph("*")

        _, marker_element, _ = _find_insert_marker(doc)

        self.assertIs(marker_element, body_marker._element)

    def test_copy_to_placeholder_supports_township_and_blank_values(self) -> None:
        classified = Document()
        classified.add_paragraph("抄送：*")
        self.assertTrue(_replace_cc_placeholder(classified, "南部商务区"))
        self.assertEqual(classified.paragraphs[0].text, "抄送：南部商务区")

        unclassified = Document()
        unclassified.add_paragraph("抄送：*")
        self.assertTrue(_replace_cc_placeholder(unclassified))
        self.assertEqual(unclassified.paragraphs[0].text, "抄送：")

    def test_rewrite_marker_is_persisted_in_document_properties(self) -> None:
        doc = Document()

        _mark_rewritten_notice(doc)
        _mark_rewritten_notice(doc)

        markers = {item.strip() for item in doc.core_properties.comments.split(";")}
        self.assertIn(REWRITTEN_NOTICE_METADATA, markers)
        self.assertEqual(doc.core_properties.comments.count(REWRITTEN_NOTICE_METADATA), 1)


if __name__ == "__main__":
    unittest.main()
