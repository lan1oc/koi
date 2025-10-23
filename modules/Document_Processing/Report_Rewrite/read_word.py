#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word文档内容读取工具
支持读取.docx格式的Word文档
"""

import sys
from docx import Document

# 设置Windows控制台编码为UTF-8
if sys.platform == 'win32':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass


def read_word_document(file_path):
    """
    读取Word文档的内容
    
    参数:
        file_path: Word文档的路径
        
    返回:
        文档内容的字符串
    """
    try:
        # 打开Word文档
        doc = Document(file_path)
        
        # 存储文档内容
        content = []
        
        print(f"\n正在读取文档: {file_path}")
        print("=" * 60)
        
        # 读取所有段落
        for i, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text
            if text.strip():  # 只添加非空段落
                content.append(text)
                print(f"\n段落 {i}:")
                print(text)
        
        print("\n" + "=" * 60)
        print(f"总共读取了 {len(content)} 个段落")
        
        # 返回完整内容
        return "\n".join(content)
        
    except FileNotFoundError:
        print(f"错误: 找不到文件 '{file_path}'")
        return None
    except Exception as e:
        print(f"读取文档时出错: {str(e)}")
        return None


def read_word_with_tables(file_path):
    """
    读取Word文档的内容（包含表格）
    
    参数:
        file_path: Word文档的路径
    """
    try:
        doc = Document(file_path)
        
        print(f"\n正在读取文档（包含表格）: {file_path}")
        print("=" * 60)
        
        # 读取段落
        print("\n【文档段落】")
        for i, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text
            if text.strip():
                print(f"\n段落 {i}: {text}")
        
        # 读取表格
        if doc.tables:
            print("\n" + "=" * 60)
            print("【文档表格】")
            for table_num, table in enumerate(doc.tables, 1):
                print(f"\n表格 {table_num}:")
                for row_num, row in enumerate(table.rows, 1):
                    row_data = [cell.text for cell in row.cells]
                    print(f"  行 {row_num}: {' | '.join(row_data)}")
        
        print("\n" + "=" * 60)
        
    except FileNotFoundError:
        print(f"错误: 找不到文件 '{file_path}'")
    except Exception as e:
        print(f"读取文档时出错: {str(e)}")


def save_word_to_txt(word_file, output_file=None):
    """
    将Word文档内容保存为文本文件
    
    参数:
        word_file: Word文档的路径
        output_file: 输出文本文件的路径（可选，默认为原文件名.txt）
    """
    try:
        doc = Document(word_file)
        
        # 如果没有指定输出文件，使用默认名称
        if output_file is None:
            output_file = word_file.rsplit('.', 1)[0] + '.txt'
        
        # 提取所有段落
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        
        # 写入文本文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        print(f"成功将内容保存到: {output_file}")
        return True
        
    except Exception as e:
        print(f"保存文件时出错: {str(e)}")
        return False


if __name__ == "__main__":
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("=" * 60)
        print("Word文档读取工具")
        print("=" * 60)
        print("\n使用方法:")
        print("  python read_word.py <Word文档路径>")
        print("  python read_word.py <Word文档路径> --table")
        print("  python read_word.py <Word文档路径> --save [输出文件]")
        print("\n功能说明:")
        print("  默认: 读取并显示所有非空段落")
        print("  --table: 读取并显示段落和表格")
        print("  --save: 将内容保存为文本文件")
        print("\n示例:")
        print("  python read_word.py document.docx")
        print("  python read_word.py document.docx --table")
        print("  python read_word.py document.docx --save")
        print("  python read_word.py document.docx --save output.txt")
        print("\n通报改写功能:")
        print("  请使用 rewrite_report.py 脚本")
        print("  python rewrite_report.py <源通报文档>")
        print("=" * 60)
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # 检查是否需要读取表格
    if len(sys.argv) > 2 and sys.argv[2] == "--table":
        read_word_with_tables(file_path)
    # 检查是否需要保存为文本文件
    elif len(sys.argv) > 2 and sys.argv[2] == "--save":
        output_file = sys.argv[3] if len(sys.argv) > 3 else None
        save_word_to_txt(file_path, output_file)
    else:
        # 普通读取
        read_word_document(file_path)
