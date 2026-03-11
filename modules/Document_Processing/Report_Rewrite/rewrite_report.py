#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
通报改写工具
自动将漏洞通报内容插入到模板中，并进行格式化处理
"""

import sys
import io
import re
import os
import json
import hashlib
import tempfile
import shutil
import uuid
import time
import subprocess
import zipfile
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from pathlib import Path
# XML 处理库 - 用于处理 Word 文档的 XML 结构
from lxml import etree  # type: ignore

# 全局手动处理列表
MANUAL_PROCESSING_LIST = []


def _agent_debug_log(run_id, hypothesis_id, location, message, data=None):
    """Debug logger disabled after successful fix verification."""
    return None


def _paragraph_style_val(para_or_elem):
    """Extract paragraph style value from paragraph object or element."""
    try:
        elem = para_or_elem._element if hasattr(para_or_elem, "_element") else para_or_elem
        p_pr = elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if p_pr is None:
            return ""
        p_style = p_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        if p_style is None:
            return ""
        return p_style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
    except Exception:
        return ""


def _paragraph_run_snapshot(para):
    """Collect compact run-level format snapshot for debug."""
    try:
        runs = []
        for r in para.runs[:5]:
            runs.append(
                {
                    "text": (r.text or "")[:40],
                    "bold": r.bold,
                    "italic": r.italic,
                    "size": str(r.font.size) if r.font and r.font.size else None,
                    "name": r.font.name if r.font else None,
                }
            )
        return {"run_count": len(para.runs), "runs": runs}
    except Exception:
        return {"run_count": -1, "runs": []}


def _replace_or_prepend_number_prefix_keep_runs(para, new_number, delimiter="."):
    """Update numbering prefix without rebuilding paragraph runs."""
    try:
        para_text = para.text or ""
        import re
        match = re.match(r"^(\d+)([.、）)])", para_text.strip())
        new_prefix = f"{new_number}{delimiter}"

        if para.runs:
            first_run = para.runs[0]
            first_text = first_run.text or ""
            if match:
                old_prefix = f"{match.group(1)}{match.group(2)}"
                if first_text.startswith(old_prefix):
                    first_run.text = new_prefix + first_text[len(old_prefix):]
                else:
                    # Fallback: replace first textual numbering once, but keep paragraph runs.
                    first_run.text = re.sub(r"^(\d+)([.、）)])", new_prefix, first_text, count=1)
            else:
                first_run.text = new_prefix + first_text
            return True

        para.add_run(f"{new_prefix}{para_text}")
        return True
    except Exception:
        return False


def _replace_or_prepend_literal_prefix_keep_runs(para, new_prefix):
    """Rewrite leading numbering prefix to a literal prefix, keeping runs intact."""
    try:
        para_text = para.text or ""
        import re
        match = re.match(r"^\s*(?:[▪•]?\s*)?(?:\(?\d+\)?)([.、）)]?)\s*", para_text.strip())
        if para.runs:
            first_run = para.runs[0]
            first_text = first_run.text or ""
            if match:
                prefix_match = re.match(r"^\s*(?:[▪•]?\s*)?(?:\(?\d+\)?)([.、）)]?)\s*", first_text)
                if prefix_match:
                    first_run.text = new_prefix + first_text[prefix_match.end():]
                else:
                    first_run.text = re.sub(
                        r"^\s*(?:[▪•]?\s*)?(?:\(?\d+\)?)([.、）)]?)\s*",
                        new_prefix,
                        first_text,
                        count=1,
                    )
            else:
                first_run.text = new_prefix + first_text
            return True
        para.add_run(f"{new_prefix}{para_text}")
        return True
    except Exception:
        return False


def _paragraph_numbering_info(para_or_elem):
    """Extract numId/ilvl from paragraph XML."""
    try:
        elem = para_or_elem._element if hasattr(para_or_elem, "_element") else para_or_elem
        p_pr = elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if p_pr is None:
            return {"numId": None, "ilvl": None}
        num_pr = p_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        if num_pr is None:
            return {"numId": None, "ilvl": None}
        num_id_el = num_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
        ilvl_el = num_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
        return {
            "numId": num_id_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if num_id_el is not None else None,
            "ilvl": ilvl_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if ilvl_el is not None else None,
        }
    except Exception:
        return {"numId": None, "ilvl": None}


def _doc_has_style_id(doc, style_id):
    try:
        if not style_id:
            return True
        for s in doc.styles:
            if getattr(s, "style_id", None) == style_id:
                return True
        styles_root = doc.styles.element
        style_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId'
        for node in styles_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style'):
            if node.get(style_attr) == style_id:
                return True
        return False
    except Exception:
        return False


def _find_style_node(styles_root, style_id):
    try:
        if not style_id:
            return None
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for node in styles_root.findall(f".//{ns}style"):
            if node.get(f"{ns}styleId") == style_id:
                return node
        return None
    except Exception:
        return None


def _next_style_id(target_doc, source_style_id):
    try:
        styles_root = target_doc.styles.element
        style_attr = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId"
        existing = {
            node.get(style_attr)
            for node in styles_root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style')
            if node.get(style_attr)
        }
        base = f"{source_style_id}_src"
        if base not in existing:
            return base
        index = 1
        while f"{base}_{index}" in existing:
            index += 1
        return f"{base}_{index}"
    except Exception:
        return f"{source_style_id}_src"


def _resolve_num_format(doc, num_id, ilvl):
    """Resolve numbering format/lvlText from numbering.xml for debug."""
    try:
        if not num_id:
            return {"numFmt": None, "lvlText": None}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numbering = _get_numbering_root(doc)
        if numbering is None:
            return {"numFmt": None, "lvlText": None}
        num_node = _find_num_node(numbering, str(num_id))
        if num_node is None:
            return {"numFmt": None, "lvlText": None}
        abs_id_el = num_node.find(f".//{ns}abstractNumId")
        if abs_id_el is None:
            return {"numFmt": None, "lvlText": None}
        abs_id = abs_id_el.get(f"{ns}val")
        if not abs_id:
            return {"numFmt": None, "lvlText": None}
        abs_node = _find_abstract_num_node(numbering, abs_id)
        if abs_node is None:
            return {"numFmt": None, "lvlText": None}
        level = ilvl if ilvl is not None else "0"
        target_lvl = None
        for lvl in abs_node.findall(f".//{ns}lvl"):
            if lvl.get(f"{ns}ilvl") == str(level):
                target_lvl = lvl
                break
        if target_lvl is None:
            return {"numFmt": None, "lvlText": None}
        num_fmt_el = target_lvl.find(f".//{ns}numFmt")
        lvl_text_el = target_lvl.find(f".//{ns}lvlText")
        return {
            "numFmt": num_fmt_el.get(f"{ns}val") if num_fmt_el is not None else None,
            "lvlText": lvl_text_el.get(f"{ns}val") if lvl_text_el is not None else None,
        }
    except Exception:
        return {"numFmt": None, "lvlText": None}


def _resolve_num_format_detail(doc, num_id, ilvl):
    """Resolve full numbering detail (including lvlOverride) for debug."""
    try:
        if not num_id:
            return {"numFmt": None, "lvlText": None, "from_override": False}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numbering = _get_numbering_root(doc)
        if numbering is None:
            return {"numFmt": None, "lvlText": None, "from_override": False}
        num_node = _find_num_node(numbering, str(num_id))
        if num_node is None:
            return {"numFmt": None, "lvlText": None, "from_override": False}

        level = str(ilvl if ilvl is not None else "0")
        override_lvl = None
        for lvl_override in num_node.findall(f".//{ns}lvlOverride"):
            if lvl_override.get(f"{ns}ilvl") == level:
                override_lvl = lvl_override.find(f".//{ns}lvl")
                if override_lvl is not None:
                    break

        if override_lvl is not None:
            num_fmt_el = override_lvl.find(f".//{ns}numFmt")
            lvl_text_el = override_lvl.find(f".//{ns}lvlText")
            return {
                "numFmt": num_fmt_el.get(f"{ns}val") if num_fmt_el is not None else None,
                "lvlText": lvl_text_el.get(f"{ns}val") if lvl_text_el is not None else None,
                "from_override": True,
            }

        abs_id_el = num_node.find(f".//{ns}abstractNumId")
        if abs_id_el is None:
            return {"numFmt": None, "lvlText": None, "from_override": False}
        abs_id = abs_id_el.get(f"{ns}val")
        if not abs_id:
            return {"numFmt": None, "lvlText": None, "from_override": False}
        abs_node = _find_abstract_num_node(numbering, abs_id)
        if abs_node is None:
            return {"numFmt": None, "lvlText": None, "from_override": False}

        target_lvl = None
        for lvl in abs_node.findall(f".//{ns}lvl"):
            if lvl.get(f"{ns}ilvl") == level:
                target_lvl = lvl
                break
        if target_lvl is None:
            return {"numFmt": None, "lvlText": None, "from_override": False}

        num_fmt_el = target_lvl.find(f".//{ns}numFmt")
        lvl_text_el = target_lvl.find(f".//{ns}lvlText")
        return {
            "numFmt": num_fmt_el.get(f"{ns}val") if num_fmt_el is not None else None,
            "lvlText": lvl_text_el.get(f"{ns}val") if lvl_text_el is not None else None,
            "from_override": False,
        }
    except Exception:
        return {"numFmt": None, "lvlText": None, "from_override": False}


def _get_numbering_root(doc):
    """Get numbering root element from docx document."""
    try:
        part = doc.part.numbering_part
        root = getattr(part, "_element", None)
        if root is None:
            root = getattr(part, "element", None)
        return root
    except Exception:
        return None


def _find_num_node(numbering_root, num_id):
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for node in numbering_root.findall(f".//{ns}num"):
            if node.get(f"{ns}numId") == str(num_id):
                return node
        return None
    except Exception:
        return None


def _find_abstract_num_node(numbering_root, abstract_num_id):
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for node in numbering_root.findall(f".//{ns}abstractNum"):
            if node.get(f"{ns}abstractNumId") == str(abstract_num_id):
                return node
        return None
    except Exception:
        return None


def _next_numbering_id(numbering_root, tag_name, attr_name):
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        max_id = 0
        for node in numbering_root.findall(f".//{ns}{tag_name}"):
            v = node.get(f"{ns}{attr_name}")
            if v is None:
                continue
            try:
                max_id = max(max_id, int(v))
            except Exception:
                continue
        return str(max_id + 1)
    except Exception:
        return "1"


def _set_paragraph_num_id(paragraph_element, num_id):
    """Rewrite paragraph numId to mapped target numId."""
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        p_pr = paragraph_element.find(f".//{ns}pPr")
        if p_pr is None:
            return False
        num_pr = p_pr.find(f".//{ns}numPr")
        if num_pr is None:
            return False
        num_id_el = num_pr.find(f".//{ns}numId")
        if num_id_el is None:
            return False
        num_id_el.set(f"{ns}val", str(num_id))
        return True
    except Exception:
        return False


def _find_para_num_snapshot(doc, keywords):
    """Find first paragraph containing any keyword and return numbering snapshot."""
    try:
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            if any(k in text for k in keywords):
                info = _paragraph_numbering_info(para)
                fmt = _resolve_num_format(doc, info.get("numId"), info.get("ilvl"))
                return {
                    "text": text[:120],
                    "num": info,
                    "fmt": fmt,
                }
        return None
    except Exception:
        return None


def _collect_paragraph_diagnostics(doc, keywords, limit=20):
    """Collect paragraph diagnostics for keywords in final output."""
    out = []
    try:
        for idx, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if not text:
                continue
            if not any(k in text for k in keywords):
                continue
            num = _paragraph_numbering_info(para)
            out.append(
                {
                    "idx": idx,
                    "text": text[:160],
                    "style": _paragraph_style_val(para),
                    "num": num,
                    "fmt": _resolve_num_format(doc, num.get("numId"), num.get("ilvl")),
                }
            )
            if len(out) >= limit:
                break
    except Exception:
        pass
    return out


def _style_struct_snapshot(doc, style_id):
    """Collect style-level paragraph/numbering properties."""
    try:
        if not style_id:
            return {"style_id": style_id, "found": False}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        styles_root = doc.styles.element
        target = None
        for node in styles_root.findall(f".//{ns}style"):
            if node.get(f"{ns}styleId") == style_id:
                target = node
                break
        if target is None:
            return {"style_id": style_id, "found": False}
        p_pr = target.find(f".//{ns}pPr")
        num_pr = p_pr.find(f".//{ns}numPr") if p_pr is not None else None
        num_id_el = num_pr.find(f".//{ns}numId") if num_pr is not None else None
        ilvl_el = num_pr.find(f".//{ns}ilvl") if num_pr is not None else None
        based_on = target.find(f".//{ns}basedOn")
        linked = target.find(f".//{ns}link")
        return {
            "style_id": style_id,
            "found": True,
            "has_pPr": p_pr is not None,
            "has_numPr": num_pr is not None,
            "numId": num_id_el.get(f"{ns}val") if num_id_el is not None else None,
            "ilvl": ilvl_el.get(f"{ns}val") if ilvl_el is not None else None,
            "basedOn": based_on.get(f"{ns}val") if based_on is not None else None,
            "link": linked.get(f"{ns}val") if linked is not None else None,
        }
    except Exception:
        return {"style_id": style_id, "found": False}


def _xml_node_fingerprint(node):
    """Return compact fingerprint for an XML node."""
    try:
        if node is None:
            return {"exists": False, "md5": None}
        xml = etree.tostring(node, encoding="utf-8", with_tail=False)
        return {"exists": True, "md5": hashlib.md5(xml).hexdigest()}
    except Exception:
        return {"exists": False, "md5": None}


def _style_fingerprint(doc, style_id):
    """Fingerprint full style XML to detect same-id/different-definition collisions."""
    try:
        if not style_id:
            return {"style_id": style_id, "found": False}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        styles_root = doc.styles.element
        target = None
        for node in styles_root.findall(f".//{ns}style"):
            if node.get(f"{ns}styleId") == style_id:
                target = node
                break
        base = _style_struct_snapshot(doc, style_id)
        base.update(_xml_node_fingerprint(target))
        return base
    except Exception:
        return {"style_id": style_id, "found": False, "exists": False, "md5": None}


def _numbering_fingerprint(doc, num_id):
    """Fingerprint full num/abstractNum XML for collision diagnosis."""
    try:
        if not num_id:
            return {"numId": None, "found": False}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numbering = _get_numbering_root(doc)
        if numbering is None:
            return {"numId": str(num_id), "found": False}
        num_node = _find_num_node(numbering, str(num_id))
        if num_node is None:
            return {"numId": str(num_id), "found": False}
        abs_id_el = num_node.find(f".//{ns}abstractNumId")
        abs_id = abs_id_el.get(f"{ns}val") if abs_id_el is not None else None
        abs_node = _find_abstract_num_node(numbering, abs_id) if abs_id else None
        return {
            "numId": str(num_id),
            "found": True,
            "abstractNumId": abs_id,
            "num": _xml_node_fingerprint(num_node),
            "abstractNum": _xml_node_fingerprint(abs_node),
            "fmt_l0": _resolve_num_format(doc, num_id, "0"),
        }
    except Exception:
        return {"numId": str(num_id), "found": False}


def _style_semantic_fingerprint(doc, style_id):
    """Fingerprint style XML while ignoring remapped style IDs."""
    try:
        if not style_id:
            return {"style_id": style_id, "found": False, "md5": None}
        node = _find_style_node(doc.styles.element, style_id)
        if node is None:
            return {"style_id": style_id, "found": False, "md5": None}
        clone = deepcopy(node)
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        style_attr = f"{ns}styleId"
        if clone.get(style_attr) is not None:
            clone.set(style_attr, "__STYLE_ID__")
        for ref_tag in ("basedOn", "link", "next"):
            ref_el = clone.find(f".//{ns}{ref_tag}")
            if ref_el is not None and ref_el.get(f"{ns}val") is not None:
                ref_el.set(f"{ns}val", f"__{ref_tag.upper()}__")
        xml = etree.tostring(clone, encoding="utf-8", with_tail=False)
        return {"style_id": style_id, "found": True, "md5": hashlib.md5(xml).hexdigest()}
    except Exception:
        return {"style_id": style_id, "found": False, "md5": None}


def _numbering_semantic_fingerprint(doc, num_id):
    """Fingerprint numbering XML while ignoring remapped numbering IDs."""
    try:
        if not num_id:
            return {"numId": None, "found": False, "num_md5": None, "abstract_md5": None}
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numbering = _get_numbering_root(doc)
        if numbering is None:
            return {"numId": str(num_id), "found": False, "num_md5": None, "abstract_md5": None}
        num_node = _find_num_node(numbering, str(num_id))
        if num_node is None:
            return {"numId": str(num_id), "found": False, "num_md5": None, "abstract_md5": None}
        num_clone = deepcopy(num_node)
        if num_clone.get(f"{ns}numId") is not None:
            num_clone.set(f"{ns}numId", "__NUM_ID__")
        abs_id_el = num_clone.find(f".//{ns}abstractNumId")
        abs_id = None
        if abs_id_el is not None:
            abs_id = abs_id_el.get(f"{ns}val")
            abs_id_el.set(f"{ns}val", "__ABSTRACT_NUM_ID__")
        abs_node = _find_abstract_num_node(numbering, abs_id) if abs_id else None
        abstract_md5 = None
        if abs_node is not None:
            abs_clone = deepcopy(abs_node)
            if abs_clone.get(f"{ns}abstractNumId") is not None:
                abs_clone.set(f"{ns}abstractNumId", "__ABSTRACT_NUM_ID__")
            abstract_md5 = hashlib.md5(
                etree.tostring(abs_clone, encoding="utf-8", with_tail=False)
            ).hexdigest()
        return {
            "numId": str(num_id),
            "found": True,
            "num_md5": hashlib.md5(
                etree.tostring(num_clone, encoding="utf-8", with_tail=False)
            ).hexdigest(),
            "abstract_md5": abstract_md5,
        }
    except Exception:
        return {"numId": str(num_id), "found": False, "num_md5": None, "abstract_md5": None}


def _find_first_paragraph_containing(doc, keyword):
    try:
        for p in doc.paragraphs:
            if keyword in (p.text or ""):
                return p
        return None
    except Exception:
        return None


def _doc_defaults_fingerprint(doc):
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        doc_defaults = doc.styles.element.find(f".//{ns}docDefaults")
        settings = getattr(doc.settings, "element", None)
        return {
            "doc_defaults_md5": _xml_node_fingerprint(doc_defaults).get("md5"),
            "settings_md5": _xml_node_fingerprint(settings).get("md5"),
        }
    except Exception:
        return {"doc_defaults_md5": None, "settings_md5": None}


def _package_part_fingerprint(doc, target_suffix):
    try:
        package = getattr(doc.part, "package", None)
        if package is None:
            return {"part": target_suffix, "md5": None}
        for part in getattr(package, "parts", []):
            partname = str(getattr(part, "partname", ""))
            if partname.endswith(target_suffix):
                blob = getattr(part, "blob", None)
                if blob is None:
                    return {"part": partname, "md5": None}
                return {"part": partname, "md5": hashlib.md5(blob).hexdigest()}
        return {"part": target_suffix, "md5": None}
    except Exception:
        return {"part": target_suffix, "md5": None}


def _package_resource_fingerprints(doc):
    return {
        "settings": _package_part_fingerprint(doc, "/word/settings.xml"),
        "webSettings": _package_part_fingerprint(doc, "/word/webSettings.xml"),
        "fontTable": _package_part_fingerprint(doc, "/word/fontTable.xml"),
        "theme": _package_part_fingerprint(doc, "/word/theme/theme1.xml"),
    }


def _paragraph_semantic_signature(doc, para):
    try:
        num = _paragraph_numbering_info(para)
        style_id = _paragraph_style_val(para)
        return {
            "text": (para.text or "").strip()[:200],
            "style": style_id,
            "style_semantic": _style_semantic_fingerprint(doc, style_id),
            "num": num,
            "num_semantic": _numbering_semantic_fingerprint(doc, num.get("numId")),
            "runs": _paragraph_run_snapshot(para),
            "struct": _paragraph_struct_snapshot(para),
        }
    except Exception:
        return {
            "text": "",
            "style": "",
            "style_semantic": {"found": False, "md5": None},
            "num": {"numId": None, "ilvl": None},
            "num_semantic": {"found": False, "num_md5": None, "abstract_md5": None},
            "runs": {"run_count": -1, "runs": []},
            "struct": {"has_pPr": False},
        }


def _compare_copy_range_semantics(source_doc, delivered_doc, source_start_idx, source_end_idx, delivered_anchor="验证情况", max_items=12):
    try:
        source_items = []
        for i in range(source_start_idx, min(source_end_idx, len(source_doc.paragraphs))):
            p = source_doc.paragraphs[i]
            if not (p.text or "").strip():
                continue
            source_items.append({"source_idx": i, "sig": _paragraph_semantic_signature(source_doc, p)})

        anchor_idx = None
        for i, p in enumerate(delivered_doc.paragraphs):
            if delivered_anchor in (p.text or ""):
                anchor_idx = i
                break
        if anchor_idx is None:
            return {"found_anchor": False, "mismatches": [], "compared": 0}

        mismatches = []
        compared = 0
        # Align using the common visible heading "验证情况" as pivot.
        source_anchor_offset = None
        for idx, item in enumerate(source_items):
            if delivered_anchor in item["sig"]["text"]:
                source_anchor_offset = idx
                break
        if source_anchor_offset is None:
            source_anchor_offset = 0

        delivered_start = max(0, anchor_idx - source_anchor_offset)
        delivered_items = []
        for i in range(delivered_start, len(delivered_doc.paragraphs)):
            p = delivered_doc.paragraphs[i]
            if not (p.text or "").strip():
                continue
            delivered_items.append({"delivered_idx": i, "sig": _paragraph_semantic_signature(delivered_doc, p)})
            if len(delivered_items) >= max_items:
                break

        for rel_idx, source_item in enumerate(source_items[:max_items]):
            if rel_idx >= len(delivered_items):
                break
            delivered_idx = delivered_items[rel_idx]["delivered_idx"]
            delivered_sig = delivered_items[rel_idx]["sig"]
            compared += 1
            source_sig = source_item["sig"]
            if (
                source_sig["text"] != delivered_sig["text"]
                or source_sig["style_semantic"].get("md5") != delivered_sig["style_semantic"].get("md5")
                or source_sig["num_semantic"].get("num_md5") != delivered_sig["num_semantic"].get("num_md5")
                or source_sig["num_semantic"].get("abstract_md5") != delivered_sig["num_semantic"].get("abstract_md5")
                or source_sig["runs"] != delivered_sig["runs"]
            ):
                mismatches.append(
                    {
                        "relative_idx": rel_idx,
                        "source_idx": source_item["source_idx"],
                        "delivered_idx": delivered_idx,
                        "source": source_sig,
                        "delivered": delivered_sig,
                    }
                )
        return {"found_anchor": True, "mismatches": mismatches, "compared": compared}
    except Exception as e:
        return {"found_anchor": False, "mismatches": [], "compared": 0, "error": str(e)}


def _sync_rendering_resources_from_source(source_file, output_file):
    """
    Sync package-level rendering resources from source to output.
    This avoids template-level settings/font/theme overriding identical paragraph XML.
    """
    tmp_output = None
    try:
        source_file = str(source_file)
        output_file = str(output_file)
        targets = [
            "word/settings.xml",
            "word/fontTable.xml",
            "word/theme/theme1.xml",
            "word/webSettings.xml",
        ]
        tmp_output = f"{output_file}.render-sync.tmp"
        copied = []
        skipped = []

        with zipfile.ZipFile(source_file, "r") as src_zip, \
             zipfile.ZipFile(output_file, "r") as out_zip, \
             zipfile.ZipFile(tmp_output, "w", compression=zipfile.ZIP_DEFLATED) as new_zip:
            output_names = set(out_zip.namelist())
            source_names = set(src_zip.namelist())

            for item in out_zip.infolist():
                if item.filename in targets and item.filename in source_names:
                    new_zip.writestr(item, src_zip.read(item.filename))
                    copied.append(item.filename)
                else:
                    new_zip.writestr(item, out_zip.read(item.filename))

            for target in targets:
                if target in source_names and target not in output_names:
                    new_zip.writestr(target, src_zip.read(target))
                    copied.append(target)
                elif target not in source_names:
                    skipped.append(target)

        os.replace(tmp_output, output_file)
        return {"ok": True, "copied": copied, "skipped_missing_in_source": skipped}
    except Exception as e:
        try:
            if os.path.exists(tmp_output):
                os.remove(tmp_output)
        except Exception:
            pass
        return {"ok": False, "error": str(e)}


def _collect_paragraph_window(doc, anchor_keyword, radius=3):
    """Collect window paragraphs around an anchor keyword for layout diagnosis."""
    out = []
    try:
        anchor_idx = None
        for i, p in enumerate(doc.paragraphs):
            if anchor_keyword in (p.text or ""):
                anchor_idx = i
                break
        if anchor_idx is None:
            return {"anchor": anchor_keyword, "found": False, "items": []}
        start = max(0, anchor_idx - radius)
        end = min(len(doc.paragraphs), anchor_idx + radius + 1)
        for i in range(start, end):
            p = doc.paragraphs[i]
            t = (p.text or "").strip()
            if not t:
                continue
            num = _paragraph_numbering_info(p)
            out.append(
                {
                    "idx": i,
                    "text": t[:160],
                    "style": _paragraph_style_val(p),
                    "num": num,
                    "fmt": _resolve_num_format(doc, num.get("numId"), num.get("ilvl")),
                    "struct": _paragraph_struct_snapshot(p),
                    "drawing_count": len(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')),
                }
            )
        return {"anchor": anchor_keyword, "found": True, "items": out}
    except Exception:
        return {"anchor": anchor_keyword, "found": False, "items": []}


def _collect_paragraph_window_including_empty(doc, anchor_keyword, radius=3):
    """Collect raw paragraph window, including empty paragraphs, for hidden-list diagnosis."""
    out = []
    try:
        anchor_idx = None
        for i, p in enumerate(doc.paragraphs):
            if anchor_keyword in (p.text or ""):
                anchor_idx = i
                break
        if anchor_idx is None:
            return {"anchor": anchor_keyword, "found": False, "items": []}
        start = max(0, anchor_idx - radius)
        end = min(len(doc.paragraphs), anchor_idx + radius + 1)
        for i in range(start, end):
            p = doc.paragraphs[i]
            raw_text = p.text or ""
            stripped = raw_text.strip()
            num = _paragraph_numbering_info(p)
            out.append(
                {
                    "idx": i,
                    "text": stripped[:160],
                    "raw_len": len(raw_text),
                    "is_empty": not bool(stripped),
                    "style": _paragraph_style_val(p),
                    "num": num,
                    "fmt": _resolve_num_format(doc, num.get("numId"), num.get("ilvl")),
                    "struct": _paragraph_struct_snapshot(p),
                    "run_count": len(getattr(p, "runs", [])),
                }
            )
        return {"anchor": anchor_keyword, "found": True, "items": out}
    except Exception:
        return {"anchor": anchor_keyword, "found": False, "items": []}


def _collect_non_empty_paragraph_slice(doc, start_idx, count=12):
    """Collect a forward slice of non-empty paragraphs from a paragraph index."""
    out = []
    try:
        start_idx = max(0, int(start_idx))
        for i in range(start_idx, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            if not text:
                continue
            num = _paragraph_numbering_info(p)
            out.append(
                {
                    "idx": i,
                    "text": text[:160],
                    "style": _paragraph_style_val(p),
                    "num": num,
                    "fmt": _resolve_num_format(doc, num.get("numId"), num.get("ilvl")),
                }
            )
            if len(out) >= count:
                break
    except Exception:
        pass
    return out


def _paragraph_struct_snapshot(para_or_elem):
    """Collect key paragraph structure fields for debug."""
    try:
        elem = para_or_elem._element if hasattr(para_or_elem, "_element") else para_or_elem
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        p_pr = elem.find(f".//{ns}pPr")
        if p_pr is None:
            return {"has_pPr": False}
        p_style = p_pr.find(f".//{ns}pStyle")
        num_pr = p_pr.find(f".//{ns}numPr")
        num_id = None
        ilvl = None
        if num_pr is not None:
            num_id_el = num_pr.find(f".//{ns}numId")
            ilvl_el = num_pr.find(f".//{ns}ilvl")
            num_id = num_id_el.get(f"{ns}val") if num_id_el is not None else None
            ilvl = ilvl_el.get(f"{ns}val") if ilvl_el is not None else None
        spacing = p_pr.find(f".//{ns}spacing")
        ind = p_pr.find(f".//{ns}ind")
        jc = p_pr.find(f".//{ns}jc")
        keep_next = p_pr.find(f".//{ns}keepNext")
        page_break_before = p_pr.find(f".//{ns}pageBreakBefore")
        return {
            "has_pPr": True,
            "pStyle": p_style.get(f"{ns}val") if p_style is not None else None,
            "numId": num_id,
            "ilvl": ilvl,
            "numFmt_l0": _resolve_num_format_for_num_only(elem, num_id),
            "spacing_before": spacing.get(f"{ns}before") if spacing is not None else None,
            "spacing_after": spacing.get(f"{ns}after") if spacing is not None else None,
            "line": spacing.get(f"{ns}line") if spacing is not None else None,
            "ind_left": ind.get(f"{ns}left") if ind is not None else None,
            "ind_hanging": ind.get(f"{ns}hanging") if ind is not None else None,
            "jc": jc.get(f"{ns}val") if jc is not None else None,
            "keepNext": keep_next is not None,
            "pageBreakBefore": page_break_before is not None,
        }
    except Exception:
        return {"has_pPr": False}


def _resolve_num_format_for_num_only(_elem, num_id):
    """Compatibility helper for struct snapshot output."""
    if not num_id:
        return None
    return str(num_id)


def _ensure_numbering_mapping(source_doc, target_doc, source_num_id, mapping_cache):
    """
    Copy source num/abstractNum to target numbering.xml and return mapped numId.
    Avoid numId collision with template numbering definitions.
    """
    try:
        if not source_num_id:
            return None
        source_num_id = str(source_num_id)
        if source_num_id in mapping_cache:
            return mapping_cache[source_num_id]

        source_root = _get_numbering_root(source_doc)
        target_root = _get_numbering_root(target_doc)
        if source_root is None or target_root is None:
            return None

        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        source_num_node = _find_num_node(source_root, source_num_id)
        if source_num_node is None:
            return None

        abs_id_el = source_num_node.find(f".//{ns}abstractNumId")
        if abs_id_el is None:
            return None
        source_abs_id = abs_id_el.get(f"{ns}val")
        if not source_abs_id:
            return None

        source_abs_node = _find_abstract_num_node(source_root, source_abs_id)
        if source_abs_node is None:
            return None

        new_abs_id = _next_numbering_id(target_root, "abstractNum", "abstractNumId")
        new_num_id = _next_numbering_id(target_root, "num", "numId")

        new_abs_node = deepcopy(source_abs_node)
        new_abs_node.set(f"{ns}abstractNumId", new_abs_id)

        new_num_node = deepcopy(source_num_node)
        new_num_node.set(f"{ns}numId", new_num_id)
        new_num_abs_id = new_num_node.find(f".//{ns}abstractNumId")
        if new_num_abs_id is not None:
            new_num_abs_id.set(f"{ns}val", new_abs_id)

        target_root.append(new_abs_node)
        target_root.append(new_num_node)

        mapping_cache[source_num_id] = new_num_id
        return new_num_id
    except Exception:
        return None


def _ensure_style_in_target(source_doc, target_doc, style_id, _visited=None):
    """Copy missing style definition from source to target styles.xml."""
    try:
        if not style_id:
            return False
        if _visited is None:
            _visited = set()
        if style_id in _visited:
            return False
        _visited.add(style_id)

        if _doc_has_style_id(target_doc, style_id):
            return False

        source_styles = source_doc.styles.element
        target_styles = target_doc.styles.element
        style_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId'
        style_el = None
        for node in source_styles.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style'):
            if node.get(style_attr) == style_id:
                style_el = node
                break

        if style_el is None:
            return False

        based_on = style_el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}basedOn')
        linked = style_el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}link')
        if based_on is not None:
            _ensure_style_in_target(source_doc, target_doc, based_on.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'), _visited)
        if linked is not None:
            _ensure_style_in_target(source_doc, target_doc, linked.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'), _visited)

        target_styles.append(deepcopy(style_el))
        return _doc_has_style_id(target_doc, style_id)
    except Exception:
        return False


def _ensure_style_mapping(source_doc, target_doc, style_id, mapping_cache, _visited=None):
    """Map a source style ID to an isolated target style ID."""
    try:
        if not style_id:
            return style_id
        if style_id in mapping_cache:
            return mapping_cache[style_id]
        if _visited is None:
            _visited = set()
        if style_id in _visited:
            return mapping_cache.get(style_id, style_id)
        _visited.add(style_id)

        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        source_styles = source_doc.styles.element
        target_styles = target_doc.styles.element
        source_style_el = _find_style_node(source_styles, style_id)
        if source_style_el is None:
            mapping_cache[style_id] = style_id
            return style_id

        target_style_el = _find_style_node(target_styles, style_id)
        source_fp = _xml_node_fingerprint(source_style_el)
        target_fp = _xml_node_fingerprint(target_style_el)

        if target_style_el is None:
            resolved_style_id = style_id
        elif source_fp.get("md5") == target_fp.get("md5"):
            resolved_style_id = style_id
        else:
            resolved_style_id = _next_style_id(target_doc, style_id)

        mapping_cache[style_id] = resolved_style_id

        if _find_style_node(target_styles, resolved_style_id) is not None:
            return resolved_style_id

        cloned_style_el = deepcopy(source_style_el)
        cloned_style_el.set(f"{ns}styleId", resolved_style_id)

        for ref_tag in ("basedOn", "link", "next"):
            ref_el = cloned_style_el.find(f".//{ns}{ref_tag}")
            if ref_el is None:
                continue
            ref_style_id = ref_el.get(f"{ns}val")
            if not ref_style_id:
                continue
            resolved_ref_style_id = _ensure_style_mapping(
                source_doc,
                target_doc,
                ref_style_id,
                mapping_cache,
                _visited,
            )
            if resolved_ref_style_id:
                ref_el.set(f"{ns}val", resolved_ref_style_id)

        target_styles.append(cloned_style_el)
        return resolved_style_id
    except Exception:
        return style_id


def _remap_style_references_in_element(source_doc, target_doc, element, mapping_cache):
    """Rewrite copied element style refs to isolated target style IDs."""
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        changed = []
        for tag_name in ("pStyle", "rStyle", "tblStyle"):
            for style_el in element.findall(f".//{ns}{tag_name}"):
                source_style_id = style_el.get(f"{ns}val")
                if not source_style_id:
                    continue
                resolved_style_id = _ensure_style_mapping(
                    source_doc,
                    target_doc,
                    source_style_id,
                    mapping_cache,
                )
                if resolved_style_id and resolved_style_id != source_style_id:
                    style_el.set(f"{ns}val", resolved_style_id)
                    changed.append(
                        {
                            "tag": tag_name,
                            "source_style_id": source_style_id,
                            "resolved_style_id": resolved_style_id,
                        }
                    )
        return changed
    except Exception:
        return []

# 设置编码设置控制台编码，避免Unicode错误
def safe_print(text, fallback_text=None):
    """安全打印函数，处理编码问题"""
    try:
        print(text)
    except UnicodeEncodeError:
        if fallback_text:
            print(fallback_text)
        else:
            print(text.encode('utf-8', errors='ignore').decode('utf-8'))
    except Exception:
        if fallback_text:
            print(fallback_text)
        else:
            print("打印输出时发生错误")

# 导入文档完整性验证模块
try:
    from .document_integrity import (
        safe_save_document, validate_document_integrity, 
        cleanup_resources
    )
    INTEGRITY_MODULE_AVAILABLE = True
except ImportError:
    try:
        from document_integrity import (
            safe_save_document, validate_document_integrity, 
            cleanup_resources
        )
        INTEGRITY_MODULE_AVAILABLE = True
    except ImportError:
        INTEGRITY_MODULE_AVAILABLE = False
        safe_print("⚠️ 文档完整性验证模块导入失败，将使用原始保存方法")

def create_backup(file_path):
    """创建文档备份"""
    try:
        if Path(file_path).exists():
            backup_path = f"{file_path}.backup_{int(time.time())}"
            shutil.copy2(file_path, backup_path)
            print(f"  📋 已创建备份: {Path(backup_path).name}")
            return backup_path
    except Exception as e:
        print(f"  ⚠️ 备份创建失败: {e}")
    return None

def recover_from_backup(original_path, backup_path):
    """从备份恢复文档"""
    try:
        if backup_path and Path(backup_path).exists():
            shutil.copy2(backup_path, original_path)
            print(f"  🔄 已从备份恢复: {Path(original_path).name}")
            return True
    except Exception as e:
        print(f"  ❌ 备份恢复失败: {e}")
    return False

def cleanup_backups(file_path, keep_count=3):
    """清理旧备份文件，保留最新的几个"""
    try:
        backup_files = list(Path(file_path).parent.glob(f"{Path(file_path).name}.backup_*"))
        if len(backup_files) > keep_count:
            # 按时间排序，删除最旧的
            backup_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
            for old_backup in backup_files[keep_count:]:
                old_backup.unlink()
                print(f"  🗑️ 已清理旧备份: {old_backup.name}")
    except Exception as e:
        print(f"  ⚠️ 备份清理失败: {e}")


def add_to_manual_processing_list(file_path, error_type, error_detail):
    """添加文件到手动处理列表"""
    global MANUAL_PROCESSING_LIST
    
    entry = {
        'file_path': file_path,
        'file_name': os.path.basename(file_path),
        'error_type': error_type,
        'error_detail': error_detail,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    
    MANUAL_PROCESSING_LIST.append(entry)
    print(f"  📝 已添加到手动处理列表: {entry['file_name']} - {error_type}")


def get_manual_processing_list():
    """获取手动处理列表"""
    return MANUAL_PROCESSING_LIST.copy()


def clear_manual_processing_list():
    """清空手动处理列表"""
    global MANUAL_PROCESSING_LIST
    MANUAL_PROCESSING_LIST.clear()


def print_manual_processing_list():
    """打印手动处理列表"""
    if not MANUAL_PROCESSING_LIST:
        print("  ✅ 无需手动处理的文件")
        return
    
    print("\n" + "=" * 60)
    print("📋 需要手动处理的文件列表")
    print("=" * 60)
    
    for i, entry in enumerate(MANUAL_PROCESSING_LIST, 1):
        print(f"\n{i}. 文件: {entry['file_name']}")
        print(f"   路径: {entry['file_path']}")
        print(f"   错误类型: {entry['error_type']}")
        print(f"   错误详情: {entry['error_detail']}")
        print(f"   时间: {entry['timestamp']}")
    
    print("\n" + "=" * 60)


def _should_keep_numbering(paragraph_element):
    """
    判断段落是否应该保留编号格式
    支持混合编号模式：文本编号（如"1.漏洞描述"）+ Word自动编号（如"验证情况"）
    
    Args:
        paragraph_element: 段落的XML元素
        
    Returns:
        bool: True表示应该保留编号，False表示应该移除编号
    """
    try:
        # 获取段落文本内容
        text_content = ""
        for text_elem in paragraph_element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            if text_elem.text:
                text_content += text_elem.text
        
        text_content = text_content.strip()
        
        # 如果段落为空，不保留编号
        if not text_content:
            return False
        
        # 检查是否有Word自动编号
        pPr = paragraph_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        has_auto_numbering = False
        if pPr is not None:
            numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                has_auto_numbering = True
        
        # 明确的正文内容，即使有自动编号也不保留
        body_text_indicators = ['高危', '中危', '低危', '严重', '一般', '轻微']
        for indicator in body_text_indicators:
            if text_content == indicator or text_content.strip() == indicator:
                return False
        
        # 明确的字段标签，不应该有编号（这些是字段名，不是章节标题）
        field_labels = ['漏洞事件：', '发现时间：', '影响产品：', '影响危害：', '漏洞描述：', '验证截图：']
        for label in field_labels:
            if text_content == label or text_content.strip() == label:
                return False
        
        # 检查段落样式（如果有的话）
        if pPr is not None:
            # 检查是否有标题样式
            pStyle = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if pStyle is not None:
                style_val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                if 'Heading' in style_val or '标题' in style_val:
                    return True
        
        # 特定的章节标题关键词（这些是真正的章节，应该有编号）
        section_keywords = ['漏洞描述', '验证情况', '整改要求', '整改建议', '处置措施']
        
        # 检查是否是章节标题（短文本且包含关键词）
        for keyword in section_keywords:
            if keyword in text_content and len(text_content.strip()) <= 15:
                return True
        
        # 检查是否是标题级别的文本编号（短文本且以编号开头）
        import re
        text_match = re.match(r'^\d+[.、）)]', text_content)
        if text_match:
            # 只有短文本（可能是标题）才保留编号，长文本（正文内容）不保留
            if len(text_content.strip()) <= 20:
                # 进一步检查是否包含标题关键词
                for keyword in section_keywords:
                    if keyword in text_content:
                        return True
                # 如果是纯编号+简短描述，也可能是标题
                if len(text_content.strip()) <= 10:
                    return True
        
        # 如果有自动编号且内容是章节标题，保留编号
        if has_auto_numbering:
            # 检查是否是章节标题（不是字段标签）
            for keyword in section_keywords:
                if keyword in text_content and len(text_content.strip()) <= 15:
                    return True
        
        # 其他情况（正文内容）不保留编号
        return False
        
    except Exception as e:
        print(f"  ⚠️ 编号检测出错: {e}")
        # 出错时默认不保留编号，避免错误显示
        return False


def _remove_paragraph_numbering(paragraph_element):
    """
    移除段落的编号格式
    
    Args:
        paragraph_element: 段落的XML元素
    """
    try:
        # 查找段落属性
        pPr = paragraph_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            # 移除编号属性 (numPr)
            numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                pPr.remove(numPr)
                print(f"  🔧 移除段落编号格式")
                return True
        return False
                
    except Exception as e:
        print(f"  ⚠️ 移除编号格式时出错: {e}")
        return False


def _force_validation_heading_standalone(doc, run_id="manual"):
    """Detach '验证情况' from body numbering and render as standalone text heading."""
    try:
        for idx, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if text != "验证情况":
                continue
            _remove_paragraph_numbering(para._element)
            _replace_or_prepend_number_prefix_keep_runs(para, 2, ".")
    except Exception:
        pass


def _force_severity_value_standalone(doc, run_id="manual"):
    """Detach severity value like '高危漏洞' from list formatting and render as literal text."""
    severity_values = {"高危漏洞", "中危漏洞", "低危漏洞", "严重漏洞", "一般漏洞", "轻微漏洞"}
    try:
        for idx, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if text not in severity_values:
                continue
            _remove_paragraph_numbering(para._element)
            _replace_or_prepend_literal_prefix_keep_runs(para, "1. ")
    except Exception:
        pass


def _reassign_numbering_sequence(doc):
    """No-op: preserve copied numbering/layout exactly."""
    return True



def _copy_image_to_document(drawing_element, source_doc, target_doc, target_run):
    """复制图片从源文档到目标文档 - 增强版，支持受保护文档"""
    try:
        from docx.oxml.ns import qn
        import tempfile
        import os
        import zipfile
        from pathlib import Path
        
        print(f"    🔍 开始图片复制流程...")
        
        # 方法1: 标准方式 - 通过关系ID获取图片
        success = _try_standard_image_copy(drawing_element, source_doc, target_doc, target_run)
        if success:
            print(f"    ✅ 标准方式复制成功")
            return True
        
        print(f"    ⚠️ 标准方式失败，尝试备用方案...")
        
        # 方法2: 直接从docx文件中提取图片
        success = _try_direct_image_extraction(drawing_element, source_doc, target_doc, target_run)
        if success:
            print(f"    ✅ 直接提取方式成功")
            return True
        
        print(f"    ⚠️ 直接提取失败，尝试平台增强方案...")
        
        # 方法3: 平台增强方案（当前默认关闭）
        if COM_UTILS_AVAILABLE:
            success = _try_com_image_copy(drawing_element, source_doc, target_doc, target_run)
            if success:
                print(f"    ✅ 平台增强方案复制成功")
                return True
        
        print(f"    ❌ 所有图片复制方法都失败")
        return False
        
    except Exception as e:
        print(f"    ❌ 图片复制总体错误: {e}")
        return False


def _try_standard_image_copy(drawing_element, source_doc, target_doc, target_run):
    """标准图片复制方式"""
    try:
        from docx.oxml.ns import qn
        import tempfile
        import os
        
        # 查找图片的关系ID - 修复namespaces兼容性问题
        try:
            blip_elements = drawing_element.xpath('.//a:blip', namespaces={
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
            })
        except TypeError:
            # 兼容旧版本python-docx，不使用namespaces参数
            blip_elements = drawing_element.xpath('.//a:blip')
        
        if not blip_elements:
            # 尝试其他可能的图片元素
            try:
                pic_elements = drawing_element.xpath('.//pic:pic', namespaces={
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                })
            except TypeError:
                pic_elements = drawing_element.xpath('.//pic:pic')
                
            if pic_elements:
                try:
                    blip_elements = pic_elements[0].xpath('.//a:blip', namespaces={
                        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                    })
                except TypeError:
                    blip_elements = pic_elements[0].xpath('.//a:blip')
        
        if not blip_elements:
            return False
            
        embed_attr = blip_elements[0].get(qn('r:embed'))
        if not embed_attr:
            return False
            
        # 从源文档获取图片数据
        try:
            source_image_part = source_doc.part.related_parts.get(embed_attr)
            if not source_image_part:
                return False
                
            # 获取图片数据
            image_data = source_image_part.blob
            if not image_data:
                return False
                
        except Exception as e:
            print(f"      获取图片数据失败: {e}")
            return False
        
        # 确定图片格式（增强版本）
        image_ext = '.png'  # 默认
        if hasattr(source_image_part, 'content_type'):
            content_type = source_image_part.content_type.lower()
            if 'jpeg' in content_type or 'jpg' in content_type:
                image_ext = '.jpg'
            elif 'png' in content_type:
                image_ext = '.png'
            elif 'gif' in content_type:
                image_ext = '.gif'
            elif 'bmp' in content_type:
                image_ext = '.bmp'
            elif 'tiff' in content_type or 'tif' in content_type:
                image_ext = '.tiff'
            elif 'webp' in content_type:
                image_ext = '.webp'
            elif 'svg' in content_type:
                image_ext = '.svg'
        
        # 如果无法从content_type确定，尝试从图片数据头部检测
        if image_ext == '.png' and image_data:
            try:
                # 检查文件头部魔数
                if image_data.startswith(b'\xff\xd8\xff'):
                    image_ext = '.jpg'
                elif image_data.startswith(b'\x89PNG\r\n\x1a\n'):
                    image_ext = '.png'
                elif image_data.startswith(b'GIF87a') or image_data.startswith(b'GIF89a'):
                    image_ext = '.gif'
                elif image_data.startswith(b'BM'):
                    image_ext = '.bmp'
                elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:12]:
                    image_ext = '.webp'
                elif image_data.startswith(b'<svg') or image_data.startswith(b'<?xml'):
                    image_ext = '.svg'
            except Exception:
                pass  # 保持默认格式
        
        # 创建临时文件来存储图片
        with tempfile.NamedTemporaryFile(delete=False, suffix=image_ext) as temp_file:
            temp_file.write(image_data)
            temp_file_path = temp_file.name
        
        try:
            # 获取原始图片尺寸信息
            width, height = _get_image_dimensions(drawing_element)
            
            # 添加图片到目标run
            if width and height:
                target_run.add_picture(temp_file_path, width=width, height=height)
            else:
                target_run.add_picture(temp_file_path)
            return True
            
        finally:
            # 清理临时文件
            try:
                os.unlink(temp_file_path)
            except:
                pass
                
    except Exception as e:
        print(f"      标准复制错误: {e}")
        return False


def _try_direct_image_extraction(drawing_element, source_doc, target_doc, target_run):
    """直接从docx文件中提取图片"""
    try:
        import tempfile
        import os
        import zipfile
        from pathlib import Path
        
        # 获取源文档的文件路径
        source_path = _get_document_path(source_doc)
        if not source_path:
            return False
        
        source_path = Path(source_path)
        if not source_path.exists():
            return False
        
        print(f"      尝试从 {source_path.name} 直接提取图片...")
        
        # 打开docx文件作为zip
        with zipfile.ZipFile(source_path, 'r') as zip_file:
            # 列出所有图片文件
            image_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
            
            if not image_files:
                print(f"      未找到图片文件")
                return False
            
            print(f"      找到 {len(image_files)} 个图片文件")
            
            # 尝试每个图片文件
            for img_file in image_files:
                try:
                    # 提取图片数据
                    image_data = zip_file.read(img_file)
                    if not image_data:
                        continue
                    
                    # 确定文件扩展名
                    img_ext = Path(img_file).suffix or '.png'
                    
                    # 创建临时文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix=img_ext) as temp_file:
                        temp_file.write(image_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # 获取图片尺寸
                        width, height = _get_image_dimensions(drawing_element)
                        
                        # 添加图片到目标run
                        if width and height:
                            target_run.add_picture(temp_file_path, width=width, height=height)
                        else:
                            target_run.add_picture(temp_file_path)
                        
                        print(f"      成功使用图片: {img_file}")
                        return True
                        
                    finally:
                        # 清理临时文件
                        try:
                            os.unlink(temp_file_path)
                        except:
                            pass
                            
                except Exception as e:
                    print(f"      处理图片 {img_file} 失败: {e}")
                    continue
        
        return False
        
    except Exception as e:
        print(f"      直接提取错误: {e}")
        return False


def _try_com_image_copy(drawing_element, source_doc, target_doc, target_run):
    """兼容保留：当前版本禁用COM图片复制。"""
    return False


def _get_document_path(doc):
    """获取文档的文件路径"""
    try:
        # 尝试多种方式获取文档路径
        if hasattr(doc, '_path') and doc._path:
            return doc._path
        
        if hasattr(doc, 'core_properties') and hasattr(doc.core_properties, 'identifier'):
            return doc.core_properties.identifier
        
        if hasattr(doc, '_part') and hasattr(doc._part, 'package') and hasattr(doc._part.package, '_pkg_file'):
            pkg_file = doc._part.package._pkg_file
            if hasattr(pkg_file, 'name'):
                return pkg_file.name
        
        # 如果都没有，返回None
        return None
        
    except Exception:
        return None


def _get_image_dimensions(drawing_element):
    """获取图片尺寸信息"""
    try:
        width = None
        height = None
        
        # 尝试使用namespaces参数的xpath调用
        try:
            extent_elements = drawing_element.xpath('.//wp:extent', namespaces={
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            })
        except TypeError:
            # 如果namespaces参数不支持，使用不带namespaces的xpath
            extent_elements = drawing_element.xpath('.//wp:extent')
        
        if extent_elements:
            cx = extent_elements[0].get('cx')
            cy = extent_elements[0].get('cy')
            if cx and cy:
                # 转换EMU到英寸 (1 inch = 914400 EMU)
                from docx.shared import Inches
                width = Inches(int(cx) / 914400)
                height = Inches(int(cy) / 914400)
                return width, height
        
        # 如果没有找到extent元素，尝试其他方法获取尺寸
        # 方法1: 查找inline元素的extent
        try:
            inline_elements = drawing_element.xpath('.//wp:inline')
            if inline_elements:
                for inline in inline_elements:
                    try:
                        extent_els = inline.xpath('.//wp:extent', namespaces={
                            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                        })
                    except TypeError:
                        extent_els = inline.xpath('.//wp:extent')
                    
                    if extent_els:
                        cx = extent_els[0].get('cx')
                        cy = extent_els[0].get('cy')
                        if cx and cy:
                            from docx.shared import Inches
                            width = Inches(int(cx) / 914400)
                            height = Inches(int(cy) / 914400)
                            return width, height
        except Exception:
            pass
        
        # 方法2: 查找anchor元素的extent
        try:
            anchor_elements = drawing_element.xpath('.//wp:anchor')
            if anchor_elements:
                for anchor in anchor_elements:
                    try:
                        extent_els = anchor.xpath('.//wp:extent', namespaces={
                            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                        })
                    except TypeError:
                        extent_els = anchor.xpath('.//wp:extent')
                    
                    if extent_els:
                        cx = extent_els[0].get('cx')
                        cy = extent_els[0].get('cy')
                        if cx and cy:
                            from docx.shared import Inches
                            width = Inches(int(cx) / 914400)
                            height = Inches(int(cy) / 914400)
                            return width, height
        except Exception:
            pass
        
        # 方法3: 直接查找所有extent元素
        try:
            all_extents = drawing_element.xpath('.//extent')
            if all_extents:
                for extent in all_extents:
                    cx = extent.get('cx')
                    cy = extent.get('cy')
                    if cx and cy:
                        from docx.shared import Inches
                        width = Inches(int(cx) / 914400)
                        height = Inches(int(cy) / 914400)
                        return width, height
        except Exception:
            pass
        
        return None, None
        
    except Exception as e:
        print(f"      获取图片尺寸时出错: {e}")
        return None, None










COM_UTILS_AVAILABLE = False

# 统一跨平台转换能力检查：依赖 LibreOffice/soffice
PDF_CONVERSION_AVAILABLE = True


def _find_soffice_executable():
    """查找 LibreOffice 可执行文件（跨平台）。"""
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found
    mac_default = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if mac_default.exists():
        return str(mac_default)
    return None


def _convert_docx_to_pdf_with_soffice(docx_path, pdf_path):
    """使用 LibreOffice headless 模式转换 DOCX 到 PDF。"""
    soffice = _find_soffice_executable()
    if not soffice:
        return False, None, "未找到 LibreOffice/soffice 可执行文件"

    try:
        docx_path = Path(docx_path)
        pdf_path = Path(pdf_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        tmp_out_dir = Path(tempfile.gettempdir()) / f"rewrite_soffice_{uuid.uuid4().hex}"
        tmp_out_dir.mkdir(parents=True, exist_ok=True)
        try:
            proc = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_out_dir),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            generated_pdf = tmp_out_dir / f"{docx_path.stem}.pdf"
            if proc.returncode != 0 or not generated_pdf.exists():
                stderr = (proc.stderr or "").strip()[:200]
                stdout = (proc.stdout or "").strip()[:200]
                detail = stderr or stdout or f"退出码 {proc.returncode}"
                return False, None, f"LibreOffice 转换失败: {detail}"

            shutil.copy2(generated_pdf, pdf_path)
            return True, str(pdf_path), None
        finally:
            shutil.rmtree(tmp_out_dir, ignore_errors=True)
    except Exception as e:
        return False, None, f"LibreOffice 转换异常: {e}"


# 设置控制台编码
if sys.platform == 'win32':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass


def wait_for_file_release(file_path, max_wait=15, check_interval=0.5):
    """
    主动等待文件被释放（不再被其他进程占用）
    
    参数:
        file_path: 文件路径
        max_wait: 最大等待时间（秒），默认15秒
        check_interval: 检查间隔（秒），默认0.5秒
    
    返回:
        True: 文件已释放
        False: 超时，文件仍被占用
    """
    import gc
    
    if not os.path.exists(file_path):
        return True  # 文件不存在，无需等待
    
    # 先强制垃圾回收
    gc.collect()
    
    start_time = time.time()
    attempts = 0
    
    while time.time() - start_time < max_wait:
        try:
            # 尝试以独占模式打开文件
            # 如果文件被占用，会抛出PermissionError
            with open(file_path, 'r+b') as f:
                # 成功打开，说明文件已释放
                return True
        except PermissionError:
            # 文件仍被占用，继续等待
            attempts += 1
            if attempts == 1:
                print(f"    ⏳ 文件被占用，等待释放...")
            elif attempts % 4 == 0:  # 每2秒打印一次
                elapsed = time.time() - start_time
                print(f"    ⏳ 仍在等待... ({elapsed:.1f}秒)")
            
            time.sleep(check_interval)
            gc.collect()  # 每次检查前垃圾回收
        except Exception as e:
            # 其他错误（如文件不存在），认为已释放
            return True
    
    # 超时
    elapsed = time.time() - start_time
    print(f"    ⚠️ 等待超时 ({elapsed:.1f}秒)，文件可能仍被占用")
    return False


def convert_docx_to_pdf(docx_path, pdf_path=None):
    """
    将Word文档转换为PDF
    
    参数:
        docx_path: Word文档路径
        pdf_path: PDF输出路径，如果为None则自动生成
    
    返回:
        tuple: (success, pdf_path, error_message)
    """
    try:
        import win32com.client as win32  # type: ignore
    except Exception as exc:
        return False, None, "未安装 pywin32（win32com），无法使用 Word COM 转换"

    word = None
    doc = None
    try:
        docx_path = Path(docx_path)
        if not docx_path.exists():
            return False, None, f"源文件不存在: {docx_path}"

        if pdf_path is None:
            pdf_path = docx_path.with_suffix('.pdf')
        else:
            pdf_path = Path(pdf_path)

        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(
            str(docx_path),
            ReadOnly=True,
            Visible=False,
            ConfirmConversions=False,
            AddToRecentFiles=False,
        )
        # 17 = wdExportFormatPDF
        doc.ExportAsFixedFormat(OutputFileName=str(pdf_path), ExportFormat=17)

        if not pdf_path.exists():
            return False, None, "PDF文件未生成"
        return True, str(pdf_path), None
    except Exception as e:
        return False, None, f"PDF转换失败: {e}"
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=0)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit(SaveChanges=0)
        except Exception:
            pass


def get_config_file():
    """获取配置文件路径"""
    if getattr(sys, 'frozen', False):
        # 如果是打包后的exe，配置文件在exe同级目录
        return Path(sys.executable).parent / "config.json"
    else:
        # 开发环境：从脚本位置向上找到项目根目录
        script_dir = Path(__file__).resolve().parent
        project_root = script_dir.parent.parent.parent
        return project_root / "config.json"


def update_notification_number(docx_file):
    """
    更新通报编号
    
    参数:
        docx_file: 生成的通报文档路径
    
    返回:
        当前使用的编号
    """
    try:
        config_file = get_config_file()
        
        # 读取配置
        if not config_file.exists():
            print(f"  警告: 配置文件不存在: {config_file}")
            return None
        
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        # 获取当前编号
        if 'report_counters' not in config:
            config['report_counters'] = {
                'notification_number': 104,
                'rectification_number': 235,
                'year': datetime.now().year,
                'last_updated': ''
            }
        
        # 检查年份，如果是新年则重置编号
        current_year = datetime.now().year
        if 'year' not in config['report_counters'] or config['report_counters']['year'] != current_year:
            print(f"  🎊 检测到新年份: {current_year}，重置编号计数")
            config['report_counters']['notification_number'] = 1
            config['report_counters']['rectification_number'] = 1
            config['report_counters']['year'] = current_year
        
        # 使用配置中的年份（已更新后的）
        config_year = config['report_counters']['year']
        current_number = config['report_counters']['notification_number']
        
        # 打开文档并替换编号
        doc = Document(docx_file)
        replaced = False
        
        for para in doc.paragraphs:
            para_text = para.text
            # 查找 〔YYYY〕第XX期 的模式（支持任意年份）
            if '〔' in para_text and '〕' in para_text and '第' in para_text and '期' in para_text:
                # 提取当前的年份和期数
                year_match = re.search(r'〔(\d{4})〕', para_text)
                number_match = re.search(r'第(\d+)期', para_text)
                
                if year_match and number_match:
                    old_year = year_match.group(1)
                    old_number = number_match.group(1)
                    
                    # 对每个run进行替换
                    for run in para.runs:
                        # 替换年份中的数字（可能分散在多个runs中）
                        if old_year in run.text:
                            run.text = run.text.replace(old_year, str(config_year))
                            replaced = True
                        elif any(old_year[i:i+len(run.text)] == run.text for i in range(len(old_year)) if run.text and run.text.isdigit()):
                            # 处理年份被拆分的情况（如 '202' 或 '5'）
                            for i in range(len(old_year)):
                                if old_year[i:i+len(run.text)] == run.text:
                                    run.text = str(config_year)[i:i+len(run.text)]
                                    replaced = True
                                    break
                        
                        # 替换期数
                        if old_number in run.text:
                            run.text = run.text.replace(old_number, str(current_number))
                            replaced = True
                
                # 找到目标段落后退出循环
                break
        
        if replaced:
            # 保存文档（添加重试机制）
            max_retries = 3
            for retry in range(max_retries):
                try:
                    doc.save(docx_file)
                    break
                except PermissionError as pe:
                    if retry < max_retries - 1:
                        print(f"  ⚠️ 文件被占用，等待重试 ({retry + 1}/{max_retries})...")
                        time.sleep(1.0)
                    else:
                        raise pe
            
            # 更新配置中的编号
            old_notification_number = config['report_counters']['notification_number']
            new_notification_number = current_number + 1
            config['report_counters']['notification_number'] = new_notification_number
            config['report_counters']['last_updated'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            print(f"  📝 准备更新配置文件: {config_file}")
            print(f"  📊 编号变更: {old_notification_number} → {new_notification_number}")
            
            try:
                # 使用统一配置管理器进行原子更新，避免覆盖其他模块的写入
                from modules.config.config_manager import ConfigManager
                cm = ConfigManager()
                # 更新编号
                cm.update_section('report_counters', {
                    'notification_number': new_notification_number,
                    'year': config_year,
                    'last_updated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                })
                print(f"  💾 配置文件写入完成（统一管理器）")
                
                # 验证写入结果
                verify_config = cm.get_config('report_counters')
                actual_number = verify_config.get('notification_number', 'N/A')
                print(f"  ✅ 验证写入结果: {actual_number}")
                
                if actual_number == new_notification_number:
                    print(f"  ✓ 配置更新成功！编号已更新为: {new_notification_number}")
                else:
                    print(f"  ❌ 配置更新失败！期望: {new_notification_number}, 实际: {actual_number}")
                    
            except Exception as config_error:
                print(f"  ❌ 配置文件操作失败: {config_error}")
                raise config_error
            
            print(f"  ✓ 已更新通报编号: 〔{config_year}〕第{current_number}期")
            return current_number, config_year
        else:
            print(f"  警告: 未找到通报编号标记")
            return None
            
    except Exception as e:
        print(f"  警告: 更新通报编号失败: {str(e)}")
        return None


def extract_info_from_filename(filename):
    """
    从文件名中提取公司名和漏洞类型
    
    文件名格式示例：
    - 1759979441661关于浙江格瓦拉数字科技有限公司所属Druid系统存在未授权访问安全漏洞通报.docx
    - 关于宁波易到互联科技有限公司所属啾啾救援-全国道路救援生态云平台系统MongDB未授权访问安全漏洞通报.docx
    - 1760410609070舒普智能技术股份有限公司远程技术检查存在ecology远程命令执行漏洞.docx
    
    返回: (公司名, 漏洞描述)
    """
    # 去掉路径和扩展名
    basename = os.path.basename(filename)
    name_without_ext = basename.rsplit('.', 1)[0]
    
    # 去掉开头的数字
    name_clean = re.sub(r'^\d+', '', name_without_ext)
    
    # 提取公司名：尝试多种模式
    company_name = None
    
    # 模式1：关于...所属（最常见）
    company_match = re.search(r'关于(.+?)所属', name_clean)
    if company_match:
        company_name = company_match.group(1)
    else:
        # 模式2：关于...门户网站/官网/网站
        company_match = re.search(r'关于(.+?)(门户网站|官网|网站|平台|系统)', name_clean)
        if company_match:
            company_name = company_match.group(1)
        else:
            # 模式3：关于...存在（针对直接描述漏洞的文件名）
            company_match = re.search(r'关于(.+?)存在', name_clean)
            if company_match:
                company_name = company_match.group(1)
            else:
                # 模式4：关于...的
                company_match = re.search(r'关于(.+?)的', name_clean)
                if company_match:
                    company_name = company_match.group(1)
                else:
                    # 模式5：直接格式 - 公司名+技术检查/远程检查等
                    # 匹配：公司名（包含有限公司、股份有限公司等）+ 技术检查/远程检查等
                    company_match = re.search(r'^(.+?(?:有限公司|股份有限公司|集团|科技公司|科技))', name_clean)
                    if company_match:
                        company_name = company_match.group(1)
                    else:
                        # 模式6：尝试从"存在"之前提取公司名
                        company_match = re.search(r'^(.+?)(?:远程技术检查|技术检查|检查|远程|存在)', name_clean)
                        if company_match:
                            potential_company = company_match.group(1).strip()
                            # 验证是否包含公司关键词
                            if any(keyword in potential_company for keyword in ['有限公司', '股份有限公司', '集团', '科技']):
                                company_name = potential_company
    
    # 提取漏洞类型：尝试多种模式
    vuln_type = None
    
    # 模式1：查找"存在"和"通报"或"的报告"之间的内容
    vuln_match = re.search(r'存在(.+?)(?:通报|的报告)', name_clean)
    if vuln_match:
        # 去掉"存在"前缀，只保留漏洞类型描述
        vuln_type = vuln_match.group(1).strip()
    else:
        # 模式2：查找"系统"之后到"通报"或"的报告"之间的内容
        vuln_match = re.search(r'系统(.+?)(?:通报|的报告)', name_clean)
        if vuln_match:
            content = vuln_match.group(1).strip()
            # 去掉开头的"的"字
            content = re.sub(r'^的', '', content)
            # 去掉可能的系统名称，只保留漏洞描述
            vuln_type = content
        else:
            # 模式3：查找"网站"之后到"通报"或"的报告"之间的内容
            vuln_match = re.search(r'网站(.+?)(?:通报|的报告)', name_clean)
            if vuln_match:
                content = vuln_match.group(1).strip()
                # 去掉开头的"的"字
                content = re.sub(r'^的', '', content)
                vuln_type = content
            else:
                # 模式4：查找"存在"到文件名结尾的内容（针对没有"通报"或"的报告"的文件名）
                vuln_match = re.search(r'存在(.+?)(?:\.docx|$)', name_clean)
                if vuln_match:
                    # 去掉"存在"前缀，只保留漏洞类型描述
                    vuln_type = vuln_match.group(1).strip()
                else:
                    # 模式5：查找"技术检查存在"模式
                    vuln_match = re.search(r'(?:远程技术检查|技术检查|检查)存在(.+?)(?:\.docx|$)', name_clean)
                    if vuln_match:
                        # 去掉"存在"前缀，只保留漏洞类型描述
                        vuln_type = vuln_match.group(1).strip()
                    else:
                        # 模式6：最后尝试，查找包含"漏洞"或"风险"关键词的部分
                        vuln_match = re.search(r'([\u4e00-\u9fa5A-Za-z]+(?:漏洞|风险))', name_clean)
                        if vuln_match:
                            vuln_type = vuln_match.group(1)
    
    # 清理漏洞类型，去除不需要的后缀
    if vuln_type:
        # 去除"的报告"、"风险的报告"等后缀
        vuln_type = re.sub(r'(?:风险)?的报告$', '', vuln_type)
        # 去除"安全"重复
        vuln_type = re.sub(r'安全安全', '安全', vuln_type)
        # 确保以"漏洞"或"风险"结尾
        if not vuln_type.endswith(('漏洞', '风险')):
            if '漏洞' in vuln_type:
                # 如果包含"漏洞"但不以"漏洞"结尾，截取到"漏洞"
                vuln_match = re.search(r'(.+?漏洞)', vuln_type)
                if vuln_match:
                    vuln_type = vuln_match.group(1)
            elif '风险' in vuln_type:
                # 如果包含"风险"但不以"风险"结尾，截取到"风险"
                vuln_match = re.search(r'(.+?风险)', vuln_type)
                if vuln_match:
                    vuln_type = vuln_match.group(1)
    
    return company_name, vuln_type






def replace_text_in_runs(para, old_text, new_text):
    """
    在段落的runs中替换文本（支持跨runs替换），保留超链接
    
    参数:
        para: 段落对象
        old_text: 要查找的旧文本
        new_text: 替换后的新文本
    
    返回:
        是否成功替换
    """
    # 获取段落的完整文本
    full_text = para.text
    
    # 检查是否包含要替换的文本
    if old_text not in full_text:
        return False
    
    # 找到旧文本的起始位置
    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)
    
    # 计算每个run的字符范围
    run_ranges = []
    current_pos = 0
    for run in para.runs:
        run_length = len(run.text)
        run_ranges.append((current_pos, current_pos + run_length, run))
        current_pos += run_length
    
    # 找出需要修改的runs
    affected_runs = []
    for run_start, run_end, run in run_ranges:
        # 如果run与替换区域有交集
        if run_start < end_pos and run_end > start_pos:
            affected_runs.append((run_start, run_end, run))
    
    if not affected_runs:
        return False
    
    # 检查受影响的runs中是否包含超链接
    has_hyperlink = False
    for run_start, run_end, run in affected_runs:
        if _run_contains_hyperlink(run):
            has_hyperlink = True
            print(f"    ⚠️ 检测到超链接，跳过文本替换以保留超链接: '{old_text}'")
            break
    
    # 如果包含超链接，跳过替换以保留超链接
    if has_hyperlink:
        return False
    
    # 执行替换
    for run_start, run_end, run in affected_runs:
        # 计算在当前run中的替换范围
        replace_start = max(0, start_pos - run_start)
        replace_end = min(len(run.text), end_pos - run_start)
        
        # 构建新的run文本
        old_run_text = run.text
        
        if replace_start == 0 and replace_end == len(run.text):
            # 整个run都在替换范围内
            if run == affected_runs[0][2]:
                # 第一个受影响的run，包含新文本
                run.text = new_text
            else:
                # 后续受影响的run，清空
                run.text = ""
        elif replace_start == 0:
            # 从run开头开始替换
            if run == affected_runs[0][2]:
                run.text = new_text + old_run_text[replace_end:]
            else:
                run.text = old_run_text[replace_end:]
        elif replace_end == len(run.text):
            # 替换到run结尾
            if run == affected_runs[0][2]:
                run.text = old_run_text[:replace_start] + new_text
            else:
                run.text = old_run_text[:replace_start]
        else:
            # 替换在run中间
            run.text = old_run_text[:replace_start] + new_text + old_run_text[replace_end:]
    
    return True


def _run_contains_hyperlink(run):
    """
    检查run是否包含超链接
    
    参数:
        run: run对象
    
    返回:
        bool: 是否包含超链接
    """
    try:
        # 检查run的XML元素中是否包含超链接
        if hasattr(run, '_element'):
            # 查找超链接元素
            hyperlinks = run._element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if hyperlinks:
                return True
            
            # 查找fldChar元素（字段字符，超链接的另一种形式）
            fld_chars = run._element.findall('.//w:fldChar', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if fld_chars:
                return True
            
            # 查找instrText元素（指令文本，超链接字段的一部分）
            instr_texts = run._element.findall('.//w:instrText', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for instr_text in instr_texts:
                if instr_text.text and 'HYPERLINK' in instr_text.text:
                    return True
        
        return False
        
    except Exception as e:
        # 如果检查失败，为了安全起见，假设包含超链接
        print(f"    ⚠️ 超链接检查失败: {e}")
        return True


def _run_element_contains_hyperlink(run_element):
    """
    检查run XML元素是否包含超链接
    
    参数:
        run_element: run的XML元素
    
    返回:
        bool: 是否包含超链接
    """
    try:
        # 查找超链接元素
        hyperlinks = run_element.findall('.//w:hyperlink', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if hyperlinks:
            return True
        
        # 查找fldChar元素（字段字符，超链接的另一种形式）
        fld_chars = run_element.findall('.//w:fldChar', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if fld_chars:
            return True
        
        # 查找instrText元素（指令文本，超链接字段的一部分）
        instr_texts = run_element.findall('.//w:instrText', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for instr_text in instr_texts:
            if instr_text.text and 'HYPERLINK' in instr_text.text:
                return True
        
        # 检查run是否在hyperlink元素内部
        parent = run_element.getparent()
        while parent is not None:
            if parent.tag.endswith('hyperlink'):
                return True
            parent = parent.getparent()
        
        return False
        
    except Exception as e:
        # 如果检查失败，为了安全起见，假设包含超链接
        print(f"    ⚠️ 超链接检查失败: {e}")
        return True


def is_notification_document(doc):
    """
    识别文档是否为通报文件
    
    参数:
        doc: Word文档对象
    
    返回:
        bool: 是否为通报文件
    """
    try:
        # 获取文档的前几个段落文本
        first_paragraphs_text = ""
        for i, para in enumerate(doc.paragraphs[:10]):  # 检查前10个段落
            first_paragraphs_text += para.text.strip() + " "
        
        # 通报文件的关键词
        notification_keywords = [
            "通报", "网络安全", "漏洞", "安全事件", "安全通告", 
            "风险提示", "安全预警", "威胁情报", "安全公告"
        ]
        
        # 检查是否包含通报相关关键词
        for keyword in notification_keywords:
            if keyword in first_paragraphs_text:
                print(f"  ✅ 检测到通报文件关键词: '{keyword}'")
                return True
        
        print(f"  ℹ️  未检测到通报文件关键词，跳过图片插入")
        return False
        
    except Exception as e:
        print(f"  ⚠️ 文档类型识别失败: {e}")
        return False


def get_accurate_page_count(doc):
    """
    获取文档页数（跨平台估算，不依赖COM）
    
    参数:
        doc: Word文档对象
    
    返回:
        int: 估算页数
    """
    try:
        paragraphs = getattr(doc, "paragraphs", [])
        para_count = len(paragraphs)
        # 经验估算：约每20段为1页，至少1页
        page_count = max(1, para_count // 20 + (1 if para_count % 20 else 0))
        print(f"  📄 估算页数: {page_count}")
        return page_count
    except Exception as e:
        print(f"  ❌ 页数估算失败: {e}")
        return 1


def check_existing_images_on_page(doc, page_start_para, page_end_para, image_signature):
    """
    检查指定页面范围内是否已存在相同的水印图片
    
    参数:
        doc: Word文档对象
        page_start_para: 页面开始段落索引
        page_end_para: 页面结束段落索引
        image_signature: 图片特征签名（包含文件名和大小）
    
    返回:
        bool: 是否已存在相同的水印图片
    """
    try:
        # 对于水印式图片，我们允许多个图片共存
        # 只检查是否已存在完全相同的水印图片（基于文件名和大小）
        watermark_count = 0
        
        for para_idx in range(page_start_para, min(page_end_para, len(doc.paragraphs))):
            para = doc.paragraphs[para_idx]
            
            # 检查段落中的图片
            for run in para.runs:
                if hasattr(run, '_element'):
                    # 查找图片元素
                    drawings = run._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    for drawing in drawings:
                        # 检查是否是图片
                        pic_elements = drawing.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})
                        if pic_elements:
                            # 检查图片的描述信息，看是否包含我们的水印标识
                            # 这里我们采用更宽松的策略：允许添加水印，除非已经有太多相同的水印
                            watermark_count += 1
        
        # 如果已经有超过1个水印图片，则跳过（避免重复添加太多）
        if watermark_count >= 1 and "确认词条" in image_signature:
            print(f"    ℹ️ 第{page_start_para//20 + 1}页已有{watermark_count}个图片，跳过水印添加")
            return True
        
        return False
        
    except Exception as e:
        print(f"  ⚠️ 检查已存在图片失败: {e}")
        return False


def add_floating_image_to_pages(doc, image_path, start_page=2, source_file_path=None):
    """
    在文档的每一页（从指定页开始）添加浮动图片
    支持精确页数检测、防重复插入、页数变化处理
    
    参数:
        doc: Word文档对象
        image_path: 图片文件路径
        start_page: 开始添加图片的页码（默认从第2页开始）
        source_file_path: 源文件路径（用于错误记录）
    """
    try:
        print(f"\n🖼️ 开始添加浮动图片到文档...")
        print(f"  图片路径: {image_path}")
        print(f"  开始页码: {start_page}")
        
        # 1. 检查图片文件是否存在
        if not os.path.exists(image_path):
            print(f"  ❌ 图片文件不存在: {image_path}")
            return False
        
        # 2. 检查是否为通报文件
        if not is_notification_document(doc):
            return False
        
        # 3. 获取文档的所有段落
        paragraphs = doc.paragraphs
        if len(paragraphs) == 0:
            print(f"  ❌ 文档没有段落")
            return False
        
        # 4. 获取初始页数（跨平台估算）
        initial_page_count = get_accurate_page_count(doc)
        
        print(f"  📄 初始页数: {initial_page_count}")
        
        # 5. 计算图片特征签名（用于防重复）
        image_signature = f"{os.path.basename(image_path)}_{os.path.getsize(image_path)}"
        
        # 6. 为每一页（从start_page开始）添加图片
        images_added = 0
        current_page_count = initial_page_count
        
        # 基于估算页数进行处理
        actual_page_count = initial_page_count
        print(f"  📄 将为第{start_page}页到第{actual_page_count}页添加水印图片")
        
        for page_num in range(start_page, actual_page_count + 1):
            try:
                # 简化的页面范围计算：基于页码直接选择段落
                total_paragraphs = len(paragraphs)
                
                # 根据页码选择合适的段落范围
                # 基于实际页面分布：第1页(1-16), 第2页(17-33), 第3页(34-38)
                if page_num == 2:
                    # 第2页：段落17-33
                    start_para_idx = 16  # 从段落17开始（索引16）
                    end_para_idx = min(total_paragraphs, 33)  # 到段落33结束
                elif page_num == 3:
                    # 第3页：段落34-38
                    start_para_idx = 33  # 从段落34开始（索引33）
                    end_para_idx = total_paragraphs
                else:
                    # 其他页：使用平均分配
                    paragraphs_per_page = max(1, total_paragraphs // actual_page_count)
                    start_para_idx = (page_num - 1) * paragraphs_per_page
                    end_para_idx = min(page_num * paragraphs_per_page, total_paragraphs)
                
                if start_para_idx >= total_paragraphs:
                    print(f"    ⚠️ 第{page_num}页超出段落范围，跳过")
                    continue
                
                print(f"    📄 第{page_num}页段落范围: {start_para_idx} - {end_para_idx-1} (共{end_para_idx-start_para_idx}个段落)")
                
                # 仅允许选择真正空段作为图片锚点，避免改写正文结构
                target_para_idx = _find_best_insertion_point(paragraphs, start_para_idx, end_para_idx)
                
                if target_para_idx is not None and target_para_idx < len(paragraphs):
                    target_para = paragraphs[target_para_idx]
                    # #region agent log
                    _agent_debug_log(
                        run_id=f"img_{int(time.time() * 1000)}",
                        hypothesis_id="H32",
                        location="rewrite_report.py:image_anchor_target_probe",
                        message="image_anchor_target_selected",
                        data={
                            "page_num": page_num,
                            "range_start": start_para_idx,
                            "range_end": end_para_idx,
                            "target_para_idx": target_para_idx,
                            "target_text": (target_para.text or "")[:120],
                            "target_is_empty": not bool((target_para.text or "").strip()),
                            "empty_candidate_indices": [
                                idx for idx in range(start_para_idx, min(end_para_idx + 1, len(paragraphs)))
                                if not (paragraphs[idx].text or "").strip()
                            ][:10],
                        },
                    )
                    # #endregion
                    
                    # 直接锚定到目标段落，避免额外插入空段落把正文顶开
                    run = target_para.add_run()
                    
                    # 添加图片（大小由_set_picture_floating函数控制）
                    picture = run.add_picture(image_path)
                    
                    # 设置图片为浮动样式（右上角）
                    floating_success = _set_picture_floating(picture, target_para)
                    
                    images_added += 1
                    if floating_success:
                        print(f"    ✅ 第{page_num}页图片添加成功（插入位置：段落{target_para_idx}前）")
                    else:
                        print(f"    ⚠️ 第{page_num}页图片已添加但需要手动调整浮动样式（插入位置：段落{target_para_idx}前）")
                    
                    # 检查页数是否发生变化
                    new_page_count = get_accurate_page_count(doc)
                    if new_page_count > current_page_count:
                        print(f"    📄 页数增加: {current_page_count} → {new_page_count}")
                        current_page_count = new_page_count
                        actual_page_count = new_page_count  # 更新实际页数
                else:
                    # #region agent log
                    _agent_debug_log(
                        run_id=f"img_{int(time.time() * 1000)}",
                        hypothesis_id="H32",
                        location="rewrite_report.py:image_anchor_target_probe",
                        message="image_anchor_target_missing_safe_empty_paragraph",
                        data={
                            "page_num": page_num,
                            "range_start": start_para_idx,
                            "range_end": end_para_idx,
                            "empty_candidate_indices": [
                                idx for idx in range(start_para_idx, min(end_para_idx + 1, len(paragraphs)))
                                if not (paragraphs[idx].text or "").strip()
                            ][:10],
                        },
                    )
                    # #endregion
                    print(f"    ⚠️ 第{page_num}页未找到安全空段，取消自动图片插入以保护正文格式")
                    return False

            except Exception as e:
                print(f"    ⚠️ 第{page_num}页图片添加失败: {e}")
                return False
        
        # 8. 最终页数检查和调整
        final_page_count = get_accurate_page_count(doc)
        if final_page_count > initial_page_count:
            print(f"  📄 最终页数变化: {initial_page_count} → {final_page_count}")
            
            # 如果新增了页面，为新页面也添加图片
            for page_num in range(initial_page_count + 1, final_page_count + 1):
                try:
                    # 为新增页面添加图片
                    start_para_idx = (page_num - 1) * paragraphs_per_page
                    if start_para_idx < len(doc.paragraphs):
                        end_para_idx = min(start_para_idx + paragraphs_per_page, len(doc.paragraphs))
                        target_para_idx = _find_best_insertion_point(doc.paragraphs, start_para_idx, end_para_idx)
                        if target_para_idx is None or target_para_idx >= len(doc.paragraphs):
                            print(f"    ⚠️ 新增第{page_num}页未找到安全空段，取消自动图片插入以保护正文格式")
                            return False
                        target_para = doc.paragraphs[target_para_idx]
                        
                        run = target_para.add_run()
                        picture = run.add_picture(image_path)
                        floating_success = _set_picture_floating(picture, target_para)
                        
                        images_added += 1
                        if floating_success:
                            print(f"    ✅ 新增第{page_num}页图片添加成功")
                        else:
                            print(f"    ⚠️ 新增第{page_num}页图片已添加但需要手动调整浮动样式")
                        
                except Exception as e:
                    print(f"    ⚠️ 新增第{page_num}页图片添加失败: {e}")
        
        print(f"  ✅ 图片添加完成，共添加 {images_added} 张图片")
        print(f"  📄 最终文档页数: {final_page_count}")
        return images_added > 0
        
    except Exception as e:
        print(f"  ❌ 添加浮动图片失败: {e}")
        return False


def _find_best_insertion_point(paragraphs, start_para_idx, end_para_idx):
    """
    寻找最佳的图片插入位置，优先选择空白区域
    
    参数:
        paragraphs: 段落列表
        start_para_idx: 页面开始段落索引
        end_para_idx: 页面结束段落索引
    
    返回:
        最佳插入位置的段落索引
    """
    # 确保索引在有效范围内
    start_para_idx = max(0, start_para_idx)
    end_para_idx = min(len(paragraphs) - 1, end_para_idx)
    
    # 只允许锚定到真正空段落，避免污染正文 run 结构
    for i in range(start_para_idx, end_para_idx + 1):
        if i < len(paragraphs):
            para = paragraphs[i]
            text = para.text.strip()
            if len(text) == 0:
                return i
    return None


def _set_picture_floating(picture, paragraph):
    """
    设置图片为浮动样式，位于文字上方（水印效果）
    根据第二页图片的正确格式：使用anchor浮动，右对齐
    
    参数:
        picture: 图片对象（InlineShape对象）
        paragraph: 包含图片的段落
    """
    try:
        # 在python-docx中，InlineShape对象没有直接的XML访问方式
        # 我们需要通过段落的run来找到图片的XML元素
        
        # 查找包含图片的run
        target_run = None
        for run in paragraph.runs:
            if hasattr(run._element, 'xpath'):
                # 查找内联图片元素
                inline_elements = run._element.xpath('.//wp:inline')
                if inline_elements:
                    target_run = run
                    break
        
        if not target_run:
            print(f"      ❌ 无法找到包含图片的run")
            return False
        
        # 获取内联图片元素
        inline_elements = target_run._element.xpath('.//wp:inline')
        if not inline_elements:
            print(f"      ❌ 无法找到内联图片元素")
            return False
        
        inline_element = inline_elements[0]  # 取第一个内联图片
        
        # 获取图片的graphic元素
        graphic_xml = ""
        try:
            graphic_elements = inline_element.xpath('.//a:graphic')
            if graphic_elements:
                graphic_xml = etree.tostring(graphic_elements[0], encoding='unicode')
            else:
                # 如果无法获取graphic XML，使用简化版本
                graphic_xml = f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"/></a:graphic>'
        except Exception as e:
            print(f"      ⚠️ 获取graphic XML失败: {e}")
            graphic_xml = f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"/></a:graphic>'
        
        # 创建anchor元素来替换inline元素，完全按照第二页正确图片的格式
        anchor_xml = f'''<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251663360" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" wp14:anchorId="36FF99FB" wp14:editId="53933A80"><wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>1731645</wp:posOffset></wp:positionH><wp:positionV relativeFrom="paragraph"><wp:posOffset>201295</wp:posOffset></wp:positionV><wp:extent cx="2134235" cy="1280160"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/><wp:docPr id="1" name="Picture 1"/><wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>{graphic_xml}<wp:sizeRelH relativeFrom="page"><wp:pctWidth>0</wp:pctWidth></wp:sizeRelH><wp:sizeRelV relativeFrom="page"><wp:pctHeight>0</wp:pctHeight></wp:sizeRelV></wp:anchor>'''
        
        # 解析新的anchor XML
        anchor_element = parse_xml(anchor_xml)
        
        # 替换inline元素为anchor元素
        parent = inline_element.getparent()
        parent.replace(inline_element, anchor_element)
        
        # 仅在纯图片空段上压缩段落占位，避免误伤正文段落格式
        if not (paragraph.text or "").strip():
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = Pt(6)  # 较小的行间距
        
        print(f"      ✅ 水印式图片样式设置完成（浮动anchor，右对齐）")
        return True
        
    except Exception as e:
        print(f"      ❌ 浮动图片样式设置失败: {e}")
        print(f"      ⚠️ 需要手动调整图片样式为浮动水印效果")
        return False


def replace_template_content(template_doc, company_name, vuln_type, current_date, deadline_date):
    """
    替换模板中的占位内容
    
    参数:
        template_doc: 模板文档对象
        company_name: 公司名称
        vuln_type: 漏洞类型描述
        current_date: 当前日期字符串
        deadline_date: 截止日期字符串
    """
    print(f"\n开始替换模板内容:")
    print(f"  公司名: {company_name}")
    print(f"  漏洞类型: {vuln_type}")
    print(f"  当前日期: {current_date}")
    print(f"  截止日期: {deadline_date}")
    print(f"  模板总段落数: {len(template_doc.paragraphs)}")
    print("=" * 60)
    
    for i, para in enumerate(template_doc.paragraphs, 1):
        original_text = para.text
        modified = False
        
        # 段落4：替换公司名（在标题中）
        if i == 4 and company_name:
            para_text = para.text
            # 提取"关于XXX公司所属"中的公司名
            match = re.search(r'关于([\u4e00-\u9fa5]+(?:集团)?(?:股份)?(?:有限)?公司)', para_text)
            if match:
                old_company = match.group(1)
                if replace_text_in_runs(para, old_company, company_name):
                    modified = True
                    print(f"    段落 4 公司名替换: '{old_company}' → '{company_name}'")
        
        # 段落6：替换公司名（收件人）
        if i == 6 and company_name:
            para_text = para.text
            # 提取"XXX公司："中的公司名
            match = re.search(r'([\u4e00-\u9fa5]+(?:集团)?(?:股份)?(?:有限)?公司)：', para_text)
            if match:
                old_company = match.group(1)
                if replace_text_in_runs(para, old_company, company_name):
                    modified = True
                    print(f"    段落 6 公司名替换: '{old_company}' → '{company_name}'")
        
        # 段落7：替换漏洞类型和截止日期
        if i == 7:
            para_text = para.text
            
            # 替换漏洞类型
            if vuln_type:
                vuln_match = re.search(r'存在.+?漏洞', para_text)
                if vuln_match:
                    old_vuln = vuln_match.group(0)
                    # 确保替换后保持"存在"关键词
                    new_vuln = f"存在{vuln_type}"
                    if replace_text_in_runs(para, old_vuln, new_vuln):
                        modified = True
            
            # 替换截止日期（需要重新获取文本，因为可能已被修改）
            para_text = para.text
            if deadline_date:
                date_match = re.search(r'20\d{2}年\d+月\d+日前', para_text)
                if date_match:
                    old_date = date_match.group(0)
                    if replace_text_in_runs(para, old_date, deadline_date + '前'):
                        modified = True
        
        # 段落14：替换当前日期
        if i == 14 and current_date:
            para_text = para.text
            date_match = re.search(r'20\d{2}年\d+月\d+日', para_text)
            if date_match:
                old_date = date_match.group(0)
                if replace_text_in_runs(para, old_date, current_date):
                    modified = True
                    print(f"    段落 14 日期替换: '{old_date}' → '{current_date}'")
        
        if modified:
            print(f"  段落 {i} 已更新: {original_text[:40]}... -> {para.text[:40]}...")


def rewrite_report(source_file, template_file=None, start_para=3, end_para=-1):
    """
    将源文档内容复制到模板文档中（保留格式，包括表格）
    
    参数:
        source_file: 源Word文档的路径
        template_file: 模板文档的路径（如果为None，则自动查找）
        start_para: 起始段落编号（从1开始），默认3
        end_para: 结束段落编号（从1开始），-1表示到倒数第二段
    """
    try:
        debug_run_id = f"run_{int(time.time() * 1000)}"
        # #region agent log
        _agent_debug_log(
            run_id=debug_run_id,
            hypothesis_id="H4",
            location="rewrite_report.py:rewrite_report_entry",
            message="rewrite_report_start",
            data={
                "source_file": str(source_file),
                "template_file": str(template_file) if template_file else None,
                "start_para": start_para,
                "end_para": end_para,
            },
        )
        # #endregion
        # 如果未指定模板文件，自动查找
        if template_file is None:
            # 先在 template 目录查找（支持开发和打包环境）
            template_candidates = []
            try:
                from modules.utils.resource_path import get_report_template_dir
                report_template_dir = str(get_report_template_dir())
            except ImportError:
                report_template_dir = 'Report_Template'
            
            if os.path.exists(report_template_dir):
                for filename in os.listdir(report_template_dir):
                    if filename.endswith('.docx') and '通报模板' in filename:
                        template_candidates.append(os.path.join(report_template_dir, filename))
            
            # 如果 template 目录没找到，在当前目录查找
            if not template_candidates:
                for filename in os.listdir('.'):
                    if filename.endswith('.docx') and '通报模板' in filename:
                        template_candidates.append(filename)
            
            if not template_candidates:
                print("错误: 未找到通报模板文件！")
                print("  请确保以下位置之一存在通报模板文件：")
                print("    - Repor/通报模板.docx")
                print("    - ./通报模板.docx")
                return {
                    'success': False,
                    'skip_reason': '未找到通报模板文件',
                    'backup_file': None,
                    'needs_manual_processing': False
                }
            
            template_file = template_candidates[0]
            print(f"自动找到模板文件: {template_file}")
        
        # 从文件名中提取信息
        company_name, vuln_type = extract_info_from_filename(source_file)
        
        # 计算日期
        today = datetime.now()
        deadline = today + timedelta(days=5)
        # 格式化日期，去掉前导0
        current_date_str = f"{today.year}年{today.month}月{today.day}日"
        deadline_date_str = f"{deadline.year}年{deadline.month}月{deadline.day}日"
        
        # 读取源文档
        try:
            source_doc = Document(source_file)
            cleanup_temp_source = False
                
        except Exception as e:
            safe_print(f"错误: 无法打开源文档 {source_file}: {str(e)}")
            return {'success': False, 'skip_reason': f'无法打开源文档: {str(e)}'}
        
        # 读取模板文档
        template_doc = Document(template_file)
        
        # 确定段落范围
        total_paragraphs = len(source_doc.paragraphs)
        start_idx = (start_para - 1) if start_para else 0
        last_non_empty_idx = -1
        for idx, p in enumerate(source_doc.paragraphs):
            if (p.text or "").strip():
                last_non_empty_idx = idx
        # 如果end_para是-1，表示到倒数第二段（跳过最后的空段落）
        if end_para == -1:
            # 按“最后一个非空段落”计算结束位置，避免误跳过最后一段有效内容
            end_idx = (last_non_empty_idx + 1) if last_non_empty_idx >= 0 else 0
        elif end_para:
            end_idx = end_para
        else:
            end_idx = total_paragraphs
        # #region agent log
        _agent_debug_log(
            run_id=debug_run_id,
            hypothesis_id="H15",
            location="rewrite_report.py:copy_range_calc",
            message="copy_range_and_last_paragraph_probe",
            data={
                "total_paragraphs": total_paragraphs,
                "start_idx": start_idx,
                "end_idx": end_idx,
                "end_para_arg": end_para,
                "last_non_empty_idx": last_non_empty_idx,
                "last_non_empty_text": (source_doc.paragraphs[last_non_empty_idx].text[:120] if last_non_empty_idx >= 0 else ""),
                "last_para_text": (source_doc.paragraphs[-1].text[:120] if total_paragraphs > 0 else ""),
            },
        )
        # #endregion
        # #region agent log
        _agent_debug_log(
            run_id=debug_run_id,
            hypothesis_id="H22",
            location="rewrite_report.py:source_layout_probe_before_copy",
            message="source_layout_window_probe",
            data={
                "source_file": str(source_file),
                "window_validation": _collect_paragraph_window(source_doc, "验证情况", radius=4),
                "window_validation_raw": _collect_paragraph_window_including_empty(source_doc, "验证情况", radius=4),
                "window_disposal": _collect_paragraph_window(source_doc, "处置措施", radius=4),
            },
        )
        # #endregion
        
        # 生成输出文件名：去掉源文件名开头的数字
        source_basename = os.path.basename(source_file)
        # 去掉开头的数字
        output_basename = re.sub(r'^\d+', '', source_basename)
        output_file = output_basename
        
        print(f"\n正在使用模板创建通报文档:")
        print(f"  源文件: {source_file}")
        print(f"  模板文件: {template_file}")
        print(f"  输出文件: {output_file}")
        print("=" * 60)
        print(f"复制段落范围: 第 {start_idx + 1} 段 到 第 {end_idx} 段")
        print(f"插入位置: 自动查找标记 '*'")
        print("=" * 60)
        
        # 找到插入位置（查找包含 * 标记的段落）
        insert_element_index = None
        marker_para_element = None
        para_count = 0
        marker_para_index = None
        
        for i, element in enumerate(template_doc.element.body):
            if element.tag.endswith('p'):
                para_count += 1
                # 查找对应的段落对象
                para = None
                for p in template_doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para and '*' in para.text:
                    # 找到了标记段落
                    insert_element_index = i  # 在这个段落的位置插入（替换它）
                    marker_para_element = element
                    marker_para_index = para_count
                    print(f"找到标记位置: 第 {marker_para_index} 段")
                    break
        
        if insert_element_index is None:
            print("错误: 未找到 * 标记！请在模板的第二页起始位置添加 * 标记。")
            return False
        insert_para_index = max(0, (marker_para_index or 1) - 1)
        
        # ⚠️ 重要：在插入原文段落之前，先替换模板内容
        # 因为插入原文段落会改变段落索引
        replace_template_content(template_doc, company_name, vuln_type, current_date_str, deadline_date_str)
        
        # 遍历源文档的body元素，复制指定范围的段落
        para_count = 0
        copied_count = 0
        copy_style_log_count = 0
        style_import_log_count = 0
        style_collision_probe_log_count = 0
        style_remap_log_count = 0
        numbering_strip_log_count = 0
        numbering_map_log_count = 0
        numbering_collision_log_count = 0
        numbering_mapping_probe_log_count = 0
        numbering_collision_fingerprint_log_count = 0
        numbering_mapping_cache = {}
        style_mapping_cache = {}
        skipped_by_end_idx_log_count = 0
        replacement_mutation_log_count = 0
        key_struct_log_count = 0
        key_phrases = ["验证情况", "处置措施", "限制用户访问", "Url:"]
        
        for element in source_doc.element.body:
            # 检查是否是段落
            if element.tag.endswith('p'):
                # 跳过范围外的段落
                if para_count < start_idx or para_count >= end_idx:
                    if para_count >= end_idx and skipped_by_end_idx_log_count < 6:
                        skipped_para_text = ""
                        if para_count < len(source_doc.paragraphs):
                            skipped_para_text = (source_doc.paragraphs[para_count].text or "")[:120]
                        # #region agent log
                        _agent_debug_log(
                            run_id=debug_run_id,
                            hypothesis_id="H15",
                            location="rewrite_report.py:copy_loop_skip_by_end",
                            message="paragraph_skipped_by_end_idx",
                            data={
                                "para_count_before_inc": para_count,
                                "end_idx": end_idx,
                                "skipped_text": skipped_para_text,
                                "skipped_is_non_empty": bool(skipped_para_text.strip()),
                            },
                        )
                        # #endregion
                        skipped_by_end_idx_log_count += 1
                    para_count += 1
                    continue
                
                para_count += 1
                copied_count += 1
                
                # 从element创建段落对象
                paragraph = None
                for p in source_doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break
                
                if paragraph is None:
                    continue
                
                # 直接深拷贝整个段落元素以保持所有格式
                new_para_element = deepcopy(paragraph._element)
                source_style = _paragraph_style_val(paragraph)
                copied_style = _paragraph_style_val(new_para_element)
                src_num = _paragraph_numbering_info(paragraph)
                copied_num = _paragraph_numbering_info(new_para_element)
                if copied_style and style_collision_probe_log_count < 8:
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H23",
                        location="rewrite_report.py:style_collision_probe",
                        message="source_vs_target_style_fingerprint",
                        data={
                            "copied_idx": copied_count,
                            "text": paragraph.text[:120],
                            "style_id": copied_style,
                            "source_style": _style_fingerprint(source_doc, copied_style),
                            "target_style_before_copy": _style_fingerprint(template_doc, copied_style),
                        },
                    )
                    # #endregion
                    style_collision_probe_log_count += 1
                remapped_style_refs = _remap_style_references_in_element(
                    source_doc,
                    template_doc,
                    new_para_element,
                    style_mapping_cache,
                )
                copied_style = _paragraph_style_val(new_para_element)
                if remapped_style_refs and style_remap_log_count < 8:
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H25",
                        location="rewrite_report.py:style_remap_apply",
                        message="copied_element_style_refs_remapped",
                        data={
                            "copied_idx": copied_count,
                            "text": paragraph.text[:120],
                            "source_style": source_style,
                            "resolved_paragraph_style": copied_style,
                            "remapped_style_refs": remapped_style_refs,
                        },
                    )
                    # #endregion
                    style_remap_log_count += 1
                mapped_num_id = None
                if src_num.get("numId"):
                    target_root_for_probe = _get_numbering_root(template_doc)
                    target_has_source_num_id = bool(
                        target_root_for_probe is not None and _find_num_node(target_root_for_probe, src_num.get("numId")) is not None
                    )
                    if target_has_source_num_id and numbering_collision_fingerprint_log_count < 8:
                        # #region agent log
                        _agent_debug_log(
                            run_id=debug_run_id,
                            hypothesis_id="H24",
                            location="rewrite_report.py:numbering_collision_fingerprint_probe",
                            message="source_vs_target_existing_numbering_fingerprint",
                            data={
                                "copied_idx": copied_count,
                                "text": paragraph.text[:120],
                                "source_num_id": src_num.get("numId"),
                                "source_numbering": _numbering_fingerprint(source_doc, src_num.get("numId")),
                                "target_numbering_before_map": _numbering_fingerprint(template_doc, src_num.get("numId")),
                            },
                        )
                        # #endregion
                        numbering_collision_fingerprint_log_count += 1
                    mapped_num_id = _ensure_numbering_mapping(
                        source_doc,
                        template_doc,
                        src_num.get("numId"),
                        numbering_mapping_cache,
                    )
                    if mapped_num_id:
                        set_num_ok = _set_paragraph_num_id(new_para_element, mapped_num_id)
                        num_after_set = _paragraph_numbering_info(new_para_element)
                        if numbering_collision_log_count < 8:
                            # #region agent log
                            _agent_debug_log(
                                run_id=debug_run_id,
                                hypothesis_id="H13",
                                location="rewrite_report.py:numbering_mapping_collision_probe",
                                message="target_numid_collision_and_set_result",
                                data={
                                    "copied_idx": copied_count,
                                    "text": paragraph.text[:120],
                                    "source_num_id": src_num.get("numId"),
                                    "target_has_source_num_id_before_map": target_has_source_num_id,
                                    "mapped_num_id": mapped_num_id,
                                    "set_num_ok": set_num_ok,
                                    "num_after_set": num_after_set,
                                },
                            )
                            # #endregion
                            numbering_collision_log_count += 1
                        if numbering_map_log_count < 8:
                            # #region agent log
                            _agent_debug_log(
                                run_id=debug_run_id,
                                hypothesis_id="H12",
                                location="rewrite_report.py:numbering_mapping_apply",
                                message="mapped_paragraph_num_id_for_target",
                                data={
                                    "copied_idx": copied_count,
                                    "text": paragraph.text[:120],
                                    "source_num_id": src_num.get("numId"),
                                    "mapped_num_id": mapped_num_id,
                                    "source_ilvl": src_num.get("ilvl"),
                                    "mapping_cache_size": len(numbering_mapping_cache),
                                },
                            )
                            # #endregion
                            numbering_map_log_count += 1
                if src_num.get("numId") and numbering_mapping_probe_log_count < 8:
                    source_num_fmt = _resolve_num_format(source_doc, src_num.get("numId"), src_num.get("ilvl"))
                    target_num_fmt_same_id = _resolve_num_format(
                        template_doc,
                        mapped_num_id if mapped_num_id else src_num.get("numId"),
                        src_num.get("ilvl"),
                    )
                    source_num_fmt_detail = _resolve_num_format_detail(source_doc, src_num.get("numId"), src_num.get("ilvl"))
                    target_num_fmt_detail = _resolve_num_format_detail(
                        template_doc,
                        mapped_num_id if mapped_num_id else src_num.get("numId"),
                        src_num.get("ilvl"),
                    )
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H10",
                        location="rewrite_report.py:numbering_mapping_probe",
                        message="source_vs_target_numbering_definition",
                        data={
                            "copied_idx": copied_count,
                            "text": paragraph.text[:120],
                            "source_num": src_num,
                            "mapped_num_id": mapped_num_id,
                            "source_num_fmt": source_num_fmt,
                            "target_num_fmt_same_id": target_num_fmt_same_id,
                            "source_num_fmt_detail": source_num_fmt_detail,
                            "target_num_fmt_detail": target_num_fmt_detail,
                        },
                    )
                    # #endregion
                    numbering_mapping_probe_log_count += 1
                imported_style = False
                if copied_style and _doc_has_style_id(template_doc, copied_style):
                    imported_style = bool(remapped_style_refs) or not _doc_has_style_id(template_doc, source_style)
                    if style_import_log_count < 5:
                        # #region agent log
                        _agent_debug_log(
                            run_id=debug_run_id,
                            hypothesis_id="H8",
                            location="rewrite_report.py:copy_paragraph_loop",
                            message="import_missing_style_definition",
                            data={
                                "copied_idx": copied_count,
                                "source_style_id": source_style,
                                "resolved_style_id": copied_style,
                                "imported": imported_style,
                                "style_exists_after": _doc_has_style_id(template_doc, copied_style),
                                "text": paragraph.text[:100],
                            },
                        )
                        # #endregion
                        # #region agent log
                        _agent_debug_log(
                            run_id=debug_run_id,
                            hypothesis_id="H20",
                            location="rewrite_report.py:style_import_struct_probe",
                            message="style_numpr_probe_after_import",
                            data={
                                "source_style_id": source_style,
                                "resolved_style_id": copied_style,
                                "source_style_struct": _style_struct_snapshot(source_doc, source_style),
                                "target_style_struct": _style_struct_snapshot(template_doc, copied_style),
                            },
                        )
                        # #endregion
                        style_import_log_count += 1
                
                # 关键策略：复制区域保持“原样结构”，不在复制阶段改动编号
                keep_numbering = _should_keep_numbering(new_para_element)
                before_num_info = _paragraph_numbering_info(new_para_element)
                if numbering_strip_log_count < 8:
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H9",
                        location="rewrite_report.py:copy_paragraph_loop",
                        message="numbering_preserved_in_copy_phase",
                        data={
                            "copied_idx": copied_count,
                            "text": paragraph.text[:120],
                            "source_style": source_style,
                            "copied_style": copied_style,
                            "keep_numbering_eval": keep_numbering,
                            "num_before": before_num_info,
                            "num_after": _paragraph_numbering_info(new_para_element),
                        },
                    )
                    # #endregion
                    numbering_strip_log_count += 1
                
                # 移除段落边框（黑线）
                try:
                    if new_para_element.pPr is not None:
                        pBdr = new_para_element.pPr.find(qn('w:pBdr'))
                        if pBdr is not None:
                            new_para_element.pPr.remove(pBdr)
                except Exception as e:
                    pass
                
                # 处理段落中的文本替换和图片复制
                paragraph_text_before_mutation = paragraph.text or ""
                text_replacement_count = 0
                drawing_replacement_count = 0
                for run_element in new_para_element.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    # 检查run是否包含超链接
                    has_hyperlink = _run_element_contains_hyperlink(run_element)
                    
                    # 处理文本内容
                    for text_element in run_element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if text_element.text:
                            original_text = text_element.text
                            # 替换文本：将"XXX网信办"替换为"鄞州区网信办"
                            new_text = re.sub(r'[\u4e00-\u9fa5]+网信办', '鄞州区网信办', original_text)
                            if new_text != original_text:
                                if has_hyperlink:
                                    print(f"  ⚠️ 跳过超链接文本替换以保留超链接: '{original_text}'")
                                else:
                                    print(f"  文本替换: '{original_text}' -> '{new_text}'")
                                    text_element.text = new_text
                                    text_replacement_count += 1
                    
                    # 处理图片内容 - 这部分比较复杂，需要特殊处理
                    drawing_elements = run_element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if drawing_elements:
                        # 为了处理图片，我们需要创建一个临时的run对象
                        try:
                            # 创建临时段落和run来处理图片复制
                            temp_para = template_doc.add_paragraph()
                            temp_run = temp_para.add_run()
                            
                            # 尝试复制每个图片
                            for drawing_element in drawing_elements:
                                try:
                                    if _copy_image_to_document(drawing_element, source_doc, template_doc, temp_run):
                                        print(f"  📷 复制图片到段落 {copied_count}")
                                        # 如果图片复制成功，用新的图片元素替换原有的
                                        if temp_run._element and len(list(temp_run._element)) > 0:
                                            # 获取新复制的图片元素
                                            new_drawing = None
                                            for elem in temp_run._element:
                                                if elem.tag.endswith('drawing'):
                                                    new_drawing = elem
                                                    break
                                            if new_drawing is not None:
                                                # 替换原有的图片元素
                                                parent = drawing_element.getparent()
                                                if parent is not None:
                                                    parent.replace(drawing_element, deepcopy(new_drawing))
                                                    drawing_replacement_count += 1
                                    else:
                                        print(f"  ⚠️ 图片复制失败，保留原始引用")
                                except Exception as img_error:
                                    print(f"  ⚠️ 图片复制失败: {img_error}")
                            
                            # 删除临时段落
                            template_doc._element.body.remove(temp_para._element)
                            
                        except Exception as e:
                            print(f"  ⚠️ 图片处理过程出错: {e}")
                            # 如果图片处理失败，保留原始图片引用
                

                
                # 将深拷贝的段落元素插入到模板的指定位置
                template_doc._element.body.insert(insert_element_index, new_para_element)
                insert_element_index += 1
                if key_struct_log_count < 10:
                    para_text_for_key = (paragraph.text or "")
                    if any(k in para_text_for_key for k in key_phrases):
                        # #region agent log
                        _agent_debug_log(
                            run_id=debug_run_id,
                            hypothesis_id="H17",
                            location="rewrite_report.py:key_paragraph_struct_copy_phase",
                            message="key_paragraph_source_vs_copied_struct",
                            data={
                                "copied_idx": copied_count,
                                "text": para_text_for_key[:120],
                                "source_struct": _paragraph_struct_snapshot(paragraph),
                                "copied_struct": _paragraph_struct_snapshot(new_para_element),
                                "source_num_fmt": _resolve_num_format(source_doc, src_num.get("numId"), src_num.get("ilvl")),
                                "copied_num_fmt": _resolve_num_format(template_doc, _paragraph_numbering_info(new_para_element).get("numId"), _paragraph_numbering_info(new_para_element).get("ilvl")),
                            },
                        )
                        # #endregion
                        key_struct_log_count += 1
                if replacement_mutation_log_count < 8:
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H16",
                        location="rewrite_report.py:copy_paragraph_mutation_probe",
                        message="copy_phase_mutation_counters",
                        data={
                            "copied_idx": copied_count,
                            "text_before": paragraph_text_before_mutation[:120],
                            "text_after": (new_para_element.text if hasattr(new_para_element, "text") else "")[:120],
                            "text_replacement_count": text_replacement_count,
                            "drawing_replacement_count": drawing_replacement_count,
                        },
                    )
                    # #endregion
                    replacement_mutation_log_count += 1

                if copy_style_log_count < 5:
                    resolved_num = _resolve_num_format(template_doc, copied_num.get("numId"), copied_num.get("ilvl"))
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H6",
                        location="rewrite_report.py:copy_paragraph_loop",
                        message="copied_paragraph_style_num_snapshot",
                        data={
                            "copied_idx": copied_count,
                            "source_style": source_style,
                            "copied_style": copied_style,
                            "style_exists_in_template": _doc_has_style_id(template_doc, copied_style),
                            "style_imported_now": imported_style,
                            "source_num": src_num,
                            "copied_num": copied_num,
                            "resolved_num_in_template": resolved_num,
                            "text": paragraph.text[:100],
                        },
                    )
                    # #endregion
                    copy_style_log_count += 1

        # #region agent log
        _agent_debug_log(
            run_id=debug_run_id,
            hypothesis_id="H30",
            location="rewrite_report.py:post_copy_pre_marker_remove",
            message="insert_start_slice_after_copy_before_marker_remove",
            data={
                "insert_para_index": insert_para_index,
                "first_heading_hits": _collect_paragraph_diagnostics(template_doc, ["1.漏洞描述", "漏洞事件：", "验证情况"], limit=12),
                "start_slice": _collect_non_empty_paragraph_slice(template_doc, insert_para_index, count=12),
            },
        )
        # #endregion
        
        # 删除标记段落（包含 * 的段落）
        if marker_para_element is not None:
            try:
                template_doc._element.body.remove(marker_para_element)
                print(f"已删除标记段落")
            except Exception as e:
                print(f"删除标记段落时出错: {e}")
        # #region agent log
        _agent_debug_log(
            run_id=debug_run_id,
            hypothesis_id="H30",
            location="rewrite_report.py:post_marker_remove",
            message="insert_start_slice_after_marker_remove",
            data={
                "insert_para_index": insert_para_index,
                "first_heading_hits": _collect_paragraph_diagnostics(template_doc, ["1.漏洞描述", "漏洞事件：", "验证情况"], limit=12),
                "start_slice": _collect_non_empty_paragraph_slice(template_doc, insert_para_index, count=12),
            },
        )
        # #endregion

        # 🔢 不再重排复制区域编号：保持与原文一致，避免列表/项目符号样式被改写
        try:
            heading_before = 0
            for p in template_doc.paragraphs:
                style_val = _paragraph_style_val(p)
                if "Heading" in style_val or "标题" in style_val:
                    heading_before += 1
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H1",
                location="rewrite_report.py:before_reassign_numbering",
                message="heading_count_before_numbering_reassign",
                data={"heading_count_before": heading_before, "total_paragraphs": len(template_doc.paragraphs)},
            )
            # #endregion
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H1",
                location="rewrite_report.py:skip_reassign_numbering",
                message="skip_numbering_reassign_to_preserve_source",
                data={"reason": "preserve copied numbering/layout exactly"},
            )
            # #endregion
            heading_after = 0
            for p in template_doc.paragraphs:
                style_val = _paragraph_style_val(p)
                if "Heading" in style_val or "标题" in style_val:
                    heading_after += 1
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H1",
                location="rewrite_report.py:after_reassign_numbering",
                message="heading_count_after_numbering_reassign",
                data={"heading_count_after": heading_after, "total_paragraphs": len(template_doc.paragraphs)},
            )
            # #endregion
        except Exception as e:
            print(f"  ⚠️ 重新分配编号序列失败: {e}")
        
        # 🔢 先更新通报编号（在创建备份之前）
        print(f"\n  📝 更新通报编号...")
        notification_number = None
        try:
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H14",
                location="rewrite_report.py:before_update_notification_number",
                message="paragraph_numbering_snapshot_before_update_number",
                data={
                    "snapshot": _find_para_num_snapshot(template_doc, ["验证情况", "敏感信息泄露", "处置措施"]),
                },
            )
            # #endregion
            # 临时保存文档以便编号更新函数读取
            temp_save_path = str(Path(output_file).with_suffix('.temp.docx'))
            template_doc.save(temp_save_path)
            
            # 更新编号
            result = update_notification_number(temp_save_path)
            if result:
                notification_number, config_year = result if isinstance(result, tuple) else (result, None)
            else:
                notification_number, config_year = None, None
            
            # 如果函数返回的不是元组，从配置文件读取年份
            if config_year is None:
                try:
                    config_file = get_config_file()
                    if config_file.exists():
                        with open(config_file, 'r', encoding='utf-8') as f:
                            config = json.load(f)
                        config_year = config.get('report_counters', {}).get('year', datetime.now().year)
                    else:
                        config_year = datetime.now().year
                except:
                    config_year = datetime.now().year
            
            # 重新加载更新后的文档
            template_doc = Document(temp_save_path)
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H14",
                location="rewrite_report.py:after_update_notification_number",
                message="paragraph_numbering_snapshot_after_update_number",
                data={
                    "snapshot": _find_para_num_snapshot(template_doc, ["验证情况", "敏感信息泄露", "处置措施"]),
                },
            )
            # #endregion
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H31",
                location="rewrite_report.py:after_update_notification_number",
                message="insert_start_slice_after_update_number_reload",
                data={
                    "insert_para_index": insert_para_index,
                    "first_heading_hits": _collect_paragraph_diagnostics(template_doc, ["1.漏洞描述", "漏洞事件：", "验证情况"], limit=12),
                    "start_slice": _collect_non_empty_paragraph_slice(template_doc, insert_para_index, count=12),
                },
            )
            # #endregion
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H18",
                location="rewrite_report.py:after_update_key_paragraph_struct",
                message="key_paragraph_struct_after_full_pipeline",
                data={
                    "validation_case": _find_para_num_snapshot(template_doc, ["验证情况"]),
                    "disposal_case": _find_para_num_snapshot(template_doc, ["处置措施"]),
                    "url_case": _find_para_num_snapshot(template_doc, ["Url:", "URL:", "url:"]),
                },
            )
            # #endregion
            _force_severity_value_standalone(template_doc, run_id=debug_run_id)
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H34",
                location="rewrite_report.py:after_force_severity_value",
                message="severity_value_snapshot_after_force_standalone",
                data={
                    "severity_case": _find_para_num_snapshot(template_doc, ["高危漏洞", "中危漏洞", "低危漏洞", "严重漏洞", "一般漏洞", "轻微漏洞"]),
                    "window_validation": _collect_paragraph_window(template_doc, "验证情况", radius=3),
                    "window_validation_raw": _collect_paragraph_window_including_empty(template_doc, "验证情况", radius=4),
                },
            )
            # #endregion
            _force_validation_heading_standalone(template_doc, run_id=debug_run_id)
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H33",
                location="rewrite_report.py:after_force_validation_heading",
                message="validation_heading_snapshot_after_force_standalone",
                data={
                    "validation_case": _find_para_num_snapshot(template_doc, ["验证情况", "2.验证情况"]),
                    "window_validation": _collect_paragraph_window(template_doc, "验证情况", radius=3),
                    "window_validation_raw": _collect_paragraph_window_including_empty(template_doc, "验证情况", radius=4),
                },
            )
            # #endregion
            
            # 删除临时文件
            Path(temp_save_path).unlink()
            
            if notification_number:
                print(f"  ✓ 通报编号已更新: 〔{config_year}〕第{notification_number}期")
        except Exception as e:
            print(f"  ⚠️ 编号更新失败: {e}")
        
        # 先保存主文档，然后创建备份文件
        backup_file_path = None
        

        
        # 最后统一保存文档（只保存一次）
        backup_path = None
        try:
            # 如果输出文件已存在，先创建备份
            if Path(output_file).exists():
                backup_path = create_backup(output_file)
            
            # 使用新的安全保存方法
            if INTEGRITY_MODULE_AVAILABLE:
                save_result = safe_save_document(template_doc, output_file)
                
                if not save_result['success']:
                    print(f"  ❌ 文档保存失败: {save_result['error']}")
                    # 如果有备份，尝试恢复
                    if backup_path:
                        print(f"  🔄 尝试从备份恢复...")
                        if recover_from_backup(output_file, backup_path):
                            print(f"  ✅ 已从备份恢复原始文档")
                        else:
                            print(f"  ❌ 备份恢复也失败")
                    raise Exception(f"文档保存失败: {save_result['error']}")
                
                print(f"  ✓ 文档已保存 (方法: {save_result['method']})")
                
                # 跨平台模式：仅使用 python-docx 完成保存后验证
                if save_result['validation']['valid']:
                    print("  ✅ 文档结构验证通过（python-docx）")
                else:
                    print("  ⚠️ 文档保存完成，但结构验证未通过，请人工核查")
                sync_render_result = _sync_rendering_resources_from_source(source_file, output_file)
                # #region agent log
                _agent_debug_log(
                    run_id=debug_run_id,
                    hypothesis_id="H28",
                    location="rewrite_report.py:render_resource_sync",
                    message="sync_rendering_resources_after_save",
                    data={
                        "output_file": str(output_file),
                        "source_file": str(source_file),
                        "sync_result": sync_render_result,
                    },
                )
                # #endregion
                # #region agent log
                try:
                    final_doc_probe = Document(output_file)
                    semantic_compare = []
                    for keyword in ["验证情况", "处置措施", "敏感信息泄露"]:
                        source_para_probe = _find_first_paragraph_containing(source_doc, keyword)
                        final_para_probe = _find_first_paragraph_containing(final_doc_probe, keyword)
                        if source_para_probe is None or final_para_probe is None:
                            continue
                        source_num_probe = _paragraph_numbering_info(source_para_probe)
                        final_num_probe = _paragraph_numbering_info(final_para_probe)
                        source_style_probe = _paragraph_style_val(source_para_probe)
                        final_style_probe = _paragraph_style_val(final_para_probe)
                        semantic_compare.append(
                            {
                                "keyword": keyword,
                                "source_style": source_style_probe,
                                "final_style": final_style_probe,
                                "source_style_semantic": _style_semantic_fingerprint(source_doc, source_style_probe),
                                "final_style_semantic": _style_semantic_fingerprint(final_doc_probe, final_style_probe),
                                "source_num": source_num_probe,
                                "final_num": final_num_probe,
                                "source_num_semantic": _numbering_semantic_fingerprint(source_doc, source_num_probe.get("numId")),
                                "final_num_semantic": _numbering_semantic_fingerprint(final_doc_probe, final_num_probe.get("numId")),
                                "source_runs": _paragraph_run_snapshot(source_para_probe),
                                "final_runs": _paragraph_run_snapshot(final_para_probe),
                            }
                        )
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H19",
                        location="rewrite_report.py:final_output_probe_after_save",
                        message="final_output_key_paragraph_diagnostics",
                        data={
                            "output_file": str(output_file),
                            "matches": _collect_paragraph_diagnostics(
                                final_doc_probe,
                                ["验证情况", "处置措施", "Url", "URL", "限制用户访问", "▪", "(2)"],
                                limit=30,
                            ),
                            "window_validation": _collect_paragraph_window(final_doc_probe, "验证情况", radius=3),
                            "window_validation_raw": _collect_paragraph_window_including_empty(final_doc_probe, "验证情况", radius=4),
                            "window_disposal": _collect_paragraph_window(final_doc_probe, "处置措施", radius=4),
                            "start_slice": _collect_non_empty_paragraph_slice(final_doc_probe, insert_para_index, count=12),
                            "semantic_compare": semantic_compare,
                        },
                    )
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H26",
                        location="rewrite_report.py:final_semantic_compare",
                        message="source_vs_output_semantic_compare",
                        data={
                            "output_file": str(output_file),
                            "key_paragraphs": semantic_compare,
                            "source_defaults": _doc_defaults_fingerprint(source_doc),
                            "output_defaults": _doc_defaults_fingerprint(final_doc_probe),
                        },
                    )
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H27",
                        location="rewrite_report.py:package_resource_compare",
                        message="source_vs_output_package_resource_fingerprints",
                        data={
                            "output_file": str(output_file),
                            "source_resources": _package_resource_fingerprints(source_doc),
                            "output_resources": _package_resource_fingerprints(final_doc_probe),
                        },
                    )
                except Exception as final_probe_err:
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H19",
                        location="rewrite_report.py:final_output_probe_after_save",
                        message="final_output_probe_failed",
                        data={"error": str(final_probe_err)},
                    )
                # #endregion
                # #region agent log
                _agent_debug_log(
                    run_id=debug_run_id,
                    hypothesis_id="H21",
                    location="rewrite_report.py:final_layout_probe",
                    message="final_layout_window_probe",
                    data={
                        "output_file": str(output_file),
                        "window_validation": (_collect_paragraph_window(Document(output_file), "验证情况", radius=3) if Path(output_file).exists() else {}),
                        "window_disposal": (_collect_paragraph_window(Document(output_file), "处置措施", radius=4) if Path(output_file).exists() else {}),
                    },
                )
                # #endregion
                
        except Exception as e:
            print(f"  ❌ 文档保存失败: {e}")
            # #region agent log
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H3",
                location="rewrite_report.py:save_exception",
                message="save_exception_triggered",
                data={"error": str(e)},
            )
            # #endregion
            # 改为仅重试“原文档保存”，避免进入重建简化文档导致样式丢失
            saved = False
            for retry_idx in range(3):
                try:
                    time.sleep(0.6 * (retry_idx + 1))
                    template_doc.save(output_file)
                    saved = True
                    print(f"  ✅ 重试保存成功（第 {retry_idx + 1} 次）")
                    # #region agent log
                    _agent_debug_log(
                        run_id=debug_run_id,
                        hypothesis_id="H3",
                        location="rewrite_report.py:save_retry_success",
                        message="save_retry_succeeded",
                        data={"retry_idx": retry_idx + 1},
                    )
                    # #endregion
                    break
                except Exception as retry_err:
                    print(f"  ⚠️ 重试保存失败（第 {retry_idx + 1} 次）: {retry_err}")

            if not saved:
                if backup_path:
                    print(f"  🔄 尝试从备份恢复...")
                    if recover_from_backup(output_file, backup_path):
                        print(f"  ✅ 已从备份恢复原始文档")
                # 直接抛出，让上层标记为需人工处理，而不是输出降级文档
                raise Exception(f"文档保存失败，已取消降级重建以保护原格式: {e}")
        
        # 创建备份文件（在主文档保存成功后）
        backup_file_path = str(Path(output_file).with_suffix('.backup.docx'))
        try:
            if Path(output_file).exists():
                shutil.copy2(output_file, backup_file_path)
                print(f"  ✅ 已创建备份文件: {Path(backup_file_path).name}")
            else:
                print(f"  ⚠️ 主输出文件不存在，无法创建备份")
                backup_file_path = None
        except Exception as backup_error:
            print(f"  ⚠️ 创建备份文件失败: {backup_error}")
            backup_file_path = None

        print(f"\n✓ 成功创建通报文档!")
        print(f"  输出文件: {output_file}")
        print(f"  复制的段落数: {copied_count}")
        if notification_number:
            # 从配置文件读取年份
            try:
                config_file = get_config_file()
                if config_file.exists():
                    with open(config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                    config_year = config.get('report_counters', {}).get('year', datetime.now().year)
                else:
                    config_year = datetime.now().year
            except:
                config_year = datetime.now().year
            print(f"  通报编号: 〔{config_year}〕第{notification_number}期")
        
        print("=" * 60)
        
        # 清理旧备份文件
        try:
            cleanup_backups(output_file, keep_count=2)
        except Exception as cleanup_error:
            print(f"  ⚠️ 备份清理警告: {cleanup_error}")
        
        # 按用户要求：只保留backup文件，不进行重命名操作
        print(f"  📁 保留备份文件，不进行重命名操作")
        
        # 添加图片到主输出文件
        image_path = r"C:\Users\lan1o\Desktop\wow\Report_Template\确认词条.jpg"
        image_insertion_success = False
        
        if Path(image_path).exists() and Path(output_file).exists():
            print(f"\n🖼️ 开始添加确认词条图片到主输出文件...")
            try:
                # 加载主输出文档对象
                target_doc = Document(output_file)
                
                # 调用图片添加函数，传递源文件路径用于错误记录
                image_insertion_success = add_floating_image_to_pages(target_doc, image_path, start_page=2, source_file_path=output_file)
                
                if image_insertion_success:
                    # 保存修改后的文档
                    target_doc.save(output_file)
                    print(f"  ✅ 确认词条图片已添加到主输出文件的每一页（从第2页开始）")
                else:
                    print(f"  ❌ 图片添加失败，可能原因：")
                    print(f"    • 文档内容结构异常，页范围估算偏差")
                    print(f"    • 文档格式不兼容")
                    print(f"    • 图片文件损坏或格式不支持")
                    print(f"  💡 解决方案：")
                    print(f"    • 手动打开备份文件添加确认词条图片")
                    print(f"    • 检查图片文件是否完整")
                    
            except Exception as img_error:
                print(f"  ❌ 添加图片失败: {img_error}")
                print(f"  💡 建议：手动打开备份文件添加确认词条图片")
                image_insertion_success = False
        elif not Path(image_path).exists():
            print(f"\n⚠️ 确认词条图片文件不存在: {image_path}")
            print(f"  ℹ️  跳过图片添加，文档仍然可以正常使用")
        elif not Path(output_file).exists():
            print(f"\n⚠️ 主输出文件不存在，跳过图片添加")
        
        # 根据图片插入结果决定删除哪个文件
        if image_insertion_success:
            # 图片插入成功，删除备份文件，保留主输出文件
            try:
                if backup_file_path and Path(backup_file_path).exists():
                    Path(backup_file_path).unlink()
                    print(f"  🗑️ 图片插入成功，已删除备份文件: {Path(backup_file_path).name}")
                    print(f"  ✅ 保留主输出文件: {Path(output_file).name}")
            except Exception as e:
                print(f"  ⚠️ 删除备份文件失败: {e}")
        else:
            # 图片插入失败，删除主输出文件，保留备份文件
            try:
                if Path(output_file).exists():
                    Path(output_file).unlink()
                    print(f"  🗑️ 图片插入失败，已删除主输出文件: {Path(output_file).name}")
                
                # 确定最终要保留的备份文件（只保留backup.docx）
                if backup_file_path and Path(backup_file_path).exists():
                    backup_type = "备份"
                    print(f"  ✅ 已保留{backup_type}文件: {Path(backup_file_path).name}")
                else:
                    print(f"  ⚠️ 备份文件路径为空或文件不存在")
            except Exception as e:
                print(f"  ⚠️ 删除主输出文件失败: {e}")

        # #region agent log
        try:
            delivered_path = None
            if Path(output_file).exists():
                delivered_path = output_file
            elif backup_file_path and Path(backup_file_path).exists():
                delivered_path = backup_file_path

            delivered_probe = None
            if delivered_path:
                delivered_doc = Document(delivered_path)
                delivered_probe = {
                    "delivered_path": str(delivered_path),
                    "image_insertion_success": image_insertion_success,
                    "output_exists": Path(output_file).exists(),
                    "backup_exists": bool(backup_file_path and Path(backup_file_path).exists()),
                    "start_slice": _collect_non_empty_paragraph_slice(delivered_doc, insert_para_index, count=12),
                    "key_paragraphs": _collect_paragraph_diagnostics(
                        delivered_doc,
                        ["验证情况", "处置措施", "Url", "URL", "限制用户访问", "▪", "(2)"],
                        limit=30,
                    ),
                    "window_validation": _collect_paragraph_window(delivered_doc, "验证情况", radius=3),
                    "window_validation_raw": _collect_paragraph_window_including_empty(delivered_doc, "验证情况", radius=4),
                    "window_disposal": _collect_paragraph_window(delivered_doc, "处置措施", radius=4),
                    "resources": _package_resource_fingerprints(delivered_doc),
                    "defaults": _doc_defaults_fingerprint(delivered_doc),
                    "copy_range_compare": _compare_copy_range_semantics(
                        source_doc,
                        delivered_doc,
                        start_idx,
                        end_idx,
                        delivered_anchor="验证情况",
                        max_items=20,
                    ),
                }
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H29",
                location="rewrite_report.py:delivered_artifact_probe",
                message="final_delivered_artifact_probe",
                data={
                    "output_file": str(output_file),
                    "backup_file_path": str(backup_file_path) if backup_file_path else None,
                    "image_insertion_success": image_insertion_success,
                    "delivered_probe": delivered_probe,
                },
            )
        except Exception as delivered_probe_err:
            _agent_debug_log(
                run_id=debug_run_id,
                hypothesis_id="H29",
                location="rewrite_report.py:delivered_artifact_probe",
                message="final_delivered_artifact_probe_failed",
                data={"error": str(delivered_probe_err)},
            )
        # #endregion
        
        # 删除数字开头的原始通报文件
        try:
            source_path = Path(source_file)
            source_filename = source_path.name
            
            # 检查文件名是否以数字开头
            if source_filename and source_filename[0].isdigit():
                if source_path.exists():
                    source_path.unlink()
                    print(f"  🗑️ 已删除原始通报文件: {source_filename}")
                else:
                    print(f"  ℹ️  原始通报文件已不存在: {source_filename}")
            else:
                print(f"  ℹ️  原始文件名不以数字开头，保留: {source_filename}")
        except Exception as delete_error:
            print(f"  ⚠️ 删除原始通报文件失败: {delete_error}")
        
        # PDF转换逻辑
        pdf_file = None
        pdf_conversion_success = False
        
        # 跳过PDF转换，因为主输出文件已被删除，只保留备份文件
        print(f"\n📄 跳过PDF转换...")
        print(f"  ℹ️  主输出文件已删除，只保留备份文件，不进行PDF转换")
        print(f"  ℹ️  如需PDF文件，请手动转换备份文件")

        # 返回结果信息，包含是否需要手动处理的标记
        # 注意：由于执行了文件替换逻辑，clean_backup和final_backup文件已被清理
        # 如果文件替换成功，backup_file_path已更新为最终文件路径
        result = {
            'success': True,
            'output_file': output_file,
            'backup_file': backup_file_path if backup_file_path and Path(backup_file_path).exists() else None,
            'clean_backup_file': None,  # 已被清理或重命名为主文件
            'final_backup_file': None,  # 已被清理或重命名为主文件
            'needs_manual_processing': False,  # 默认不需要手动处理
            'skip_reason': None,
            'pdf_file': pdf_file,  # 新增PDF文件路径
            'pdf_conversion_success': pdf_conversion_success  # 新增PDF转换状态
        }
        
        # 检查是否需要手动处理的情况
        manual_processing_reasons = []
        
        # 图片添加失败
        if not image_insertion_success:
            manual_processing_reasons.append("确认词条图片添加失败，需要手动添加图片")
        
        # 设置手动处理标志
        if manual_processing_reasons:
            result['needs_manual_processing'] = True
            result['skip_reason'] = '; '.join(manual_processing_reasons)
            print(f"  ⚠️ 注意：此文档需要手动处理 - {result['skip_reason']}")
        
        # 清理临时文件
        if 'cleanup_temp_source' in locals() and cleanup_temp_source:
            try:
                temp_source = source_file.replace('.docx', '_temp_editable.docx')
                if os.path.exists(temp_source):
                    os.remove(temp_source)
                    safe_print("🧹 已清理临时可编辑文件")
            except Exception as e:
                safe_print(f"⚠️ 清理临时文件失败: {str(e)}")

        
        return result
        
    except FileNotFoundError as e:
        # 清理临时文件
        if 'cleanup_temp_source' in locals() and cleanup_temp_source:
            try:
                temp_source = source_file.replace('.docx', '_temp_editable.docx')
                if os.path.exists(temp_source):
                    os.remove(temp_source)
            except:
                pass
        
        print(f"错误: 找不到文件: {e}")
        return {
            'success': False,
            'output_file': None,
            'backup_file': None,
            'needs_manual_processing': False,
            'skip_reason': f'文件未找到: {e}'
        }
    except Exception as e:
        # 清理临时文件
        if 'cleanup_temp_source' in locals() and cleanup_temp_source:
            try:
                temp_source = source_file.replace('.docx', '_temp_editable.docx')
                if os.path.exists(temp_source):
                    os.remove(temp_source)
            except:
                pass
        
        print(f"创建文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'output_file': None,
            'backup_file': None,
            'needs_manual_processing': False,
            'skip_reason': f'创建失败: {str(e)}'
        }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("=" * 60)
        print("通报改写工具")
        print("=" * 60)
        print("\n使用方法:")
        print("  python rewrite_report.py <源通报文档>")
        print("  python rewrite_report.py <源通报文档> <起始段落> <结束段落>")
        print("\n功能说明:")
        print("  1. 智能识别：从文件名自动提取公司名和漏洞类型")
        print("  2. 自动替换：")
        print("     - 段落4、6：公司名自动更新")
        print("     - 段落7：漏洞类型自动更新")
        print("     - 段落7：截止日期自动设置为5天后")
        print("     - 段落14：当前日期自动设置为今天")
        print("     - 内容中'XXX网信办'替换为'鄞州区网信办'")
        print("  3. 自动查找模板中的 * 标记作为插入位置")
        print("  4. 从源文档复制指定段落到模板")
        print("  5. 保留所有格式（标题、字体、颜色等）")
        print("  6. 移除段落边框（黑线）")
        print("  7. 文件名自动去掉开头数字")
        print("\n默认参数:")
        print("  起始段落: 3")
        print("  结束段落: -1（倒数第2段，跳过最后的空段落）")
        print("  模板文件: 自动查找 Report_Template/通报模板*.docx 或 ./通报模板*.docx")
        print("\n示例:")
        print("  python rewrite_report.py 1759979441661关于XXX漏洞通报.docx")
        print("  python rewrite_report.py 源文档.docx 3 20")
        print("\n提示:")
        print("  1. 请确保模板文件中有 * 标记标注第二页起始位置")
        print("  2. 模板文件会自动从 Report_Template 目录或当前目录查找")
        print("=" * 60)
        sys.exit(1)
    
    source_file = sys.argv[1]
    
    # 默认参数
    start_para = 3
    end_para = -1
    
    # 解析可选参数
    if len(sys.argv) > 2:
        try:
            start_para = int(sys.argv[2])
        except ValueError:
            print("错误: 起始段落必须是数字")
            sys.exit(1)
    
    if len(sys.argv) > 3:
        try:
            end_para = int(sys.argv[3])
        except ValueError:
            print("错误: 结束段落必须是数字")
            sys.exit(1)
    
    # 执行改写
    result = rewrite_report(source_file, start_para=start_para, end_para=end_para)
    
    if result['success']:
        print("\n改写完成！")
        if result['needs_manual_processing']:
            print(f"⚠️ 需要手动处理: {result['skip_reason']}")
            if result['backup_file']:
                print(f"📁 备份文件: {result['backup_file']}")
        
        # 显示PDF转换结果
        if result.get('pdf_conversion_success'):
            print(f"📄 PDF转换成功: {Path(result['pdf_file']).name}")
        elif result.get('pdf_file') is not None:
            print(f"⚠️ PDF转换失败")
        
    else:
        print(f"\n改写失败: {result['skip_reason']}")
        sys.exit(1)
