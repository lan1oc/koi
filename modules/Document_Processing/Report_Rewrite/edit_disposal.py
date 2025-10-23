#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
处置文件模板处理工具
1. 如果文件夹中没有处置文件，则复制模板
2. 如果已有处置文件，则修改段落4为"鄞州区网信办："
"""

import sys
import os
from pathlib import Path
from docx import Document

# 设置Windows控制台编码为UTF-8
if sys.platform == 'win32':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass


def find_disposal_file(directory: Path = None) -> Path:
    """
    查找处置文件（处置文件模板或漏洞处置文件模板）
    
    参数:
        directory: 要查找的目录，默认为当前目录
    
    返回:
        Path对象，如果找到文件；否则返回None
    """
    if directory is None:
        directory = Path.cwd()
    
    # 查找所有.docx文件
    for file in directory.glob("*.docx"):
        # 检查文件名是否包含"处置"和"模板"
        if '处置' in file.name and ('模板' in file.name or '处置文件' in file.name):
            return file
    
    return None


def copy_disposal_template(template_file: str, target_directory: Path = None) -> bool:
    """
    复制处置文件模板到目标目录
    
    参数:
        template_file: 模板文件路径
        target_directory: 目标目录，默认为当前目录
    
    返回:
        bool: 成功返回True，失败返回False
    """
    try:
        if not os.path.exists(template_file):
            print(f"❌ 模板文件不存在: {template_file}")
            return False
        
        if target_directory is None:
            target_directory = Path.cwd()
        
        # 生成目标文件名（去掉文件名开头的数字）
        template_path = Path(template_file)
        template_name = template_path.name
        
        # 去掉文件名开头的数字
        import re
        clean_name = re.sub(r'^\d+', '', template_name)
        if not clean_name:
            clean_name = template_name
        
        target_file = target_directory / clean_name
        
        # 如果目标文件已存在，不复制
        if target_file.exists():
            print(f"ℹ️  处置文件已存在: {target_file.name}")
            return True
        
        # 复制文件
        import shutil
        shutil.copy2(template_file, target_file)
        print(f"✅ 已复制处置文件模板: {clean_name}")
        
        return True
        
    except Exception as e:
        print(f"❌ 复制模板时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def edit_disposal_document(disposal_file: Path = None) -> bool:
    """
    编辑处置文件，将段落4修改为"鄞州区网信办："
    
    参数:
        disposal_file: 处置文件路径，默认为自动查找
    
    返回:
        bool: 成功返回True，失败返回False
    """
    try:
        # 如果没有指定文件，自动查找
        if disposal_file is None:
            disposal_file = find_disposal_file()
        
        if disposal_file is None:
            print("⚠️ 未找到处置文件")
            return False
        
        if not disposal_file.exists():
            print(f"❌ 处置文件不存在: {disposal_file}")
            return False
        
        print(f"\n正在编辑处置文件:")
        print(f"  文件: {disposal_file.name}")
        print("=" * 60)
        
        # 读取文档
        doc = Document(disposal_file)
        
        # 检查段落数量
        if len(doc.paragraphs) < 4:
            print(f"⚠️ 文档段落数量不足（只有{len(doc.paragraphs)}个段落，需要至少4个）")
            return False
        
        # 获取段落4（索引3，因为从0开始）
        # 注意：段落0通常是空段落或标题行，实际的"段落4"可能是索引3或4
        # 根据read_word.py的输出，段落4显示的是"××网信办："
        target_para_idx = 3  # 段落索引从0开始，段落4是索引3
        
        # 检查段落内容
        para = doc.paragraphs[target_para_idx]
        original_text = para.text.strip()
        
        print(f"  段落 {target_para_idx + 1} 原文: {original_text}")
        
        # 检查是否已经是目标内容
        if original_text == "鄞州区网信办：":
            print("  ℹ️  段落内容已经正确，无需修改")
            return True
        
        # 替换内容 - 在每个run中查找并替换，保留格式
        replaced = False
        for run in para.runs:
            if '网信办' in run.text:
                # 替换为目标内容
                run.text = run.text.replace(run.text.strip(), '鄞州区网信办：')
                replaced = True
                break
        
        # 如果没有在run中找到，直接替换整个段落
        if not replaced:
            # 清空段落的所有runs
            for run in para.runs:
                run.text = ''
            # 添加新内容
            if para.runs:
                para.runs[0].text = '鄞州区网信办：'
            else:
                para.add_run('鄞州区网信办：')
        
        print(f"  段落 {target_para_idx + 1} 新文: {para.text}")
        
        # 保存文档
        doc.save(disposal_file)
        
        print(f"\n✅ 成功编辑处置文件!")
        print(f"  输出文件: {disposal_file.name}")
        print("=" * 60)
        
        return True
        
    except Exception as e:
        print(f"❌ 编辑文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def process_disposal(template_file: str = None, target_directory: Path = None) -> bool:
    """
    处理处置文件的完整流程
    1. 检查目标目录是否已有处置文件
    2. 如果没有，复制模板
    3. 如果有，修改段落4
    
    参数:
        template_file: 模板文件路径
        target_directory: 目标目录，默认为当前目录
    
    返回:
        bool: 成功返回True，失败返回False
    """
    try:
        if target_directory is None:
            target_directory = Path.cwd()
        
        print("\n" + "=" * 60)
        print("处置文件处理工具")
        print("=" * 60)
        print(f"目标目录: {target_directory}")
        
        # 1. 查找现有的处置文件
        disposal_file = find_disposal_file(target_directory)
        
        if disposal_file is None:
            # 2. 如果没有处置文件，复制模板
            print("\n📋 未找到处置文件，开始复制模板...")
            
            if template_file is None:
                # 尝试查找默认模板位置
                possible_templates = [
                    "Report_Template/处置文件模板.docx",
                    "../../../Report_Template/处置文件模板.docx",
                ]
                
                # 也尝试查找带数字前缀的模板
                template_dir = Path("Report_Template")
                if template_dir.exists():
                    for file in template_dir.glob("*处置*.docx"):
                        possible_templates.insert(0, str(file))
                
                for tmpl in possible_templates:
                    if os.path.exists(tmpl):
                        template_file = tmpl
                        break
                
                if template_file is None:
                    print("❌ 未找到处置文件模板！")
                    print("请指定模板文件路径，或将模板放在 Report_Template/ 目录下")
                    return False
            
            # 复制模板
            success = copy_disposal_template(template_file, target_directory)
            if not success:
                return False
            
            # 重新查找复制后的文件
            disposal_file = find_disposal_file(target_directory)
        else:
            print(f"\n📄 找到处置文件: {disposal_file.name}")
        
        # 3. 编辑处置文件
        print("\n✏️  开始编辑处置文件...")
        success = edit_disposal_document(disposal_file)
        
        if success:
            print("\n✅ 处置文件处理完成！")
        else:
            print("\n❌ 处置文件处理失败")
        
        return success
        
    except Exception as e:
        print(f"❌ 处理过程出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    print("=" * 60)
    print("处置文件模板处理工具")
    print("=" * 60)
    
    # 解析命令行参数
    template_file = None
    target_directory = None
    
    if len(sys.argv) >= 2:
        template_file = sys.argv[1]
    
    if len(sys.argv) >= 3:
        target_directory = Path(sys.argv[2])
    
    # 执行处理
    success = process_disposal(template_file, target_directory)
    
    if not success:
        print("\n使用方法:")
        print("  方法1: 在目标目录运行: python edit_disposal.py")
        print("  方法2: 指定模板: python edit_disposal.py <模板路径>")
        print("  方法3: 指定模板和目标目录: python edit_disposal.py <模板路径> <目标目录>")
        print("\n功能说明:")
        print("  1. 检查目标目录是否存在处置文件")
        print("  2. 如果不存在，复制模板到目标目录")
        print("  3. 如果存在，编辑段落4为'鄞州区网信办：'")
        print("=" * 60)
        sys.exit(1)

