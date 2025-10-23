#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
授权委托书编辑工具
自动将通报文件名填入授权委托书模板的*标记处
"""

import sys
import os
import re
from docx import Document

# 设置Windows控制台编码为UTF-8
if sys.platform == 'win32':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass


def edit_authorization(report_file, template_file=None):
    """
    编辑授权委托书，将通报文件名填入*标记处
    
    参数:
        report_file: 通报文档的路径（用于提取文件名）
        template_file: 授权委托书模板的路径
    """
    try:
        # 如果没有指定模板，尝试查找
        if template_file is None:
            # 尝试多个可能的模板位置
            possible_templates = [
                "Report_Template/授权委托书（执法调查类）.docx",
                "授权委托书（执法调查类）.docx",
                "授权委托书.docx",
                "Report_Template/授权委托书.docx",
            ]
            for tmpl in possible_templates:
                if os.path.exists(tmpl):
                    template_file = tmpl
                    break
            
            if template_file is None:
                print("错误: 找不到授权委托书模板文件！")
                print("请将模板文件放在以下位置之一：")
                for tmpl in possible_templates:
                    print(f"  - {tmpl}")
                return False
        
        # 从通报文档中读取标题（更可靠的方法，避免文件名编码问题）
        try:
            report_doc = Document(report_file)
            # 查找标题（通常在前几个非空段落中）
            report_title = None
            for para in report_doc.paragraphs[:10]:
                text = para.text.strip()
                if text and '关于' in text and '通报' in text:
                    report_title = text
                    break
            
            if not report_title:
                # 如果没找到，尝试从文件名提取
                report_basename = os.path.basename(report_file)
                report_name = report_basename.rsplit('.', 1)[0]
                report_title = re.sub(r'^\d+', '', report_name)
            
            report_name_clean = report_title
        except Exception as e:
            # 如果读取失败，使用文件名
            print(f"  注意: 无法读取通报文档内容，使用文件名: {e}")
            report_basename = os.path.basename(report_file)
            report_name = report_basename.rsplit('.', 1)[0]
            report_name_clean = re.sub(r'^\d+', '', report_name)
        
        # 读取模板
        doc = Document(template_file)
        
        print(f"\n正在编辑授权委托书:")
        print(f"  模板文件: {template_file}")
        print(f"  通报文件名: {report_name_clean}")
        print("=" * 60)
        
        # 查找并替换*标记
        found_marker = False
        replaced_count = 0
        
        for i, para in enumerate(doc.paragraphs, 1):
            if '*' in para.text:
                found_marker = True
                original_text = para.text
                
                # 在每个run中查找并替换*，保留格式
                for run in para.runs:
                    if '*' in run.text:
                        # 只替换*字符，保留run的所有格式
                        run.text = run.text.replace('*', report_name_clean)
                        replaced_count += 1
                
                print(f"  段落 {i} 已替换:")
                print(f"    原文: {original_text}")
                print(f"    新文: {para.text}")
                print(f"    格式: 完全保留（包括下划线、字体等）")
        
        if not found_marker:
            print("  警告: 未找到 * 标记！")
            print("  请在模板中添加 * 标记作为占位符。")
            return False
        
        # 生成输出文件名 - 保持模板原名
        template_basename = os.path.basename(template_file)
        # 如果是.doc格式，转换为.docx；如果已经是.docx则保持不变
        if template_basename.endswith('.doc') and not template_basename.endswith('.docx'):
            output_file = template_basename + 'x'
        else:
            output_file = template_basename
        
        # 保存文档
        doc.save(output_file)
        
        print(f"\n✓ 成功生成授权委托书!")
        print(f"  输出文件: {output_file}")
        print(f"  替换次数: {replaced_count}")
        print(f"  模板格式: 完全保留")
        print("=" * 60)
        
        return True
        
    except FileNotFoundError as e:
        print(f"错误: 找不到文件: {e}")
        return False
    except Exception as e:
        print(f"编辑文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    print("=" * 60)
    print("授权委托书编辑工具")
    print("=" * 60)
    
    # 自动查找通报文档（避免中文文件名编码问题）
    report_file = None
    
    if len(sys.argv) >= 2:
        # 如果提供了参数，使用参数
        report_file = sys.argv[1]
    else:
        # 自动查找当前目录的通报文档（使用os.listdir避免编码问题）
        possible_reports = []
        for filename in os.listdir('.'):
            if filename.endswith('.docx'):
                # 排除模板和授权委托书（注意"授权委托书"而不是单独的"授权"，因为可能有"未授权"）
                if 'template' not in filename and '授权委托书' not in filename and '模板' not in filename:
                    # 优先选择以"关于"开头或包含"通报"的文件
                    if filename.startswith('关于') or '通报' in filename:
                        possible_reports.append(filename)
        
        if possible_reports:
            report_file = possible_reports[0]
            print(f"\n自动找到通报文档: {report_file}")
        else:
            print("\n未找到通报文档！")
            print("\n使用方法:")
            print("  方法1: 将通报文档放在当前目录，运行: python edit_authorization.py")
            print("  方法2: 指定文件: python edit_authorization.py <通报文档路径>")
            print("\n功能说明:")
            print("  1. 从通报文档内容中读取标题")
            print("  2. 在授权委托书模板的 * 标记处填入通报标题")
            print("  3. 生成授权委托书文件")
            print("\n重要提示:")
            print("  1. 请确保模板文件中有 * 标记作为占位符")
            print("  2. 模板文件必须是 .docx 格式")
            print("=" * 60)
            sys.exit(1)
    
    # 执行编辑
    success = edit_authorization(report_file)
    
    if success:
        print("\n编辑完成！")
    else:
        print("\n编辑失败，请检查错误信息。")
        sys.exit(1)

