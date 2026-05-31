#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
责令整改通知书编辑工具
自动替换公司名、漏洞类型和日期
"""

import sys
import io
import os
import re
import json
import shutil
import tempfile
from datetime import datetime
from docx import Document
from pathlib import Path

# 设置Windows控制台编码为UTF-8
if sys.platform == 'win32':
    try:
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass


def get_config_file():
    """获取配置文件路径"""
    try:
        from modules.config.config_manager import ConfigManager

        return Path(ConfigManager().config_file_path)
    except Exception:
        env_data_dir = os.environ.get("KOI_USER_DATA_DIR")
        if env_data_dir:
            return Path(env_data_dir).resolve() / "config.json"
        env_app_dir = os.environ.get("KOI_APP_DIR")
        if env_app_dir:
            return Path(env_app_dir).resolve() / "config.json"
        if getattr(sys, 'frozen', False):
            app_dir = Path(sys.executable).parent
            if app_dir.name.lower() in {"koi-backend", "koi_backend"}:
                app_dir = app_dir.parent
            return app_dir / "config.json"
        script_dir = Path(__file__).resolve().parent
        project_root = script_dir.parent.parent.parent
        return project_root / "config.json"


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_document_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def _replace_paragraph_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _replace_in_text_nodes(text_nodes, pattern, replacement):
    texts = [node.text or "" for node in text_nodes]
    full_text = "".join(texts)
    match = pattern.search(full_text)
    if not match:
        return False

    start, end = match.span()
    offset = 0
    start_index = None
    end_index = None
    start_inner = 0
    end_inner = 0
    for index, text in enumerate(texts):
        next_offset = offset + len(text)
        if start_index is None and start <= next_offset:
            start_index = index
            start_inner = max(0, start - offset)
        if end_index is None and end <= next_offset:
            end_index = index
            end_inner = max(0, end - offset)
            break
        offset = next_offset

    if start_index is None or end_index is None:
        return False

    if start_index == end_index:
        text_nodes[start_index].text = (
            texts[start_index][:start_inner]
            + replacement
            + texts[start_index][end_inner:]
        )
        return True

    text_nodes[start_index].text = texts[start_index][:start_inner] + replacement
    for index in range(start_index + 1, end_index):
        text_nodes[index].text = ""
    text_nodes[end_index].text = texts[end_index][end_inner:]
    return True


def _replace_pattern_in_paragraph_text_nodes(para, pattern_text, replacement):
    try:
        text_nodes = para._element.findall(
            ".//w:t",
            {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if not text_nodes:
            return False
        return _replace_in_text_nodes(text_nodes, re.compile(pattern_text), replacement)
    except Exception:
        return False


def _replace_text_in_docx_xml(docx_file, pattern_text, replacement):
    from lxml import etree  # type: ignore
    import zipfile

    docx_path = Path(docx_file)
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".number.tmp")
    pattern = re.compile(pattern_text)
    changed = False
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"

    try:
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    try:
                        root = etree.fromstring(data)
                        text_nodes = root.findall(f".//{w_ns}")
                        if text_nodes and _replace_in_text_nodes(text_nodes, pattern, replacement):
                            data = etree.tostring(
                                root,
                                encoding="UTF-8",
                                xml_declaration=data.lstrip().startswith(b"<?xml"),
                            )
                            changed = True
                    except Exception:
                        pass
                target.writestr(info, data)
        if changed:
            os.replace(tmp_path, docx_path)
        else:
            tmp_path.unlink(missing_ok=True)
        return changed
    except Exception:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
        return False


def _parse_unavailable_numbers(value) -> set[int]:
    """
    解析不可用编号配置。
    支持：
    - list[int]
    - "170,171"
    - "170-175" / "170~175"
    - 中英文逗号、分号、空格分隔
    """
    out: set[int] = set()
    if value is None:
        return out
    if isinstance(value, list):
        for v in value:
            try:
                n = int(str(v).strip())
                if n > 0:
                    out.add(n)
            except Exception:
                continue
        return out

    s = str(value).strip()
    if not s:
        return out
    s = s.replace("，", ",").replace("；", ";").replace("~", "-")
    parts = re.split(r"[,\s;]+", s)
    for p in parts:
        p = (p or "").strip()
        if not p:
            continue
        if "-" in p:
            a, b = p.split("-", 1)
            try:
                start = int(a.strip())
                end = int(b.strip())
                if start <= 0 or end <= 0:
                    continue
                if start > end:
                    start, end = end, start
                # 防御：避免超大区间导致卡死
                if end - start > 200000:
                    continue
                for n in range(start, end + 1):
                    out.add(n)
            except Exception:
                continue
        else:
            try:
                n = int(p)
                if n > 0:
                    out.add(n)
            except Exception:
                continue
    return out


def _next_available_number(n: int, unavailable: set[int]) -> int:
    """从 n 开始（含 n）寻找第一个不在 unavailable 的正整数。"""
    cur = int(n) if n is not None else 1
    if cur <= 0:
        cur = 1
    for _ in range(500000):
        if cur not in unavailable:
            return cur
        cur += 1
    return cur


def _remaining_unavailable_numbers(unavailable: set[int], next_number: int) -> list[int]:
    return sorted(number for number in unavailable if number >= int(next_number or 1))


def _get_report_counters(config):
    counters = config.get('report_counters') if isinstance(config, dict) else None
    return counters if isinstance(counters, dict) else {}


def _ensure_report_counters(config):
    counters = _get_report_counters(config)
    if not counters:
        counters = {
            'notification_number': 1,
            'rectification_number': 1,
            'year': datetime.now().year,
            'last_updated': ''
        }
    counters.setdefault('notification_number', 1)
    counters.setdefault('rectification_number', 1)
    counters.setdefault('year', datetime.now().year)
    counters.setdefault('last_updated', '')
    config['report_counters'] = counters
    return counters


def _prepare_template_for_docx(template_file):
    template_path = Path(template_file)
    if template_path.suffix.lower() != '.doc':
        return str(template_path), None

    temp_dir = Path(tempfile.mkdtemp(prefix='koi_rect_template_'))
    converted_path = temp_dir / f"{template_path.stem}.docx"
    word = None
    doc = None
    com_initialized = False
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore

        try:
            pythoncom.CoInitialize()
            com_initialized = True
        except Exception:
            com_initialized = False

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(
            str(template_path),
            ReadOnly=True,
            Visible=False,
            ConfirmConversions=False,
            AddToRecentFiles=False,
        )
        try:
            doc.SaveAs2(str(converted_path), FileFormat=16)
        except Exception:
            doc.SaveAs(str(converted_path), 16)
    except Exception as exc:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise RuntimeError(f"无法将 .doc 模板转换为 .docx: {exc}") from exc
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=0)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit(SaveChanges=0)
        except Exception:
            pass
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    if not converted_path.exists():
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise RuntimeError("无法将 .doc 模板转换为 .docx: 未生成转换文件")
    return str(converted_path), str(temp_dir)


def update_rectification_number(docx_file):
    """
    更新责令整改编号
    
    参数:
        docx_file: 生成的责令整改通知书路径
    
    返回:
        当前使用的编号
    """
    try:
        config_file = get_config_file()
        
        # 读取配置
        if not config_file.exists():
            print(f"  警告: 配置文件不存在: {config_file}")
            return None
        
        with open(config_file, 'r', encoding='utf-8-sig') as f:
            config = json.load(f)
        
        # 获取当前编号
        counters = _ensure_report_counters(config)
        
        # 检查年份，如果是新年则重置编号
        current_year = datetime.now().year
        if 'year' not in counters or counters['year'] != current_year:
            print(f"  🎊 检测到新年份: {current_year}，重置编号计数")
            counters['notification_number'] = 1
            counters['rectification_number'] = 1
            counters['year'] = current_year
        
        # 不可用编号（命中则自动跳过）
        unavailable = _parse_unavailable_numbers(
            counters.get('unavailable_rectification_numbers', None)
        )
        # 兼容旧字段：unavailable_numbers
        if not unavailable:
            unavailable = _parse_unavailable_numbers(counters.get('unavailable_numbers', []))

        # 使用配置中的年份（已更新后的）
        config_year = counters['year']
        current_number = counters['rectification_number']
        current_number = _next_available_number(current_number, unavailable)
        
        # 打开文档并替换编号
        doc = Document(docx_file)
        replaced = False
        
        for para in _iter_document_paragraphs(doc):
            para_text = para.text
            # 查找 鄞网办责字[YYYY]XXX号 的模式（支持任意年份）
            if '鄞网办责字' in para_text and '[' in para_text and ']' in para_text and '号' in para_text:
                # 提取当前的年份和编号
                year_match = re.search(r'\[(\d{4})\]', para_text)
                number_match = re.search(r'\](\d+)号', para_text)
                
                if year_match and number_match:
                    replaced = _replace_pattern_in_paragraph_text_nodes(
                        para,
                        r'鄞网办责字\[\d{4}\]\d+号',
                        f'鄞网办责字[{config_year}]{current_number}号',
                    )
                
                # 找到目标段落后退出循环
                break
        
        if not replaced:
            replaced = _replace_text_in_docx_xml(
                docx_file,
                r'鄞网办责字\[\d{4}\]\d+号',
                f'鄞网办责字[{config_year}]{current_number}号',
            )

        if replaced:
            # 保存文档
            if any(
                re.search(r'鄞网办责字\[\d{4}\]\d+号', para.text or "")
                for para in _iter_document_paragraphs(doc)
            ):
                doc.save(docx_file)
            
            # 使用统一配置管理器进行原子更新，避免覆盖其他模块的写入
            # 兼容在压缩包子目录中执行时的导入路径问题
            try:
                from modules.config.config_manager import ConfigManager
            except ImportError:
                # 尝试从项目根目录导入
                import sys
                project_root = get_config_file().parent
                if str(project_root) not in sys.path:
                    sys.path.insert(0, str(project_root))
                from modules.config.config_manager import ConfigManager
            
            # 传递具体的配置文件路径
            cm = ConfigManager(str(get_config_file()))
            next_rectification_number = _next_available_number(current_number + 1, unavailable)
            cm.update_section('report_counters', {
                'rectification_number': next_rectification_number,
                'unavailable_rectification_numbers': _remaining_unavailable_numbers(unavailable, next_rectification_number),
                'year': config_year,
                'last_updated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
            
            print(f"  ✓ 已更新责令整改编号: 鄞网办责字[{config_year}]{current_number}号")
            return current_number, config_year
        else:
            print(f"  警告: 未找到责令整改编号标记")
            return None
            
    except Exception as e:
        print(f"  警告: 更新责令整改编号失败: {str(e)}")
        return None


def extract_info_from_filename(filename):
    """
    从文件名中提取公司名和漏洞类型
    
    文件名格式示例：
    - 关于浙江格瓦拉数字科技有限公司所属Druid系统存在未授权访问安全漏洞通报.docx
    - 1760410609070舒普智能技术股份有限公司远程技术检查存在ecology远程命令执行漏洞.docx
    
    返回: (公司名, 漏洞描述)
    """
    basename = os.path.basename(filename)
    name_without_ext = basename.rsplit('.', 1)[0]
    name_without_ext = re.sub(r'_\d{8,}$', '', name_without_ext).strip()
    name_clean = re.sub(r'^\d+', '', name_without_ext).strip()

    try:
        from modules.Document_Processing.Report_Rewrite import group_folders as gf
        company_name = gf.normalize_company(name_clean)
        company_suffix_pattern = gf.company_suffix_regex()
    except Exception:
        company_name = None
        company_suffix_pattern = r'(?:股份有限公司|有限责任公司|责任有限公司|有限公司|集团公司|集团|公司)'
    report_tail_pattern = r'(?:的预警通报|预警通报|的通报|通报|的报告(?:（.*?）)?|报告(?:（.*?）)?)'

    if not company_name:
        for pattern in [
            rf'关于(.+?{company_suffix_pattern})',
            r'关于(.+?)所属',
            r'关于(.+?)(门户网站|官网|网站|平台|系统)',
            r'关于(.+?)存在',
            r'关于(.+?)的',
            rf'^(.+?{company_suffix_pattern})',
            r'^(.+?)(?:远程技术检查|技术检查|检查|远程|存在)',
        ]:
            company_match = re.search(pattern, name_clean)
            if company_match:
                candidate = company_match.group(1).strip()
                if candidate:
                    company_name = candidate
                    break

    vuln_type = None
    for pattern in [
        rf'关于.+?((?:疑似感染|发现|遭受|发生).+?)(?:{report_tail_pattern})\s*$',
        rf'关于.+?{company_suffix_pattern}(.+?)(?:{report_tail_pattern})\s*$',
        r'存在(.+?)(?:通报|的报告)',
        r'系统(.+?)(?:通报|的报告)',
        r'网站(.+?)(?:通报|的报告)',
        r'存在(.+?)(?:_\d+)?(?:\.docx|$)',
        r'(?:远程技术检查|技术检查|检查)存在(.+?)(?:_\d+)?(?:\.docx|$)',
        r'([\u4e00-\u9fa5A-Za-z]+(?:漏洞|风险|事件))',
    ]:
        vuln_match = re.search(pattern, name_clean)
        if vuln_match:
            vuln_type = vuln_match.group(1).strip()
            break
    
    # 后处理：清理提取的漏洞类型
    if vuln_type:
        vuln_type = re.sub(r'^疑似', '', vuln_type).strip()
        # 去掉开头的"存在"（如果有的话）
        vuln_type = re.sub(r'^存在', '', vuln_type)
        
        # 去掉结尾的"的报告"、"风险的报告"等
        vuln_type = re.sub(r'(?:风险)?的报告$', '', vuln_type)
        # 去掉预警通报后缀
        vuln_type = re.sub(r'(?:的)?预警通报$', '', vuln_type)
        vuln_type = re.sub(r'的报告（.*?）$', '', vuln_type)
        vuln_type = re.sub(r'报告（.*?）$', '', vuln_type)
        vuln_type = re.sub(r'的报告$', '', vuln_type)
        vuln_type = re.sub(r'报告$', '', vuln_type)
        vuln_type = re.sub(r'的通报$', '', vuln_type)
        vuln_type = re.sub(r'通报$', '', vuln_type)
        if vuln_type.startswith('关于') and company_name and company_name in vuln_type:
            vuln_type = vuln_type.split(company_name, 1)[-1].strip()
        
        # 修正重复的"安全"
        vuln_type = re.sub(r'安全安全', '安全', vuln_type)
        
        # 确保以"漏洞"或"风险"结尾
        if not (vuln_type.endswith('漏洞') or vuln_type.endswith('风险')):
            if '风险' in vuln_type:
                vuln_type = re.sub(r'风险.*$', '风险', vuln_type)
            elif '漏洞' in vuln_type:
                vuln_type = re.sub(r'漏洞.*$', '漏洞', vuln_type)
            elif '事件' in vuln_type:
                vuln_type = re.sub(r'事件.*$', '事件', vuln_type)
        
        vuln_type = vuln_type.strip()
    
    return company_name, vuln_type


def extract_info_from_document(doc_file):
    """
    从通报文档中读取内容来提取信息（备用方案）
    """
    try:
        doc = Document(doc_file)
        # 读取文档内容，尝试从内容中提取信息
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # 从内容中提取公司名
        company_match = re.search(r'关于(.+?)所属', full_text)
        company_name = company_match.group(1) if company_match else None
        
        # 从内容中提取漏洞类型
        vuln_match = re.search(r'(存在.+?安全漏洞)', full_text)
        vuln_type = vuln_match.group(1) if vuln_match else None
        
        return company_name, vuln_type
    except Exception as e:
        print(f"从文档内容提取信息时出错: {e}")
        return None, None


def edit_rectification(
    report_file=None,
    template_file=None,
    company_name=None,
    vuln_type=None,
    company_name_override=None,
    vuln_type_override=None,
):
    """
    编辑责令整改通知书
    
    参数:
        report_file: 通报文档路径（如果为None，则自动查找）
        template_file: 责令整改模板路径（如果为None，则自动查找）
        company_name: 兼容旧参数，明确指定公司名
        vuln_type: 兼容旧参数，明确指定漏洞类型
        company_name_override: 新参数，明确指定公司名（优先）
        vuln_type_override: 新参数，明确指定漏洞类型（优先）
    """
    print("=" * 60)
    print("责令整改通知书编辑工具")
    print("=" * 60)
    
    # 如果未指定通报文件且未提供明确信息，自动查找
    if report_file is None and (not company_name or not vuln_type):
        possible_reports = []
        for filename in os.listdir('.'):
            if filename.endswith('.docx'):
                # 排除模板和其他文件
                if 'Report_Template' not in filename and '授权委托书' not in filename and '模板' not in filename and '责令整改' not in filename:
                    # 优先选择以"关于"开头或包含"通报"的文件
                    if filename.startswith('关于') or '通报' in filename:
                        possible_reports.append(filename)
        
        if possible_reports:
            report_file = possible_reports[0]
            print(f"\n自动找到通报文档: {report_file}")
        else:
            print("\n未找到通报文档！")
            print("\n使用方法:")
            print("  方法1: 将通报文档放在当前目录，运行: python edit_rectification.py")
            print("  方法2: 指定文件: python edit_rectification.py <通报文档路径>")
            print("\n功能说明:")
            print("  1. 从通报文档中提取公司名和漏洞类型")
            print("  2. 在责令整改模板中替换公司名、漏洞类型和日期")
            print("  3. 生成责令整改通知书文件")
            print("\n重要提示:")
            print("  1. 模板文件会从 template 目录自动查找")
            print("  2. 模板文件必须是 .docx 格式")
            print("=" * 60)
            return False
    
    # 如果未指定模板文件，自动查找
    if template_file is None:
        template_candidates = []
        
        # 先在 template 目录查找（支持开发和打包环境）
        try:
            from modules.utils.resource_path import get_report_template_dir
            report_template_dir = str(get_report_template_dir())
        except ImportError:
            report_template_dir = 'Report_Template'
        
        if os.path.exists(report_template_dir):
            for filename in os.listdir(report_template_dir):
                if filename.lower().endswith(('.doc', '.docx')) and ('责令整改' in filename or '整改通知' in filename):
                    template_candidates.append(os.path.join(report_template_dir, filename))
        
        # 如果 template 目录没找到，在当前目录查找
        if not template_candidates:
            for filename in os.listdir('.'):
                if filename.lower().endswith(('.doc', '.docx')) and ('责令整改' in filename or '整改通知' in filename):
                    template_candidates.append(filename)
        
        if not template_candidates:
            print("\n错误: 未找到责令整改模板文件！")
            print("  请确保以下位置之一存在责令整改模板文件：")
            print("    - template/责令整改*.docx 或 .doc")
            print("    - ./责令整改*.docx 或 .doc")
            return False
        
        template_file = template_candidates[0]

    original_template_file = template_file
    template_cleanup_dir = None
    try:
        template_file, template_cleanup_dir = _prepare_template_for_docx(template_file)
    except Exception as exc:
        print(f"\n错误: {exc}")
        return False
    
    # 如果没有明确提供信息，则从文件名提取
    if not company_name or not vuln_type:
        extracted_company, extracted_vuln = extract_info_from_filename(report_file)
        
        if not company_name:
            company_name = extracted_company
        if not vuln_type:
            vuln_type = extracted_vuln
    
    # 如果从文件名提取失败，尝试从文档内容提取
    if (not company_name or not vuln_type) and report_file:
        print("从文件名提取信息失败，尝试从文档内容提取...")
        company_name_doc, vuln_type_doc = extract_info_from_document(report_file)
        if not company_name:
            company_name = company_name_doc
        if not vuln_type:
            vuln_type = vuln_type_doc

    
    if company_name_override:
        company_name = str(company_name_override).strip()

    if vuln_type_override:
        vuln_type = str(vuln_type_override).strip()

    if not company_name:
        print("\n警告: 无法提取公司名！")
        company_name = "【公司名】"
    
    if not vuln_type:
        print("\n警告: 无法提取漏洞类型！")
        vuln_type = "【漏洞类型】"
    
    # 获取当前日期
    today = datetime.now()
    current_date = f"{today.year}年{today.month}月{today.day}日"
    
    # 获取模板文件名（用于生成输出文件名）
    template_basename = Path(original_template_file).with_suffix('.docx').name
    
    print(f"\n正在编辑责令整改通知书:")
    print(f"  模板文件: {template_file}")
    print(f"  公司名: {company_name}")
    print(f"  漏洞类型: {vuln_type}")
    print(f"  日期: {current_date}")
    print("=" * 60)
    
    try:
        # 加载模板文档
        doc = Document(template_file)
        
        replacements = 0
        
        # 遍历所有段落
        for para_idx, para in enumerate(doc.paragraphs, 1):
            original_text = para.text
            modified = False
            
            # 首先检查整个段落是否包含日期（因为日期可能被分成多个run）
            date_pattern = r'20\d{2}\s*年\s*\d+\s*月\s*\d+\s*日'
            if re.search(date_pattern, para.text):
                # 替换整个段落的日期
                new_para_text = re.sub(date_pattern, current_date, para.text)
                if new_para_text != para.text and para.runs:
                    # 清空所有run，只保留第一个
                    for run in para.runs[1:]:
                        run.text = ''
                    para.runs[0].text = new_para_text
                    modified = True
            else:
                # 如果没有日期，按run处理公司名和漏洞类型
                for run in para.runs:
                    run_text = run.text
                    new_text = run_text
                    
                    # 替换公司名（查找任何公司名模式）
                    if company_name and '有限公司' in run_text:
                        # 替换任何公司名为实际公司名
                        new_text = re.sub(r'[\u4e00-\u9fa5]+有限公司', company_name, new_text)
                    
                    # 替换漏洞类型
                    if vuln_type and ('存在' in run_text and '漏洞' in run_text):
                        # 替换任何漏洞描述为实际漏洞类型
                        # vuln_type现在已经不包含"存在"前缀，需要确保不重复添加
                        # 如果vuln_type已经包含"存在"，直接使用；否则添加"存在"
                        if vuln_type.startswith('存在'):
                            new_text = re.sub(r'存在.+?漏洞', vuln_type, new_text)
                        else:
                            new_text = re.sub(r'存在.+?漏洞', f'存在{vuln_type}', new_text)
                    
                    # 如果有修改，更新run的文本
                    if new_text != run_text:
                        run.text = new_text
                        modified = True
            
            if modified:
                replacements += 1
                print(f"  段落 {para_idx} 已替换:")
                print(f"    原文: {original_text}")
                print(f"    新文: {para.text}")
                print()
        
        # 生成输出文件名（保持模板原文件名）
        output_file = template_basename
        
        # 保存文档
        doc.save(output_file)
        
        # 更新责令整改编号
        result = update_rectification_number(output_file)
        if result:
            rectification_number, config_year = result if isinstance(result, tuple) else (result, None)
        else:
            rectification_number, config_year = None, None
        
        # 如果函数返回的不是元组，从配置文件读取年份
        if config_year is None:
            try:
                config_file = get_config_file()
                if config_file.exists():
                    with open(config_file, 'r', encoding='utf-8') as f:
                        config = json.load(f)
                    config_year = _get_report_counters(config).get('year', datetime.now().year)
                else:
                    config_year = datetime.now().year
            except:
                config_year = datetime.now().year
        
        print("=" * 60)
        print(f"✓ 成功生成责令整改通知书!")
        print(f"  输出文件: {output_file}")
        print(f"  替换次数: {replacements} 个段落")
        if rectification_number:
            print(f"  文号: 鄞网办责字[{config_year}]{rectification_number}号")
        print("=" * 60)
        
        return True
        
    except FileNotFoundError as e:
        print(f"\n错误: 找不到文件: {e}")
        return False
    except ValueError as e:
        print(f"\n错误: {e}")
        return False
    except Exception as e:
        print(f"\n编辑文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        if template_cleanup_dir:
            shutil.rmtree(template_cleanup_dir, ignore_errors=True)


if __name__ == "__main__":
    report_file = None
    
    if len(sys.argv) >= 2:
        # 如果提供了参数，使用参数
        report_file = sys.argv[1]
    
    # 执行编辑
    success = edit_rectification(report_file)
    
    if success:
        print("\n编辑完成！")
    else:
        print("\n编辑失败，请检查错误信息。")
        sys.exit(1)
