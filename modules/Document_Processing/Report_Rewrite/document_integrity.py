"""
文档完整性验证和修复模块
提供安全的文档保存、验证和修复功能
"""

import os
import sys
import time
import shutil
import tempfile
import gc
from pathlib import Path
from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH

def safe_print(text, fallback_text=None):
    """安全打印函数，处理编码问题"""
    try:
        print(text)
    except UnicodeEncodeError:
        if fallback_text:
            print(fallback_text)
        else:
            print(text.encode('utf-8', errors='ignore').decode('utf-8'))

def validate_document_integrity(doc_path, min_paragraphs=5, min_size_kb=10):
    """
    验证文档完整性
    
    Args:
        doc_path: 文档路径
        min_paragraphs: 最小段落数
        min_size_kb: 最小文件大小(KB)
    
    Returns:
        dict: 验证结果
    """
    result = {
        'valid': False,
        'file_exists': False,
        'size_ok': False,
        'readable': False,
        'paragraph_count': 0,
        'file_size': 0,
        'error': None
    }
    
    try:
        # 检查文件是否存在
        if not Path(doc_path).exists():
            result['error'] = "文件不存在"
            return result
        result['file_exists'] = True
        
        # 检查文件大小
        file_size = Path(doc_path).stat().st_size
        result['file_size'] = file_size
        if file_size < min_size_kb * 1024:
            result['error'] = f"文件大小异常: {file_size}字节"
            return result
        result['size_ok'] = True
        
        # 尝试用python-docx打开文档
        doc = Document(doc_path)
        result['readable'] = True
        
        # 检查段落数
        paragraph_count = len(doc.paragraphs)
        result['paragraph_count'] = paragraph_count
        if paragraph_count < min_paragraphs:
            result['error'] = f"段落数异常: {paragraph_count}"
            return result
        
        # 检查是否有实际内容
        content_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
        if content_paragraphs < 3:
            result['error'] = f"有效内容段落过少: {content_paragraphs}"
            return result
        
        result['valid'] = True
        return result
        
    except Exception as e:
        result['error'] = f"文档验证失败: {str(e)}"
        return result

def safe_save_document(doc, output_path, max_retries=3):
    """
    安全保存文档，包含重试和验证机制
    
    Args:
        doc: Document对象
        output_path: 输出路径
        max_retries: 最大重试次数
    
    Returns:
        dict: 保存结果
    """
    result = {
        'success': False,
        'method': None,
        'error': None,
        'validation': None
    }
    
    for attempt in range(max_retries):
        try:
            safe_print(f"  📝 尝试保存文档 (第{attempt + 1}次)...")
            
            # 直接保存
            doc.save(output_path)
            
            # 等待文件完全写入
            time.sleep(1.0)
            
            # 验证保存结果
            validation = validate_document_integrity(output_path)
            result['validation'] = validation
            
            if validation['valid']:
                result['success'] = True
                result['method'] = 'direct_save'
                safe_print(f"  ✅ 文档保存成功 ({validation['paragraph_count']}个段落，{validation['file_size']}字节)")
                return result
            else:
                safe_print(f"  ⚠️ 保存验证失败: {validation['error']}")
                if attempt < max_retries - 1:
                    time.sleep(0.5)
                    continue
                    
        except Exception as e:
            safe_print(f"  ❌ 保存失败: {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(0.5)
                continue
            result['error'] = str(e)
    
    # 如果直接保存失败，尝试修复保存
    return repair_and_save_document(doc, output_path)

def repair_and_save_document(original_doc, output_path):
    """
    修复并保存文档
    
    Args:
        original_doc: 原始Document对象
        output_path: 输出路径
    
    Returns:
        dict: 保存结果
    """
    result = {
        'success': False,
        'method': None,
        'error': None,
        'validation': None
    }
    
    try:
        safe_print(f"  🔧 尝试修复文档...")
        
        # 创建新文档
        repaired_doc = Document()
        
        # 复制段落内容，避免复杂格式
        copied_count = 0
        for para in original_doc.paragraphs:
            if para.text.strip():  # 只复制有内容的段落
                new_para = repaired_doc.add_paragraph()
                new_para.text = para.text
                
                # 尝试保留基本格式，但忽略错误
                try:
                    if para.style:
                        new_para.style = para.style
                except:
                    pass
                
                try:
                    new_para.alignment = para.alignment
                except:
                    pass
                
                copied_count += 1
        
        # 如果没有复制到内容，添加默认内容
        if copied_count == 0:
            default_para = repaired_doc.add_paragraph()
            default_para.text = "文档内容修复失败，请检查原始文件。"
            default_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存修复后的文档
        temp_file = tempfile.mktemp(suffix='.docx')
        repaired_doc.save(temp_file)
        
        # 验证临时文件
        validation = validate_document_integrity(temp_file)
        if validation['valid']:
            shutil.move(temp_file, output_path)
            result['success'] = True
            result['method'] = 'repaired_save'
            result['validation'] = validation
            safe_print(f"  ✅ 修复保存成功 ({validation['paragraph_count']}个段落)")
            return result
        else:
            safe_print(f"  ❌ 修复后验证失败: {validation['error']}")
            if Path(temp_file).exists():
                os.remove(temp_file)
    
    except Exception as e:
        safe_print(f"  ❌ 文档修复失败: {str(e)}")
        result['error'] = str(e)
    
    # 最后的降级策略：创建简化文档
    return create_fallback_document(original_doc, output_path)

def create_fallback_document(original_doc, output_path):
    """
    创建降级文档（纯文本版本）
    
    Args:
        original_doc: 原始Document对象
        output_path: 输出路径
    
    Returns:
        dict: 保存结果
    """
    result = {
        'success': False,
        'method': 'fallback_save',
        'error': None,
        'validation': None
    }
    
    try:
        safe_print(f"  🆘 创建降级文档...")
        
        fallback_doc = Document()
        
        # 添加标题
        title_para = fallback_doc.add_paragraph()
        title_para.text = "网络安全漏洞通报"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加说明
        note_para = fallback_doc.add_paragraph()
        note_para.text = "注意：此文档为简化版本，不包含复杂格式。"
        
        # 添加基本内容
        content_added = False
        for para in original_doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 5:  # 过滤掉太短的段落
                new_para = fallback_doc.add_paragraph()
                new_para.text = text
                content_added = True
        
        if not content_added:
            fallback_doc.add_paragraph("原始文档内容无法读取，请检查源文件。")
        
        # 保存降级文档
        fallback_doc.save(output_path)
        
        # 验证降级文档
        validation = validate_document_integrity(output_path, min_paragraphs=2, min_size_kb=5)
        result['validation'] = validation
        
        if validation['valid']:
            result['success'] = True
            safe_print(f"  ✅ 降级文档创建成功")
            safe_print(f"  ⚠️ 注意：此文档为简化版本，不包含复杂格式")
        else:
            result['error'] = f"降级文档验证失败: {validation['error']}"
            safe_print(f"  ❌ {result['error']}")
        
        return result
        
    except Exception as e:
        result['error'] = f"降级文档创建失败: {str(e)}"
        safe_print(f"  ❌ {result['error']}")
        return result



def cleanup_resources():
    """清理资源，确保文件句柄释放"""
    try:
        gc.collect()
        time.sleep(0.5)
    except:
        pass

def wait_for_file_release(file_path, max_wait=10):
    """等待文件释放"""
    for i in range(max_wait):
        try:
            # 尝试以写模式打开文件
            with open(file_path, 'a'):
                pass
            return True
        except:
            time.sleep(0.5)
    return False