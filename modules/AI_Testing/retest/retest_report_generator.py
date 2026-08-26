#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
复测报告生成工具
基于Word文档漏洞扫描结果，自动生成复测报告
"""

import os
import re
import shutil
from pathlib import Path
from typing import Any, List, Dict, Optional, Tuple
from docx import Document
from docx.shared import Inches

from modules.AI_Testing.retest.word_vulnerability_scanner import is_generated_retest_report_path
from modules.Document_Processing.Report_Rewrite.notice_name_utils import normalize_issue_type


SCREENSHOT_INSERT_WIDTH_INCHES = 4.2
MAX_SCREENSHOT_PART_HEIGHT_INCHES = 5.0
MIN_SHRUNK_SCREENSHOT_WIDTH_INCHES = 3.2


class RetestReportGenerator:
    """复测报告生成器"""
    
    def __init__(
        self,
        target_dir: str,
        template_path: str,
        output_dir: str | None = None,
        screenshot_path: Optional[str] = None,
        screenshot_sections: Optional[List[Dict[str, Any]]] = None,
    ):
        """
        初始化报告生成器
        
        Args:
            target_dir: 要扫描的通报文档目录
            template_path: 复测模板文件路径
            output_dir: 输出目录（默认与原通报文件同目录）
        """
        self.target_dir = Path(target_dir)
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir) if output_dir else None
        # 复测截图路径（可选，用于在模板正文中插入证明图片）
        self.screenshot_path = Path(screenshot_path) if screenshot_path else None
        self.screenshot_sections: List[Tuple[str, Path]] = []
        for item in screenshot_sections or []:
            if not isinstance(item, dict):
                continue
            section_path = str(item.get("path") or item.get("screenshot_path") or item.get("screenshotPath") or "").strip()
            if not section_path:
                continue
            self.screenshot_sections.append((str(item.get("caption") or "").strip(), Path(section_path)))
        if self.screenshot_sections and self.screenshot_path is None:
            self.screenshot_path = self.screenshot_sections[0][1]
        self._screenshot_insert_items: Optional[List[Tuple[Path, float]]] = None
        self._temp_screenshot_parts: List[Path] = []
        
        if self.output_dir is not None:
            self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # URL正则表达式（复用原扫描器的逻辑）
        self.url_pattern = re.compile(
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        )
        
        # IP地址正则表达式
        self.ip_pattern = re.compile(
            r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b'
        )
    
    def find_word_files(self) -> List[Path]:
        """
        递归查找所有Word文档
        
        Returns:
            Word文档路径列表
        """
        word_files = []
        for ext in ['*.docx', '*.doc']:
            word_files.extend(self.target_dir.rglob(ext))
        
        # 过滤掉临时文件和非通报文件
        filtered_files = []
        for file in word_files:
            filename = file.name
            
            # 跳过Word临时文件
            if filename.startswith('~$'):
                continue
            
            # 跳过隐藏文件
            if filename.startswith('.'):
                continue
            if is_generated_retest_report_path(file):
                continue
            
            # 跳过非通报文件
            non_report_keywords = [
                '模板', 'template', '处置文件', '示例', '样例', 
                'example', '说明', 'readme', '备份'
            ]
            
            is_non_report = False
            filename_lower = filename.lower()
            for keyword in non_report_keywords:
                if keyword in filename_lower:
                    is_non_report = True
                    break
            
            if is_non_report:
                continue
            
            filtered_files.append(file)
        
        return filtered_files
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """从Word文档提取文本内容"""
        try:
            # 显式转换为 str，满足类型检查器的签名要求
            doc = Document(str(file_path))
            text_content = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            print(f"[错误] 读取文件 {file_path} 失败: {e}")
            return ""
    
    def extract_urls_and_ips(self, text: str, file_path: Optional[Path] = None) -> Dict[str, List[str]]:
        """从文本中提取URL和IP地址（复用原扫描器逻辑）"""
        urls = []
        ips = []
        
        lines = text.split('\n')
        
        url_markers = [
            'URl:', 'URL:', 'url:', 'Url:',
            'URl：', 'URL：', 'url：', 'Url：',
            '网址:', '地址:',
            '网址：', '地址：',
        ]
        
        # 1) 首选：在"验证情况"之后查找
        verify_idx = None
        for i, line in enumerate(lines):
            if '验证情况' in line:
                verify_idx = i
                break
        
        if verify_idx is not None:
            for line in lines[verify_idx + 1: verify_idx + 15]:
                if not line.strip():
                    continue
                
                if any(marker in line for marker in url_markers):
                    line_urls = self.url_pattern.findall(line)
                    if line_urls:
                        urls.append(line_urls[0])
                        break
                
                line_urls = self.url_pattern.findall(line)
                if line_urls:
                    urls.append(line_urls[0])
                    break
        
        # 2) 全文查找URL标记
        if not urls:
            for line in lines:
                if any(marker in line for marker in url_markers):
                    line_urls = self.url_pattern.findall(line)
                    if line_urls:
                        urls.append(line_urls[0])
                        break
        
        # 3) 全文兜底提取
        if not urls:
            all_urls = self.url_pattern.findall(text)
            if all_urls:
                urls.append(all_urls[0])
        
        # 4) 识别domain行
        if not urls:
            domain_pattern = re.compile(r'(?:domain|Domain|DOMAIN|域名)\s*[:：]\s*([^\s]+)')
            for line in lines:
                m = domain_pattern.search(line)
                if m:
                    domain_str = m.group(1).strip()
                    if domain_str:
                        urls.append(domain_str)
                        break
        
        # 5) 从底层XML提取
        if not urls and file_path is not None:
            try:
                doc = Document(str(file_path))
                raw_xml = doc._element.xml
                extra_urls = self.url_pattern.findall(raw_xml)
                if extra_urls:
                    urls.append(extra_urls[0])
            except Exception:
                pass
        
        # 6) 提取IP
        if not urls:
            ips.extend(self.ip_pattern.findall(text))
            ips = [ip for ip in ips if self._is_valid_ip(ip)]
        
        return {
            'urls': list(set(urls)),
            'ips': list(set(ips))
        }
    
    def _is_valid_ip(self, ip: str) -> bool:
        """验证IP地址是否有效"""
        parts = ip.split('.')
        if len(parts) != 4:
            return False
        try:
            return all(0 <= int(part) <= 255 for part in parts)
        except ValueError:
            return False
    
    def extract_vulnerability_from_filename(self, filename: str) -> List[str]:
        """从文件名中提取漏洞类型"""
        found_vulns = []
        normalized_name = Path(filename).stem
        normalized_name = re.sub(r'[\s_-]*(?:19|20)\d{6,12}$', '', normalized_name)
        
        patterns = [
            r'存在漏洞的(.+?)(?:漏洞|通报|的)',
            r'存在(.+?)的安全漏洞',
            r'存在(.+?)安全漏洞',
            r'存在(.+?)的漏洞',
            r'存在(.+?)漏洞',
            r'存在(.+?)(?:安全问题|安全隐患|风险)$',
            r'存在(.+)$',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, normalized_name)
            if matches:
                for match in matches:
                    vuln = match.strip() if isinstance(match, str) else match[0].strip()
                    
                    # 清理噪音词
                    for noise in ['问题']:
                        vuln = vuln.replace(noise, '').strip()
                    vuln = re.sub(r'[\s_-]*(?:19|20)\d{6,12}$', '', vuln).strip()
                    vuln = re.sub(r'(?:的)?通报$', '', vuln).strip()
                    vuln = normalize_issue_type(vuln) or vuln
                    
                    # 过滤无效结果
                    invalid_words = ['存在', '所属', '官网', '网站', '系统', '平台', '域名', '通报']
                    if vuln and len(vuln) > 2 and vuln not in invalid_words:
                        found_vulns.append(vuln)
                break
        
        deduped = []
        seen = set()
        for item in found_vulns:
            key = item.lower()
            if key in seen:
                continue
            seen.add(key)
            deduped.append(item)
        return deduped
    
    def extract_title_from_filename(self, filename: str) -> str:
        """
        从通报文件名提取标题
        
        Args:
            filename: 文件名（如：关于XXX公司所属网站存在XXX漏洞的通报.docx）
            
        Returns:
            标题（如：关于XXX公司所属网站存在XXX漏洞）
        """
        # 去掉文件扩展名
        title = filename.replace('.docx', '').replace('.doc', '')
        
        # 去掉"的通报"后缀
        title = re.sub(r'的通报$', '', title)
        
        return title
    
    def scan_document(self, file_path: Path) -> Dict:
        """
        扫描单个文档，提取信息
        
        Args:
            file_path: 文档路径
            
        Returns:
            扫描结果字典
        """
        print(f"[扫描] {file_path.name}")
        
        # 提取文本
        text = self.extract_text_from_docx(file_path)
        
        if not text:
            return {
                'file': file_path,
                'status': 'failed',
                'title': '',
                'vulnerability_type': '',
                'url': ''
            }
        
        # 提取标题
        title = self.extract_title_from_filename(file_path.name)
        
        # 提取漏洞类型
        vulnerability_types = self.extract_vulnerability_from_filename(file_path.name)
        vulnerability_type = vulnerability_types[0] if vulnerability_types else '未知漏洞'
        
        # 提取URL/IP
        url_ip_data = self.extract_urls_and_ips(text, file_path)
        
        # 优先使用URL，其次使用IP
        if url_ip_data['urls']:
            url = url_ip_data['urls'][0]
        elif url_ip_data['ips']:
            url = url_ip_data['ips'][0]
        else:
            url = '未检测到URL'
        
        return {
            'file': file_path,
            'status': 'success',
            'title': title,
            'vulnerability_type': vulnerability_type,
            'url': url
        }
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """
        在段落中替换文本，保留原有格式
        
        Args:
            paragraph: 段落对象
            old_text: 要替换的文本
            new_text: 新文本
        """
        if old_text in paragraph.text:
            # 先保存替换后的完整文本
            replaced_text = paragraph.text.replace(old_text, new_text)
            
            # 获取第一个run的格式（如果存在）
            if paragraph.runs:
                # 保存第一个run的格式
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                font_underline = first_run.font.underline
                font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
                
                # 清空所有runs
                for run in paragraph.runs:
                    run.text = ''
                
                # 添加新文本，应用原有格式
                new_run = paragraph.add_run(replaced_text)
                new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                new_run.font.bold = font_bold
                new_run.font.italic = font_italic
                new_run.font.underline = font_underline
                if font_color:
                    new_run.font.color.rgb = font_color
            else:
                # 如果没有runs，直接替换
                paragraph.text = replaced_text
    
    def _replace_text_in_cell(self, cell, old_text: str, new_text: str):
        """
        在单元格中替换文本，保留原有格式
        
        Args:
            cell: 单元格对象
            old_text: 要替换的文本
            new_text: 新文本
        """
        for paragraph in cell.paragraphs:
            if old_text in paragraph.text:
                # 先保存替换后的完整文本
                replaced_text = paragraph.text.replace(old_text, new_text)
                
                # 获取第一个run的格式（如果存在）
                if paragraph.runs:
                    # 保存第一个run的格式
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    font_bold = first_run.font.bold
                    font_italic = first_run.font.italic
                    font_underline = first_run.font.underline
                    font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
                    
                    # 清空所有runs
                    for run in paragraph.runs:
                        run.text = ''
                    
                    # 添加新文本，应用原有格式
                    new_run = paragraph.add_run(replaced_text)
                    new_run.font.name = font_name
                    if font_size:
                        new_run.font.size = font_size
                    new_run.font.bold = font_bold
                    new_run.font.italic = font_italic
                    new_run.font.underline = font_underline
                    if font_color:
                        new_run.font.color.rgb = font_color
                else:
                    # 如果没有runs，直接替换
                    paragraph.text = replaced_text

    def _cleanup_temp_screenshot_parts(self) -> None:
        for part_path in self._temp_screenshot_parts:
            try:
                if part_path.exists():
                    part_path.unlink()
            except Exception:
                pass
        self._temp_screenshot_parts = []
        self._screenshot_insert_items = None

    def _screenshot_insert_width(self, image_width: int, image_height: int) -> float:
        if image_width <= 0 or image_height <= 0:
            return SCREENSHOT_INSERT_WIDTH_INCHES
        rendered_height = image_height * SCREENSHOT_INSERT_WIDTH_INCHES / image_width
        if rendered_height <= MAX_SCREENSHOT_PART_HEIGHT_INCHES:
            return SCREENSHOT_INSERT_WIDTH_INCHES
        # 稍微超高的图优先缩小，避免切出多张很短的图；太窄会影响证据可读性，所以设置下限。
        shrunk_width = MAX_SCREENSHOT_PART_HEIGHT_INCHES * image_width / image_height
        return max(MIN_SHRUNK_SCREENSHOT_WIDTH_INCHES, min(SCREENSHOT_INSERT_WIDTH_INCHES, shrunk_width))

    def _prepare_screenshot_insert_items_for_path(self, image_path: Path | None) -> List[Tuple[Path, float]]:
        if image_path is None:
            return []
        if not image_path.exists():
            print(f"[警告] 复测截图不存在: {image_path}")
            return []

        try:
            from PIL import Image

            with Image.open(image_path) as image:
                image_width, image_height = image.size
                insert_width = self._screenshot_insert_width(image_width, image_height)
                rendered_height = image_height * insert_width / max(1, image_width)
                if rendered_height <= MAX_SCREENSHOT_PART_HEIGHT_INCHES:
                    return [(image_path, insert_width)]

                max_part_height_px = max(
                    480,
                    int(image_width * MAX_SCREENSHOT_PART_HEIGHT_INCHES / SCREENSHOT_INSERT_WIDTH_INCHES),
                )
                output_dir = image_path.parent
                suffix = image_path.suffix.lower() or ".png"
                image_format = "JPEG" if suffix in {".jpg", ".jpeg"} else "PNG"
                parts: List[Tuple[Path, float]] = []
                for part_index, top in enumerate(range(0, image_height, max_part_height_px), 1):
                    bottom = min(image_height, top + max_part_height_px)
                    crop = image.crop((0, top, image_width, bottom))
                    if image_format == "JPEG" and crop.mode not in {"RGB", "L"}:
                        crop = crop.convert("RGB")
                    part_path = output_dir / f"{image_path.stem}_part_{part_index:02d}{suffix}"
                    crop.save(part_path, image_format)
                    self._temp_screenshot_parts.append(part_path)
                    parts.append((part_path, SCREENSHOT_INSERT_WIDTH_INCHES))
                print(f"[信息] 复测截图过长，已拆分为 {len(parts)} 张证据图插入报告")
                return parts or [(image_path, insert_width)]
        except Exception as e:
            print(f"[警告] 复测截图切分失败，改为原图插入: {e}")
            return [(image_path, SCREENSHOT_INSERT_WIDTH_INCHES)]

    def _prepare_screenshot_insert_items(self) -> List[Tuple[Path, float]]:
        if self._screenshot_insert_items is None:
            self._screenshot_insert_items = self._prepare_screenshot_insert_items_for_path(self.screenshot_path)
        return self._screenshot_insert_items

    def _insert_screenshot_in_paragraph(self, paragraph, intro_text: str = "") -> bool:
        """在段落中插入复测说明和截图，长图会先缩小或拆分，成功返回 True。"""
        intro_lines = [line.rstrip() for line in str(intro_text or "").splitlines()]
        intro_lines = [line for line in intro_lines if line.strip()]
        if self.screenshot_sections:
            has_existing_section = any(path.exists() for _caption, path in self.screenshot_sections)
            if not has_existing_section:
                return False
        else:
            insert_items = self._prepare_screenshot_insert_items()
            if not insert_items and not intro_lines:
                return False

        if self.screenshot_sections:
            sections = list(self.screenshot_sections)
        else:
            sections = [("\n".join(intro_lines), self.screenshot_path)] if self.screenshot_path is not None else []
        if not sections and not intro_lines:
            return False

        for run in paragraph.runs:
            run.text = ""

        try:
            for section_index, (caption, image_path) in enumerate(sections):
                if section_index > 0:
                    paragraph.add_run().add_break()
                caption_lines = [line.rstrip() for line in str(caption or "").splitlines() if line.strip()]
                for line in caption_lines:
                    run = paragraph.add_run(line)
                    run.add_break()
                insert_items = self._prepare_screenshot_insert_items_for_path(image_path)
                if caption_lines and insert_items:
                    paragraph.add_run().add_break()
                for image_index, (part_path, width_inches) in enumerate(insert_items):
                    if image_index > 0:
                        paragraph.add_run().add_break()
                    run = paragraph.add_run()
                    run.add_picture(str(part_path), width=Inches(width_inches))
                    run.add_break()
            return True
        except Exception as e:
            print(f"[警告] 插入复测截图失败: {e}")
            return False

    def _safe_filename_part(self, value: str, fallback: str = "复测报告") -> str:
        text = str(value or "").strip() or fallback
        text = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", text)
        text = re.sub(r"\s+", " ", text).strip(" .")
        if not text:
            text = fallback
        return text[:120]

    def _fallback_scan_result(self, scan_result: Dict) -> Dict:
        file_path = Path(scan_result.get("file") or self.target_dir / "通报.docx")
        title = self.extract_title_from_filename(file_path.name) if file_path.name else "通报"
        vulns = self.extract_vulnerability_from_filename(file_path.name)
        return {
            "file": file_path,
            "status": "success",
            "title": title or file_path.stem or "通报",
            "vulnerability_type": vulns[0] if vulns else "详见复测结论",
            "url": "详见复测结果截图",
        }
    
    def generate_report(self, scan_result: Dict) -> Optional[Path]:
        """
        根据扫描结果生成复测报告
        
        Args:
            scan_result: 扫描结果字典
            
        Returns:
            生成的报告文件路径
        """
        if scan_result['status'] != 'success':
            print(f"[警告] {getattr(scan_result.get('file'), 'name', scan_result.get('file'))} - 原通报扫描失败，使用文件名和复测截图生成报告")
            scan_result = self._fallback_scan_result(scan_result)
        
        # Agent 可传 staging 目录，待当前 turn 仍有效时再提交到原目录。
        source_dir = self.output_dir or scan_result['file'].parent
        source_dir.mkdir(parents=True, exist_ok=True)

        output_filename = f"{self._safe_filename_part(scan_result['title'])}_复测报告.docx"
        output_path = source_dir / output_filename
        
        try:
            shutil.copy(self.template_path, output_path)
            
            # 打开复制的模板
            doc = Document(output_path)
            
            # 标记：标题是否已替换、截图是否已插入
            title_replaced = False
            image_inserted = False

            # 替换段落中的占位符：
            # - 第一次遇到包含"*"的段落，用于替换标题
            # - 如果存在仅包含"*"的段落，且提供了截图路径，则在该段落插入截图
            for para in doc.paragraphs:
                text = para.text or ""

                # 1) 先处理截图占位（正文通常是单独一行的"*"）
                if (
                    not image_inserted
                    and self.screenshot_path is not None
                    and text.strip() == "*"
                ):
                    if self._insert_screenshot_in_paragraph(para, str(scan_result.get("report_detail") or "")):
                        image_inserted = True
                        continue

                # 2) 标题占位：第一个包含"*"的段落视为标题位置
                if not title_replaced and "*" in text:
                    self._replace_text_in_paragraph(para, "*", scan_result["title"])
                    title_replaced = True
                    continue

            # 替换表格中的*标记（漏洞名称和详情）- 保留格式
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    for cell_index, cell in enumerate(cells):
                        if '*' in cell.text:
                            if (
                                not image_inserted
                                and self.screenshot_path is not None
                                and cell.text.strip() == "*"
                                and cell_index not in (1, 2)
                            ):
                                for para in cell.paragraphs:
                                    if (para.text or "").strip() == "*" and self._insert_screenshot_in_paragraph(para, str(scan_result.get("report_detail") or "")):
                                        image_inserted = True
                                        break
                                if image_inserted:
                                    continue

                            # 检查是否是漏洞名称列
                            # 根据位置判断要填入的内容
                            # 假设表格结构：文件号 | 漏洞名称 | 漏洞详情 | 复测结论
                            if cell_index == 1:  # 漏洞名称
                                self._replace_text_in_cell(cell, '*', scan_result['vulnerability_type'])
                            elif cell_index == 2:  # 漏洞详情
                                self._replace_text_in_cell(cell, '*', scan_result['url'])
            
            # 保存文档
            doc.save(output_path)
            
            print(f"[成功] 生成报告: {output_filename}")
            return output_path
            
        except Exception as e:
            print(f"[错误] 生成报告失败 {scan_result.get('file')}: {e}")
            return None
        finally:
            self._cleanup_temp_screenshot_parts()
    
    def generate_all_reports(self):
        """扫描所有文档并生成复测报告"""
        if not self.target_dir.exists():
            print(f"[错误] 目录不存在: {self.target_dir}")
            return
        
        if not self.template_path.exists():
            print(f"[错误] 模板文件不存在: {self.template_path}")
            return
        
        word_files = self.find_word_files()
        print(f"\n[信息] 找到 {len(word_files)} 个通报文档\n")
        
        if not word_files:
            print("[信息] 未找到任何通报文档")
            return
        
        success_count = 0
        fail_count = 0
        
        for file_path in word_files:
            scan_result = self.scan_document(file_path)
            report_path = self.generate_report(scan_result)
            
            if report_path:
                success_count += 1
            else:
                fail_count += 1
        
        print("\n" + "="*80)
        print("报告生成完成")
        print("="*80)
        print(f"  总文档数: {len(word_files)}")
        print(f"  成功生成: {success_count}")
        print(f"  失败数量: {fail_count}")
        print(f"  说明: 复测报告已生成在各通报文档所在目录")
        print("="*80)


def main():
    """主函数"""
    import sys
    
    print("="*80)
    print("复测报告生成工具")
    print("="*80 + "\n")
    
    # 获取目标目录
    if len(sys.argv) > 1:
        target_dir = sys.argv[1]
    else:
        target_dir = input("请输入通报文档目录路径: ").strip()
    
    if not target_dir:
        print("[错误] 未提供目录路径")
        return
    
    # 获取模板路径
    if len(sys.argv) > 2:
        template_path = sys.argv[2]
    else:
        template_path = input("请输入复测模板路径 (默认: Report_Template/复测模板.docx): ").strip()
        if not template_path:
            template_path = "Report_Template/复测模板.docx"
    
    # 获取输出目录
    if len(sys.argv) > 3:
        output_dir = sys.argv[3]
    else:
        output_dir = input("请输入输出目录 (默认: retest_reports): ").strip()
        if not output_dir:
            output_dir = None
    
    # 创建生成器并执行
    generator = RetestReportGenerator(target_dir, template_path, output_dir)
    generator.generate_all_reports()


if __name__ == "__main__":
    main()
