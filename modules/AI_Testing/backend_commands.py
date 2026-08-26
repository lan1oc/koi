#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AI testing backend commands and retest agent runtime."""

import base64
import contextlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import traceback
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List
from urllib.parse import urlsplit, urlunsplit

import requests

from modules.AI_Testing.retest.llm_provider_catalog import (
    OPENROUTER_DEFAULT_BASE_URL,
    OPENROUTER_FREE_MODEL,
    PROVIDER_AUTO,
    SUPPORTED_LLM_PROVIDERS,
    infer_llm_provider,
    llm_provider_default_name,
    llm_provider_defaults,
    llm_provider_label,
    llm_provider_options,
    normalize_llm_base_url,
    normalize_provider_id,
)
from modules.AI_Testing.retest.retest_http_evidence import repair_utf8_mojibake
from modules.AI_Testing.retest.word_vulnerability_scanner import is_generated_retest_report_path

WORD_SUFFIXES = {".doc", ".docx"}
RETEST_AI_PROVIDERS = SUPPORTED_LLM_PROVIDERS
OPENROUTER_FREE_LIMITS = {
    "requests_per_minute": 20,
    "daily_without_credits": 50,
    "daily_with_credits": 1000,
    "credits_threshold_usd": 10,
}

_RETEST_CANCEL_DIR_NAME = ".retest-control"


def _retest_cancel_key(value: Any) -> str:
    text = re.sub(r"[^A-Za-z0-9._-]", "_", str(value or "").strip())[:120]
    return text or "unknown"


def _retest_cancel_dir() -> Path:
    configured = str(os.environ.get("KOI_USER_DATA_DIR") or "").strip()
    base_dir = Path(configured).expanduser() if configured else Path(__file__).resolve().parents[2]
    return base_dir / _RETEST_CANCEL_DIR_NAME


def _agent_session_store_root() -> Path:
    configured = str(os.environ.get("KOI_USER_DATA_DIR") or "").strip()
    return Path(configured).expanduser() if configured else _project_root()


def _retest_cancel_marker(kind: str, value: Any) -> Path | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    return _retest_cancel_dir() / f"{kind}-{_retest_cancel_key(raw)}.stop"


def _retest_cancel_marker_time_ns(marker: Path) -> int:
    try:
        raw = marker.read_text(encoding="ascii").strip()
        if raw:
            try:
                return int(raw)
            except ValueError:
                # Compatibility with 3.1.3/early 3.1.4 Python markers.
                return int(float(raw) * 1_000_000_000)
    except (OSError, ValueError, OverflowError):
        pass
    try:
        return marker.stat().st_mtime_ns
    except OSError:
        return 0


def _retest_cancel_requested(session_id: Any = "", task_id: Any = "", newer_than_ns: int = 0) -> bool:
    for marker in (
        _retest_cancel_marker("session", session_id),
        _retest_cancel_marker("task", task_id),
    ):
        if marker is None or not marker.is_file():
            continue
        if newer_than_ns > 0 and _retest_cancel_marker_time_ns(marker) < newer_than_ns:
            continue
        return True
    return False


def _write_retest_cancel_marker(kind: str, value: Any) -> bool:
    marker = _retest_cancel_marker(kind, value)
    if marker is None:
        return False
    try:
        marker.parent.mkdir(parents=True, exist_ok=True)
        marker.write_text(str(time.time_ns()), encoding="ascii")
        return True
    except OSError:
        return False


def _clear_retest_cancel_marker(kind: str, value: Any) -> bool:
    """Remove a stale stop marker before starting a new execution epoch."""
    marker = _retest_cancel_marker(kind, value)
    if marker is None or not marker.exists():
        return True
    try:
        marker.unlink()
        return True
    except OSError:
        return False

AI_TESTING_COMMANDS = {
    "doc.agent.message",
    "doc.agent.status",
    "doc.agent.stop",
    "doc.agent.approval.respond",
    "doc.agent.auto_approval.set",
    "doc.agent.auto_approval.status",
    "doc.agent.operation.status",
    "doc.agent.operation.stop",
    "doc.agent.tools",
    "doc.retest.run",
    "doc.retest.list_files",
    "doc.retest.run_one",
    "doc.retest.run_one.start",
    "doc.retest.run_one.status",
    "doc.retest.run_one.stop",
    "doc.retest.confirmation.respond",
    "doc.retest.event_stream.info",
    "doc.retest.agent.start",
    "doc.retest.agent.message",
    "doc.retest.agent.status",
    "doc.retest.agent.stop",
    "doc.retest.agent_chat",
    "doc.retest.session.compact",
    "doc.retest.ai_config.get",
    "doc.retest.ai_config.set",
    "doc.retest.ai_config.test",
    "doc.retest.ai_config.key_status",
    "doc.retest.tools.list",
    "doc.retest.tools.status",
    "doc.retest.tools.install",
    "doc.retest.tools.install.status",
    "doc.retest.generate_reports_with_screenshot",
    "doc.retest.open_output",
}

class _ProgressTee(io.TextIOBase):
    def __init__(self, original: Any, progress: "NoticeProgress"):
        self.original = original
        self.progress = progress
        self._buffer = ""

    @property
    def encoding(self) -> str:
        return getattr(self.original, "encoding", "utf-8")

    def writable(self) -> bool:
        return True

    def write(self, text: str) -> int:
        value = str(text)
        self._buffer += value
        while "\n" in self._buffer:
            line, self._buffer = self._buffer.split("\n", 1)
            if line.strip():
                self.progress.log(line.rstrip())
        if self.original is not None and self.original is not sys.stdout:
            try:
                self.original.write(value)
            except Exception:
                pass
        return len(value)

    def flush(self) -> None:
        if self._buffer.strip():
            self.progress.log(self._buffer.rstrip())
        self._buffer = ""
        if self.original is not None and self.original is not sys.stdout:
            try:
                self.original.flush()
            except Exception:
                pass


class NoticeProgress:
    def __init__(self, total: int = 0):
        self.lock = threading.RLock()
        self.logs: List[str] = []
        self.progress = 0
        self.message = "等待开始..."
        self.processed = 0
        self.total = max(0, int(total or 0))

    def snapshot(self) -> Dict[str, Any]:
        with self.lock:
            return {
                "progress": self.progress,
                "message": self.message,
                "logs": list(self.logs),
                "processed": self.processed,
                "total_reports": self.total,
            }

    def log(self, message: Any) -> None:
        text = str(message or "").strip("\r\n")
        if not text:
            return
        with self.lock:
            self.logs.append(text)
            self.message = text

    def extend(self, messages: Iterable[Any]) -> None:
        for message in messages:
            self.log(message)

    def set(self, progress: int | float | None = None, message: Any | None = None) -> None:
        with self.lock:
            if progress is not None:
                self.progress = max(0, min(100, int(round(float(progress)))))
            if message is not None:
                self.message = str(message)

    def set_total(self, total: int) -> None:
        with self.lock:
            self.total = max(0, int(total or 0))

    def set_processed(self, processed: int, total: int | None = None, message: Any | None = None) -> None:
        with self.lock:
            self.processed = max(0, int(processed or 0))
            if total is not None:
                self.total = max(0, int(total or 0))
            total_value = self.total
            if total_value > 0:
                self.progress = max(self.progress, min(98, 20 + int(self.processed / total_value * 75)))
            if message is not None:
                self.message = str(message)


class ProgressLogList(list):
    def __init__(self, progress: NoticeProgress | None = None):
        super().__init__()
        self.progress = progress

    def append(self, item: Any) -> None:
        super().append(str(item))
        if self.progress:
            self.progress.log(item)

    def extend(self, items: Iterable[Any]) -> None:
        for item in items:
            self.append(item)

    def extend_silent(self, items: Iterable[Any]) -> None:
        for item in items:
            super().append(str(item))


class RetestTaskProgress(NoticeProgress):
    def __init__(self, total: int = 0):
        super().__init__(total)
        self.trace_events: List[Dict[str, Any]] = []
        self.session_id = ""
        self.task_id = ""
        self.stop_requested = False
        self.cancel_epoch_ns = time.time_ns()
        self.resume_snapshot: Dict[str, Any] = {}

    def snapshot(self) -> Dict[str, Any]:
        with self.lock:
            return {
                "progress": self.progress,
                "message": self.message,
                "logs": list(self.logs),
                "processed": self.processed,
                "total_reports": self.total,
                "trace_events": list(self.trace_events),
                "stop_requested": self.stop_requested,
                "resume_snapshot": dict(self.resume_snapshot),
            }

    def delta_snapshot(self, log_offset: int = 0, trace_event_offset: int = 0) -> Dict[str, Any]:
        with self.lock:
            logs = list(self.logs)
            trace_events = list(self.trace_events)
            safe_log_offset = max(0, min(len(logs), int(log_offset or 0)))
            safe_trace_event_offset = max(0, min(len(trace_events), int(trace_event_offset or 0)))
            return {
                "progress": self.progress,
                "message": self.message,
                "logs": logs[safe_log_offset:],
                "log_count": len(logs),
                "processed": self.processed,
                "total_reports": self.total,
                "trace_events": trace_events[safe_trace_event_offset:],
                "trace_event_count": len(trace_events),
                "stop_requested": self.stop_requested,
                "resume_snapshot": dict(self.resume_snapshot),
            }

    def checkpoint(self, snapshot: Dict[str, Any] | None) -> None:
        if not isinstance(snapshot, dict) or not snapshot:
            return
        with self.lock:
            self.resume_snapshot = _json_safe_clone(snapshot)
        source_file = str(snapshot.get("source_file") or "")
        stage = str(snapshot.get("stage") or snapshot.get("resume_stage") or "checkpoint")
        self.event(_retest_trace_event(
            "status",
            "复测断点已保存",
            "已保存当前文件断点；停止或中断后会从此处继续，不重复已完成步骤。",
            "ok",
            source_file=source_file,
            metadata={
                "phase": stage,
                "sourceFileName": Path(source_file).name if source_file else "",
                "resumeSnapshot": _json_safe_clone(snapshot),
            },
        ))

    def request_stop(self, message: str = "复测已停止，可继续") -> None:
        with self.lock:
            self.stop_requested = True
            self.message = message
            self.logs.append(message)

    def should_stop(self) -> bool:
        with self.lock:
            if self.stop_requested:
                return True
            if _retest_cancel_requested(self.session_id, self.task_id, newer_than_ns=self.cancel_epoch_ns):
                # Latch the signal locally. A later run ignores this older
                # marker by epoch, while this task remains cancelled.
                self.stop_requested = True
                self.message = "复测已停止，可继续"
                return True
            return False

    def event(self, event: Dict[str, Any] | None) -> None:
        if not isinstance(event, dict):
            return
        with self.lock:
            self.trace_events.append(dict(event))
            title = str(event.get("title") or "").strip()
            if title:
                self.message = title
        try:
            from modules.backend_api.retest_event_stream import publish_retest_event

            publish_retest_event({
                "type": "retest_trace_event",
                "session_id": self.session_id,
                "task_id": self.task_id,
                "event": event,
            })
        except Exception:
            pass

_RETEST_TASKS: Dict[str, Dict[str, Any]] = {}
_RETEST_TASK_LOCK = threading.RLock()
_RETEST_TOOL_INSTALL_TASKS: Dict[str, Dict[str, Any]] = {}
_RETEST_TOOL_INSTALL_LOCK = threading.RLock()
_RETEST_AGENT_RUNNERS: Dict[str, "RetestAgentRunner"] = {}
_RETEST_AGENT_LOCK = threading.RLock()
_RETEST_AGENT_MAX_RUNNERS = 40
_RETEST_AGENT_IDLE_TTL_SECONDS = 6 * 60 * 60
_RETEST_AGENT_MAX_PENDING_MESSAGES = 8

# ---- 人在回路确认（probe 含本机破坏性操作时，暂停等用户批准）----
# confirmation_id -> {"event": threading.Event, "decision": "approve"|"reject"|"", "note": str}
_RETEST_CONFIRMATIONS: Dict[str, Dict[str, Any]] = {}
_RETEST_CONFIRM_LOCK = threading.RLock()


def _register_confirmation(confirmation_id: str) -> threading.Event:
    evt = threading.Event()
    with _RETEST_CONFIRM_LOCK:
        _RETEST_CONFIRMATIONS[confirmation_id] = {"event": evt, "decision": "", "note": ""}
    return evt


def _resolve_confirmation(confirmation_id: str, decision: str, note: str = "") -> bool:
    with _RETEST_CONFIRM_LOCK:
        entry = _RETEST_CONFIRMATIONS.get(confirmation_id)
        if not entry:
            return False
        entry["decision"] = "approve" if str(decision).lower() in {"approve", "yes", "allow", "true", "1"} else "reject"
        entry["note"] = str(note or "")
        entry["event"].set()
    return True


def _read_confirmation(confirmation_id: str) -> Dict[str, Any]:
    with _RETEST_CONFIRM_LOCK:
        entry = _RETEST_CONFIRMATIONS.get(confirmation_id) or {}
        return {"decision": entry.get("decision") or "", "note": entry.get("note") or ""}


def _discard_confirmation(confirmation_id: str) -> None:
    with _RETEST_CONFIRM_LOCK:
        _RETEST_CONFIRMATIONS.pop(confirmation_id, None)


def _retest_request_confirmation(progress: "RetestTaskProgress | None", request: Dict[str, Any]) -> Dict[str, Any]:
    """暂停并向 UI 推送确认卡片，阻塞等待用户批准/拒绝本机破坏性操作。

    request: {"operation": str, "detail": str, "script": str, "matched": str}
    返回:    {"decision": "approve"|"reject", "note": str}
    用户点停止 / 超时（默认 300s）均按拒绝处理，且原因回传给模型促其改写脚本。
    """
    confirmation_id = uuid.uuid4().hex
    evt = _register_confirmation(confirmation_id)
    operation = str(request.get("operation") or "本机敏感操作")
    matched = str(request.get("matched") or "")
    detail = str(request.get("detail") or "")
    script = str(request.get("script") or "")
    try:
        if progress is not None:
            progress.event(_retest_trace_event(
                "confirmation_request",
                f"需要你确认：{operation}",
                f"复测脚本包含可能影响本机电脑的操作（{matched}）。\n{detail}\n\n"
                f"批准则按原脚本执行；拒绝则不在你本机执行该操作，并让模型改写脚本。",
                "warn",
                metadata={
                    "confirmationId": confirmation_id,
                    "operation": operation,
                    "matched": matched,
                    "script": script[:4000],
                    "requiresUserDecision": True,
                },
            ))
        # 阻塞等待用户决定；期间若用户点了停止，立即按拒绝放行循环。
        waited = 0.0
        while not evt.wait(timeout=1.0):
            waited += 1.0
            if progress is not None and progress.should_stop():
                _discard_confirmation(confirmation_id)
                return {"decision": "reject", "note": "用户已停止复测"}
            if waited >= 300:
                _discard_confirmation(confirmation_id)
                return {"decision": "reject", "note": "确认超时（300s），默认拒绝"}
        outcome = _read_confirmation(confirmation_id)
        _discard_confirmation(confirmation_id)
        return {"decision": outcome.get("decision") or "reject", "note": outcome.get("note") or ""}
    except Exception:
        _discard_confirmation(confirmation_id)
        return {"decision": "reject", "note": "确认流程异常，默认拒绝"}




def is_ai_testing_command(command: str | None) -> bool:
    return str(command or "") in AI_TESTING_COMMANDS


def handle_ai_testing_command(command: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    if command == "doc.agent.message":
        return _doc_agent_message(payload)
    if command == "doc.agent.status":
        return _doc_agent_status(payload)
    if command == "doc.agent.stop":
        return _doc_agent_stop(payload)
    if command == "doc.agent.approval.respond":
        return _doc_agent_approval_respond(payload)
    if command == "doc.agent.auto_approval.set":
        return _doc_agent_auto_approval_set(payload)
    if command == "doc.agent.auto_approval.status":
        return _doc_agent_auto_approval_status(payload)
    if command == "doc.agent.operation.status":
        return _doc_agent_operation_status(payload)
    if command == "doc.agent.operation.stop":
        return _doc_agent_operation_stop(payload)
    if command == "doc.agent.tools":
        return _doc_agent_tools(payload)
    if command == "doc.retest.run":
        return _doc_retest_run(payload)
    if command == "doc.retest.list_files":
        return _doc_retest_list_files(payload)
    if command == "doc.retest.run_one":
        return _doc_retest_run_one(payload)
    if command == "doc.retest.run_one.start":
        return _doc_retest_run_one_start(payload)
    if command == "doc.retest.run_one.status":
        return _doc_retest_run_one_status(payload)
    if command == "doc.retest.run_one.stop":
        return _doc_retest_run_one_stop(payload)
    if command == "doc.retest.confirmation.respond":
        return _doc_retest_confirmation_respond(payload)
    if command == "doc.retest.event_stream.info":
        return _doc_retest_event_stream_info(payload)
    if command == "doc.retest.agent.start":
        return _doc_retest_agent_start(payload)
    if command == "doc.retest.agent.message":
        return _doc_retest_agent_message(payload)
    if command == "doc.retest.agent.status":
        return _doc_retest_agent_status(payload)
    if command == "doc.retest.agent.stop":
        return _doc_retest_agent_stop(payload)
    if command == "doc.retest.agent_chat":
        return _doc_retest_agent_chat(payload)
    if command == "doc.retest.session.compact":
        return _doc_retest_session_compact(payload)
    if command == "doc.retest.ai_config.get":
        return _doc_retest_ai_config_get(payload)
    if command == "doc.retest.ai_config.set":
        return _doc_retest_ai_config_set(payload)
    if command == "doc.retest.ai_config.test":
        return _doc_retest_ai_config_test(payload)
    if command == "doc.retest.ai_config.key_status":
        return _doc_retest_ai_key_status(payload)
    if command == "doc.retest.tools.list":
        return _doc_retest_tools_list(payload)
    if command == "doc.retest.tools.status":
        return _doc_retest_tools_status(payload)
    if command == "doc.retest.tools.install":
        return _doc_retest_tools_install(payload)
    if command == "doc.retest.tools.install.status":
        return _doc_retest_tools_install_status(payload)
    if command == "doc.retest.generate_reports_with_screenshot":
        return _doc_retest_generate_reports_with_screenshot(payload)
    if command == "doc.retest.open_output":
        return _doc_retest_open_output(payload)
    raise ValueError(f"未知 AI 测试命令: {command}")

def _required_text(payload: Dict[str, Any], key: str, message: str) -> str:
    value = str(payload.get(key) or "").strip()
    if not value:
        raise ValueError(message)
    return value

def _path_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    if not text:
        return []
    return [item.strip() for item in text.split(";") if item.strip()]


def _as_record(value: Any) -> Dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_text_list(value: Any, limit: int = 0) -> List[str]:
    if not isinstance(value, list):
        return []
    items = [str(item or "").strip() for item in value]
    items = [item for item in items if item]
    return items[-limit:] if limit and len(items) > limit else items


def _is_generated_retest_report_value(value: Any) -> bool:
    text = str(value or "").strip()
    return bool(text) and is_generated_retest_report_path(Path(text))


def _source_notice_name(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    name = Path(text).name
    return "" if _is_generated_retest_report_value(name) else name


def _source_notice_paths(values: Iterable[Any]) -> List[str]:
    paths: List[str] = []
    seen: set[str] = set()
    for value in values:
        text = str(value or "").strip()
        if not text or _is_generated_retest_report_value(text):
            continue
        key = str(Path(text)).lower()
        if key in seen:
            continue
        seen.add(key)
        paths.append(text)
    return paths


def _strip_retest_report_marker(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    previous = ""
    while previous != text:
        previous = text
        text = re.sub(r"(?i)(?:[\s_\-]*复测报告|[\s_\-]*retest\s*report)\s*$", "", text).strip(" _-")
    return text


def _notice_identity_keys(value: Any) -> List[str]:
    name = Path(str(value or "")).name.strip()
    if not name:
        return []
    stem = Path(name).stem
    candidates = [stem, _strip_retest_report_marker(stem)]
    for item in list(candidates):
        text = str(item or "").strip()
        for suffix in ("的通报", "通报"):
            if text.endswith(suffix) and len(text) > len(suffix):
                candidates.append(text[: -len(suffix)].strip())
    keys: List[str] = []
    seen: set[str] = set()
    for item in candidates:
        normalized = re.sub(r"\s+", "", str(item or "").strip().lower())
        if normalized and normalized not in seen:
            seen.add(normalized)
            keys.append(normalized)
    return keys


def _existing_retest_report_evidence(target_dir: Path, source_files: Iterable[Path]) -> List[Dict[str, str]]:
    source_paths = [Path(item) for item in source_files]
    by_parent: Dict[tuple[str, str], Path] = {}
    by_key: Dict[str, List[Path]] = {}
    for source_file in source_paths:
        parent_key = _retest_target_key(source_file.parent)
        for key in _notice_identity_keys(source_file.name):
            by_parent.setdefault((parent_key, key), source_file)
            by_key.setdefault(key, []).append(source_file)

    evidence_by_source: Dict[str, Dict[str, str]] = {}
    if not target_dir.exists():
        return []
    for root, _, files in os.walk(target_dir):
        root_path = Path(root)
        parent_key = _retest_target_key(root_path)
        for filename in files:
            report_path = root_path / filename
            if report_path.suffix.lower() not in WORD_SUFFIXES or not is_generated_retest_report_path(report_path):
                continue
            matched_source: Path | None = None
            for key in _notice_identity_keys(filename):
                matched_source = by_parent.get((parent_key, key))
                if matched_source is None:
                    candidates = by_key.get(key) or []
                    if len(candidates) == 1:
                        matched_source = candidates[0]
                if matched_source is not None:
                    break
            if matched_source is None:
                continue
            source_key = _retest_target_key(matched_source)
            evidence_by_source.setdefault(
                source_key,
                {
                    "source_file": str(matched_source),
                    "source_file_name": matched_source.name,
                    "report_path": str(report_path),
                    "report_file_name": report_path.name,
                },
            )

    ordered: List[Dict[str, str]] = []
    for source_file in source_paths:
        item = evidence_by_source.get(_retest_target_key(source_file))
        if item:
            ordered.append(item)
    return ordered


def _as_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(value))
    except Exception:
        return default


def _as_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on", "enable", "enabled", "approve", "approved"}:
        return True
    if text in {"0", "false", "no", "n", "off", "disable", "disabled", "reject", "rejected"}:
        return False
    return default


def _truncate_agent_context(value: Any, limit: int = 2000) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + f"\n...[已截断 {len(text) - limit} 字]"


def _frontend_context_memory_text(value: Any, limit: int = 16000) -> str:
    context = _as_record(value)
    if not context:
        return ""
    session = _as_record(context.get("session"))
    progress_evidence = _as_record(context.get("progressEvidence"))
    conversation = context.get("conversation") if isinstance(context.get("conversation"), list) else []
    recent_events = context.get("recentEvents") if isinstance(context.get("recentEvents"), list) else []
    lines: List[str] = ["[会话恢复记忆：来自前端持久化上下文，恢复/继续时必须优先使用]"]
    if session:
        memory_markdown = _truncate_agent_context(session.get("memoryMarkdown"), 10000)
        if memory_markdown:
            lines.append("AI 语义压缩记忆（必须优先使用，不能忽略）:\n" + memory_markdown)
        lines.append(f"会话: {session.get('title') or session.get('sessionId') or ''}")
        lines.append(f"状态: {session.get('status') or ''}；进度: {session.get('progress') or 0}%")
        if session.get("targetDir"):
            lines.append(f"通报目录: {session.get('targetDir')}")
        resume_state = _as_record(session.get("resumeState"))
        if resume_state:
            source_files = _as_text_list(resume_state.get("sourceFiles"))
            lines.append(
                "断点: "
                f"canContinue={bool(resume_state.get('canContinue'))}, "
                f"nextIndex={resume_state.get('nextIndex')}, total={len(source_files)}, "
                f"reason={resume_state.get('blockedReason') or session.get('status') or ''}"
            )
        result_text = _truncate_agent_context(session.get("resultText"), 2200)
        if result_text:
            lines.append("最近复测结果:\n" + result_text)
        log_tail = _as_text_list(session.get("logTail"), 30)
        if log_tail:
            lines.append("最近日志:\n" + "\n".join(_truncate_agent_context(item, 420) for item in log_tail[-30:]))

    completed_names = _as_text_list(progress_evidence.get("completedFileNames"), 200)
    latest_source_name = str(progress_evidence.get("latestSourceFileName") or "").strip()
    completed_count_hint = max(0, _as_int(progress_evidence.get("completedCountHint"), 0))
    next_index_hint = max(0, _as_int(progress_evidence.get("nextIndexHint"), 0))
    next_source_file_name = Path(str(progress_evidence.get("nextSourceFileName") or "")).name
    if completed_count_hint or next_index_hint or next_source_file_name:
        lines.append(
            "Frontend recovery numeric hints: "
            f"completedCountHint={completed_count_hint}, "
            f"nextIndexHint={next_index_hint}, "
            f"nextSourceFileName={next_source_file_name or ''}. "
            "Treat completedCountHint/nextIndexHint as weak metadata only; do not skip notices from numeric hints alone. "
            "Skip only when completedFileNames, disk report evidence, or an exact nextSourceFileName confirms the queue position."
        )
    if completed_names or latest_source_name:
        lines.append(
            "前端恢复进度证据: "
            f"已完成文件 {len(completed_names)} 个"
            + (f"，最近处理 {latest_source_name}" if latest_source_name else "")
        )
        if completed_names:
            lines.append("已完成文件名:\n" + "\n".join(completed_names[-100:]))

    if conversation:
        lines.append("最近对话轮次:")
        for item in conversation[-8:]:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "agent")
            title = str(item.get("title") or role)
            content = _truncate_agent_context(item.get("content"), 900)
            if content:
                lines.append(f"- {role}/{title}: {content}")
    elif recent_events:
        lines.append("最近事件:")
        for item in recent_events[-16:]:
            if not isinstance(item, dict):
                continue
            lines.append(
                f"- {item.get('timestamp') or ''} {item.get('type') or ''}/{item.get('title') or ''}: "
                f"{_truncate_agent_context(item.get('content'), 500)}"
            )
    return _truncate_agent_context("\n".join(line for line in lines if str(line).strip()), limit)


def _retest_target_key(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return os.path.normcase(os.path.abspath(os.path.expanduser(text)))
    except Exception:
        return text.replace("\\", "/").rstrip("/").lower()

_MODEL_REPRODUCED_VALUES = {
    "reproduced",
    "reproducible",
    "unfixed",
    "not_fixed",
    "notfixed",
    "risk",
    "vulnerable",
    "可复现",
    "未修复",
}
_MODEL_CLEAN_VALUES = {
    "not_reproduced",
    "notreproduced",
    "not_reproducible",
    "notreproducible",
    "fixed",
    "clean",
    "pass",
    "passed",
    "已修复",
    "复测通过",
    "不可复现",
}
_MODEL_CLEAN_PATTERNS = (
    r"\bnot[\s_-]?reproduced\b",
    r"\bnot[\s_-]?reproducible\b",
    r"不可复现",
    r"未能?复现",
    r"未见复现",
    r"未发现复现",
    r"无从复现",
    r"无法复现",
    r"未形成可复现证据",
    r"未形成.*复现证据",
    r"没有.*复现证据",
    r"缺乏.*复现证据",
    r"目标.*不可达",
    r"未能验证",
    r"复测通过",
    r"已修复",
)
_MODEL_REPRODUCED_PATTERNS = (
    r"(?<!not[\s_-])\breproduced\b",
    r"\breproducible\b",
    r"仍可复现",
    r"可以复现",
    r"可复现",
    r"未修复",
    r"漏洞仍然成立",
    r"风险仍然存在",
)
_MODEL_VERDICT_FIELD_KEYS = (
    "verdict", "final_verdict", "model_verdict", "reproduction_status",
    "fix_status", "status", "复现状态", "判定", "判断", "结论状态", "是否复现",
)
_MODEL_BOOL_FIELD_KEYS = (
    "reproduced", "is_reproduced", "reproducible", "复现", "是否复现", "漏洞是否复现",
)
_MODEL_TEXT_FIELD_KEYS = (
    "conclusion", "result", "message", "reason", "notes", "summary", "rationale",
    "analysis", "final_answer", "answer", "agent_message", "AGENT_MESSAGE", "decision",
    "judgement", "judgment", "raw_verdict", "_raw_model_text", "结论", "复测结论",
    "判定", "判断", "原因", "说明", "分析",
)
_MODEL_CONTAINER_KEYS = (
    "analysis", "report_analysis", "plan", "retest_plan", "judgement", "judgment",
    "result", "data", "output", "json", "json_result", "final", "final_result",
    "response", "answer", "final_answer", "content",
)


def _canonical_model_verdict(value: Any) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return ""
    return re.sub(r"[\s\-]+", "_", text)


def _model_verdict_from_texts(values: Iterable[Any]) -> str:
    texts = [str(value or "").strip() for value in values if str(value or "").strip()]
    for text in texts:
        canonical = _canonical_model_verdict(text)
        if canonical in _MODEL_REPRODUCED_VALUES:
            return "reproduced"
        if canonical in _MODEL_CLEAN_VALUES:
            return "not_reproduced"
    combined = "\n".join(texts)
    if not combined:
        return ""
    for pattern in _MODEL_CLEAN_PATTERNS:
        if re.search(pattern, combined, flags=re.IGNORECASE):
            return "not_reproduced"
    for pattern in _MODEL_REPRODUCED_PATTERNS:
        if re.search(pattern, combined, flags=re.IGNORECASE):
            return "reproduced"
    return ""


def _normalized_model_key(value: Any) -> str:
    return re.sub(r"[\s_\-]+", "", str(value or "").strip().lower())


def _model_response_dicts(response: Dict[str, Any] | None) -> List[Dict[str, Any]]:
    if not isinstance(response, dict):
        return []
    container_keys = {_normalized_model_key(key) for key in _MODEL_CONTAINER_KEYS}
    queue: List[tuple[Dict[str, Any], int]] = [(response, 0)]
    out: List[Dict[str, Any]] = []
    seen: set[int] = set()
    while queue and len(out) < 24:
        item, depth = queue.pop(0)
        marker = id(item)
        if marker in seen:
            continue
        seen.add(marker)
        out.append(item)
        if depth >= 4:
            continue
        for key, value in item.items():
            if _normalized_model_key(key) in container_keys and isinstance(value, dict):
                queue.append((value, depth + 1))
    return out


def _model_response_values(response: Dict[str, Any] | None, keys: Iterable[str]) -> List[Any]:
    wanted = {_normalized_model_key(key) for key in keys}
    values: List[Any] = []
    for item in _model_response_dicts(response):
        for key, value in item.items():
            if _normalized_model_key(key) not in wanted:
                continue
            if isinstance(value, (dict, list)):
                continue
            values.append(value)
    return values


def _model_response_bool(response: Dict[str, Any] | None, keys: Iterable[str]) -> bool | None:
    for value in _model_response_values(response, keys):
        if isinstance(value, bool):
            return value
        text = str(value or "").strip().lower()
        if text in {"true", "yes", "1", "是", "可复现", "已复现", "复现"}:
            return True
        if text in {"false", "no", "0", "否", "不可复现", "未复现", "未能复现"}:
            return False
    return None


def _model_verdict_from_judgement(judgement: Dict[str, Any] | None, fallback: Any = "") -> str:
    """Return only the model's explicit verdict; never infer from tool counts."""
    source = judgement if isinstance(judgement, dict) else {}
    verdict = _model_verdict_from_texts(_model_response_values(source, _MODEL_VERDICT_FIELD_KEYS) + [fallback])
    if verdict:
        return verdict
    reproduced = _model_response_bool(source, _MODEL_BOOL_FIELD_KEYS)
    if reproduced is not None:
        return "reproduced" if reproduced else "not_reproduced"
    verdict = _model_verdict_from_texts(_model_response_values(source, _MODEL_TEXT_FIELD_KEYS))
    if verdict:
        return verdict
    return ""


def _model_verdict_from_result_data(data: Dict[str, Any] | None) -> str:
    source = data if isinstance(data, dict) else {}
    judgement = source.get("ai_judgement") if isinstance(source.get("ai_judgement"), dict) else {}
    return _model_verdict_from_judgement(judgement, source.get("final_verdict"))


def _model_reproduced_from_result_data(data: Dict[str, Any] | None) -> bool | None:
    verdict = _model_verdict_from_result_data(data)
    if verdict == "reproduced":
        return True
    if verdict == "not_reproduced":
        return False
    return None


def _failure_dicts(failures: Iterable[tuple[Path, str]]) -> List[Dict[str, str]]:
    return [{"file": str(src), "name": src.name, "reason": str(reason)} for src, reason in failures]


def _captured_lines(buffer: io.StringIO) -> List[str]:
    return [line for line in buffer.getvalue().splitlines() if line.strip()]


def _store_captured_lines(logs: List[str], lines: Iterable[Any]) -> None:
    if isinstance(logs, ProgressLogList):
        logs.extend_silent(lines)
    else:
        logs.extend(str(line) for line in lines)


def _append_log(logs: List[str], message: Any, progress: NoticeProgress | None = None) -> None:
    text = str(message or "")
    logs.append(text)
    if progress:
        progress.log(text)


def _extend_logs(logs: List[str], messages: Iterable[Any], progress: NoticeProgress | None = None) -> None:
    for message in messages:
        _append_log(logs, message, progress)


def _run_with_progress_capture(func: Any, progress: NoticeProgress | None = None) -> tuple[Any, List[str]]:
    buffer = io.StringIO()
    stream: Any = _ProgressTee(buffer, progress) if progress else buffer
    with contextlib.redirect_stdout(stream):
        result = func()
    try:
        stream.flush()
    except Exception:
        pass
    return result, _captured_lines(buffer)


def _call_with_progress_capture(func: Any, progress: NoticeProgress | None = None) -> tuple[Any, List[str], BaseException | None]:
    buffer = io.StringIO()
    stream: Any = _ProgressTee(buffer, progress) if progress else buffer
    result = None
    error: BaseException | None = None
    with contextlib.redirect_stdout(stream):
        try:
            result = func()
        except BaseException as exc:
            error = exc
    try:
        stream.flush()
    except Exception:
        pass
    return result, _captured_lines(buffer), error

def _open_path_in_system(path: Path) -> tuple[bool, str | None]:
    try:
        if os.name == "nt":
            os.startfile(str(path))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
        else:
            subprocess.Popen(["xdg-open", str(path)])
        return True, None
    except Exception as exc:
        return False, str(exc)

def _project_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _agent_session_id(payload: Dict[str, Any], default_prefix: str = "agent") -> str:
    raw = (
        payload.get("session_id")
        or payload.get("agent_session_id")
        or payload.get("sessionId")
        or payload.get("agentSessionId")
        or ""
    )
    session_id = str(raw or "").strip()
    return session_id or f"{default_prefix}-{uuid.uuid4().hex[:10]}"


def _agent_target_dir_from_payload(payload: Dict[str, Any]) -> str:
    for key in ("target_dir", "targetDir"):
        value = str(payload.get(key) or "").strip()
        if value:
            return value
    context = _as_record(payload.get("frontend_context"))
    session = _as_record(context.get("session"))
    resume_state = _as_record(session.get("resumeState"))
    progress_evidence = _as_record(context.get("progressEvidence"))
    for value in (
        session.get("targetDir"),
        resume_state.get("targetDir"),
        progress_evidence.get("targetDir"),
    ):
        text = str(value or "").strip()
        if text:
            return text
    return ""


def _agent_workspace_root_from_payload(payload: Dict[str, Any]) -> str:
    for key in ("workspace_root", "workspaceRoot"):
        value = str(payload.get(key) or "").strip()
        if value:
            return value
    context = _as_record(payload.get("frontend_context"))
    session = _as_record(context.get("session"))
    for value in (
        session.get("workspaceRoot"),
        context.get("workspaceRoot"),
        context.get("workspace_root"),
    ):
        text = str(value or "").strip()
        if text:
            return text
    return ""


def _agent_workspace_root(payload: Dict[str, Any] | None = None, session_id: str = "") -> Path:
    payload = payload if isinstance(payload, dict) else {}
    workspace_root = _agent_workspace_root_from_payload(payload)
    if workspace_root:
        candidate = Path(workspace_root).expanduser()
        try:
            if candidate.exists() and candidate.is_dir():
                return candidate.resolve()
        except Exception:
            pass
    clean_id = str(session_id or _agent_session_id(payload)).strip()
    if clean_id:
        store_root = _agent_session_store_root()
        target_key = _retest_target_key(_agent_target_dir_from_payload(payload))
        safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", clean_id or "agent").strip("._")
        session_path = store_root / ".koi_agent_sessions" / f"{safe or 'agent'}.json"
        try:
            raw = json.loads(session_path.read_text(encoding="utf-8"))
            saved_root = str(raw.get("workspace_root") or "").strip() if isinstance(raw, dict) else ""
            if saved_root and _retest_target_key(saved_root) != target_key:
                candidate = Path(saved_root).expanduser()
                if candidate.exists() and candidate.is_dir():
                    return candidate.resolve()
        except Exception:
            pass
    return _project_root()


def _publish_agent_runtime_event(session_id: str, event: Dict[str, Any]) -> None:
    try:
        from modules.backend_api.retest_event_stream import publish_retest_event

        publish_retest_event({
            "type": "retest_trace_event",
            "session_id": session_id,
            "task_id": "agent",
            "event": event,
        })
    except Exception:
        pass


def _make_hybrid_agent_runtime(session_id: str, mode: str = "hybrid", payload: Dict[str, Any] | None = None):
    from modules.AI_Testing.hybrid_agent_runtime import HybridAgentRuntime

    clean_id = str(session_id or "").strip() or f"agent-{uuid.uuid4().hex[:10]}"
    return HybridAgentRuntime(
        clean_id,
        _agent_workspace_root(payload, clean_id),
        publish=lambda event: _publish_agent_runtime_event(clean_id, event),
        mode=mode,
        store_root=_agent_session_store_root(),
    )


def _extract_agent_path(message: str, payload: Dict[str, Any]) -> str:
    for key in ("path", "file", "file_path", "filePath", "target"):
        value = str(payload.get(key) or "").strip()
        if value:
            return value
    text = str(message or "")
    for pattern in (
        r"`([^`]+)`",
        r'"([^"]+)"',
        r"'([^']+)'",
        r"([A-Za-z0-9_.:/\\-]+\.(?:py|ts|tsx|js|jsx|rs|md|json|toml|css|html|yml|yaml|lock|ps1|txt))",
    ):
        match = re.search(pattern, text)
        if match:
            return str(match.group(1)).strip()
    return ""


def _extract_agent_query(message: str, payload: Dict[str, Any]) -> str:
    for key in ("query", "q", "pattern", "keyword"):
        value = str(payload.get(key) or "").strip()
        if value:
            return value
    text = str(message or "").strip()
    quoted = re.search(r"`([^`]+)`|\"([^\"]+)\"|'([^']+)'", text)
    if quoted:
        return next((group for group in quoted.groups() if group), "").strip()
    for marker in ("搜索", "查找", "search", "find", "rg"):
        idx = text.lower().find(marker)
        if idx >= 0:
            return text[idx + len(marker):].strip(" :：，,")
    return text[:120]


def _agent_tool_request(message: str, payload: Dict[str, Any]) -> tuple[str, Dict[str, Any]]:
    raw_tool = payload.get("tool") or payload.get("tool_name") or payload.get("action") or payload.get("name")
    args = payload.get("args") or payload.get("arguments") or payload.get("tool_args") or {}
    if raw_tool:
        return str(raw_tool), args if isinstance(args, dict) else {}

    text = str(message or "")
    lower = text.lower()
    path = _extract_agent_path(text, payload)
    if "run_command" in lower or "执行命令" in text or "跑命令" in text:
        return "run_command", {"command": str(payload.get("command") or text).strip()}
    if "apply_patch" in lower or "写文件" in text or "修改文件" in text:
        return "apply_patch", {"patch": str(payload.get("patch") or text).strip()}
    if "run_tests" in lower or "跑测试" in text or "测试" in text and ("执行" in text or "运行" in text):
        return "run_tests", {"command": str(payload.get("command") or "").strip()}
    if "build_project" in lower or "构建" in text or "build" in lower:
        return "build_project", {"command": str(payload.get("command") or "").strip()}
    if "diff" in lower or ("git" in lower and ("状态" in text or "status" in lower)):
        return "inspect_git_diff", {}
    if "summarize_file" in lower or "总结文件" in text or "概括文件" in text:
        return "summarize_file", {"path": path}
    if "read_file" in lower or "读取" in text or "打开" in text or ("看看" in text and path):
        return "read_file", {"path": path}
    if "search_code" in lower or "搜索" in text or "查找" in text or "rg " in lower:
        return "search_code", {
            "query": _extract_agent_query(text, payload),
            "path": str(payload.get("path") or ""),
            "max_matches": payload.get("max_matches") or 80,
        }
    return "workspace_tree", {
        "path": path or str(payload.get("path") or ""),
        "max_entries": payload.get("max_entries") or 120,
    }


def _runtime_for_agent_approval(approval_id: str) -> tuple[Any | None, str]:
    session_dir = _agent_session_store_root() / ".koi_agent_sessions"
    if not session_dir.exists():
        return None, ""
    for path in session_dir.glob("*.json"):
        try:
            raw = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        approvals = raw.get("approvals") if isinstance(raw, dict) else {}
        if isinstance(approvals, dict) and approval_id in approvals:
            session_id = str(raw.get("id") or path.stem)
            runtime = _make_hybrid_agent_runtime(session_id, payload=raw if isinstance(raw, dict) else None)
            return runtime, session_id
    return None, ""


def _resolve_agent_approval_from_store(approval_id: str, decision: str, note: str = "") -> tuple[bool, str]:
    runtime, session_id = _runtime_for_agent_approval(approval_id)
    if not runtime:
        return False, ""
    return runtime.resolve_approval(approval_id, decision, note), session_id


def _doc_agent_tools(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import HybridWorkspaceTools

    session_id = _agent_session_id(payload)
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    return {
        "success": True,
        "session_id": session_id,
        "auto_approve": bool(runtime.session.auto_approve),
        "workspace_root": str(runtime.workspace_root),
        "tools": HybridWorkspaceTools(runtime).tool_specs(),
    }


def _start_auto_approved_agent_operation(runtime: Any, session_id: str, result: Dict[str, Any]) -> Dict[str, Any]:
    approval_id = str(result.get("approval_id") or "").strip()
    if not approval_id:
        return result
    note = "Auto approval enabled for this Hybrid Agent session."
    if not runtime.resolve_approval(approval_id, "approve", note):
        return {
            **result,
            "success": False,
            "running": False,
            "status": "failed",
            "message": "Auto approval could not resolve the approval request.",
            "final_message": "Auto approval could not resolve the approval request.",
            "agent_session": runtime.snapshot(),
        }
    operation = runtime.operation_for_approval(approval_id)
    operation_id = operation.id if operation else str(result.get("operation_id") or "")
    try:
        operation_id = _start_agent_operation_worker(session_id, approval_id)
    except Exception as exc:
        runtime.record_status(
            "Auto approval operation start failed",
            str(exc),
            "error",
            metadata={
                "phase": "auto_approval",
                "approvalId": approval_id,
                "operationId": operation_id,
                "autoApproved": True,
                "agentRuntime": True,
            },
        )
        return {
            **result,
            "success": False,
            "running": False,
            "status": "failed",
            "operation_id": operation_id,
            "message": f"Auto-approved operation failed to start: {exc}",
            "final_message": f"Auto-approved operation failed to start: {exc}",
            "agent_session": runtime.snapshot(),
        }
    return {
        **result,
        "success": True,
        "blocked": False,
        "running": True,
        "status": "running",
        "approval_id": approval_id,
        "operation_id": operation_id,
        "message": result.get("message") or f"Auto-approved operation started: {operation_id}",
        "final_message": result.get("final_message") or f"Auto-approved operation started: {operation_id}",
        "agent_session": runtime.snapshot(),
    }


def _doc_agent_message(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import HybridAgentLoop
    from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient, load_retest_prompt

    session_id = _agent_session_id(payload)
    cancel_epoch_ns = time.time_ns()
    message = str(payload.get("message") or payload.get("content") or "").strip()
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    if "auto_approve" in payload or "autoApprove" in payload:
        requested_auto_approve = _as_bool(payload.get("auto_approve", payload.get("autoApprove")), False)
        if requested_auto_approve != bool(runtime.session.auto_approve):
            runtime.set_auto_approve(requested_auto_approve, "Updated from agent message payload.")
    frontend_memory = _frontend_context_memory_text(payload.get("frontend_context"), 12000)
    if frontend_memory:
        runtime.set_compact_memory(frontend_memory)
    try:
        ai_config = _ensure_retest_ai_ready("hybrid_agent")
        prompt = load_retest_prompt("hybrid_agent_system")
        client = RetestLLMClient({**ai_config, "_dialogue_stream": True})
        loop = HybridAgentLoop(
            runtime,
            client,
            prompt,
            stop_check=lambda: _retest_cancel_requested(session_id, newer_than_ns=cancel_epoch_ns),
        )
        result = loop.run(message)
        if result.get("auto_approved") and runtime.session.auto_approve and result.get("approval_id"):
            result = _start_auto_approved_agent_operation(runtime, session_id, result)
        return {"session_id": session_id, **result}
    except RetestAIBlockedError as exc:
        runtime.record_status(_ai_blocked_title(exc), str(exc), "warn", metadata={"phase": exc.stage, "blockedByAiConfig": True})
        return {
            "success": False,
            "session_id": session_id,
            "message": str(exc),
            "final_message": str(exc),
            "blocked": True,
            "blocked_by_ai_config": True,
            "blocked_stage": exc.stage,
            "blocked_title": _ai_blocked_title(exc),
            "agent_session": runtime.snapshot(),
        }
    except Exception as exc:
        if _is_ai_runtime_block_message(exc):
            blocked = RetestAIBlockedError(str(exc), "hybrid_agent")
            runtime.record_status(_ai_blocked_title(blocked), str(blocked), "warn", metadata={"phase": blocked.stage, "blockedByAiConfig": True})
            return {
                "success": False,
                "session_id": session_id,
                "message": str(blocked),
                "final_message": str(blocked),
                "blocked": True,
                "blocked_by_ai_config": True,
                "blocked_stage": blocked.stage,
                "blocked_title": _ai_blocked_title(blocked),
                "agent_session": runtime.snapshot(),
            }
        runtime.record_status("Agent 执行失败", str(exc), "error", metadata={"phase": "hybrid_agent"})
        return {
            "success": False,
            "session_id": session_id,
            "message": f"Agent 执行失败: {exc}",
            "final_message": f"Agent 执行失败: {exc}",
            "blocked": False,
            "agent_session": runtime.snapshot(),
        }


def _doc_agent_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import running_agent_operations_snapshot

    session_id = _agent_session_id(payload, default_prefix="agent")
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    runtime.mark_stale_operations()
    snapshot = runtime.snapshot()
    live_operations = running_agent_operations_snapshot(session_id)
    return {
        "success": True,
        "session_id": session_id,
        "auto_approve": bool(runtime.session.auto_approve),
        "workspace_root": str(runtime.workspace_root),
        "agent_session": snapshot,
        "operations": list((snapshot.get("operations") or {}).values()) if isinstance(snapshot.get("operations"), dict) else [],
        "running_operations": live_operations,
    }


def _doc_agent_auto_approval_set(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = _agent_session_id(payload, default_prefix="agent")
    if not session_id:
        return {"success": False, "message": "missing session_id"}
    enabled = _as_bool(payload.get("enabled", payload.get("auto_approve", payload.get("autoApprove"))), False)
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    runtime.set_auto_approve(enabled, str(payload.get("note") or ""))
    return {
        "success": True,
        "session_id": session_id,
        "auto_approve": bool(runtime.session.auto_approve),
        "workspace_root": str(runtime.workspace_root),
        "agent_session": runtime.snapshot(),
    }


def _doc_agent_auto_approval_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = _agent_session_id(payload, default_prefix="agent")
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    return {
        "success": True,
        "session_id": session_id,
        "auto_approve": bool(runtime.session.auto_approve),
        "workspace_root": str(runtime.workspace_root),
        "agent_session": runtime.snapshot(),
    }


def _doc_agent_stop_legacy_removed(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = _agent_session_id(payload, default_prefix="agent")
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    runtime.record_status("Agent 已停止", "当前通用 Agent 会话已收到停止指令。", "warn", metadata={"phase": "stop"})
    runtime.finish_run(None, "stopped")
    return {"success": True, "session_id": session_id, "message": "Agent 已停止", "agent_session": runtime.snapshot()}


def _doc_agent_stop(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import cancel_agent_operations

    session_id = _agent_session_id(payload, default_prefix="agent")
    _write_retest_cancel_marker("session", session_id)
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    cancelled = cancel_agent_operations(session_id)
    snapshot = runtime.snapshot()
    for operation in (snapshot.get("operations") or {}).values() if isinstance(snapshot.get("operations"), dict) else []:
        if isinstance(operation, dict) and str(operation.get("status") or "") not in {"completed", "failed", "rejected", "cancelled", "stale"}:
            runtime.mark_operation_cancel_requested(str(operation.get("id") or ""), "Agent session stop requested by user.")
    runtime.record_status(
        "Agent stopped",
        "Generic Agent stop requested. Running commands were cancelled when possible.",
        "warn",
        metadata={"phase": "stop", "cancelledOperations": cancelled, "agentRuntime": True},
    )
    runtime.finish_run(None, "stopped")
    return {
        "success": True,
        "session_id": session_id,
        "message": "Agent stopped",
        "cancelled_operations": cancelled,
        "agent_session": runtime.snapshot(),
    }


def _doc_agent_operation_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import running_agent_operations_snapshot

    session_id = _agent_session_id(payload, default_prefix="agent")
    operation_id = str(payload.get("operation_id") or payload.get("operationId") or "").strip()
    if not operation_id:
        return {"success": False, "message": "missing operation_id"}
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    operation = runtime.operation_snapshot(operation_id)
    if not operation:
        return {"success": False, "session_id": session_id, "operation_id": operation_id, "message": "operation not found"}
    running = [
        item for item in running_agent_operations_snapshot(session_id)
        if item.get("operation_id") == operation_id
    ]
    return {
        "success": True,
        "session_id": session_id,
        "operation_id": operation_id,
        "operation": operation,
        "running_operation": running[0] if running else None,
        "agent_session": runtime.snapshot(),
    }


def _doc_agent_operation_stop(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.hybrid_agent_runtime import cancel_agent_operations

    session_id = _agent_session_id(payload, default_prefix="agent")
    operation_id = str(payload.get("operation_id") or payload.get("operationId") or "").strip()
    if not operation_id:
        return {"success": False, "message": "missing operation_id"}
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
    cancelled = cancel_agent_operations(session_id=session_id, operation_id=operation_id)
    marked = runtime.mark_operation_cancel_requested(operation_id, "Operation cancellation requested by user.")
    if marked:
        runtime.record_status(
            "Agent operation cancelled",
            f"Cancellation requested for {operation_id}.",
            "warn",
            metadata={"phase": "operation_stop", "operationId": operation_id, "cancelledOperations": cancelled, "agentRuntime": True},
        )
    return {
        "success": marked,
        "session_id": session_id,
        "operation_id": operation_id,
        "cancelled_operations": cancelled,
        "operation": runtime.operation_snapshot(operation_id),
        "agent_session": runtime.snapshot(),
        "message": "operation stop requested" if marked else "operation not found",
    }


def _start_agent_operation_worker(session_id: str, approval_id: str) -> str:
    runtime = _make_hybrid_agent_runtime(session_id)
    operation = runtime.operation_for_approval(approval_id)
    if not operation:
        raise RuntimeError("operation not found for approval")
    thread = threading.Thread(
        target=_agent_operation_worker,
        args=(session_id, approval_id),
        daemon=True,
        name=f"koi-agent-operation-{operation.id}",
    )
    thread.start()
    return operation.id


def _agent_operation_worker(session_id: str, approval_id: str) -> None:
    runtime = _make_hybrid_agent_runtime(session_id)
    try:
        result = runtime.execute_approved_operation(approval_id)
        _agent_operation_reflect(runtime, result)
    except Exception as exc:
        runtime.record_status(
            "Agent operation failed",
            str(exc),
            "error",
            metadata={"phase": "operation", "approvalId": approval_id, "agentRuntime": True},
        )
        runtime.record_chat(
            "Agent",
            f"Approved operation failed: {exc}",
            "error",
            metadata={"phase": "operation", "approvalId": approval_id, "agentRuntime": True},
        )
        runtime.finish_run(None, "failed")


def _agent_operation_reflect(runtime: Any, result: Dict[str, Any]) -> None:
    if _agent_operation_reflect_with_model(runtime, result):
        return


def _agent_operation_reflect_with_model(runtime: Any, result: Dict[str, Any]) -> bool:
    status = str(result.get("status") or "")
    tool_name = str(result.get("tool_name") or "operation")
    summary = str(result.get("summary") or result.get("message") or "")
    operation_id = str(result.get("operation_id") or "")
    approval_id = str(result.get("approval_id") or "")
    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient, load_retest_prompt

        ai_config = _ensure_retest_ai_ready("hybrid_agent_continuation")
        prompt = load_retest_prompt("hybrid_agent_system")
        client = RetestLLMClient({**ai_config, "_dialogue_stream": True})
        messages = [
            {
                "role": "system",
                "content": (
                    f"{prompt}\n\n"
                    "You are resuming after one user-approved operation. "
                    "Do not call tools. Do not invent results. "
                    "Give a brief reflection and final response based only on the recorded observation."
                ),
            },
            *runtime.conversation_messages(30),
            {
                "role": "user",
                "content": (
                    "Approved operation observation:\n"
                    f"tool={tool_name}\nstatus={status}\noperation_id={operation_id}\napproval_id={approval_id}\n"
                    f"summary={summary}\n"
                    f"raw_output_excerpt={str(result.get('raw_output') or '')[:6000]}"
                ),
            },
        ]
        reply = client.chat(messages, [])
        thinking = str(reply.get("thinking") or "").strip()
        content = str(reply.get("content") or "").strip()
        tool_calls = [item for item in (reply.get("tool_calls") or []) if isinstance(item, dict)]
        if thinking:
            runtime.record_thought(
                "Agent reflection",
                thinking[:3000],
                metadata={"phase": "operation_reflection", "operationId": operation_id, "approvalId": approval_id},
            )
        if tool_calls:
            runtime.record_thought(
                "Agent reflection",
                "The continuation model attempted to call tools after an approved operation; tool calls were ignored.",
                metadata={"phase": "operation_reflection", "operationId": operation_id, "approvalId": approval_id},
            )
        if not content:
            content = f"{tool_name} finished with status={status}. The captured output is available in the operation timeline."
        runtime.record_chat(
            "Agent",
            content[:6000],
            "ok" if status == "completed" else ("warn" if status == "cancelled" else "error"),
            metadata={"phase": "operation_final", "operationId": operation_id, "approvalId": approval_id},
        )
        runtime.finish_latest_run("completed" if status == "completed" else status or "completed")
        return True
    except RetestAIBlockedError as exc:
        runtime.record_status(
            _ai_blocked_title(exc),
            str(exc),
            "warn",
            metadata={
                "phase": exc.stage,
                "operationId": operation_id,
                "approvalId": approval_id,
                "blockedByAiConfig": True,
                "agentRuntime": True,
            },
        )
        runtime.finish_latest_run("blocked_by_ai_config")
        return True
    except Exception as exc:
        if _is_ai_runtime_block_message(exc):
            blocked = RetestAIBlockedError(str(exc), "hybrid_agent_continuation")
            runtime.record_status(
                _ai_blocked_title(blocked),
                str(blocked),
                "warn",
                metadata={
                    "phase": blocked.stage,
                    "operationId": operation_id,
                    "approvalId": approval_id,
                    "blockedByAiConfig": True,
                    "agentRuntime": True,
                },
            )
        else:
            runtime.record_status(
                "Agent continuation blocked",
                f"Approved operation finished, but the model continuation failed: {exc}",
                "warn",
                metadata={"phase": "hybrid_agent_continuation", "operationId": operation_id, "approvalId": approval_id, "agentRuntime": True},
            )
        runtime.finish_latest_run("blocked_by_ai_config")
        return True


def _doc_agent_approval_respond_legacy_removed(payload: Dict[str, Any]) -> Dict[str, Any]:
    approval_id = str(payload.get("approval_id") or payload.get("approvalId") or payload.get("confirmation_id") or payload.get("confirmationId") or "").strip()
    if not approval_id:
        return {"success": False, "message": "缺少 approval_id"}
    decision = str(payload.get("decision") or "").strip().lower()
    note = str(payload.get("note") or "")
    if decision not in {"approve", "reject", "yes", "no", "allow", "deny"}:
        return {"success": False, "message": "decision 必须是 approve 或 reject", "approval_id": approval_id}
    session_id = str(payload.get("session_id") or payload.get("sessionId") or payload.get("agent_session_id") or "").strip()
    if session_id:
        runtime = _make_hybrid_agent_runtime(session_id, payload=payload)
        if runtime.resolve_approval(approval_id, decision, note):
            return {"success": True, "session_id": session_id, "approval_id": approval_id, "decision": decision}
    ok, resolved_session_id = _resolve_agent_approval_from_store(approval_id, decision, note)
    if ok:
        return {"success": True, "session_id": resolved_session_id, "approval_id": approval_id, "decision": decision}
    if _resolve_confirmation(approval_id, decision, note):
        return {"success": True, "confirmation_id": approval_id, "decision": decision}
    return {"success": False, "message": "审批请求不存在或已超时", "approval_id": approval_id}


def _doc_agent_approval_respond(payload: Dict[str, Any]) -> Dict[str, Any]:
    approval_id = str(
        payload.get("approval_id")
        or payload.get("approvalId")
        or payload.get("confirmation_id")
        or payload.get("confirmationId")
        or ""
    ).strip()
    if not approval_id:
        return {"success": False, "message": "missing approval_id"}
    decision = str(payload.get("decision") or "").strip().lower()
    note = str(payload.get("note") or "")
    if decision not in {"approve", "reject", "yes", "no", "allow", "deny"}:
        return {"success": False, "message": "decision must be approve or reject", "approval_id": approval_id}

    session_id = str(payload.get("session_id") or payload.get("sessionId") or payload.get("agent_session_id") or "").strip()
    runtime = _make_hybrid_agent_runtime(session_id, payload=payload) if session_id else None
    if runtime and not runtime.approval_request(approval_id):
        runtime = None
    if not runtime:
        runtime, session_id = _runtime_for_agent_approval(approval_id)
    if runtime:
        approved = decision in {"approve", "yes", "allow"}
        if not runtime.resolve_approval(approval_id, "approve" if approved else "reject", note):
            return {"success": False, "message": "approval request could not be resolved", "approval_id": approval_id}
        operation = runtime.operation_for_approval(approval_id)
        operation_id = operation.id if operation else ""
        if approved:
            try:
                operation_id = _start_agent_operation_worker(session_id, approval_id)
            except Exception as exc:
                runtime.record_status(
                    "Agent operation start failed",
                    str(exc),
                    "error",
                    metadata={"phase": "operation_start", "approvalId": approval_id, "operationId": operation_id},
                )
                return {
                    "success": False,
                    "session_id": session_id,
                    "approval_id": approval_id,
                    "operation_id": operation_id,
                    "status": "failed",
                    "message": str(exc),
                    "agent_session": runtime.snapshot(),
                }
            return {
                "success": True,
                "session_id": session_id,
                "approval_id": approval_id,
                "operation_id": operation_id,
                "decision": "approve",
                "status": "running",
                "agent_session": runtime.snapshot(),
            }
        return {
            "success": True,
            "session_id": session_id,
            "approval_id": approval_id,
            "operation_id": operation_id,
            "decision": "reject",
            "status": "rejected",
            "agent_session": runtime.snapshot(),
        }

    if _resolve_confirmation(approval_id, decision, note):
        return {"success": True, "confirmation_id": approval_id, "decision": decision}
    return {"success": False, "message": "approval request not found or expired", "approval_id": approval_id}


def _template_dir() -> Path:
    try:
        from modules.utils.resource_path import get_report_template_dir

        return get_report_template_dir()
    except Exception:
        return _project_root() / "Report_Template"


def _find_template(keyword: str) -> Path | None:
    template_root = _template_dir()
    if not template_root.exists():
        return None
    candidates = sorted(
        item
        for item in template_root.iterdir()
        if item.is_file() and item.suffix.lower() in {".doc", ".docx"}
    )
    return next((item for item in candidates if keyword in item.name), None)


def _retest_template_path() -> Path:
    template_root = _template_dir()
    for filename in ("复测模板.docx", "复测模板.doc"):
        candidate = template_root / filename
        if candidate.exists():
            return candidate
    found = _find_template("复测")
    return found or (template_root / "复测模板.docx")


def _retest_disposal_template_path() -> Path:
    template_root = _template_dir()
    candidate = template_root / "漏洞隐患处置文件.docx"
    if candidate.exists():
        return candidate
    found = _find_template("漏洞隐患处置文件")
    return found or candidate


def _is_disposal_document(path: Path) -> bool:
    name = path.name
    if path.name.startswith("~$") or path.suffix.lower() not in {".doc", ".docx", ".pdf"}:
        return False
    return "处置" in name


def _is_disposal_word_template(path: Path) -> bool:
    if path.name.startswith("~$") or path.suffix.lower() not in {".doc", ".docx"}:
        return False
    name = path.name
    return "处置" in name and ("模板" in name or "处置文件" in name or "处置报告" in name)


def _find_existing_disposal_word_template(directory: Path) -> Path | None:
    candidates = [
        item
        for item in directory.glob("*.docx")
        if _is_disposal_word_template(item)
    ]
    candidates.extend(
        item
        for item in directory.glob("*.doc")
        if _is_disposal_word_template(item)
    )
    candidates = sorted(candidates, key=lambda item: ("模板" not in item.name, str(item).lower()))
    return candidates[0] if candidates else None


def _unique_retest_disposal_output_path(directory: Path, preferred_name: str = "漏洞隐患处置文件.docx") -> Path:
    base = directory / preferred_name
    if not base.exists():
        return base
    stem = base.stem
    suffix = base.suffix
    for index in range(2, 1000):
        candidate = directory / f"{stem} ({index}){suffix}"
        if not candidate.exists():
            return candidate
    return directory / f"{stem}_{int(time.time())}{suffix}"

def _normalize_retest_ai_provider(value: Any) -> str:
    provider = normalize_provider_id(value)
    if provider == PROVIDER_AUTO:
        return PROVIDER_AUTO
    return provider or "openai"


def _retest_ai_provider_default_name(provider: str) -> str:
    return llm_provider_default_name(_normalize_retest_ai_provider(provider))


def _retest_ai_provider_label(provider: str) -> str:
    return llm_provider_label(_normalize_retest_ai_provider(provider))


def _retest_ai_provider_defaults(provider: str) -> Dict[str, Any]:
    return llm_provider_defaults(_normalize_retest_ai_provider(provider))


def _infer_retest_ai_provider(source: Dict[str, Any], fallback_name: str = "") -> str:
    provider = source.get("provider") if isinstance(source, dict) else ""
    base_url = source.get("base_url") if isinstance(source, dict) else ""
    model = source.get("model") if isinstance(source, dict) else ""
    name = source.get("name") if isinstance(source, dict) else fallback_name
    return infer_llm_provider(provider, base_url, model, name)


def _resolve_retest_ai_profile_provider(profile: Dict[str, Any], keep_empty_auto: bool = False) -> Dict[str, Any]:
    raw_provider = _normalize_retest_ai_provider(profile.get("provider") or "")
    has_hint = bool(str(profile.get("base_url") or "").strip() or str(profile.get("model") or "").strip())
    if raw_provider == PROVIDER_AUTO and keep_empty_auto and not has_hint:
        profile["provider"] = PROVIDER_AUTO
        return profile
    base_url_correction = normalize_llm_base_url(
        raw_provider,
        profile.get("base_url") or "",
        profile.get("model") or "",
        profile.get("name") or "",
    )
    provider = infer_llm_provider(
        base_url_correction.get("provider") or raw_provider,
        base_url_correction.get("base_url") or profile.get("base_url") or "",
        profile.get("model") or "",
        profile.get("name") or "",
    )
    defaults = _retest_ai_provider_defaults(provider)
    profile["provider"] = provider
    if str(profile.get("base_url") or "").strip():
        profile["base_url"] = str(base_url_correction.get("base_url") or "").strip()
    else:
        profile["base_url"] = str(defaults.get("base_url") or "")
    if not str(profile.get("model") or "").strip():
        profile["model"] = str(defaults.get("model") or "")
    return profile


def _default_retest_ai_profile(profile_id: str = "default", name: str | None = None, provider: str = "openai") -> Dict[str, Any]:
    provider = _normalize_retest_ai_provider(provider)
    defaults = _retest_ai_provider_defaults(provider)
    return {
        "id": profile_id,
        "name": name or _retest_ai_provider_default_name(provider),
        "provider": provider,
        "base_url": defaults["base_url"],
        "api_key": "",
        "model": defaults["model"],
        "temperature": 0.1,
        "max_tokens": defaults["max_tokens"],
        "context_window": defaults["context_window"],
        "last_updated": "",
    }


def _default_retest_ai_config() -> Dict[str, Any]:
    return {
        "enabled": False,
        "active_profile_id": "auto",
        "profiles": [
            _default_retest_ai_profile("auto", "自动识别", PROVIDER_AUTO),
            _default_retest_ai_profile(),
            _default_retest_ai_profile("openrouter-free", "OpenRouter 免费路由", "openrouter"),
        ],
        "last_updated": "",
    }


def _retest_ai_now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _sanitize_retest_ai_profile_id(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"[^A-Za-z0-9_-]+", "-", text).strip("-")
    return text[:64] or f"profile-{uuid.uuid4().hex[:8]}"


def _retest_ai_float(value: Any, default: float = 0.1) -> float:
    try:
        return max(0.0, min(2.0, float(value)))
    except Exception:
        return default


def _retest_ai_int(value: Any, default: int = 1600) -> int:
    try:
        return max(128, min(65536, int(value)))
    except Exception:
        return default


def _retest_ai_context_int(value: Any, default: int = 128000) -> int:
    try:
        return max(4096, min(2000000, int(value)))
    except Exception:
        return default


def _normalize_retest_ai_profile(raw: Dict[str, Any] | None, fallback_id: str = "default", fallback_name: str = "默认 OpenAI") -> Dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    profile_id = _sanitize_retest_ai_profile_id(source.get("id") or fallback_id)
    provider = _normalize_retest_ai_provider(source.get("provider") or "openai")
    profile = _default_retest_ai_profile(fallback_id, fallback_name, provider)
    defaults = _retest_ai_provider_defaults(provider)
    profile.update({
        "id": profile_id,
        "name": str(source.get("name") or fallback_name or profile_id).strip()[:80] or profile_id,
        "provider": provider,
        "base_url": str(source.get("base_url") or defaults["base_url"]).strip(),
        "api_key": str(source.get("api_key") or "").strip(),
        "model": str(source.get("model") or defaults["model"]).strip(),
        "temperature": _retest_ai_float(source.get("temperature"), 0.1),
        "max_tokens": _retest_ai_int(source.get("max_tokens"), defaults["max_tokens"]),
        "context_window": _retest_ai_context_int(source.get("context_window"), defaults["context_window"]),
        "last_updated": str(source.get("last_updated") or ""),
    })
    return _resolve_retest_ai_profile_provider(profile, keep_empty_auto=True)


def _normalize_retest_ai_config(raw: Dict[str, Any] | None) -> Dict[str, Any]:
    source = raw if isinstance(raw, dict) else {}
    store = _default_retest_ai_config()
    store["enabled"] = _payload_bool(source.get("enabled"), False)
    store["last_updated"] = str(source.get("last_updated") or "")

    raw_profiles = source.get("profiles") if isinstance(source.get("profiles"), list) else None
    if raw_profiles is None:
        if not any(str(source.get(key) or "").strip() for key in ("provider", "base_url", "api_key", "model", "name")):
            return store
        # 兼容旧版单配置：provider/base_url/api_key/model 直接挂在 retest_ai_agent 下。
        migrated = _normalize_retest_ai_profile(source, "default", str(source.get("name") or "默认 OpenAI"))
        store["profiles"] = [_default_retest_ai_profile("auto", "自动识别", PROVIDER_AUTO), migrated]
        if migrated.get("provider") != "openrouter":
            store["profiles"].append(_default_retest_ai_profile("openrouter-free", "OpenRouter 免费路由", "openrouter"))
        store["active_profile_id"] = migrated["id"]
        return store

    profiles: List[Dict[str, Any]] = []
    used_ids: set[str] = set()
    for index, item in enumerate(raw_profiles):
        if not isinstance(item, dict):
            continue
        fallback_id = "default" if index == 0 else f"profile-{index + 1}"
        profile = _normalize_retest_ai_profile(item, fallback_id, str(item.get("name") or f"配置 {index + 1}"))
        base_id = profile["id"]
        if base_id in used_ids:
            suffix = 2
            while f"{base_id}-{suffix}" in used_ids:
                suffix += 1
            profile["id"] = f"{base_id}-{suffix}"
        used_ids.add(profile["id"])
        profiles.append(profile)

    if not profiles:
        profiles = [_default_retest_ai_profile()]
        used_ids = {"default"}
    if not any(profile.get("provider") == PROVIDER_AUTO for profile in profiles):
        profile_id = "auto"
        suffix = 2
        while profile_id in used_ids:
            profile_id = f"auto-{suffix}"
            suffix += 1
        profiles.insert(0, _default_retest_ai_profile(profile_id, "自动识别", PROVIDER_AUTO))
        used_ids.add(profile_id)
    if not any(profile.get("provider") == "openrouter" for profile in profiles):
        profile_id = "openrouter-free"
        suffix = 2
        while profile_id in used_ids:
            profile_id = f"openrouter-free-{suffix}"
            suffix += 1
        profiles.append(_default_retest_ai_profile(profile_id, "OpenRouter 免费路由", "openrouter"))
        used_ids.add(profile_id)
    active_id = _sanitize_retest_ai_profile_id(source.get("active_profile_id") or profiles[0]["id"])
    if active_id not in used_ids:
        active_id = profiles[0]["id"]
    store["profiles"] = profiles
    store["active_profile_id"] = active_id
    return store


def _load_retest_ai_store() -> Dict[str, Any]:
    from modules.config.config_manager import ConfigManager

    manager = ConfigManager()
    config = manager.load_config()
    raw = config.get("retest_ai_agent") if isinstance(config.get("retest_ai_agent"), dict) else {}
    return _normalize_retest_ai_config(raw)


def _active_retest_ai_profile(store: Dict[str, Any]) -> Dict[str, Any]:
    active_id = str(store.get("active_profile_id") or "")
    for profile in store.get("profiles") or []:
        if isinstance(profile, dict) and profile.get("id") == active_id:
            return profile
    profiles = store.get("profiles") or []
    return profiles[0] if profiles and isinstance(profiles[0], dict) else _default_retest_ai_profile()


def _load_retest_ai_config() -> Dict[str, Any]:
    store = _load_retest_ai_store()
    active_profile = _active_retest_ai_profile(store)
    runtime = dict(active_profile)
    runtime["enabled"] = bool(store.get("enabled"))
    runtime["active_profile_id"] = active_profile.get("id")
    return runtime


class RetestAIBlockedError(RuntimeError):
    def __init__(self, message: str, stage: str = "config", resume_snapshot: Dict[str, Any] | None = None):
        super().__init__(message)
        self.stage = stage
        self.resume_snapshot = resume_snapshot if isinstance(resume_snapshot, dict) else {}


def _json_safe_clone(value: Any) -> Any:
    try:
        return json.loads(json.dumps(value, ensure_ascii=False, default=str))
    except Exception:
        if isinstance(value, dict):
            return {str(key): _json_safe_clone(item) for key, item in value.items()}
        if isinstance(value, list):
            return [_json_safe_clone(item) for item in value]
        return str(value)


def _judgement_resume_snapshot(
    source_file: str,
    scan_result: Dict[str, Any] | None,
    result_data: Dict[str, Any] | None,
) -> Dict[str, Any]:
    result_source = result_data if isinstance(result_data, dict) else {}
    return {
        "schema": "retest_judgement_resume.v1",
        "stage": "judgement",
        "source_file": source_file or str(result_source.get("file") or ""),
        "scan_result": _json_safe_clone(scan_result or {}),
        "result_data": _json_safe_clone(result_data or {}),
    }


def _report_resume_snapshot(
    source_file: str,
    summary: str,
    result_data: Dict[str, Any] | None,
    report_summary: str = "",
) -> Dict[str, Any]:
    result_source = result_data if isinstance(result_data, dict) else {}
    return {
        "schema": "retest_report_resume.v1",
        "stage": "report",
        "source_file": source_file or str(result_source.get("file") or ""),
        "summary": str(summary or ""),
        "report_summary": str(report_summary or ""),
        "result_data": _json_safe_clone(result_data or {}),
    }


def _result_resume_snapshot(
    source_file: str,
    summary: str,
    result_data: Dict[str, Any] | None,
) -> Dict[str, Any]:
    result_source = result_data if isinstance(result_data, dict) else {}
    return {
        "schema": "retest_result_resume.v1",
        "stage": "result",
        "source_file": source_file or str(result_source.get("file") or ""),
        "summary": str(summary or ""),
        "result_data": _json_safe_clone(result_data or {}),
    }


def _execution_resume_snapshot(
    source_file: str,
    scan_result: Dict[str, Any] | None,
    valid_urls: List[str] | None,
    retest_results: List[Dict[str, Any]] | None,
    next_url_index: int = 0,
    use_ai: bool = True,
    context_supported: bool = False,
) -> Dict[str, Any]:
    urls = [str(item) for item in (valid_urls or [])]
    completed = [item for item in (retest_results or []) if isinstance(item, dict)]
    bounded_next = max(0, min(int(next_url_index or 0), len(urls)))
    return {
        "schema": "retest_execution_resume.v1",
        "stage": "execution",
        "source_file": source_file,
        "scan_result": _json_safe_clone(scan_result or {}),
        "valid_urls": _json_safe_clone(urls),
        "retest_results": _json_safe_clone(completed),
        "next_url_index": bounded_next,
        "completed_url_count": min(bounded_next, len(completed)),
        "total_url_count": len(urls),
        "use_ai": bool(use_ai),
        "context_supported": bool(context_supported),
    }


def _ai_blocked_title(exc: RetestAIBlockedError) -> str:
    message = str(exc)
    lowered = message.lower()
    if "超时" in message:
        return "模型响应超时"
    if "HTTP 429" in message or "限流" in message or "并发" in message:
        return "模型并发/限流"
    if exc.stage == "probe_repair":
        return "Python 探针修复待继续"
    if exc.stage == "report":
        return "报告生成待重试"
    if (
        "WinError 10013" in message
        or "访问权限不允许" in message
        or "Failed to establish a new connection" in message
        or "NewConnectionError" in message
        or "NameResolutionError" in message
        or "Failed to resolve" in message
        or "connection refused" in lowered
        or "network is unreachable" in lowered
        or "name resolution" in lowered
        or "getaddrinfo failed" in lowered
        or "max retries exceeded" in lowered
        or "httpsconnectionpool" in lowered
        or "proxyerror" in lowered
    ):
        return "模型网络/权限受限"
    if (
        "模型调用失败" in message
        or "/chat/completions" in message
        or "chat/completions" in message
        or "模型接口" in message
        or "SSLError" in message
        or "SSL" in message
    ):
        return "模型调用失败"
    if exc.stage == "config" or "配置" in message or "未启用" in message:
        return "AI 配置阻塞"
    if exc.stage == "session_react":
        return "模型对话暂停"
    return "AI 会话暂停"


def _is_ai_runtime_block_message(message: Any) -> bool:
    text = str(message or "")
    lowered = text.lower()
    return (
        "模型响应超时/网络超时" in text
        or "模型调用失败" in text
        or "WinError 10013" in text
        or "访问权限不允许" in text
        or "Failed to establish a new connection" in text
        or "NewConnectionError" in text
        or "NameResolutionError" in text
        or "Failed to resolve" in text
        or "failed to resolve" in lowered
        or "name resolution" in lowered
        or "connection refused" in lowered
        or "network is unreachable" in lowered
        or "getaddrinfo failed" in lowered
        or "max retries exceeded" in lowered
        or "httpsconnectionpool" in lowered
        or "proxyerror" in lowered
        or "模型并发" in text
        or "模型接口" in text
        or "模型额度" in text
        or "HTTP 429" in text
        or "AI Agent" in text
        or "LLM" in text
        or "llm" in lowered
        or "chat/completions" in lowered
        or "/v1/chat/completions" in lowered
        or "openai" in lowered
        or "openrouter" in lowered
        or "anthropic" in lowered
    )


def _ensure_retest_ai_ready(stage: str = "config") -> Dict[str, Any]:
    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient, load_retest_prompt

        ai_config = _load_retest_ai_config()
        client = RetestLLMClient(ai_config)
    except RetestAIBlockedError:
        raise
    except Exception as exc:
        raise RetestAIBlockedError(f"读取 AI 测试配置失败: {exc}", stage) from exc

    if not ai_config.get("enabled"):
        raise RetestAIBlockedError(
            "AI 测试未启用。请先在「模型与工具」配置并启用 AI，然后回到测试工作台点击「继续测试」。",
            stage,
        )
    if not client.is_ready():
        missing = "、".join(client.missing_items() or ["必要配置"])
        raise RetestAIBlockedError(
            f"AI 测试配置不完整，缺少 {missing}。请先在「模型与工具」补全配置，然后回到测试工作台点击「继续测试」。",
            stage,
        )
    return ai_config


def _retest_payload_uses_ai(payload: Dict[str, Any]) -> bool:
    return not (
        payload.get("use_ai") is False
        or str(payload.get("use_ai") or "").strip().lower() in {"0", "false", "no", "off"}
        or str(payload.get("mode") or "").strip().lower() in {"fast", "quick", "legacy", "local"}
    )


def _ai_blocked_payload(
    exc: RetestAIBlockedError,
    source_file: str = "",
    logs: List[str] | None = None,
    trace_events: List[Dict[str, Any]] | None = None,
) -> Dict[str, Any]:
    return {
        "success": False,
        "blocked_by_ai": True,
        "blocked_by_ai_config": True,
        "blocked_stage": exc.stage,
        "blocked_title": _ai_blocked_title(exc),
        "message": str(exc),
        "source_file": source_file,
        "manual_test_required": False,
        "logs": list(logs or []),
        "trace_events": list(trace_events or []),
        "resume_snapshot": dict(exc.resume_snapshot or {}),
    }


def _safe_retest_ai_profile(profile: Dict[str, Any]) -> Dict[str, Any]:
    safe = _normalize_retest_ai_profile(profile, str(profile.get("id") or "default"), str(profile.get("name") or "默认 OpenAI"))
    api_key = str(safe.get("api_key") or "")
    safe.pop("api_key", None)
    safe["api_key_configured"] = bool(api_key)
    safe["api_key_masked"] = (api_key[:3] + "***" + api_key[-3:]) if len(api_key) >= 8 else ("***" if api_key else "")
    return safe


def _safe_retest_ai_config(config: Dict[str, Any]) -> Dict[str, Any]:
    store = _normalize_retest_ai_config(config)
    safe_profiles = [_safe_retest_ai_profile(profile) for profile in store.get("profiles") or [] if isinstance(profile, dict)]
    active_profile = next(
        (profile for profile in safe_profiles if profile.get("id") == store.get("active_profile_id")),
        safe_profiles[0] if safe_profiles else _safe_retest_ai_profile(_default_retest_ai_profile()),
    )
    return {
        "enabled": bool(store.get("enabled")),
        "active_profile_id": active_profile.get("id"),
        "active_profile": active_profile,
        "profiles": safe_profiles,
        "last_updated": str(store.get("last_updated") or ""),
        # 兼容旧前端字段，指向当前 active profile。
        "provider": active_profile.get("provider"),
        "base_url": active_profile.get("base_url"),
        "model": active_profile.get("model"),
        "temperature": active_profile.get("temperature"),
        "max_tokens": active_profile.get("max_tokens"),
        "context_window": active_profile.get("context_window"),
        "api_key_configured": active_profile.get("api_key_configured"),
        "api_key_masked": active_profile.get("api_key_masked"),
        "provider_options": llm_provider_options(include_auto=True),
    }


def _doc_retest_ai_config_get(payload: Dict[str, Any]) -> Dict[str, Any]:
    store = _load_retest_ai_store()
    return {"success": True, "message": "复测 AI Agent 配置已读取", "config": _safe_retest_ai_config(store)}


def _doc_retest_ai_config_set(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.config.config_manager import ConfigManager

    manager = ConfigManager()
    config = manager.load_config()
    current = config.get("retest_ai_agent") if isinstance(config.get("retest_ai_agent"), dict) else {}
    store = _normalize_retest_ai_config(current)
    action = str(payload.get("action") or "save_profile").strip().lower()

    if "enabled" in payload:
        store["enabled"] = _payload_bool(payload.get("enabled"), False)

    profiles = [profile for profile in (store.get("profiles") or []) if isinstance(profile, dict)]
    active_id = str(store.get("active_profile_id") or (profiles[0].get("id") if profiles else "default"))

    if action == "set_enabled":
        message = "复测 AI Agent 已启用" if store.get("enabled") else "复测 AI Agent 已关闭"
    elif action == "create_profile":
        provider = _normalize_retest_ai_provider(payload.get("provider") or "openai")
        requested_id = _sanitize_retest_ai_profile_id(payload.get("profile_id") or payload.get("name") or provider)
        used_ids = {str(profile.get("id")) for profile in profiles}
        profile_id = requested_id
        suffix = 2
        while profile_id in used_ids:
            profile_id = f"{requested_id}-{suffix}"
            suffix += 1
        name = str(payload.get("name") or _retest_ai_provider_default_name(provider)).strip()[:80]
        new_profile = _default_retest_ai_profile(profile_id, name or profile_id, provider)
        new_profile["last_updated"] = _retest_ai_now()
        profiles.append(new_profile)
        store["active_profile_id"] = profile_id
        message = "复测 AI 配置档已创建"
    elif action == "switch_profile":
        profile_id = _sanitize_retest_ai_profile_id(payload.get("profile_id") or active_id)
        if not any(profile.get("id") == profile_id for profile in profiles):
            raise ValueError("要切换的 AI 配置档不存在")
        store["active_profile_id"] = profile_id
        message = "已切换复测 AI 配置档"
    elif action == "delete_profile":
        profile_id = _sanitize_retest_ai_profile_id(payload.get("profile_id") or active_id)
        if len(profiles) <= 1:
            raise ValueError("至少需要保留一个 AI 配置档")
        original_count = len(profiles)
        profiles = [profile for profile in profiles if profile.get("id") != profile_id]
        if len(profiles) == original_count:
            raise ValueError("要删除的 AI 配置档不存在")
        if store.get("active_profile_id") == profile_id:
            store["active_profile_id"] = profiles[0].get("id")
        message = "复测 AI 配置档已删除"
    else:
        profile_id = _sanitize_retest_ai_profile_id(payload.get("profile_id") or active_id)
        target = next((profile for profile in profiles if profile.get("id") == profile_id), None)
        provider = _normalize_retest_ai_provider(payload.get("provider") or (target.get("provider") if isinstance(target, dict) else "openai"))
        if target is None:
            target = _default_retest_ai_profile(profile_id, str(payload.get("name") or _retest_ai_provider_default_name(provider)), provider)
            profiles.append(target)
        previous_provider = _normalize_retest_ai_provider(target.get("provider") or "openai")
        defaults = _retest_ai_provider_defaults(provider)
        target["provider"] = provider
        if "name" in payload:
            target["name"] = str(payload.get("name") or target.get("name") or profile_id).strip()[:80] or profile_id
        for key in ("base_url", "model"):
            incoming_value = str(payload.get(key) or "").strip() if key in payload else ""
            if key in payload:
                target[key] = incoming_value or str(defaults.get(key) or "")
            elif provider != previous_provider and not str(target.get(key) or "").strip():
                target[key] = str(defaults.get(key) or "")
        if "api_key" in payload:
            incoming_key = str(payload.get("api_key") or "")
            if incoming_key.strip():
                target["api_key"] = incoming_key.strip()
            elif _payload_bool(payload.get("clear_api_key"), False):
                target["api_key"] = ""
        if "temperature" in payload:
            target["temperature"] = _retest_ai_float(payload.get("temperature"), 0.1)
        if "max_tokens" in payload:
            target["max_tokens"] = _retest_ai_int(payload.get("max_tokens"), 800)
        if "context_window" in payload:
            target["context_window"] = _retest_ai_context_int(payload.get("context_window"), 128000)
        target["last_updated"] = _retest_ai_now()
        base_url_before_resolve = str(target.get("base_url") or "").strip()
        target = _resolve_retest_ai_profile_provider(target, keep_empty_auto=True)
        store["active_profile_id"] = profile_id
        message = f"复测 AI 配置档已保存（{_retest_ai_provider_label(target.get('provider'))}）"
        base_url_after_resolve = str(target.get("base_url") or "").strip()
        if base_url_before_resolve and base_url_after_resolve and base_url_before_resolve.rstrip("/") != base_url_after_resolve.rstrip("/"):
            message += f"，Base URL 已自动修正为 {base_url_after_resolve}"

    store["profiles"] = profiles or [_default_retest_ai_profile()]
    if not any(profile.get("id") == store.get("active_profile_id") for profile in store["profiles"]):
        store["active_profile_id"] = store["profiles"][0].get("id")
    store["last_updated"] = _retest_ai_now()
    config["retest_ai_agent"] = store
    if not manager.save_config(config):
        return {"success": False, "message": "保存复测 AI Agent 配置失败", "config": _safe_retest_ai_config(store)}
    return {"success": True, "message": message, "config": _safe_retest_ai_config(store)}


def _runtime_retest_ai_profile_from_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    store = _load_retest_ai_store()
    profiles = [profile for profile in (store.get("profiles") or []) if isinstance(profile, dict)]
    requested_id = _sanitize_retest_ai_profile_id(payload.get("profile_id") or store.get("active_profile_id") or "")
    saved_profile = next((profile for profile in profiles if profile.get("id") == requested_id), None) or _active_retest_ai_profile(store)
    runtime_profile = dict(saved_profile)

    for key in ("provider", "base_url", "model"):
        if key in payload:
            runtime_profile[key] = str(payload.get(key) or "").strip()
    if "temperature" in payload:
        runtime_profile["temperature"] = _retest_ai_float(payload.get("temperature"), 0.1)
    if "max_tokens" in payload:
        runtime_profile["max_tokens"] = _retest_ai_int(payload.get("max_tokens"), 800)
    if "context_window" in payload:
        runtime_profile["context_window"] = _retest_ai_context_int(payload.get("context_window"), 128000)
    incoming_key = str(payload.get("api_key") or "").strip()
    if incoming_key:
        runtime_profile["api_key"] = incoming_key
    elif _payload_bool(payload.get("clear_api_key"), False):
        runtime_profile["api_key"] = ""

    return _normalize_retest_ai_profile(runtime_profile, str(saved_profile.get("id") or "default"), str(saved_profile.get("name") or "AI 配置"))


def _sanitize_openrouter_key_status(value: Any, depth: int = 0) -> Any:
    if depth > 8:
        return str(value)[:500]
    if isinstance(value, dict):
        clean: Dict[str, Any] = {}
        for key, item in value.items():
            key_text = str(key)
            if key_text.lower() in {"api_key", "apikey", "key", "token", "authorization", "password", "secret"}:
                continue
            clean[key_text] = _sanitize_openrouter_key_status(item, depth + 1)
        return clean
    if isinstance(value, list):
        return [_sanitize_openrouter_key_status(item, depth + 1) for item in value[:50]]
    if isinstance(value, str):
        return value[:2000]
    return value


def _openrouter_key_status_summary(data: Any) -> Dict[str, Any]:
    source = data.get("data") if isinstance(data, dict) and isinstance(data.get("data"), dict) else data
    if not isinstance(source, dict):
        return {}
    keys = (
        "label",
        "usage",
        "limit",
        "limit_remaining",
        "limit_reset",
        "is_free_tier",
        "rate_limit",
        "rate_limit_remaining",
        "requests",
        "requests_remaining",
        "credits",
        "credit_balance",
    )
    return {key: source.get(key) for key in keys if key in source}


def _doc_retest_ai_key_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    profile = _runtime_retest_ai_profile_from_payload(payload)
    provider = _normalize_retest_ai_provider(profile.get("provider"))
    if provider != "openrouter":
        return {
            "success": False,
            "message": "当前 Key 状态查询仅支持 OpenRouter。请切换到 OpenRouter 免费路由配置后再查询。",
            "provider": provider,
            "model": profile.get("model") or "",
            "free_model_limits": OPENROUTER_FREE_LIMITS,
        }
    api_key = str(profile.get("api_key") or "").strip()
    if not api_key:
        return {
            "success": False,
            "message": "请先填写或保存 OpenRouter API Key，再查询当前限制和剩余额度。",
            "provider": provider,
            "model": profile.get("model") or "",
            "free_model_limits": OPENROUTER_FREE_LIMITS,
        }

    base_url = str(profile.get("base_url") or OPENROUTER_DEFAULT_BASE_URL).strip().rstrip("/") or OPENROUTER_DEFAULT_BASE_URL
    endpoint = f"{base_url}/key"
    started = time.time()
    try:
        response = requests.get(
            endpoint,
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=(10, 30),
        )
        elapsed_ms = int((time.time() - started) * 1000)
        if response.status_code == 429:
            return {
                "success": False,
                "message": "OpenRouter Key 状态查询被限流（HTTP 429），请稍后再试。",
                "provider": provider,
                "model": profile.get("model") or "",
                "endpoint": endpoint,
                "status_code": response.status_code,
                "elapsed_ms": elapsed_ms,
                "free_model_limits": OPENROUTER_FREE_LIMITS,
            }
        try:
            response.raise_for_status()
        except requests.HTTPError as exc:
            return {
                "success": False,
                "message": f"OpenRouter Key 状态查询失败（HTTP {response.status_code}）: {response.text[:500]}",
                "provider": provider,
                "model": profile.get("model") or "",
                "endpoint": endpoint,
                "status_code": response.status_code,
                "elapsed_ms": elapsed_ms,
                "error": str(exc),
                "free_model_limits": OPENROUTER_FREE_LIMITS,
            }
        try:
            data = response.json()
        except Exception as exc:
            return {
                "success": False,
                "message": f"OpenRouter Key 状态返回内容不是 JSON: {exc}",
                "provider": provider,
                "model": profile.get("model") or "",
                "endpoint": endpoint,
                "status_code": response.status_code,
                "elapsed_ms": elapsed_ms,
                "error": str(exc),
                "free_model_limits": OPENROUTER_FREE_LIMITS,
            }
        safe_data = _sanitize_openrouter_key_status(data)
        return {
            "success": True,
            "message": "OpenRouter Key 状态已读取。免费路由在本应用内会串行并按 20 次/分钟保护；每日额度以 OpenRouter 当前账号状态为准。",
            "provider": provider,
            "model": profile.get("model") or "",
            "endpoint": endpoint,
            "status_code": response.status_code,
            "elapsed_ms": elapsed_ms,
            "summary": _openrouter_key_status_summary(safe_data),
            "data": safe_data,
            "free_model_limits": OPENROUTER_FREE_LIMITS,
        }
    except requests.exceptions.Timeout as exc:
        return {
            "success": False,
            "message": f"OpenRouter Key 状态查询超时: {exc}",
            "provider": provider,
            "model": profile.get("model") or "",
            "endpoint": endpoint,
            "elapsed_ms": int((time.time() - started) * 1000),
            "error": str(exc),
            "free_model_limits": OPENROUTER_FREE_LIMITS,
        }
    except Exception as exc:
        return {
            "success": False,
            "message": f"OpenRouter Key 状态查询异常: {exc}",
            "provider": provider,
            "model": profile.get("model") or "",
            "endpoint": endpoint,
            "elapsed_ms": int((time.time() - started) * 1000),
            "error": str(exc),
            "free_model_limits": OPENROUTER_FREE_LIMITS,
        }


def _doc_retest_ai_config_test(payload: Dict[str, Any]) -> Dict[str, Any]:
    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient, load_retest_prompt
    except Exception as exc:
        return {"success": False, "message": f"导入 AI Agent 客户端失败: {exc}", "error": str(exc)}

    profile = _runtime_retest_ai_profile_from_payload(payload)
    client = RetestLLMClient({**profile, "enabled": True, "_non_stream_json": True})
    provider = client.provider
    model = client.model
    started = time.time()
    if not client.is_ready():
        missing = []
        if not client.api_key:
            missing.append("API Key")
        if not client.model:
            missing.append("Model")
        if provider not in RETEST_AI_PROVIDERS:
            missing.append("Provider")
        return {
            "success": False,
            "message": "模型测试失败：缺少 " + "、".join(missing or ["必要配置"]),
            "provider": provider,
            "model": model,
            "elapsed_ms": 0,
        }

    try:
        response = client.complete_json(
            load_retest_prompt("connectivity_test_system"),
            json.dumps({
                "task": "请确认当前模型 API 是否可通信。",
                "schema": {"ok": True, "message": "一句中文测试结果"},
            }, ensure_ascii=False),
        )
        elapsed_ms = int((time.time() - started) * 1000)
        reply = json.dumps(response, ensure_ascii=False)
        ok = bool(response.get("ok", True))
        message = str(response.get("message") or response.get("reply") or "模型已返回 JSON，通信正常").strip()
        return {
            "success": ok,
            "message": f"模型测试{'成功' if ok else '失败'}：{message}",
            "provider": provider,
            "model": model,
            "reply": reply,
            "elapsed_ms": elapsed_ms,
        }
    except Exception as exc:
        elapsed_ms = int((time.time() - started) * 1000)
        return {
            "success": False,
            "message": f"模型测试失败：{exc}",
            "provider": provider,
            "model": model,
            "elapsed_ms": elapsed_ms,
            "error": str(exc),
        }


def _doc_retest_tools_list(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.AI_Testing.retest.retest_tool_registry import RetestToolRegistry

    tools = RetestToolRegistry().catalog()
    categories: Dict[str, int] = {}
    for tool in tools:
        category = str(tool.get("category") or "other")
        categories[category] = categories.get(category, 0) + 1
    return {"success": True, "message": f"已加载 {len(tools)} 个复测工具", "tools": tools, "categories": categories}


def _doc_retest_tools_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    try:
        from modules.AI_Testing.retest.retest_external_tools import tool_status

        return tool_status()
    except Exception as exc:
        return {"success": False, "message": f"读取外部工具状态失败: {exc}", "logs": [traceback.format_exc()]}


def _external_tool_names_from_payload(payload: Dict[str, Any]) -> List[str] | None:
    tool_names = payload.get("tools")
    if isinstance(tool_names, str):
        tool_names = [item.strip() for item in tool_names.split(",") if item.strip()]
    if isinstance(tool_names, list):
        selected = [str(item).strip().lower() for item in tool_names if str(item).strip()]
        selected = [item for item in dict.fromkeys(selected) if item in {"nmap", "sqlmap", "ffuf"}]
        return selected or None
    return None


def _retest_tools_install_payload(task_id: str) -> Dict[str, Any]:
    with _RETEST_TOOL_INSTALL_LOCK:
        task = _RETEST_TOOL_INSTALL_TASKS.get(task_id)
        if not task:
            return {
                "success": False,
                "task_id": task_id,
                "running": False,
                "done": True,
                "message": "外部工具下载任务不存在或已过期",
                "progress": 100,
                "logs": [],
                "failures": [{"tool": "all", "reason": "task not found"}],
            }
        progress = task["progress"]
        snapshot = progress.snapshot()
        result = task.get("result") if isinstance(task.get("result"), dict) else {}
        done = bool(task.get("done"))
        running = bool(task.get("running"))
        success = bool(task.get("success")) if done else True
        response = {
            "success": success,
            "task_id": task_id,
            "running": running,
            "done": done,
            "message": str(task.get("message") or snapshot.get("message") or ""),
            "progress": 100 if done else snapshot.get("progress", 0),
            "logs": snapshot.get("logs", []),
            "log_count": len(snapshot.get("logs", [])),
            "install_progress": dict(task.get("install_progress") or {}),
            "tool_root": task.get("tool_root") or "",
            "failures": result.get("failures", []) if done else [],
            "error": task.get("error") or "",
        }
        if done:
            response["result"] = result
            response["status"] = result.get("status") if isinstance(result, dict) else None
            if isinstance(result, dict) and result.get("logs"):
                response["logs"] = result.get("logs") or []
                response["log_count"] = len(response["logs"])
        return response


def _retest_tools_install_worker(task_id: str, tool_names: List[str] | None) -> None:
    with _RETEST_TOOL_INSTALL_LOCK:
        task = _RETEST_TOOL_INSTALL_TASKS.get(task_id)
    if not task:
        return
    progress = task["progress"]

    def on_progress(event: Dict[str, Any]) -> None:
        if not isinstance(event, dict):
            return
        message = str(event.get("message") or "").strip()
        percent = event.get("overall_percent", event.get("percent"))
        with _RETEST_TOOL_INSTALL_LOCK:
            current = _RETEST_TOOL_INSTALL_TASKS.get(task_id)
            if current is not None:
                current["install_progress"] = dict(event)
                if message:
                    current["message"] = message
        progress.set(percent, message or None)
        if message and str(event.get("phase") or "") in {
            "start",
            "tool_start",
            "metadata",
            "download_done",
            "extract",
            "copy",
            "verify",
            "done",
            "failed",
            "all_done",
            "all_failed",
        }:
            progress.log(message)

    try:
        from modules.AI_Testing.retest.retest_external_tools import install_tools

        result = install_tools(tool_names, progress=on_progress)
        success = bool(result.get("success"))
        message = str(result.get("message") or ("外部工具下载完成" if success else "外部工具下载失败"))
        progress.set(100, message)
        progress.log(message)
        with _RETEST_TOOL_INSTALL_LOCK:
            task.update({
                "running": False,
                "done": True,
                "success": success,
                "message": message,
                "result": result,
                "tool_root": result.get("tool_root") or task.get("tool_root") or "",
                "error": "" if success else json.dumps(result.get("failures") or [], ensure_ascii=False),
                "finished_at": time.time(),
            })
    except Exception as exc:
        message = f"外部工具下载失败: {exc}"
        progress.set(100, message)
        progress.log(message)
        result = {
            "success": False,
            "message": message,
            "logs": [traceback.format_exc()],
            "failures": [{"tool": "all", "reason": str(exc)}],
        }
        with _RETEST_TOOL_INSTALL_LOCK:
            task.update({
                "running": False,
                "done": True,
                "success": False,
                "message": message,
                "result": result,
                "error": str(exc),
                "finished_at": time.time(),
            })


def _start_retest_tools_install_task(tool_names: List[str] | None) -> Dict[str, Any]:
    selected = tool_names or ["nmap", "sqlmap", "ffuf"]
    task_id = uuid.uuid4().hex
    progress = NoticeProgress(total=len(selected))
    progress.set(1, f"外部工具下载任务已创建: {', '.join(selected)}")
    task = {
        "task_id": task_id,
        "running": True,
        "done": False,
        "success": False,
        "message": progress.snapshot().get("message"),
        "progress": progress,
        "install_progress": {
            "phase": "start",
            "percent": 0,
            "overall_percent": 1,
            "tool_index": 0,
            "tool_count": len(selected),
            "message": progress.snapshot().get("message"),
        },
        "result": None,
        "created_at": time.time(),
        "finished_at": None,
        "error": "",
        "tools": list(selected),
    }
    with _RETEST_TOOL_INSTALL_LOCK:
        now = time.time()
        expired = [
            item_id
            for item_id, item in _RETEST_TOOL_INSTALL_TASKS.items()
            if item.get("done") and now - float(item.get("finished_at") or item.get("created_at") or now) > 3600
        ]
        for item_id in expired:
            _RETEST_TOOL_INSTALL_TASKS.pop(item_id, None)
        _RETEST_TOOL_INSTALL_TASKS[task_id] = task
    worker = threading.Thread(target=_retest_tools_install_worker, args=(task_id, list(selected)), daemon=True)
    task["thread"] = worker
    worker.start()
    return _retest_tools_install_payload(task_id)


def _doc_retest_tools_install(payload: Dict[str, Any]) -> Dict[str, Any]:
    try:
        from modules.AI_Testing.retest.retest_external_tools import install_tools

        tool_names = _external_tool_names_from_payload(payload)
        if payload.get("async") or payload.get("background"):
            return _start_retest_tools_install_task(tool_names)
        return install_tools(tool_names)
    except Exception as exc:
        return {
            "success": False,
            "message": f"一键下载外部工具失败: {exc}",
            "logs": [traceback.format_exc()],
            "failures": [{"tool": "all", "reason": str(exc)}],
        }


def _doc_retest_tools_install_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    task_id = _required_text(payload, "task_id", "缺少外部工具下载任务ID")
    return _retest_tools_install_payload(task_id)


def _valid_http_target(value: str) -> bool:
    text = str(value or "").strip().lower()
    return text.startswith(("http://", "https://")) and not text.startswith((
        "http://schemas.microsoft.com",
        "https://schemas.microsoft.com",
        "http://schemas.openxmlformats.org",
        "https://schemas.openxmlformats.org",
        "http://purl.oclc.org",
        "https://purl.oclc.org",
        "http://www.w3.org",
        "https://www.w3.org",
        "http://www.wps.cn",
        "https://www.wps.cn",
    ))


def _dedupe_http_targets(values: Iterable[Any]) -> List[str]:
    """Return stable, fragment-free HTTP targets without duplicate ReAct runs."""
    seen: set[tuple[str, str, str, str]] = set()
    targets: List[str] = []
    for value in values or []:
        raw = str(value or "").strip()
        if not _valid_http_target(raw):
            continue
        parsed = urlsplit(raw)
        scheme = parsed.scheme.lower()
        hostname = (parsed.hostname or "").lower()
        if not hostname:
            continue
        try:
            port = parsed.port
        except ValueError:
            continue
        default_port = 80 if scheme == "http" else 443
        authority_host = f"[{hostname}]" if ":" in hostname else hostname
        authority = authority_host if not port or port == default_port else f"{authority_host}:{port}"
        path = parsed.path or "/"
        key = (scheme, authority, path, parsed.query)
        if key in seen:
            continue
        seen.add(key)
        targets.append(urlunsplit((scheme, authority, path, parsed.query, "")))
    return targets


_RETEST_CONTEXT_TAG_LABELS = {
    "unauthorized": "未授权/越权",
    "directory_listing": "目录列表",
    "path_traversal": "路径遍历",
    "file_read": "任意文件读取",
    "sensitive_file": "敏感文件",
    "config_leak": "配置泄露",
    "source_leak": "源码泄露",
    "backup_file": "备份/压缩文件",
    "swagger_api": "Swagger/API",
    "phpinfo": "PHPInfo/探针",
    "js_library": "前端库版本",
    "response_header": "响应头泄露",
    "tls": "TLS/证书",
    "cors": "CORS",
    "clickjacking": "点击劫持",
    "http_methods": "HTTP方法",
    "weak_password": "弱口令",
    "sql_injection": "SQL注入",
    "xss": "XSS",
    "ssrf": "SSRF",
    "rce": "命令/代码执行",
    "file_upload": "文件上传",
}


def _format_retest_context_tags(tags: Iterable[Any]) -> str:
    labels = [_RETEST_CONTEXT_TAG_LABELS.get(str(tag), str(tag)) for tag in tags]
    return "；".join(dict.fromkeys(label for label in labels if label))


def _redact_retest_context_value(value: Any, replacements: Dict[str, str]) -> Any:
    if isinstance(value, str):
        text = value
        for secret, masked in replacements.items():
            if secret:
                text = text.replace(secret, masked)
        return text
    if isinstance(value, list):
        return [_redact_retest_context_value(item, replacements) for item in value]
    if isinstance(value, dict):
        return {
            key: _redact_retest_context_value(item, replacements)
            for key, item in value.items()
            if str(key).lower() not in {"password", "passwd", "pwd", "cookie", "authorization", "x-api-key", "token"}
        }
    return value


def _sanitize_http_request_candidates(value: Any, replacements: Dict[str, str]) -> List[Dict[str, Any]]:
    sanitized_requests: List[Dict[str, Any]] = []
    if not isinstance(value, list):
        return sanitized_requests

    for item in value:
        if not isinstance(item, dict):
            continue
        headers = item.get("headers") if isinstance(item.get("headers"), dict) else {}
        body = str(item.get("body") or "")
        safe_item = {
            "method": item.get("method"),
            "target": item.get("target"),
            "url": item.get("url"),
            "header_names": sorted(str(key) for key in headers.keys()),
            "has_body": bool(body),
            "body_line_count": item.get("body_line_count"),
            "body_size": len(body),
            "evidence_lines": item.get("evidence_lines") or [],
            "source_line": item.get("source_line"),
        }
        sanitized_requests.append(_redact_retest_context_value(safe_item, replacements))
    return sanitized_requests


def _sanitize_retest_context(context: Any) -> Any:
    if not isinstance(context, dict):
        return context

    replacements: Dict[str, str] = {}
    for item in context.get("credential_candidates") or []:
        if isinstance(item, dict):
            password = str(item.get("password") or "")
            masked = str(item.get("password_masked") or "***")
            if password:
                replacements[password] = masked

    sanitized: Dict[str, Any] = {}
    for key, value in context.items():
        if key == "http_request_candidates":
            sanitized[key] = _sanitize_http_request_candidates(value, replacements)
            continue

        if key == "credential_candidates" and isinstance(value, list):
            sanitized_credentials: List[Dict[str, Any]] = []
            for item in value:
                if not isinstance(item, dict):
                    continue
                cleaned = {k: v for k, v in item.items() if k != "password"}
                sanitized_credentials.append(_redact_retest_context_value(cleaned, replacements))
            sanitized[key] = sanitized_credentials
            continue

        if key.lower() in {"password", "passwd", "pwd", "cookie", "authorization", "x-api-key", "token"}:
            continue
        sanitized[key] = _redact_retest_context_value(value, replacements)
    return sanitized


def _sanitize_retest_scan_result(scan_result: Dict[str, Any]) -> Dict[str, Any]:
    sanitized = dict(scan_result)
    sanitized.pop("raw_text", None)
    sanitized["retest_context"] = _sanitize_retest_context(sanitized.get("retest_context") or {})
    return sanitized


def _truncate_retest_agent_text(value: Any, limit: int = 6000) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "..."


def _sanitize_retest_agent_payload(value: Any, depth: int = 0) -> Any:
    if depth > 5:
        return "..."
    if isinstance(value, dict):
        sanitized: Dict[str, Any] = {}
        for key, item in value.items():
            key_text = str(key)
            if key_text.lower() in {"password", "passwd", "pwd", "cookie", "authorization", "x-api-key", "token", "api_key"}:
                continue
            sanitized[key_text] = _sanitize_retest_agent_payload(item, depth + 1)
        return sanitized
    if isinstance(value, list):
        return [_sanitize_retest_agent_payload(item, depth + 1) for item in value[:50]]
    if isinstance(value, str):
        return _truncate_retest_agent_text(value, 1000)
    return value


def _doc_retest_session_compact(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = str(payload.get("session_id") or "").strip()
    local_memory = _truncate_retest_agent_text(payload.get("local_memory"), 16000)
    frontend_context = _sanitize_retest_agent_payload(payload.get("frontend_context") if isinstance(payload.get("frontend_context"), dict) else {})
    compact_stats = _sanitize_retest_agent_payload(payload.get("compact_stats") if isinstance(payload.get("compact_stats"), dict) else {})
    recent_events = _sanitize_retest_agent_payload(payload.get("recent_events") if isinstance(payload.get("recent_events"), list) else [])
    logs = payload.get("logs") if isinstance(payload.get("logs"), list) else []
    logs = [_truncate_retest_agent_text(item, 800) for item in logs[-80:]]
    failure_stage = "prepare"
    model_call_started = False

    try:
        failure_stage = "import_client"
        from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient

        failure_stage = "config"
        ai_config = _ensure_retest_ai_ready("session_compaction")
        failure_stage = "client"
        client = RetestLLMClient({
            **ai_config,
            "max_tokens": max(2400, _retest_ai_int(ai_config.get("max_tokens"), 2400)),
            "read_timeout": max(240, _retest_ai_int(ai_config.get("read_timeout"), 240)),
            "_non_stream_json": True,
        })
        system_prompt = (
            "你是 KOI AI 测试会话的压缩器。你的任务是把旧会话压缩成可继续执行的 Markdown 记忆。\n"
            "只根据输入中已有信息总结，不要编造不存在的复测结果、文件名、报告路径或断点。\n"
            "结构化事实优先级最高：nextIndex、sourceFiles、completedFileNames、targetDir、reports、resultText 不能被模型猜改。\n"
            "如果发现旧会话信息已缺失，要在 warning 里说明缺失范围，并在 memory_markdown 中写明只能从现存信息恢复。\n"
            "返回严格 JSON，不要 Markdown 代码块。字段："
            "{\"memory_markdown\":\"完整 Markdown\", \"brief\":\"一句话摘要\", "
            "\"warning\":\"信息缺失或风险，没有则空字符串\", \"confidence\":\"high|medium|low\"}"
        )
        user_payload = {
            "session_id": session_id,
            "local_deterministic_memory": local_memory,
            "frontend_context": frontend_context,
            "compact_stats": compact_stats,
            "recent_events": recent_events,
            "recent_logs": logs,
            "required_markdown_shape": [
                "# KOI AI 测试会话记忆",
                "目标: ...",
                "当前断点: ...",
                "交付状态: ...",
                "--------------",
                "## 值得总结的经验",
                "--------------",
                "## 上一段总结",
                "--------------",
                "## 新的内容",
                "--------------",
                "## 证据索引",
            ],
        }
        failure_stage = "model_request"
        model_call_started = True
        response = client.complete_json(system_prompt, json.dumps(user_payload, ensure_ascii=False))
        failure_stage = "parse_response"
        memory_markdown = _truncate_retest_agent_text(response.get("memory_markdown"), 24000)
        if not memory_markdown:
            raise RuntimeError("模型没有返回 memory_markdown")
        return {
            "success": True,
            "message": "AI 语义压缩完成",
            "ai_compacted": True,
            "memory_markdown": memory_markdown,
            "brief": _truncate_retest_agent_text(response.get("brief"), 1200),
            "warning": _truncate_retest_agent_text(response.get("warning"), 1200),
            "confidence": str(response.get("confidence") or "medium"),
            "provider": client.provider,
            "model": client.model,
            "model_call_started": True,
            "failure_stage": "",
        }
    except RetestAIBlockedError as exc:
        return {
            "success": False,
            "message": str(exc),
            "ai_compacted": False,
            "compact_failed": True,
            "blocked_stage": exc.stage,
            "blocked_title": _ai_blocked_title(exc),
            "failure_stage": failure_stage,
            "model_call_started": model_call_started,
        }
    except Exception as exc:
        return {
            "success": False,
            "message": f"AI 语义压缩失败: {exc}",
            "ai_compacted": False,
            "compact_failed": True,
            "blocked_stage": "session_compaction",
            "blocked_title": _ai_blocked_title(RetestAIBlockedError(str(exc), "session_compaction")),
            "failure_stage": failure_stage,
            "model_call_started": model_call_started,
        }


def _doc_retest_agent_chat(payload: Dict[str, Any]) -> Dict[str, Any]:
    message = str(payload.get("message") or "").strip()
    if not message:
        return {"success": False, "message": "请输入要追问 Agent 的内容", "reply": ""}

    summary = _truncate_retest_agent_text(payload.get("summary"), 8000)
    logs = payload.get("logs") if isinstance(payload.get("logs"), list) else []
    logs = [_truncate_retest_agent_text(item, 800) for item in logs[-60:]]
    history = payload.get("history") if isinstance(payload.get("history"), list) else []
    history = [_sanitize_retest_agent_payload(item) for item in history[-20:] if isinstance(item, dict)]
    result_data = payload.get("result_data") if isinstance(payload.get("result_data"), dict) else {}
    safe_result_data = _sanitize_retest_agent_payload(result_data)
    session_id = str(payload.get("session_id") or "").strip()
    stream_key = f"agent-chat:{session_id or 'adhoc'}:{uuid.uuid4().hex[:10]}"

    def publish_chat_event(content: str, tone: str = "info", complete: bool = False) -> None:
        if not session_id:
            return
        event = _retest_trace_event(
            "chat",
            "Agent",
            _truncate_retest_agent_text(content, 6000),
            tone,
            metadata={
                "role": "agent",
                "phase": "agent_chat",
                "modelOutput": True,
                "dialogueOutput": True,
                "streaming": not complete,
                "completeModelOutput": complete,
                "streamKey": stream_key,
            },
        )
        try:
            from modules.backend_api.retest_event_stream import publish_retest_event

            publish_retest_event({
                "type": "retest_trace_event",
                "session_id": session_id,
                "task_id": "agent-chat",
                "event": event,
            })
        except Exception:
            pass

    def visible_dialogue(raw: str) -> str:
        text = str(raw or "")
        cut_points = []
        for marker in ("```json", "```JSON", "JSON_RESULT:", "\nJSON:", "\n{"):
            index = text.find(marker)
            if index >= 0:
                cut_points.append(index)
        if text.lstrip().startswith("{"):
            return "我正在理解你的意图，并准备调度当前复测会话。"
        if cut_points:
            text = text[:min(cut_points)]
        text = re.sub(r"^\s*AGENT_MESSAGE\s*[:：]\s*", "", text, flags=re.IGNORECASE).strip()
        return text

    stream_state = {"buffer": "", "visible": "", "last_emit": 0.0, "last_emit_count": 0, "count": 0}

    def stream_callback(chunk: str) -> None:
        text = str(chunk or "")
        if not text:
            return
        stream_state["buffer"] = (str(stream_state["buffer"]) + text)[-12000:]
        stream_state["visible"] = visible_dialogue(str(stream_state["buffer"]))
        stream_state["count"] = int(stream_state["count"]) + len(text)
        now = time.time()
        if now - float(stream_state["last_emit"] or 0.0) < 0.2 and int(stream_state["count"]) - int(stream_state["last_emit_count"]) < 80:
            return
        preview = str(stream_state["visible"] or "").strip()
        if not preview:
            return
        stream_state["last_emit"] = now
        stream_state["last_emit_count"] = stream_state["count"]
        publish_chat_event(preview)

    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestLLMClient, load_retest_prompt

        ai_config = _load_retest_ai_config()
        client = RetestLLMClient({**ai_config, "_stream_callback": stream_callback, "_dialogue_stream": True})
        provider = client.provider
        model = client.model
        if not ai_config.get("enabled") or not client.is_ready():
            missing = "、".join(client.missing_items() or ["必要配置"])
            reply = f"AI Agent 未启用或配置不完整，缺少 {missing}。当前对话不会回退成本地规则判断，请先在「模型与工具」补全配置后继续。"
            return {
                "success": False,
                "message": reply,
                "reply": reply,
                "action": "none",
                "action_reason": "AI Agent 未配置完成，无法由模型判断会话动作。",
                "provider": provider,
                "model": model,
                "blocked_by_ai_config": True,
                "blocked_stage": "chat",
                "blocked_title": "AI 会话配置阻塞",
                "stream_key": stream_key,
            }

        system_prompt = load_retest_prompt("agent_chat_system")
        user_payload = {
            "schema": {
                "reply": "中文回答",
                "action": "none|continue_retest|rerun_retest",
                "generate_reports": "boolean，默认 false，只有用户明确要求报告时为 true",
                "action_reason": "为什么选择该动作",
            },
            "user_question": message,
            "session_state": {
                "can_continue": bool(payload.get("can_continue")),
                "is_running": bool(payload.get("is_running")),
                "target_dir": str(payload.get("target_dir") or ""),
            },
            "session_summary": summary,
            "recent_logs": logs,
            "chat_history": history,
            "result_data": safe_result_data,
        }
        response = client.complete_json(system_prompt, json.dumps(user_payload, ensure_ascii=False))
        reply = str(response.get("reply") or response.get("answer") or "").strip()
        if not reply:
            if response.get("action") == "continue_retest":
                reply = "我会从当前暂停断点继续复测。"
            elif response.get("action") == "rerun_retest":
                reply = "我会基于当前通报目录重新创建一轮完整复测。"
            else:
                reply = "我已读取当前复测上下文。"
        action = str(response.get("action") or "none").strip()
        if action not in {"none", "continue_retest", "rerun_retest"}:
            action = "none"
        generate_reports = bool(response.get("generate_reports")) and action in {"continue_retest", "rerun_retest"}
        publish_chat_event(reply, "ok", True)
        return {
            "success": True,
            "message": "AI Agent 已回答",
            "reply": _truncate_retest_agent_text(reply, 6000),
            "action": action,
            "generate_reports": generate_reports,
            "action_reason": _truncate_retest_agent_text(response.get("action_reason") or "", 1000),
            "provider": provider,
            "model": model,
            "stream_key": stream_key,
        }
    except Exception as exc:
        message_text = str(exc)
        if "HTTP 429" in message_text or "限流" in message_text or "并发" in message_text:
            title = "模型并发/限流，待继续"
        elif "超时" in message_text or "timeout" in message_text.lower():
            title = "模型响应超时，待继续"
        else:
            title = "AI 会话暂停"
        reply = f"{title}: {message_text}。当前不会回退成本地规则判断，网络或模型恢复后请继续。"
        publish_chat_event(reply, "warn", True)
        return {
            "success": False,
            "message": reply,
            "reply": reply,
            "action": "none",
            "action_reason": "AI Agent 会话调用失败，未使用本地规则兜底。",
            "provider": "",
            "model": "",
            "blocked_by_ai_config": True,
            "blocked_stage": "chat",
            "blocked_title": title,
            "stream_key": stream_key,
        }

def _bounded_retest_ai_config(config: Dict[str, Any], seconds: int) -> Dict[str, Any]:
    bounded = dict(config or {})
    try:
        configured_timeout = int(bounded.get("read_timeout") or seconds)
    except Exception:
        configured_timeout = seconds
    bounded["read_timeout"] = min(configured_timeout, seconds)
    bounded["max_retries"] = 0
    return bounded


def _apply_retest_ai_agent(scan_result: Dict[str, Any], logs: List[str], stream_callback: Callable[[str], None] | None = None) -> Dict[str, Any]:
    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestAIAgent
        from modules.AI_Testing.retest.retest_tool_registry import RetestToolRegistry

        ai_config = _bounded_retest_ai_config(_ensure_retest_ai_ready("planning"), 60)
        if stream_callback:
            ai_config = {**ai_config, "_stream_callback": stream_callback, "_dialogue_stream": True}
        agent = RetestAIAgent(ai_config, RetestToolRegistry())
        agent.client.set_request_deadline(60)
        try:
            advice = agent.advise(scan_result)
        finally:
            agent.client.clear_request_deadline()
        updated = agent.apply_advice(scan_result, advice)
        context = updated.get("retest_context") if isinstance(updated.get("retest_context"), dict) else {}
        agent_advice = context.get("agent_advice") if isinstance(context, dict) else {}
        if isinstance(agent_advice, dict):
            provider = agent_advice.get("provider") or ai_config.get("provider") or "openai"
            model = agent_advice.get("model") or ai_config.get("model") or ""
            prefix = f"AI Agent({provider}{('/' + model) if model else ''})"
            if agent_advice.get("used"):
                recommended = agent_advice.get("recommended_checks") or []
                logs.append(f"{prefix} 已参与复测规划，推荐工具: {', '.join(recommended) if recommended else '无'}")
        return updated
    except RetestAIBlockedError:
        raise
    except Exception as exc:
        message = str(exc)
        if "模型响应超时/网络超时" in message or "HTTP 429" in message:
            raise RetestAIBlockedError(f"AI Agent 规划阶段暂停: {message}", "planning") from exc
        raise RetestAIBlockedError(f"AI Agent 规划阶段调用异常，已暂停，可继续: {message}", "planning") from exc


def _apply_retest_ai_judgement(
    scan_result: Dict[str, Any],
    result_data: Dict[str, Any],
    logs: List[str],
    stream_callback: Callable[[str], None] | None = None,
) -> Dict[str, Any]:
    source_file = str(result_data.get("file") or "")
    resume_snapshot = _judgement_resume_snapshot(source_file, scan_result, result_data)
    try:
        from modules.AI_Testing.retest.retest_ai_agent import RetestAIAgent
        from modules.AI_Testing.retest.retest_tool_registry import RetestToolRegistry

        ai_config = _bounded_retest_ai_config(_ensure_retest_ai_ready("judgement"), 60)
        if stream_callback:
            ai_config = {**ai_config, "_stream_callback": stream_callback, "_dialogue_stream": True}
        agent = RetestAIAgent(ai_config, RetestToolRegistry())
        agent.client.set_request_deadline(60)
        try:
            judgement = agent.judge_retest(scan_result, result_data)
        finally:
            agent.client.clear_request_deadline()
        verdict = _model_verdict_from_judgement(judgement)
        if not verdict:
            raise RuntimeError("模型没有给出明确 verdict，不能由工具结果或代码兜底判定。")
        reproduced = verdict == "reproduced"
        judgement["verdict"] = verdict
        judgement["reproduced"] = reproduced
        judgement["fix_status"] = "risk" if reproduced else "clean"
        judgement["conclusion"] = judgement.get("conclusion") or ("漏洞未修复/可复现" if reproduced else "漏洞已修复/复测通过")
        result_data["ai_judgement"] = judgement
        result_data["final_verdict"] = verdict
        result_data["ai_reproduced"] = reproduced
        # Compatibility fields only mirror the model verdict.
        result_data["risk_count"] = 1 if reproduced else 0
        result_data["manual_count"] = 0
        result_data["manual_test_required"] = False
        result_data["reason"] = judgement.get("reason") or result_data.get("reason") or ""
        logs.append(
            "AI Agent 已完成最终判定: "
            + ("漏洞未修复/可复现" if reproduced else "漏洞已修复/复测通过")
        )
        return result_data
    except RetestAIBlockedError as exc:
        if not exc.resume_snapshot:
            exc.resume_snapshot = resume_snapshot
        raise
    except Exception as exc:
        message = str(exc)
        if "模型响应超时/网络超时" in message or "HTTP 429" in message:
            raise RetestAIBlockedError(f"AI Agent 判定阶段暂停: {message}", "judgement", resume_snapshot) from exc
        raise RetestAIBlockedError(f"AI Agent 判定阶段调用异常，已暂停，可继续: {message}", "judgement", resume_snapshot) from exc


def _format_retest_summary(file_path: Path, result_data: Dict[str, Any]) -> str:
    lines: List[str] = []
    lines.append(f"文件: {file_path.name}")
    scan_result = result_data.get("scan_result") or {}

    vuln_types = scan_result.get("vulnerability_types") or []
    if vuln_types:
        lines.append("通报漏洞类型: " + "；".join(vuln_types))
    context = scan_result.get("retest_context") or {}
    context_tags = context.get("issue_tags") or []
    if context_tags:
        lines.append("正文识别标签: " + _format_retest_context_tags(context_tags))
    agent_advice = context.get("agent_advice") if isinstance(context, dict) else {}
    if isinstance(agent_advice, dict) and agent_advice.get("enabled"):
        provider = agent_advice.get("provider") or ""
        model = agent_advice.get("model") or ""
        ai_state = "已完成通报阅读与复测规划" if agent_advice.get("used") else "规划未完成（已暂停等待继续）"
        lines.append(f"AI Agent: {provider}{('/' + model) if model else ''}，{ai_state}")
        recommended = agent_advice.get("recommended_checks") or []
        if recommended:
            lines.append("AI推荐工具: " + "；".join(str(item) for item in recommended[:10]))
        plan_steps = agent_advice.get("plan_steps") or []
        if plan_steps:
            lines.append("AI复测计划: " + "；".join(str(item) for item in plan_steps[:5]))
        notes = str(agent_advice.get("notes") or agent_advice.get("reason") or agent_advice.get("error") or "").strip()
        if notes:
            lines.append("AI说明: " + notes[:220])
        warnings = agent_advice.get("warnings") or []
        if warnings:
            lines.append("AI提醒: " + "；".join(str(item) for item in warnings[:4]))
    http_requests = context.get("http_request_candidates") or []
    payloads = context.get("payload_candidates") or []
    if http_requests or payloads:
        lines.append(f"正文复测线索: HTTP请求块 {len(http_requests)} 个，参数载荷 {len(payloads)} 个")
    path_candidates = context.get("path_candidates") or []
    if path_candidates:
        lines.append("正文证据路径: " + "；".join(str(item) for item in path_candidates[:6]))
    expected_markers = context.get("expected_markers") or []
    if expected_markers:
        lines.append("正文证据特征: " + "；".join(str(item) for item in expected_markers[:6]))
    credentials = context.get("credential_candidates") or []
    if credentials:
        cred_text = []
        for item in credentials[:3]:
            username = item.get("username", "")
            password = item.get("password_masked", "")
            if username or password:
                cred_text.append(f"{username}/{password}")
        if cred_text:
            lines.append("正文账号线索: " + "；".join(cred_text))
    urls = result_data.get("urls") or []
    lines.append(f"复测URL数量: {len(urls)}")
    if result_data.get("context_supported"):
        lines.append("复测策略: 已启用通报正文上下文复测")
    ai_judgement = result_data.get("ai_judgement") if isinstance(result_data.get("ai_judgement"), dict) else {}
    judgement_label = "快速规则" if result_data.get("fast_mode") or ai_judgement.get("source") == "fast_rules" else "AI"
    if ai_judgement:
        lines.append(f"{judgement_label}最终判定: " + str(ai_judgement.get("conclusion") or ""))
        if ai_judgement.get("reason"):
            lines.append(f"{judgement_label}判定理由: " + str(ai_judgement.get("reason"))[:260])
        evidence = ai_judgement.get("evidence") or []
        if evidence:
            lines.append(f"{judgement_label}关键证据: " + "；".join(str(item) for item in evidence[:6]))

    results = result_data.get("retest_results") or []
    observation_records = [
        vuln
        for item in results
        for vuln in (item.get("vulnerabilities") or [])
        if isinstance(vuln, dict)
        and not vuln.get("tool_unavailable")
    ]
    observation_count = result_data.get("observation_count")
    try:
        observation_count = int(observation_count)
    except Exception:
        observation_count = sum(_retest_observation_count(item) for item in results if isinstance(item, dict))
    if observation_count < len(observation_records):
        observation_count = len(observation_records)
    lines.append(f"工具观察记录总数: {max(0, observation_count)}")
    final_verdict = _model_verdict_from_result_data(result_data)
    _judgement = result_data.get("ai_judgement") if isinstance(result_data.get("ai_judgement"), dict) else {}
    _unreachable = bool(_judgement.get("unverified_unreachable") or result_data.get("target_unreachable"))
    if final_verdict == "reproduced":
        lines.append("复测结论: 漏洞未修复/可复现")
    elif final_verdict == "not_reproduced":
        lines.append("复测结论: 未复现：目标不可达，未能验证（建议复查）" if _unreachable else "复测结论: 漏洞已修复/复测通过")
    else:
        lines.append(f"复测结论: 等待{judgement_label}判定")
    for index, item in enumerate(results, 1):
        lines.append("")
        lines.append(f"[{index}] {item.get('url', '')}")
        if item.get("error"):
            if item.get("target_unreachable"):
                lines.append(f"    目标不可达: {item.get('error')}（当前未见可利用入口）")
            else:
                lines.append(f"    复测错误: {item.get('error')}")
            continue
        meta = item.get("request_meta") or {}
        if meta:
            if meta.get("status_code") is not None:
                lines.append(f"    响应: HTTP {meta.get('status_code')}，长度 {meta.get('content_length', '-')}, 耗时 {meta.get('elapsed_ms', '-')}ms")
            elif meta.get("error"):
                lines.append(f"    响应错误: {meta.get('error')}")
        vulnerabilities = item.get("vulnerabilities") or []
        if not vulnerabilities:
            lines.append("    未记录额外工具观察")
            if item.get("note"):
                lines.append(f"    说明: {item.get('note')}")
            continue
        for vuln in vulnerabilities:
            detail = vuln.get("detail") or ""
            evidence = vuln.get("evidence")
            lines.append(f"    [{vuln.get('severity', 'info')}] {vuln.get('type', '未知类型')} - {detail}")
            if evidence:
                lines.append(f"        证据: {evidence}")
    return "\n".join(lines)


def _report_value(data: Any, *keys: str, limit: int = 0) -> str:
    if not isinstance(data, dict):
        return ""
    for key in keys:
        value = data.get(key)
        if value is None:
            continue
        if isinstance(value, (dict, list, tuple)):
            text = json.dumps(value, ensure_ascii=False, default=str)
        else:
            text = str(value)
        text = repair_utf8_mojibake(text).strip()
        if text:
            return text[:limit] if limit and len(text) > limit else text
    return ""


def _report_dict(data: Any, *keys: str) -> Dict[str, Any]:
    if not isinstance(data, dict):
        return {}
    for key in keys:
        value = data.get(key)
        if isinstance(value, dict) and value:
            return value
    return {}


def _append_report_block(lines: List[str], label: str, text: Any, max_chars: int = 1600, max_lines: int = 18) -> None:
    value = repair_utf8_mojibake(text).replace("\r\n", "\n").replace("\r", "\n").strip()
    if not value:
        return
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", value)
    truncated = len(value) > max_chars
    if truncated:
        value = value[:max_chars].rstrip()
    block_lines = value.splitlines() or [value]
    lines.append(f"{label}:")
    for line in block_lines[:max_lines]:
        lines.append(f"  {line[:220]}")
    if truncated or len(block_lines) > max_lines:
        lines.append("  ...(已截断)")


def _format_report_meta(data: Dict[str, Any]) -> str:
    meta = _report_dict(data, "response_meta", "responseMeta", "request_meta", "requestMeta")
    if not meta:
        return ""
    parts: List[str] = []
    status = _report_value(meta, "status_code", "statusCode", "status")
    if status:
        parts.append(f"HTTP {status}")
    final_url = _report_value(meta, "final_url", "finalUrl", "url", limit=220)
    if final_url:
        parts.append(f"final_url={final_url}")
    content_type = _report_value(meta, "content_type", "contentType", "mime_type", "mimeType", limit=120)
    if content_type:
        parts.append(f"content-type={content_type}")
    content_length = _report_value(meta, "content_length", "contentLength", "length", "bytes")
    if content_length:
        parts.append(f"length={content_length}")
    elapsed = _report_value(meta, "elapsed_ms", "elapsedMs", "duration_ms", "durationMs")
    if elapsed:
        parts.append(f"elapsed={elapsed}ms")
    error = _report_value(meta, "error", limit=260)
    if error:
        parts.append(f"error={error}")
    return "，".join(parts)


def _append_report_headers(lines: List[str], headers: Dict[str, Any]) -> None:
    if not isinstance(headers, dict) or not headers:
        return
    lines.append("响应头(脱敏):")
    for index, (key, value) in enumerate(headers.items()):
        if index >= 8:
            lines.append("  ...(已截断)")
            break
        lines.append(f"  {key}: {repair_utf8_mojibake(value)[:180]}")


def _report_judgement_parts(result_data: Dict[str, Any]) -> tuple[str, str, str]:
    data = result_data if isinstance(result_data, dict) else {}
    ai_judgement = data.get("ai_judgement") if isinstance(data.get("ai_judgement"), dict) else {}
    judgement_label = "快速规则" if data.get("fast_mode") or ai_judgement.get("source") == "fast_rules" else "AI"
    final_verdict = _model_verdict_from_result_data(data)
    conclusion = ai_judgement.get("conclusion") or (
        "漏洞未修复/可复现"
        if final_verdict == "reproduced"
        else "漏洞已修复/复测通过"
        if final_verdict == "not_reproduced"
        else "模型未给出判定"
    )
    return judgement_label, repair_utf8_mojibake(conclusion), final_verdict


def _format_report_text_explanation(summary: str, result_data: Dict[str, Any]) -> str:
    data = result_data if isinstance(result_data, dict) else {}
    ai_judgement = data.get("ai_judgement") if isinstance(data.get("ai_judgement"), dict) else {}
    judgement_label, conclusion, _final_verdict = _report_judgement_parts(data)
    results = [item for item in (data.get("retest_results") or []) if isinstance(item, dict)]
    lines: List[str] = [
        "复测说明",
        f"复测结论: {conclusion}",
    ]
    reason = repair_utf8_mojibake(ai_judgement.get("reason") or data.get("reason") or "").strip()
    if reason:
        lines.append(f"{judgement_label}判定理由: {reason[:600]}")
    urls = [repair_utf8_mojibake(item).strip() for item in (data.get("urls") or []) if str(item).strip()]
    if urls:
        lines.append("通报目标: " + "；".join(urls[:6]))
    observation_count = data.get("observation_count")
    if observation_count is not None or results:
        lines.append(f"执行概况: 覆盖 {len(results)} 个目标，记录 {observation_count if observation_count is not None else '-'} 条观察。")
    if data.get("target_unreachable") or ai_judgement.get("unverified_unreachable"):
        lines.append("可达性说明: 目标不可达或未发现可利用入口，未将不可达直接等同于漏洞已修复。")
    lines.append("证据位置: HTTP 请求/响应、命中特征和工具探针输出已作为下方证据截图写入报告。")
    if len(lines) <= 3:
        fallback = repair_utf8_mojibake(summary).strip()
        if fallback:
            _append_report_block(lines, "复测摘要", fallback, max_chars=1200, max_lines=12)
    return "\n".join(lines[:28])


def _format_report_evidence_screenshot_text(summary: str, result_data: Dict[str, Any]) -> str:
    data = result_data if isinstance(result_data, dict) else {}
    lines: List[str] = [
        "复测证据",
        "HTTP请求/响应证据",
    ]
    urls = [repair_utf8_mojibake(item).strip() for item in (data.get("urls") or []) if str(item).strip()]
    if urls:
        lines.append("通报目标: " + "；".join(urls[:6]))

    results = [item for item in (data.get("retest_results") or []) if isinstance(item, dict)]
    if not results:
        lines.append("未记录结构化 HTTP 响应证据。")
    for index, item in enumerate(results[:5], 1):
        if len(lines) >= 108:
            lines.append("...(更多证据已截断)")
            break
        url = _report_value(item, "url", limit=240)
        lines.append("")
        lines.append(f"[{index}] 目标: {url or '-'}")
        if item.get("target_unreachable"):
            lines.append(f"目标状态: 不可达/未见可利用入口。{_report_value(item, 'error', limit=300)}")
        elif item.get("error"):
            lines.append(f"执行错误: {_report_value(item, 'error', limit=300)}")
        meta_line = _format_report_meta(item)
        if meta_line:
            lines.append(f"响应信息: {meta_line}")
        request_text = _report_value(item, "request_safe", "requestSafe", "request_raw", "requestRaw", limit=1400)
        _append_report_block(lines, "重放请求包", request_text, max_chars=1400, max_lines=12)
        _append_report_headers(lines, _report_dict(item, "response_headers_safe", "responseHeadersSafe"))
        response_text = _report_value(
            item,
            "response_raw_excerpt",
            "responseRawExcerpt",
            "response_body_preview",
            "responseBodyPreview",
            limit=1800,
        )
        _append_report_block(lines, "响应数据片段", response_text, max_chars=1800, max_lines=18)

        vulnerabilities = [v for v in (item.get("vulnerabilities") or []) if isinstance(v, dict)]
        if vulnerabilities:
            lines.append("工具/探针证据:")
        for vuln_index, vuln in enumerate(vulnerabilities[:4], 1):
            detail = _report_value(vuln, "detail", "evidence", limit=360)
            severity = _report_value(vuln, "severity") or "info"
            vuln_type = _report_value(vuln, "type") or "复测证据"
            lines.append(f"  {vuln_index}. [{severity}] {vuln_type}: {detail or '-'}")
            markers = vuln.get("matched_markers") or vuln.get("matchedMarkers") or []
            if isinstance(markers, list) and markers:
                lines.append("     命中特征: " + "；".join(repair_utf8_mojibake(marker) for marker in markers[:8]))
            vuln_meta = _format_report_meta(vuln)
            if vuln_meta and vuln_meta != meta_line:
                lines.append(f"     响应信息: {vuln_meta}")
            vuln_request = _report_value(vuln, "request_safe", "requestSafe", "request_raw", "requestRaw", limit=900)
            if vuln_request and vuln_request != request_text:
                _append_report_block(lines, "     证据请求包", vuln_request, max_chars=900, max_lines=7)
            vuln_response = _report_value(
                vuln,
                "response_raw_excerpt",
                "responseRawExcerpt",
                "response_body_preview",
                "responseBodyPreview",
                limit=1200,
            )
            if vuln_response and vuln_response != response_text:
                _append_report_block(lines, "     证据响应片段", vuln_response, max_chars=1200, max_lines=10)
    return "\n".join(lines[:118])


def _format_report_evidence_sections(summary: str, result_data: Dict[str, Any]) -> List[Dict[str, str]]:
    data = result_data if isinstance(result_data, dict) else {}
    results = [item for item in (data.get("retest_results") or []) if isinstance(item, dict)]
    sections: List[Dict[str, str]] = []
    if not results:
        return sections

    for index, item in enumerate(results[:5], 1):
        url = _report_value(item, "url", limit=240) or "-"
        caption_lines = [f"证据 {index}: {url}"]
        meta_line = _format_report_meta(item)
        if meta_line:
            caption_lines.append(f"响应信息: {meta_line}")
        if item.get("target_unreachable"):
            caption_lines.append(f"目标状态: 不可达/未见可利用入口。{_report_value(item, 'error', limit=300)}")
        elif item.get("error"):
            caption_lines.append(f"执行错误: {_report_value(item, 'error', limit=300)}")

        evidence_lines: List[str] = ["HTTP请求/响应证据", f"目标: {url}"]
        if meta_line:
            evidence_lines.append(f"响应信息: {meta_line}")
        request_text = _report_value(item, "request_safe", "requestSafe", "request_raw", "requestRaw", limit=1800)
        _append_report_block(evidence_lines, "重放请求包", request_text, max_chars=1800, max_lines=14)
        _append_report_headers(evidence_lines, _report_dict(item, "response_headers_safe", "responseHeadersSafe"))
        response_text = _report_value(
            item,
            "response_raw_excerpt",
            "responseRawExcerpt",
            "response_body_preview",
            "responseBodyPreview",
            limit=2400,
        )
        _append_report_block(evidence_lines, "响应数据片段", response_text, max_chars=2400, max_lines=22)

        vulnerabilities = [v for v in (item.get("vulnerabilities") or []) if isinstance(v, dict)]
        if vulnerabilities:
            evidence_lines.append("工具/探针原始证据:")
        for vuln_index, vuln in enumerate(vulnerabilities[:4], 1):
            detail = _report_value(vuln, "detail", "evidence", limit=420)
            severity = _report_value(vuln, "severity") or "info"
            vuln_type = _report_value(vuln, "type") or "复测证据"
            evidence_lines.append(f"  {vuln_index}. [{severity}] {vuln_type}: {detail or '-'}")
            markers = vuln.get("matched_markers") or vuln.get("matchedMarkers") or []
            if isinstance(markers, list) and markers:
                evidence_lines.append("     命中特征: " + "；".join(repair_utf8_mojibake(marker) for marker in markers[:8]))
            vuln_request = _report_value(vuln, "request_safe", "requestSafe", "request_raw", "requestRaw", limit=1000)
            if vuln_request and vuln_request != request_text:
                _append_report_block(evidence_lines, "     证据请求包", vuln_request, max_chars=1000, max_lines=8)
            vuln_response = _report_value(
                vuln,
                "response_raw_excerpt",
                "responseRawExcerpt",
                "response_body_preview",
                "responseBodyPreview",
                limit=1400,
            )
            if vuln_response and vuln_response != response_text:
                _append_report_block(evidence_lines, "     证据响应片段", vuln_response, max_chars=1400, max_lines=12)

        sections.append({
            "caption": "\n".join(caption_lines),
            "text": "\n".join(evidence_lines[:96]),
        })
    return sections


def _format_report_evidence_snapshot(summary: str, result_data: Dict[str, Any]) -> str:
    return _format_report_evidence_screenshot_text(summary, result_data)


def _payload_bool(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on"}:
        return True
    if text in {"0", "false", "no", "n", "off"}:
        return False
    return bool(value)


def _frontend_context_requests_reports(value: Any) -> bool:
    context = _as_record(value)
    if not context:
        return False
    session = _as_record(context.get("session"))
    resume_state = _as_record(session.get("resumeState"))
    if _payload_bool(resume_state.get("generateReports"), False) or _payload_bool(session.get("generateReports"), False):
        return True
    if _as_text_list(resume_state.get("reports")):
        return True

    target_key = _retest_target_key(session.get("targetDir") or resume_state.get("targetDir") or "")
    last_report_path = str(session.get("lastReportPath") or "").strip()
    if last_report_path and (not target_key or _retest_target_key(last_report_path) != target_key):
        return True

    positive_phrases = (
        "一键复测",
        "继续测试并生成报告",
        "继续生成报告",
        "报告已生成",
        "报告生成完成",
        "生成报告:",
        "生成报告：",
        "已为本会话已完成的通报生成报告",
    )

    def metadata_wants_reports(value: Any) -> bool:
        metadata = _as_record(value)
        return (
            _payload_bool(metadata.get("generateReports"), False)
            or _payload_bool(metadata.get("generate_reports"), False)
        )

    def record_wants_reports(item: Any) -> bool:
        record = _as_record(item)
        if not record:
            return False
        if metadata_wants_reports(record.get("metadata")):
            return True
        if _as_text_list(record.get("reports")):
            return True
        tool = _as_record(record.get("tool"))
        tool_id = str(tool.get("toolId") or "").strip().lower()
        tool_label = str(tool.get("label") or "").strip()
        if tool_id in {"generate_report", "generate_reports"} or tool_label == "生成报告":
            return True
        text = "\n".join(
            str(record.get(key) or "")
            for key in ("title", "content", "resultPreview", "failureReason")
        )
        return any(phrase in text for phrase in positive_phrases)

    recent_events = context.get("recentEvents") if isinstance(context.get("recentEvents"), list) else []
    for item in recent_events:
        if record_wants_reports(item):
            return True
    conversation = context.get("conversation") if isinstance(context.get("conversation"), list) else []
    for item in conversation:
        if record_wants_reports(item):
            return True
        tool_items = item.get("tools") if isinstance(item, dict) and isinstance(item.get("tools"), list) else []
        for tool_item in tool_items:
            if record_wants_reports(tool_item):
                return True
    text = "\n".join(str(session.get(key) or "") for key in ("status", "resultText", "memoryMarkdown"))
    return any(phrase in text for phrase in positive_phrases)


def _save_retest_screenshot_data(target_dir: Path, screenshot_data_url: str) -> Path:
    text = str(screenshot_data_url or "").strip()
    match = re.match(r"^data:image/(png|jpeg|jpg);base64,(.+)$", text, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        raise ValueError("复测截图数据格式不正确")

    image_type = match.group(1).lower()
    extension = "jpg" if image_type in {"jpeg", "jpg"} else "png"
    try:
        image_bytes = base64.b64decode(match.group(2), validate=True)
    except Exception as exc:
        raise ValueError("复测截图 base64 解码失败") from exc
    if not image_bytes:
        raise ValueError("复测截图数据为空")

    output_dir = _retest_screenshot_dir(target_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"ui_result_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{extension}"
    with open(output_path, "wb") as screenshot_file:
        screenshot_file.write(image_bytes)
    return output_path


def _save_retest_text_screenshot(target_dir: Path, text: str, title: str = "复测结果") -> Path:
    from PIL import Image, ImageDraw, ImageFont

    output_dir = _retest_screenshot_dir(target_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"agent_result_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.png"
    width = 1280
    padding = 36
    line_height = 28
    max_chars = 88

    def load_font(size: int):
        candidates = [
            Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "msyh.ttc",
            Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "simhei.ttf",
            Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "arial.ttf",
        ]
        for candidate in candidates:
            try:
                if candidate.exists():
                    return ImageFont.truetype(str(candidate), size)
            except Exception:
                continue
        return ImageFont.load_default()

    title_font = load_font(28)
    body_font = load_font(20)
    raw_lines = [repair_utf8_mojibake(title or "复测结果").strip(), ""]
    for raw_line in repair_utf8_mojibake(text).splitlines():
        line = raw_line.rstrip()
        if not line:
            raw_lines.append("")
            continue
        while len(line) > max_chars:
            raw_lines.append(line[:max_chars])
            line = line[max_chars:]
        raw_lines.append(line)
    raw_lines = raw_lines[:120]
    height = max(360, padding * 2 + 44 + len(raw_lines) * line_height)
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 86), fill="#f2f6ff")
    draw.text((padding, 24), repair_utf8_mojibake(title or "复测结果"), fill="#1f2a44", font=title_font)
    y = 104
    for line in raw_lines[2:]:
        draw.text((padding, y), line, fill="#243044", font=body_font)
        y += line_height
    image.save(output_path)
    return output_path


def _save_retest_evidence_section_screenshots(target_dir: Path, sections: List[Dict[str, str]]) -> List[Dict[str, str]]:
    saved: List[Dict[str, str]] = []
    for index, section in enumerate(sections, 1):
        if not isinstance(section, dict):
            continue
        text = repair_utf8_mojibake(section.get("text") or "").strip()
        if not text:
            continue
        path = _save_retest_text_screenshot(target_dir, text, f"复测证据 {index}")
        saved.append({
            "caption": repair_utf8_mojibake(section.get("caption") or f"证据 {index}").strip(),
            "path": str(path),
        })
    return saved


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _notice_base_date(source_file: Path) -> datetime:
    text = _extract_docx_text(source_file)
    match = re.search(r"(20\d{2})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
    if match:
        try:
            return datetime(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        except ValueError:
            pass

    name_match = re.search(r"(20\d{2})(\d{2})(\d{2})", source_file.stem)
    if name_match:
        try:
            return datetime(int(name_match.group(1)), int(name_match.group(2)), int(name_match.group(3)))
        except ValueError:
            pass

    try:
        return datetime.fromtimestamp(source_file.stat().st_mtime)
    except Exception:
        return datetime.now()


def _disposal_report_date(source_file: Path) -> datetime:
    offset_days = 5 + (sum(ord(ch) for ch in source_file.name) % 3)
    return _notice_base_date(source_file) + timedelta(days=offset_days)


def _format_chinese_date(value: datetime) -> str:
    return f"{value.year}年{value.month}月{value.day}日"


def _infer_notice_company_name(source_file: Path, scan_result: Dict[str, Any] | None = None) -> str:
    candidates = [
        Path(str((scan_result or {}).get("file") or source_file)).stem,
        str((scan_result or {}).get("title") or ""),
        source_file.parent.name,
        source_file.stem,
    ]
    for candidate in candidates:
        try:
            from modules.Document_Processing.Report_Rewrite import group_folders as gf

            company = gf.normalize_company(candidate)
        except Exception:
            company = None
        if company:
            return company
    return source_file.parent.name or source_file.stem


def _normalize_retest_issue_type(value: Any) -> str:
    try:
        from modules.Document_Processing.Report_Rewrite.notice_name_utils import normalize_issue_type

        normalized = normalize_issue_type(value)
    except Exception:
        normalized = None
    return str(normalized or value or "漏洞隐患").strip()


def _replace_paragraph_text(paragraph: Any, new_text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _replace_text_preserve_first_run(paragraph: Any, old: str, new: str) -> bool:
    if old not in (paragraph.text or ""):
        return False
    _replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def _replace_disposal_template_image(doc: Any, screenshot_path: Path, logs: List[str]) -> bool:
    from docx.shared import Inches

    image_para = None
    for para in doc.paragraphs:
        for run in list(para.runs):
            has_image = bool(run._element.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))
            if has_image:
                if image_para is None:
                    image_para = para
                run._element.getparent().remove(run._element)

    target_para = image_para
    if target_para is None:
        for index, para in enumerate(doc.paragraphs):
            text = para.text or ""
            if "已修改" in text:
                target_para = para
                break
            if "整改措施" in text or "整改结果" in text:
                for next_para in doc.paragraphs[index + 1 : index + 5]:
                    if next_para.text.strip():
                        target_para = next_para
                        break
                if target_para:
                    break

    if target_para is None:
        target_para = doc.add_paragraph()
    elif target_para.text.strip():
        target_para.add_run().add_break()
    target_para.add_run().add_picture(str(screenshot_path), width=Inches(5.2))
    logs.append(f"处置文件证据截图已替换为复测报告证据: {screenshot_path.name}")
    return True


def _fill_disposal_report_document(
    output_path: Path,
    source_file: Path,
    scan_result: Dict[str, Any],
    screenshot_path: Path,
    logs: List[str],
) -> None:
    from docx import Document

    company_name = _infer_notice_company_name(source_file, scan_result)
    issue_type = _normalize_retest_issue_type(scan_result.get("vulnerability_type"))
    if issue_type.endswith("事件"):
        issue_term = "事件"
        issue_fixed_sentence = "该事件已完成处置"
    elif issue_type.endswith(("风险", "隐患", "安全问题", "安全隐患")):
        issue_term = "隐患"
        issue_fixed_sentence = "该隐患已完成整改"
    else:
        issue_term = "漏洞"
        issue_fixed_sentence = "该漏洞已进行修补"
    report_date = _format_chinese_date(_disposal_report_date(source_file))
    doc = Document(str(output_path))

    date_pattern = re.compile(r"20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日")
    old_issue_pattern = re.compile(r"(?:信息泄露漏洞|未知漏洞|漏洞隐患|网络安全漏洞)")

    for para in doc.paragraphs:
        text = para.text or ""
        if not text:
            continue
        if "*" in text:
            _replace_text_preserve_first_run(para, "*", company_name)
            text = para.text or ""
        if "所属漏洞" in text:
            _replace_text_preserve_first_run(para, "所属漏洞", f"所属{issue_type}")
            text = para.text or ""
        if "信息泄露漏洞信息" in text:
            _replace_text_preserve_first_run(para, "信息泄露漏洞", issue_type)
            text = para.text or ""
        elif old_issue_pattern.search(text):
            _replace_paragraph_text(para, old_issue_pattern.sub(issue_type, text))
            text = para.text or ""
        if "已经确立漏洞点" in text:
            _replace_text_preserve_first_run(para, "已经确立漏洞点", f"已经确立{issue_term}点")
            text = para.text or ""
        if "该漏洞已进行修补" in text:
            _replace_text_preserve_first_run(para, "该漏洞已进行修补", issue_fixed_sentence)
            text = para.text or ""
        if date_pattern.search(text):
            _replace_paragraph_text(para, date_pattern.sub(report_date, text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text or ""
                    if "*" in text:
                        _replace_text_preserve_first_run(para, "*", company_name)
                        text = para.text or ""
                    if old_issue_pattern.search(text):
                        _replace_paragraph_text(para, old_issue_pattern.sub(issue_type, text))
                        text = para.text or ""
                    if date_pattern.search(text):
                        _replace_paragraph_text(para, date_pattern.sub(report_date, text))

    _replace_disposal_template_image(doc, screenshot_path, logs)
    doc.save(str(output_path))
    logs.append(f"处置文件已填充: {output_path.name}，单位={company_name}，问题={issue_type}，日期={report_date}")


def _convert_single_word_to_pdf(word_path: Path, logs: List[str]) -> tuple[Path | None, str | None]:
    pdf_path = word_path.with_suffix(".pdf")
    try:
        from modules.Document_Processing.doc_pdf import convert_with_word_com

        _converted, _skipped, failures = convert_with_word_com([(word_path, pdf_path)], overwrite=True)
    except Exception as exc:
        return None, str(exc)
    if failures:
        return None, "; ".join(f"{src.name}: {reason}" for src, reason in failures)
    if not pdf_path.exists():
        return None, "PDF 文件未生成"
    logs.append(f"处置文件已转为PDF: {pdf_path.name}")
    return pdf_path, None


def _delete_disposal_word_after_pdf(word_path: Path, pdf_path: Path | None, logs: List[str]) -> tuple[bool, str]:
    if pdf_path is None or not pdf_path.exists():
        return False, ""
    if word_path.suffix.lower() not in {".doc", ".docx"}:
        return False, ""
    if not word_path.exists():
        return True, ""
    try:
        word_path.unlink()
        logs.append(f"处置文件Word版已删除: {word_path.name}")
        return True, ""
    except Exception as exc:
        message = str(exc)
        logs.append(f"处置文件Word版删除失败 {word_path.name}: {message}")
        return False, message


def _prepare_retest_disposal_report(
    source_file: Path,
    scan_result: Dict[str, Any],
    screenshot_path: Path | None,
    logs: List[str],
    output_dir: Path | None = None,
) -> Dict[str, Any] | None:
    if screenshot_path is None or not screenshot_path.exists():
        logs.append(f"跳过处置文件替换: 缺少复测证据截图 ({source_file.name})")
        return None

    source_dir = source_file.parent
    target_dir = output_dir or source_dir
    target_dir.mkdir(parents=True, exist_ok=True)
    existing_template = _find_existing_disposal_word_template(source_dir)
    template_path = _retest_disposal_template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"未找到漏洞隐患处置文件模板: {template_path}")

    output_path = _unique_retest_disposal_output_path(target_dir)

    shutil.copy2(template_path, output_path)
    if existing_template:
        logs.append(f"原处置类Word文件已保留，复测处置文件另存为: {output_path.name}")
    else:
        logs.append(f"未找到处置类Word模板，已生成处置文件: {output_path.name}")

    if output_path.suffix.lower() != ".docx":
        next_output = output_path.with_suffix(".docx")
        try:
            output_path.replace(next_output)
            output_path = next_output
        except Exception:
            pass

    _fill_disposal_report_document(output_path, source_file, scan_result, screenshot_path, logs)
    pdf_path, pdf_error = _convert_single_word_to_pdf(output_path, logs)
    logs.append(f"处置文件Word版已保留: {output_path.name}")
    return {
        "source": str(source_file),
        "word": str(output_path),
        "pdf": str(pdf_path) if pdf_path else "",
        "pdf_error": pdf_error or "",
        "word_deleted": False,
        "word_delete_error": "",
        "removed_template": "",
    }


def _retest_screenshot_dir(target_dir: Path) -> Path:
    return target_dir / ".koi_retest_screenshots"


def _cleanup_retest_screenshot_dir(target_dir: Path, logs: List[str]) -> None:
    screenshot_dir = _retest_screenshot_dir(target_dir)
    if not screenshot_dir.exists():
        return

    try:
        resolved_dir = screenshot_dir.resolve()
        resolved_target = target_dir.resolve()
        if resolved_dir.name != ".koi_retest_screenshots" or resolved_dir.parent != resolved_target:
            logs.append(f"跳过异常复测截图目录清理: {resolved_dir}")
            return
        shutil.rmtree(resolved_dir)
        logs.append(f"临时复测截图目录已删除: {resolved_dir}")
    except Exception as exc:
        logs.append(f"临时复测截图目录删除失败: {screenshot_dir} -> {exc}")


def _cleanup_retest_staging_dir(staging_dir: Path, logs: List[str] | None = None) -> None:
    """Remove one agent-owned report staging directory after a turn."""
    if not staging_dir.exists():
        return
    try:
        resolved = staging_dir.resolve()
        if resolved.name in {"", ".", ".."} or resolved.parent.name != ".koi_retest_staging":
            return
        shutil.rmtree(resolved)
        try:
            resolved.parent.rmdir()
        except OSError:
            pass
        if logs is not None:
            logs.append(f"临时复测报告暂存目录已删除: {resolved}")
    except Exception as exc:
        if logs is not None:
            logs.append(f"临时复测报告暂存目录删除失败: {staging_dir} -> {exc}")


def _doc_retest_list_files(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_dir = Path(_required_text(payload, "target_dir", "请选择通报目录")).expanduser()
    if not target_dir.exists() or not target_dir.is_dir():
        return {"success": False, "message": f"通报目录不存在: {target_dir}", "logs": []}

    try:
        from modules.AI_Testing.retest.word_vulnerability_scanner import WordVulnerabilityScanner
    except Exception as exc:
        return {"success": False, "message": f"导入通报扫描器失败: {exc}", "logs": [traceback.format_exc()]}

    logs: List[str] = []
    scanner = WordVulnerabilityScanner(str(target_dir))
    buffer = io.StringIO()
    with contextlib.redirect_stdout(buffer):
        word_files = scanner.find_word_files()
    logs.extend(_captured_lines(buffer))
    report_evidence = _existing_retest_report_evidence(target_dir, word_files)
    completed_names = [item["source_file_name"] for item in report_evidence]
    next_index_hint = len(completed_names)
    next_source_file = str(word_files[next_index_hint]) if 0 <= next_index_hint < len(word_files) else ""
    logs.append(
        f"扫描完成，发现 {len(word_files)} 份原始通报文档；"
        f"从同目录复测报告识别到 {len(completed_names)} 份已完成通报"
    )
    return {
        "success": True,
        "message": (
            f"发现 {len(word_files)} 份原始通报文档"
            + (f"，已从复测报告识别 {len(completed_names)} 份已完成" if completed_names else "")
        ),
        "target_dir": str(target_dir),
        "total": len(word_files),
        "source_files": [str(file_path) for file_path in word_files],
        "completed_source_files": [item["source_file"] for item in report_evidence],
        "completed_source_file_names": completed_names,
        "existing_report_evidence": report_evidence,
        "completed_count_hint": len(completed_names),
        "next_index_hint": next_index_hint,
        "next_source_file": next_source_file,
        "next_source_file_name": Path(next_source_file).name if next_source_file else "",
        "logs": logs,
    }


def _retest_trace_event(
    event_type: str,
    title: str,
    content: str = "",
    tone: str = "info",
    tool: Dict[str, Any] | None = None,
    source_file: str | None = None,
    metadata: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    event: Dict[str, Any] = {
        "id": f"trace-{uuid.uuid4().hex[:10]}",
        "type": event_type,
        "title": title,
        "content": content,
        "tone": tone,
        "timestamp": datetime.now().strftime("%H:%M:%S"),
    }
    if tool:
        event["tool"] = tool
    if source_file:
        event["sourceFile"] = source_file
    if metadata:
        event["metadata"] = metadata
    return event


def _retest_observation_count(result_item: Dict[str, Any]) -> int:
    if not isinstance(result_item, dict):
        return 0
    if result_item.get("observation_count") is not None:
        try:
            return max(0, int(result_item.get("observation_count") or 0))
        except Exception:
            return 0
    return len([item for item in (result_item.get("vulnerabilities") or []) if isinstance(item, dict)])


def _fast_vulnerability_supports_reproduced(vuln: Dict[str, Any]) -> bool:
    if not isinstance(vuln, dict):
        return False
    if vuln.get("tool_unavailable") or vuln.get("tool_failed"):
        return False
    severity = str(vuln.get("severity") or "").strip().lower()
    if severity in {"high", "critical"}:
        return True
    text = " ".join(
        str(vuln.get(key) or "")
        for key in ("type", "detail", "evidence")
    )
    positive_markers = (
        "登录成功",
        "弱口令登录成功",
        "可读",
        "任意文件读取",
        "目录遍历",
        "目录穿越",
        "敏感文件",
        "源码泄露",
        "配置泄露",
        "sql注入",
        "xss",
        "跨站脚本",
        "open redirect",
        "开放重定向",
        "trace method enabled",
    )
    negative_markers = ("未复现", "未见复现", "访问受限", "可能已修复", "仅作为", "信息")
    normalized = text.lower()
    if any(marker.lower() in normalized for marker in positive_markers):
        return severity in {"medium", "high", "critical"} or "成功" in text or "可读" in text
    if severity == "medium" and not any(marker in text for marker in negative_markers):
        return True
    return False


def _apply_fast_retest_judgement(result_data: Dict[str, Any], logs: List[str]) -> Dict[str, Any]:
    risk_observations: List[Dict[str, Any]] = []
    all_unreachable = False
    retest_results = [item for item in (result_data.get("retest_results") or []) if isinstance(item, dict)]
    if retest_results:
        all_unreachable = all(bool(item.get("target_unreachable")) for item in retest_results)
    for item in retest_results:
        for vuln in item.get("vulnerabilities") or []:
            if isinstance(vuln, dict) and _fast_vulnerability_supports_reproduced(vuln):
                risk_observations.append(vuln)

    reproduced = bool(risk_observations)
    verdict = "reproduced" if reproduced else "not_reproduced"
    if reproduced:
        conclusion = "快速判定：漏洞未修复/可复现"
        reason = f"快速规则命中 {len(risk_observations)} 条风险观察。"
    elif all_unreachable or result_data.get("target_unreachable"):
        conclusion = "快速判定：未复现，目标不可达，未能验证（建议复查）"
        reason = "快速规则未能访问目标，当前未形成可复现证据。"
    else:
        conclusion = "快速判定：未见可复现证据/复测通过"
        reason = "快速规则复测未命中可支撑漏洞仍存在的风险观察。"
    evidence = []
    for vuln in risk_observations[:8]:
        detail = str(vuln.get("detail") or vuln.get("evidence") or "").strip()
        evidence.append(f"{vuln.get('type') or '风险观察'}: {detail}" if detail else str(vuln.get("type") or "风险观察"))

    judgement = {
        "verdict": verdict,
        "reproduced": reproduced,
        "fix_status": "risk" if reproduced else "clean",
        "conclusion": conclusion,
        "reason": reason,
        "evidence": evidence,
        "source": "fast_rules",
        "unverified_unreachable": bool((not reproduced) and (all_unreachable or result_data.get("target_unreachable"))),
    }
    result_data["ai_judgement"] = judgement
    result_data["final_verdict"] = verdict
    result_data["ai_reproduced"] = reproduced
    result_data["risk_count"] = 1 if reproduced else 0
    result_data["manual_count"] = 0
    result_data["manual_test_required"] = False
    result_data["fast_mode"] = True
    result_data["reason"] = reason
    logs.append(conclusion)
    return result_data


def _apply_decisive_reproduction_judgement(result_data: Dict[str, Any], logs: List[str]) -> Dict[str, Any]:
    """Use execution-layer-verified direct evidence without a redundant judge call."""
    decisive_items: List[Dict[str, Any]] = []
    for url_result in result_data.get("retest_results") or []:
        if not isinstance(url_result, dict) or not url_result.get("decisive_reproduction"):
            continue
        evidence = url_result.get("decisive_evidence")
        if isinstance(evidence, dict):
            decisive_items.append(evidence)
    if not decisive_items:
        return result_data

    evidence_lines: List[str] = []
    for item in decisive_items[:8]:
        title = str(item.get("type") or item.get("title") or "原通报漏洞复现证据").strip()
        proof = str(item.get("evidence") or item.get("detail") or "").strip()
        evidence_lines.append(f"{title}: {proof}" if proof else title)
    judgement = {
        "verdict": "reproduced",
        "reproduced": True,
        "fix_status": "risk",
        "confidence": "high",
        "conclusion": "漏洞未修复/可复现",
        "reason": "执行层已取得直接对应原通报漏洞的真实阳性证据，按最小充分取证原则结束。",
        "evidence": evidence_lines,
        "source": "react_decisive_evidence",
    }
    result_data["ai_judgement"] = judgement
    result_data["final_verdict"] = "reproduced"
    result_data["ai_reproduced"] = True
    result_data["risk_count"] = 1
    result_data["manual_count"] = 0
    result_data["manual_test_required"] = False
    result_data["reason"] = judgement["reason"]
    result_data["decisive_reproduction"] = True
    logs.append("ReAct 直接阳性证据已充分，跳过重复 AI 判定。")
    return result_data


def _retest_ai_advice_trace(scan_result: Dict[str, Any], source_file: Path) -> Dict[str, Any] | None:
    context = scan_result.get("retest_context") if isinstance(scan_result.get("retest_context"), dict) else {}
    advice = context.get("agent_advice") if isinstance(context, dict) else {}
    if not isinstance(advice, dict) or not advice.get("enabled"):
        return None
    provider = str(advice.get("provider") or "")
    model = str(advice.get("model") or "")
    plan_steps = [str(item) for item in (advice.get("plan_steps") or []) if str(item).strip()]
    recommended = [str(item) for item in (advice.get("recommended_checks") or []) if str(item).strip()]
    warnings = [str(item) for item in (advice.get("warnings") or []) if str(item).strip()]
    lines = [
        f"模型: {provider}{('/' + model) if model else ''}".strip(),
        f"状态: {'已完成通报阅读与复测规划' if advice.get('used') else '规划未完成（已暂停等待继续）'}",
    ]
    if plan_steps:
        lines.append("计划:\n" + "\n".join(f"- {item}" for item in plan_steps[:8]))
    if recommended:
        lines.append("推荐工具: " + "、".join(recommended))
    python_probe = advice.get("python_probe") if isinstance(advice.get("python_probe"), dict) else {}
    if python_probe and python_probe.get("script"):
        lines.append("Python 探针: 已生成受限 HTTP 探针脚本")
        if python_probe.get("reason"):
            lines.append("脚本目的: " + str(python_probe.get("reason")))
    if advice.get("notes"):
        lines.append("说明: " + str(advice.get("notes")))
    if warnings:
        lines.append("提醒:\n" + "\n".join(f"- {item}" for item in warnings[:6]))
    if advice.get("reason") and not advice.get("used"):
        lines.append("原因: " + str(advice.get("reason")))
    if advice.get("error"):
        lines.append("错误: " + str(advice.get("error")))
    return _retest_trace_event(
        "thought_summary",
        "AI 规划摘要",
        "\n".join(line for line in lines if line),
        "ok" if advice.get("used") else "warn",
        source_file=str(source_file),
        metadata={"provider": provider, "model": model, "recommended_checks": recommended, "pythonProbe": bool(python_probe and python_probe.get("script")), "phase": "planning"},
    )


def _run_retest_for_source_file(
    file_path: Path,
    payload: Dict[str, Any],
    logs: List[str],
    event_callback: Callable[[Dict[str, Any]], None] | None = None,
    stop_check: Callable[[], bool] | None = None,
    confirm_callback: Callable[[Dict[str, Any]], Dict[str, Any]] | None = None,
    checkpoint_callback: Callable[[Dict[str, Any]], None] | None = None,
) -> tuple[str, Dict[str, Any], bool, List[Dict[str, Any]]]:
    from modules.AI_Testing.retest.vulnerability_batch_scanner import VulnerabilityRetestScanner
    from modules.AI_Testing.retest.word_vulnerability_scanner import WordVulnerabilityScanner

    trace_events: List[Dict[str, Any]] = []
    round_id = str(payload.get("round_id") or f"file:{file_path.name}")
    turn_id = str(payload.get("turn_id") or "")
    source_file_name = str(payload.get("source_file_name") or file_path.name)
    use_ai = _retest_payload_uses_ai(payload)

    def emit(event: Dict[str, Any]) -> Dict[str, Any]:
        if isinstance(event, dict):
            event.setdefault("sourceFile", str(file_path))
            metadata = event.get("metadata")
            if not isinstance(metadata, dict):
                metadata = {}
                event["metadata"] = metadata
            metadata.setdefault("roundId", round_id)
            if turn_id:
                metadata.setdefault("turnId", turn_id)
            metadata.setdefault("sourceFileName", source_file_name)
            metadata.setdefault("phase", "tool" if str(event.get("type") or "").startswith("tool_") else "status")
            trace_events.append(event)
            if event_callback:
                event_callback(event)
        return event

    def checkpoint(snapshot: Dict[str, Any] | None) -> None:
        if not checkpoint_callback or not isinstance(snapshot, dict) or not snapshot:
            return
        try:
            checkpoint_callback(snapshot)
        except Exception:
            pass

    def make_model_stream_callback(title: str, phase: str) -> tuple[Callable[[str], None], Callable[[], None]]:
        state = {"buffer": "", "visible": "", "last_emit": 0.0, "last_emit_count": 0, "count": 0}
        stream_key = f"model-output:{round_id}:{phase}"

        def visible_dialogue(raw: str) -> str:
            text = str(raw or "")
            cut_points = []
            for marker in ("```json", "```JSON", "JSON_RESULT:", "\nJSON:", "\n{"):
                index = text.find(marker)
                if index >= 0:
                    cut_points.append(index)
            if text.lstrip().startswith("{"):
                return "Agent 正在生成结构化复测数据，稍后会继续执行工具。"
            if cut_points:
                text = text[:min(cut_points)]
            text = re.sub(r"^\s*AGENT_MESSAGE\s*[:：]\s*", "", text, flags=re.IGNORECASE).strip()
            return text

        def callback(chunk: str) -> None:
            text = str(chunk or "")
            if not text:
                return
            state["buffer"] = (state["buffer"] + text)[-12000:]
            state["visible"] = visible_dialogue(str(state["buffer"]))
            state["count"] = int(state["count"]) + len(text)
            now = time.time()
            if now - float(state["last_emit"] or 0.0) < 0.25 and int(state["count"]) - int(state["last_emit_count"]) < 120:
                return
            preview = str(state["visible"] or "").strip()
            if not preview:
                return
            state["last_emit"] = now
            state["last_emit_count"] = state["count"]
            emit(_retest_trace_event(
                "thought_summary",
                title,
                preview,
                "info",
                source_file=str(file_path),
                metadata={"phase": phase, "streaming": True, "modelOutput": True, "dialogueOutput": True, "streamKey": stream_key},
            ))

        def flush() -> None:
            if not state["buffer"]:
                return
            final_text = str(state["visible"] or "").strip() or "Agent 已完成本阶段模型输出，正在进入下一步。"
            emit(_retest_trace_event(
                "thought_summary",
                title.replace("对话 /", "对话完成 /"),
                final_text,
                "ok",
                source_file=str(file_path),
                metadata={"phase": phase, "streaming": False, "modelOutput": True, "dialogueOutput": True, "completeModelOutput": True, "streamKey": stream_key},
            ))

        return callback, flush

    def finish_judgement_result(
        result_data: Dict[str, Any],
        judge_started: float = 0.0,
    ) -> tuple[str, Dict[str, Any], bool, List[Dict[str, Any]]]:
        ai_judgement = result_data.get("ai_judgement") if isinstance(result_data.get("ai_judgement"), dict) else {}
        fast_mode = bool(result_data.get("fast_mode") or ai_judgement.get("source") == "fast_rules")
        decisive_mode = ai_judgement.get("source") == "react_decisive_evidence"
        model_verdict = _model_verdict_from_result_data(result_data)
        reproduced = model_verdict == "reproduced"
        trace_tone = "warn" if reproduced else ("ok" if model_verdict == "not_reproduced" else "error")
        fix_status = "risk" if reproduced else ("clean" if model_verdict == "not_reproduced" else "unknown")
        evidence_level = "confirmed" if reproduced else ("not_reproduced" if model_verdict == "not_reproduced" else "unknown")
        judge_label = "直接证据判定" if decisive_mode else ("快速规则判定" if fast_mode else "AI 结论判定")
        judge_tool_id = "react_evidence_judge" if decisive_mode else ("fast_rule_judge" if fast_mode else "llm_judge")
        judge_args_preview = (
            "mode: decisive evidence\nmodel: skipped"
            if decisive_mode
            else ("mode: fast\nmodel: skipped" if fast_mode else f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}")
        )
        judge_duration_ms = 0 if fast_mode or decisive_mode or not judge_started else int((time.time() - judge_started) * 1000)
        judgement_preview = "\n".join([
            f"verdict: {model_verdict or '-'}",
            f"conclusion: {ai_judgement.get('conclusion') or '-'}",
            f"reason: {ai_judgement.get('reason') or result_data.get('reason') or '-'}",
        ])
        emit(_retest_trace_event(
            "tool_result",
            judge_label,
            judgement_preview,
            trace_tone,
            tool={
                "toolId": judge_tool_id,
                "label": judge_label,
                "status": "completed",
                "target": file_path.name,
                "argsPreview": judge_args_preview,
                "resultPreview": judgement_preview,
                "durationMs": judge_duration_ms,
            },
            source_file=str(file_path),
            metadata={
                "provider": ai_provider,
                "model": ai_model,
                "phase": "judgement",
                "mode": "decisive_evidence" if decisive_mode else ("fast" if fast_mode else "ai"),
                "evidenceLevel": evidence_level,
                "fixStatus": fix_status,
            },
        ))
        judgement_lines = [
            f"结论: {ai_judgement.get('conclusion') or ('漏洞未修复/可复现' if reproduced else '漏洞已修复/复测通过' if model_verdict == 'not_reproduced' else '模型未给出判定')}",
            f"理由: {ai_judgement.get('reason') or result_data.get('reason') or ('直接证据已充分' if decisive_mode else '快速规则未提供额外理由' if fast_mode else 'AI 未提供额外理由')}",
        ]
        evidence = ai_judgement.get("evidence") if isinstance(ai_judgement.get("evidence"), list) else []
        if evidence:
            judgement_lines.append("证据:\n" + "\n".join(f"- {item}" for item in evidence[:8]))
        emit(_retest_trace_event(
            "thought_summary",
            "直接证据判定摘要" if decisive_mode else ("快速判定摘要" if fast_mode else "AI 判定摘要"),
            "\n".join(judgement_lines),
            trace_tone,
            source_file=str(file_path),
            metadata={
                "provider": ai_provider,
                "model": ai_model,
                "phase": "judgement",
                "mode": "decisive_evidence" if decisive_mode else ("fast" if fast_mode else "ai"),
                "fixStatus": fix_status,
                "evidenceLevel": evidence_level,
            },
        ))

        summary = _format_retest_summary(file_path, result_data)
        checkpoint(_result_resume_snapshot(str(file_path), summary, result_data))
        return summary, result_data, False, trace_events

    ai_config_for_trace: Dict[str, Any] = {}
    try:
        ai_config_for_trace = _load_retest_ai_config()
    except Exception:
        ai_config_for_trace = {}
    ai_provider = str(ai_config_for_trace.get("provider") or "")
    ai_model = str(ai_config_for_trace.get("model") or "")
    ai_enabled = bool(ai_config_for_trace.get("enabled"))

    resume_snapshot = payload.get("resume_snapshot") if isinstance(payload.get("resume_snapshot"), dict) else {}
    resume_stage = str(resume_snapshot.get("stage") or resume_snapshot.get("resume_stage") or "").strip().lower() if isinstance(resume_snapshot, dict) else ""
    resume_source = str(resume_snapshot.get("source_file") or "") if isinstance(resume_snapshot, dict) else ""
    resume_matches_source = not resume_source or _retest_target_key(resume_source) == _retest_target_key(str(file_path))
    resume_scan_result = resume_snapshot.get("scan_result") if isinstance(resume_snapshot.get("scan_result"), dict) else {}
    resume_result_data = resume_snapshot.get("result_data") if isinstance(resume_snapshot.get("result_data"), dict) else {}
    if resume_stage in {"result", "completed", "judgement_complete", "report", "report_generation"} and resume_matches_source and resume_result_data:
        result_data = dict(resume_result_data)
        result_data["file"] = str(file_path)
        summary = str(resume_snapshot.get("summary") or "").strip() or _format_retest_summary(file_path, result_data)
        logs.append(f"{file_path.name} resumed from {resume_stage or 'result'} snapshot; retest and judgement were skipped")
        emit(_retest_trace_event(
            "status",
            "Result checkpoint restored",
            "Recovered completed retest judgement without rerunning probes.",
            "ok",
            source_file=str(file_path),
            metadata={"phase": resume_stage or "result", "resumeStage": resume_stage or "result", "resumedFromSnapshot": True},
        ))
        return summary, result_data, False, trace_events
    if use_ai and resume_stage == "judgement" and resume_matches_source and resume_scan_result and resume_result_data:
        scan_result = dict(resume_scan_result)
        result_data = dict(resume_result_data)
        result_data["file"] = str(file_path)
        logs.append(f"{file_path.name} 从判定断点继续：复用已完成工具观察，仅重新调用 AI 结论判定")
        emit(_retest_trace_event(
            "status",
            "判定断点恢复",
            "已恢复上次中断前的通报解析和工具观察，本次继续只重新调用 AI 结论判定，不重复执行取证工具。",
            "ok",
            source_file=str(file_path),
            metadata={"phase": "judgement", "resumeStage": "judgement", "resumedFromSnapshot": True},
        ))
        judge_started = time.time()
        emit(_retest_trace_event(
            "tool_call",
            "AI 结论判定",
            "从判定断点继续：复用已完成工具输出并请求模型给出最终复测结论。",
            "info",
            tool={
                "toolId": "llm_judge",
                "label": "AI 结论判定",
                "status": "running",
                "target": file_path.name,
                "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}\nresume: judgement",
            },
            source_file=str(file_path),
            metadata={"provider": ai_provider, "model": ai_model, "phase": "judgement", "resumedFromSnapshot": True},
        ))
        judge_stream_callback, flush_judge_stream = make_model_stream_callback("AI Agent 对话 / 判定", "judgement")
        try:
            result_data = _apply_retest_ai_judgement(
                scan_result,
                result_data,
                logs,
                stream_callback=judge_stream_callback,
            )
            flush_judge_stream()
        except RetestAIBlockedError as exc:
            flush_judge_stream()
            if not exc.resume_snapshot:
                exc.resume_snapshot = _judgement_resume_snapshot(str(file_path), scan_result, result_data)
            blocked_preview = f"blocked_stage: {exc.stage}\nreason: {exc}"
            emit(_retest_trace_event(
                "tool_result",
                "AI 结论判定",
                blocked_preview,
                "error",
                tool={
                    "toolId": "llm_judge",
                    "label": "AI 结论判定",
                    "status": "blocked",
                    "target": file_path.name,
                    "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}\nresume: judgement",
                    "resultPreview": blocked_preview,
                    "durationMs": int((time.time() - judge_started) * 1000),
                    "failureReason": str(exc),
                },
                source_file=str(file_path),
                metadata={"provider": ai_provider, "model": ai_model, "phase": exc.stage, "blockedByAiConfig": True, "resumedFromSnapshot": True},
            ))
            raise
        return finish_judgement_result(result_data, judge_started)

    retest_scanner = VulnerabilityRetestScanner(
        timeout=int(payload.get("timeout") or 15),
        max_workers=int(payload.get("max_workers") or 5),
        trace_callback=emit,
        ai_config=ai_config_for_trace,
        stop_check=stop_check,
        confirm_callback=confirm_callback,
    )

    def execute_retest_pipeline(
        scan_result: Dict[str, Any],
        restored_valid_urls: List[str] | None = None,
        restored_retest_results: List[Dict[str, Any]] | None = None,
        start_url_index: int = 0,
        resumed_execution: bool = False,
        restored_context_supported: bool | None = None,
    ) -> tuple[str, Dict[str, Any], bool, List[Dict[str, Any]]]:
        vuln_types = scan_result.get("vulnerability_types") or []
        retest_context = scan_result.get("retest_context") or {}
        url_candidates = scan_result.get("urls") or retest_context.get("target_urls") or []
        if restored_valid_urls is not None:
            url_candidates = restored_valid_urls
        valid_urls = _dedupe_http_targets(url_candidates)
        context_supported = (
            bool(restored_context_supported)
            if restored_context_supported is not None
            else retest_scanner.context_has_retestable_signals(retest_context)
        )
        scan_result["context_supported"] = context_supported
        sanitized_scan_result = _sanitize_retest_scan_result(scan_result)

        if not valid_urls:
            reason_parts = []
            if not vuln_types:
                reason_parts.append("no vulnerability type identified")
            reason_parts.append("no usable URL extracted")
            result_data: Dict[str, Any] = {
                "file": str(file_path),
                "urls": valid_urls,
                "retest_results": [],
                "risk_count": 0,
                "manual_count": 0,
                "failed_count": 0,
                "scan_result": sanitized_scan_result,
                "manual_test_required": False,
                "reason": "; ".join(reason_parts) + "; no reproducible evidence",
                "context_supported": context_supported,
            }
            logs.append(f"{file_path.name} no retestable target found: {result_data['reason']}")
            emit(_retest_trace_event(
                "status",
                "No retestable target",
                result_data["reason"],
                "ok",
                source_file=str(file_path),
                metadata={"phase": "result", "evidenceLevel": "empty"},
            ))
        else:
            retest_results = [item for item in (restored_retest_results or []) if isinstance(item, dict)]
            restored_decisive = any(
                bool(item.get("decisive_reproduction"))
                for item in retest_results
                if isinstance(item, dict)
            )
            next_index = max(0, min(int(start_url_index or 0), len(valid_urls)))
            if restored_decisive:
                # A previous URL already proved the reported vulnerability. The
                # remaining targets were intentionally skipped and must stay skipped.
                next_index = len(valid_urls)
            elif len(retest_results) < next_index:
                next_index = len(retest_results)
            elif len(retest_results) > next_index:
                retest_results = retest_results[:next_index]

            last_execution_snapshot = _execution_resume_snapshot(
                str(file_path),
                scan_result,
                valid_urls,
                retest_results,
                next_index,
                use_ai,
                context_supported,
            )
            checkpoint(last_execution_snapshot)
            if resumed_execution:
                logs.append(f"{file_path.name} resumed execution checkpoint at URL {next_index + 1}/{len(valid_urls)}")
                emit(_retest_trace_event(
                    "status",
                    "Execution checkpoint restored",
                    f"Reused {len(retest_results)} completed URL result(s); continuing from URL {next_index + 1} of {len(valid_urls)}.",
                    "ok",
                    source_file=str(file_path),
                    metadata={
                        "phase": "execution",
                        "resumeStage": "execution",
                        "resumedFromSnapshot": True,
                        "completedUrlCount": len(retest_results),
                        "nextUrlIndex": next_index,
                        "totalUrlCount": len(valid_urls),
                    },
                ))

            try:
                for url_index in range(next_index, len(valid_urls)):
                    if callable(stop_check) and stop_check():
                        return (
                            "复测已停止，可继续",
                            {
                                "file": str(file_path),
                                "urls": valid_urls,
                                "retest_results": retest_results,
                                "scan_result": sanitized_scan_result,
                                "stopped": True,
                                "manual_test_required": False,
                            },
                            False,
                            trace_events,
                        )
                    url = valid_urls[url_index]
                    result = retest_scanner.scan_url_for_context(url, vuln_types, retest_context) if use_ai else retest_scanner.scan_url_fast_for_context(url, vuln_types, retest_context)
                    stopped_result = bool(isinstance(result, dict) and result.get("stopped"))
                    stop_requested = bool(callable(stop_check) and stop_check())
                    if stopped_result or stop_requested:
                        if use_ai and isinstance(result, dict) and result.get("probe_repair_pending"):
                            repair_resume = result.get("probe_repair_resume")
                            if isinstance(repair_resume, dict) and repair_resume:
                                retest_context["probe_repair_resume"] = {
                                    **repair_resume,
                                    "target_url": url,
                                }
                                scan_result["retest_context"] = retest_context
                                last_execution_snapshot = _execution_resume_snapshot(
                                    str(file_path),
                                    scan_result,
                                    valid_urls,
                                    retest_results,
                                    url_index,
                                    use_ai,
                                    context_supported,
                                )
                                checkpoint(last_execution_snapshot)
                        return (
                            "复测已停止，可继续",
                            {
                                "file": str(file_path),
                                "urls": valid_urls,
                                "retest_results": retest_results,
                                "scan_result": _sanitize_retest_scan_result(scan_result),
                                "resume_snapshot": last_execution_snapshot,
                                "stopped": True,
                                "manual_test_required": False,
                            },
                            False,
                            trace_events,
                        )
                    if use_ai and isinstance(result, dict) and result.get("probe_repair_paused"):
                        repair_resume = result.get("probe_repair_resume")
                        if not isinstance(repair_resume, dict):
                            repair_resume = {}
                        retest_context["probe_repair_resume"] = {
                            **repair_resume,
                            "target_url": url,
                        }
                        scan_result["retest_context"] = retest_context
                        last_execution_snapshot = _execution_resume_snapshot(
                            str(file_path),
                            scan_result,
                            valid_urls,
                            retest_results,
                            url_index,
                            use_ai,
                            context_supported,
                        )
                        checkpoint(last_execution_snapshot)
                        raise RetestAIBlockedError(
                            "Python 探针已完成多轮实质重写但仍未成功；已保留当前 URL 和错误断点。"
                            "点击继续后会从脚本修复恢复，不会重复已完成 URL，也不会把脚本失败判成漏洞已修复。",
                            "probe_repair",
                            last_execution_snapshot,
                        )
                    retest_context.pop("probe_repair_resume", None)
                    retest_results.append(result)
                    decisive = bool(isinstance(result, dict) and result.get("decisive_reproduction"))
                    checkpoint_index = len(valid_urls) if decisive else url_index + 1
                    last_execution_snapshot = _execution_resume_snapshot(
                        str(file_path),
                        scan_result,
                        valid_urls,
                        retest_results,
                        checkpoint_index,
                        use_ai,
                        context_supported,
                    )
                    checkpoint(last_execution_snapshot)
                    if decisive:
                        logs.append(f"{file_path.name} URL {url_index + 1}/{len(valid_urls)} 已取得直接阳性证据，跳过剩余 URL。")
                        break
            except RetestAIBlockedError:
                raise
            except Exception as exc:
                message = str(exc)
                if _is_ai_runtime_block_message(message):
                    raise RetestAIBlockedError(f"AI Agent execution stage paused: {message}", "execution", last_execution_snapshot) from exc
                raise

            observation_count = sum(_retest_observation_count(item) for item in retest_results)
            decisive_observed = any(
                bool(item.get("decisive_reproduction"))
                for item in retest_results
                if isinstance(item, dict)
            )
            judge_name = "direct ReAct evidence" if decisive_observed else ("AI" if use_ai else "fast rules")
            logs.append(f"{file_path.name} tool observations complete: {observation_count}; final verdict by {judge_name}")
            emit(_retest_trace_event(
                "status",
                "Tool observations complete",
                (
                    f"Recorded {observation_count} observation(s); direct evidence is sufficient, finalizing without another model call."
                    if decisive_observed
                    else f"Recorded {observation_count} observation(s); handing evidence to {judge_name} for final verdict."
                ),
                "info" if observation_count else "ok",
                source_file=str(file_path),
                metadata={"observation_count": observation_count, "phase": "result", "evidenceLevel": "observation" if observation_count else "empty"},
            ))
            all_targets_unreachable = bool(retest_results) and all(
                bool(item.get("target_unreachable") or (isinstance(item.get("request_meta"), dict) and item.get("request_meta", {}).get("error")))
                for item in retest_results
                if isinstance(item, dict)
            )
            result_data = {
                "file": str(file_path),
                "urls": valid_urls,
                "retest_results": retest_results,
                "observation_count": observation_count,
                "risk_count": 0,
                "manual_count": 0,
                "failed_count": sum(int(item.get("failed_count") or 0) for item in retest_results if isinstance(item, dict)),
                "scan_result": sanitized_scan_result,
                "manual_test_required": False,
                "target_unreachable": all_targets_unreachable,
                "reason": "",
                "context_supported": context_supported,
            }

        judge_started = 0.0
        if use_ai and any(
            bool(item.get("decisive_reproduction"))
            for item in (result_data.get("retest_results") or [])
            if isinstance(item, dict)
        ):
            result_data = _apply_decisive_reproduction_judgement(result_data, logs)
        elif use_ai:
            judge_started = time.time()
            emit(_retest_trace_event(
                "tool_call",
                "AI judgement",
                "Ask the model to read tool evidence and produce the final retest verdict.",
                "info",
                tool={
                    "toolId": "llm_judge",
                    "label": "AI judgement",
                    "status": "running",
                    "target": file_path.name,
                    "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}",
                },
                source_file=str(file_path),
                metadata={"provider": ai_provider, "model": ai_model, "phase": "judgement"},
            ))
            judge_stream_callback, flush_judge_stream = make_model_stream_callback("AI Agent 对话 / 判定", "judgement")
            try:
                result_data = _apply_retest_ai_judgement(
                    scan_result,
                    result_data,
                    logs,
                    stream_callback=judge_stream_callback,
                )
                flush_judge_stream()
            except RetestAIBlockedError as exc:
                flush_judge_stream()
                blocked_preview = f"blocked_stage: {exc.stage}\nreason: {exc}"
                emit(_retest_trace_event(
                    "tool_result",
                    "AI 结论判定",
                    blocked_preview,
                    "error",
                    tool={
                        "toolId": "llm_judge",
                        "label": "AI 结论判定",
                        "status": "blocked",
                        "target": file_path.name,
                        "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}",
                        "resultPreview": blocked_preview,
                        "durationMs": int((time.time() - judge_started) * 1000),
                        "failureReason": str(exc),
                    },
                    source_file=str(file_path),
                    metadata={"provider": ai_provider, "model": ai_model, "phase": exc.stage, "blockedByAiConfig": True},
                ))
                raise
        else:
            result_data = _apply_fast_retest_judgement(result_data, logs)
        return finish_judgement_result(result_data, judge_started if use_ai else 0.0)

    execution_resume = resume_stage in {"execution", "verification", "tool"} and resume_matches_source and resume_scan_result
    payload_explicit_ai_mode = (
        "use_ai" in payload
        or str(payload.get("mode") or "").strip().lower() in {"fast", "quick", "legacy", "local", "ai", "agent"}
    )
    if execution_resume and "use_ai" in resume_snapshot and not payload_explicit_ai_mode:
        use_ai = bool(resume_snapshot.get("use_ai"))
    if execution_resume:
        restored_urls = [str(item) for item in (resume_snapshot.get("valid_urls") or []) if str(item).strip()]
        restored_results = [item for item in (resume_snapshot.get("retest_results") or []) if isinstance(item, dict)]
        start_url_index = _as_int(resume_snapshot.get("next_url_index"), len(restored_results))
        return execute_retest_pipeline(
            dict(resume_scan_result),
            restored_valid_urls=restored_urls or None,
            restored_retest_results=restored_results,
            start_url_index=start_url_index,
            resumed_execution=True,
            restored_context_supported=bool(resume_snapshot.get("context_supported")),
        )

    scanner = WordVulnerabilityScanner(str(file_path.parent))
    emit(_retest_trace_event("status", "文档解析", f"开始解析通报文档: {file_path.name}", "info", source_file=str(file_path), metadata={"phase": "parse"}))

    buffer = io.StringIO()
    with contextlib.redirect_stdout(buffer):
        scan_result = scanner.scan_document(file_path)
    logs.extend(_captured_lines(buffer))
    frontend_memory = _frontend_context_memory_text(payload.get("frontend_context"), 16000)
    if frontend_memory:
        retest_context = scan_result.get("retest_context")
        if not isinstance(retest_context, dict):
            retest_context = {}
            scan_result["retest_context"] = retest_context
        retest_context["session_recovery_memory"] = frontend_memory
        logs.append(f"恢复会话记忆已载入: {len(frontend_memory)} chars")
        emit(_retest_trace_event(
            "status",
            "恢复会话记忆已载入",
            "已从前端持久化上下文载入 AI 语义压缩记忆，后续规划/判定会优先使用。",
            "ok",
            source_file=str(file_path),
            metadata={"phase": "frontend_context_restore", "memoryChars": len(frontend_memory)},
        ))
    emit(_retest_trace_event(
        "status",
        "文档解析完成",
        "\n".join([
            f"漏洞类型: {'、'.join(scan_result.get('vulnerability_types') or []) or '未识别'}",
            f"URL: {'、'.join(scan_result.get('urls') or []) or '未提取'}",
        ]),
        "ok",
        source_file=str(file_path),
        metadata={"phase": "parse"},
    ))
    ai_provider = str(ai_config_for_trace.get("provider") or "")
    ai_model = str(ai_config_for_trace.get("model") or "")
    ai_enabled = bool(ai_config_for_trace.get("enabled"))
    # ReAct receives the complete parsed notice and the initial HTTP response.
    # Avoid a second planning request when parsing already found a usable URL;
    # reserve the standalone planner for URL-less/ambiguous notices.
    planning_needed = use_ai and not _dedupe_http_targets(
        scan_result.get("urls") or (scan_result.get("retest_context") or {}).get("target_urls") or []
    )
    if planning_needed:
        ai_started = time.time()
        emit(_retest_trace_event(
            "tool_call",
            "AI 规划模型",
            "调用模型分析通报上下文、纠正识别结果并推荐复测工具。" if ai_enabled else "AI 未就绪，本次会话将中断并等待配置后继续。",
            "info" if ai_enabled else "warn",
            tool={
                "toolId": "llm_plan",
                "label": "AI 规划模型",
                "status": "running" if ai_enabled else "blocked",
                "target": file_path.name,
                "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}",
            },
            source_file=str(file_path),
            metadata={"provider": ai_provider, "model": ai_model, "enabled": ai_enabled, "phase": "planning"},
        ))
        plan_stream_callback, flush_plan_stream = make_model_stream_callback("AI Agent 对话 / 规划", "planning")
        try:
            scan_result = _apply_retest_ai_agent(
                scan_result,
                logs,
                stream_callback=plan_stream_callback,
            )
            flush_plan_stream()
        except RetestAIBlockedError as exc:
            flush_plan_stream()
            blocked_preview = f"blocked_stage: {exc.stage}\nreason: {exc}"
            emit(_retest_trace_event(
                "tool_result",
                "AI 规划模型",
                blocked_preview,
                "error",
                tool={
                    "toolId": "llm_plan",
                    "label": "AI 规划模型",
                    "status": "blocked",
                    "target": file_path.name,
                    "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}",
                    "resultPreview": blocked_preview,
                    "durationMs": int((time.time() - ai_started) * 1000),
                    "failureReason": str(exc),
                },
                source_file=str(file_path),
                metadata={"provider": ai_provider, "model": ai_model, "phase": exc.stage, "blockedByAiConfig": True},
            ))
            raise
        ai_trace = _retest_ai_advice_trace(scan_result, file_path)
        ai_advice = (scan_result.get("retest_context") or {}).get("agent_advice") if isinstance(scan_result.get("retest_context"), dict) else {}
        ai_used = bool(isinstance(ai_advice, dict) and ai_advice.get("used"))
        ai_error = str(ai_advice.get("error") or "") if isinstance(ai_advice, dict) else ""
        ai_reason = str(ai_advice.get("reason") or ai_advice.get("notes") or "") if isinstance(ai_advice, dict) else ""
        ai_recommended = ai_advice.get("recommended_checks") if isinstance(ai_advice, dict) else []
        ai_result_preview = "\n".join([
            f"used: {ai_used}",
            f"recommended: {', '.join(str(item) for item in (ai_recommended or [])[:12]) or '-'}",
            f"reason: {ai_error or ai_reason or '-'}",
        ])
        emit(_retest_trace_event(
            "tool_result",
            "AI 规划模型",
            ai_result_preview,
            "ok" if ai_used else ("error" if ai_error else "warn"),
            tool={
                "toolId": "llm_plan",
                "label": "AI 规划模型",
                "status": "completed" if ai_used else ("failed" if ai_error else "skipped"),
                "target": file_path.name,
                "argsPreview": f"provider: {ai_provider or '-'}\nmodel: {ai_model or '-'}",
                "resultPreview": ai_result_preview,
                "durationMs": int((time.time() - ai_started) * 1000),
            },
            source_file=str(file_path),
            metadata={"provider": ai_provider, "model": ai_model, "used": ai_used, "phase": "planning", "evidenceLevel": "ai_used" if ai_used else "ai_required"},
        ))
        if ai_trace:
            emit(ai_trace)
    elif not use_ai:
        emit(_retest_trace_event(
            "status",
            "快速复测模式",
            "本轮不调用 AI 模型，使用本地规则和通报正文线索快速复测并生成报告。",
            "info",
            source_file=str(file_path),
            metadata={"phase": "planning", "mode": "fast", "aiSkipped": True},
        ))
    else:
        emit(_retest_trace_event(
            "status",
            "跳过重复规划",
            "通报已提取到有效 URL，ReAct 将直接基于完整通报上下文制定最小复测计划，避免重复调用规划模型。",
            "ok",
            source_file=str(file_path),
            metadata={"phase": "planning", "mode": "react_inline", "planningSkipped": True},
        ))

    return execute_retest_pipeline(scan_result)

def _doc_retest_run_one(payload: Dict[str, Any], progress: RetestTaskProgress | None = None) -> Dict[str, Any]:
    source_file = Path(_required_text(payload, "source_file", "请选择通报文件")).expanduser()
    if not source_file.exists() or source_file.suffix.lower() not in WORD_SUFFIXES:
        return {"success": False, "message": f"通报文件不存在或不是 Word 文档: {source_file}", "logs": []}

    logs: List[str] = ProgressLogList(progress) if progress else []
    logs.append(f"开始复测: {source_file.name}")
    try:
        if progress:
            progress.set(8, f"开始复测: {source_file.name}")
        summary, result_data, _manual_required, trace_events = _run_retest_for_source_file(
            source_file,
            payload,
            logs,
            event_callback=progress.event if progress else None,
            stop_check=progress.should_stop if progress else None,
            confirm_callback=(lambda req: _retest_request_confirmation(progress, req)) if progress else None,
            checkpoint_callback=progress.checkpoint if progress else None,
        )
        if bool(result_data.get("stopped")) or (progress is not None and progress.should_stop()):
            return {
                "success": False,
                "stopped": True,
                "message": "复测已停止，可继续",
                "source_file": str(source_file),
                "manual_test_required": False,
                "summary": summary,
                "result_data": result_data,
                "resume_snapshot": result_data.get("resume_snapshot") or (progress.snapshot().get("resume_snapshot") if progress else None),
                "trace_events": trace_events,
                "logs": logs,
            }
        if progress:
            progress.set(92, f"复测完成: {source_file.name}")
    except RetestAIBlockedError as exc:
        blocked_title = _ai_blocked_title(exc)
        if progress:
            progress.set(progress.snapshot().get("progress", 0), str(exc))
            progress.event(_retest_trace_event(
                "status",
                blocked_title,
                str(exc),
                "warn",
                source_file=str(source_file),
                metadata={"phase": exc.stage, "blockedByAiConfig": True},
            ))
            snapshot = progress.snapshot()
            trace_events = snapshot.get("trace_events") or []
        else:
            trace_events = [_retest_trace_event(
                "status",
                blocked_title,
                str(exc),
                "warn",
                source_file=str(source_file),
                metadata={"phase": exc.stage, "blockedByAiConfig": True},
            )]
        logs.append(str(exc))
        return _ai_blocked_payload(exc, str(source_file), logs, trace_events)
    except Exception as exc:
        logs.append(traceback.format_exc())
        return {"success": False, "message": f"复测失败: {exc}", "source_file": str(source_file), "logs": logs}

    return {
        "success": True,
        "message": f"复测完成: {source_file.name}",
        "source_file": str(source_file),
        "manual_test_required": False,
        "summary": summary,
        "result_data": result_data,
        "resume_snapshot": progress.snapshot().get("resume_snapshot") if progress else None,
        "trace_events": trace_events,
        "logs": logs,
    }


def _retest_task_worker(task_id: str, payload: Dict[str, Any]) -> None:
    with _RETEST_TASK_LOCK:
        task = _RETEST_TASKS.get(task_id)
    if not task:
        return

    progress = task["progress"]
    try:
        result = _doc_retest_run_one(payload, progress=progress)
        with _RETEST_TASK_LOCK:
            if task.get("stopped") or progress.should_stop() or bool(result.get("stopped")):
                task.update({
                    "running": False,
                    "done": True,
                    "success": False,
                    "stopped": True,
                    "message": "复测已停止，可继续",
                    "result": task.get("result") or {
                        "success": False,
                        "stopped": True,
                        "message": "复测已停止，可继续",
                        "logs": progress.snapshot().get("logs", []),
                        "trace_events": progress.snapshot().get("trace_events", []),
                        "resume_snapshot": progress.snapshot().get("resume_snapshot") or None,
                        "source_file": str(payload.get("source_file") or ""),
                    },
                    "finished_at": time.time(),
                })
                return
            task.update({
                "running": False,
                "done": True,
                "success": bool(result.get("success")),
                "message": result.get("message") or ("复测完成" if result.get("success") else "复测失败"),
                "result": result,
                "finished_at": time.time(),
            })
        progress.set(100 if result.get("success") else progress.snapshot().get("progress", 0), result.get("message") or "")
    except Exception as exc:
        progress.log(_format_exception(exc))
        with _RETEST_TASK_LOCK:
            stopped = bool(task.get("stopped"))
        progress.set(progress.snapshot().get("progress", 0) if stopped else 0, "复测已停止，可继续" if stopped else f"复测失败: {exc}")
        result = {
            "success": False,
            "stopped": stopped,
            "message": "复测已停止，可继续" if stopped else f"复测失败: {exc}",
            "logs": progress.snapshot().get("logs", []),
            "trace_events": progress.snapshot().get("trace_events", []),
            "source_file": str(payload.get("source_file") or ""),
        }
        with _RETEST_TASK_LOCK:
            task.update({
                "running": False,
                "done": True,
                "success": False,
                "message": result["message"],
                "result": result,
                "error": "" if stopped else str(exc),
                "finished_at": time.time(),
            })


def _doc_retest_event_stream_info(payload: Dict[str, Any]) -> Dict[str, Any]:
    from modules.backend_api.retest_event_stream import ensure_retest_event_stream

    return ensure_retest_event_stream()


def _doc_retest_run_one_start(payload: Dict[str, Any]) -> Dict[str, Any]:
    source_file = Path(_required_text(payload, "source_file", "请选择通报文件")).expanduser()
    if not source_file.exists() or source_file.suffix.lower() not in WORD_SUFFIXES:
        return {"success": False, "message": f"通报文件不存在或不是 Word 文档: {source_file}", "logs": [], "trace_events": []}

    if _retest_payload_uses_ai(payload):
        try:
            _ensure_retest_ai_ready("config")
        except RetestAIBlockedError as exc:
            resume_snapshot = payload.get("resume_snapshot") if isinstance(payload.get("resume_snapshot"), dict) else {}
            if resume_snapshot and not exc.resume_snapshot:
                exc.resume_snapshot = dict(resume_snapshot)
            blocked_title = _ai_blocked_title(exc)
            logs = [str(exc)]
            trace_events = [_retest_trace_event(
                "status",
                blocked_title,
                str(exc),
                "warn",
                source_file=str(source_file),
                metadata={
                    "roundId": str(payload.get("round_id") or f"file:{source_file.name}"),
                    "sourceFileName": str(payload.get("source_file_name") or source_file.name),
                    "phase": "config",
                    "blockedByAiConfig": True,
                },
            )]
            return _ai_blocked_payload(exc, str(source_file), logs, trace_events)

    task_id = uuid.uuid4().hex
    session_id = str(payload.get("session_id") or "")
    progress = RetestTaskProgress(total=1)
    progress.session_id = session_id
    progress.task_id = task_id
    progress.cancel_epoch_ns = time.time_ns()
    progress.set(1, f"任务已创建: {source_file.name}")
    progress.event(_retest_trace_event(
        "status",
        "复测任务启动",
        f"后台任务已创建，准备复测: {source_file.name}",
        "info",
        source_file=str(source_file),
        metadata={"roundId": str(payload.get("round_id") or f"file:{source_file.name}"), "sourceFileName": str(payload.get("source_file_name") or source_file.name), "phase": "start"},
    ))
    task = {
        "task_id": task_id,
        "session_id": session_id,
        "running": True,
        "done": False,
        "success": False,
        "message": f"任务已创建: {source_file.name}",
        "progress": progress,
        "result": None,
        "created_at": time.time(),
        "finished_at": None,
        "error": None,
        "stopped": False,
        "source_file": str(source_file),
    }
    with _RETEST_TASK_LOCK:
        _RETEST_TASKS[task_id] = task

    worker = threading.Thread(target=_retest_task_worker, args=(task_id, dict(payload)), daemon=True)
    task["thread"] = worker
    worker.start()
    snapshot = progress.snapshot()
    return {
        "success": True,
        "task_id": task_id,
        "running": True,
        "done": False,
        "message": snapshot["message"],
        "progress": snapshot["progress"],
        "logs": snapshot["logs"],
        "log_count": len(snapshot["logs"]),
        "trace_events": snapshot["trace_events"],
        "trace_event_count": len(snapshot["trace_events"]),
        "source_file": str(source_file),
    }


def _doc_retest_run_one_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    log_offset = max(0, _as_int(payload.get("log_offset"), 0))
    trace_event_offset = max(0, _as_int(payload.get("trace_event_offset"), 0))
    wants_delta = "log_offset" in payload or "trace_event_offset" in payload
    task_id = _required_text(payload, "task_id", "缺少任务ID")
    with _RETEST_TASK_LOCK:
        task = _RETEST_TASKS.get(task_id)
        if not task:
            return {"success": False, "task_id": task_id, "done": True, "running": False, "message": "复测任务不存在或已过期", "logs": [], "trace_events": []}
        progress = task["progress"]
        snapshot = progress.delta_snapshot(log_offset, trace_event_offset) if wants_delta else progress.snapshot()
        result = task.get("result") or {}
        done = bool(task.get("done"))
        running = bool(task.get("running"))
        success = bool(task.get("success")) if done else True
        message = str(task.get("message") or snapshot.get("message") or "")
        stopped = bool(task.get("stopped") or snapshot.get("stop_requested") or result.get("stopped"))

    response = {
        "success": success,
        "task_id": task_id,
        "running": running,
        "done": done,
        "stopped": stopped,
        "message": message,
        "progress": 100 if done and success else snapshot["progress"],
        "logs": snapshot["logs"],
        "log_count": snapshot.get("log_count", len(snapshot["logs"])),
        "trace_events": snapshot["trace_events"],
        "trace_event_count": snapshot.get("trace_event_count", len(snapshot["trace_events"])),
        "source_file": result.get("source_file"),
        "manual_test_required": result.get("manual_test_required"),
        "blocked_by_ai_config": result.get("blocked_by_ai_config"),
        "blocked_stage": result.get("blocked_stage"),
        "blocked_title": result.get("blocked_title"),
        "resume_snapshot": result.get("resume_snapshot"),
        "summary": result.get("summary"),
        "result_data": result.get("result_data"),
        "error": task.get("error"),
    }
    if done:
        if wants_delta and isinstance(result, dict):
            response["result"] = {key: value for key, value in result.items() if key not in {"logs", "trace_events"}}
        else:
            response["result"] = result
    return response


def _doc_retest_confirmation_respond(payload: Dict[str, Any]) -> Dict[str, Any]:
    confirmation_id = _required_text(payload, "confirmation_id", "缺少确认ID")
    decision = str(payload.get("decision") or "").strip().lower()
    note = str(payload.get("note") or "")
    if decision not in {"approve", "reject", "yes", "no", "allow", "deny"}:
        return {"success": False, "message": "decision 必须是 approve 或 reject", "confirmation_id": confirmation_id}
    ok = _resolve_confirmation(confirmation_id, decision, note)
    if not ok:
        return {"success": False, "message": "确认请求不存在或已超时", "confirmation_id": confirmation_id}
    return {"success": True, "confirmation_id": confirmation_id, "decision": decision}


def _doc_retest_run_one_stop(payload: Dict[str, Any]) -> Dict[str, Any]:
    task_id = _required_text(payload, "task_id", "缺少任务ID")
    with _RETEST_TASK_LOCK:
        task = _RETEST_TASKS.get(task_id)
        if not task:
            return {"success": False, "task_id": task_id, "stopped": True, "done": True, "running": False, "message": "复测任务不存在或已过期", "logs": [], "trace_events": []}
        progress = task["progress"]
        _write_retest_cancel_marker("task", task_id)
        _write_retest_cancel_marker("session", str(task.get("session_id") or ""))
        progress.request_stop("复测已停止，可继续")
        progress.event(_retest_trace_event(
            "status",
            "复测已停止",
            "当前单份通报任务已收到停止指令，已保留会话断点。",
            "warn",
            source_file=str(task.get("source_file") or ""),
            metadata={"phase": "stop", "stopped": True},
        ))
        task.update({
            "running": False,
            "done": True,
            "success": False,
            "stopped": True,
            "message": "复测已停止，可继续",
            "result": {
                "success": False,
                "stopped": True,
                "message": "复测已停止，可继续",
                "logs": progress.snapshot().get("logs", []),
                "trace_events": progress.snapshot().get("trace_events", []),
                "resume_snapshot": progress.snapshot().get("resume_snapshot") or None,
                "source_file": str(task.get("source_file") or ""),
            },
            "finished_at": time.time(),
            "error": None,
        })
        snapshot = progress.snapshot()
    return {
        "success": True,
        "task_id": task_id,
        "stopped": True,
        "done": True,
        "running": False,
        "message": "复测已停止，可继续",
        "progress": snapshot["progress"],
        "logs": snapshot["logs"],
        "log_count": len(snapshot["logs"]),
        "trace_events": snapshot["trace_events"],
        "trace_event_count": len(snapshot["trace_events"]),
    }


def _completion_item_from_result(source_file: str, run_result: Dict[str, Any] | None = None, report_paths: List[str] | None = None, failure_reason: str = "") -> Dict[str, Any]:
    result_data = run_result.get("result_data") if isinstance(run_result, dict) else {}
    if not isinstance(result_data, dict):
        result_data = {}
    ai_judgement = result_data.get("ai_judgement") if isinstance(result_data.get("ai_judgement"), dict) else {}
    final_verdict = _model_verdict_from_result_data(result_data)
    reproduced = final_verdict == "reproduced"
    failed = bool(failure_reason or (isinstance(run_result, dict) and not run_result.get("success", True)))
    missing_model_verdict = not failed and final_verdict not in {"reproduced", "not_reproduced"}
    status = "failed" if failed or missing_model_verdict else ("risk" if reproduced else "clean")
    # 目标不可达属于"未能验证"，不能展示成"已修复/复测通过"。
    unreachable = bool(ai_judgement.get("unverified_unreachable") or result_data.get("target_unreachable"))
    if status == "risk":
        status_label = "漏洞未修复/可复现"
    elif status == "failed":
        status_label = "模型未给出判定" if missing_model_verdict else "执行失败"
    elif unreachable:
        status_label = "未复现：目标不可达，未能验证（建议复查）"
    else:
        status_label = "漏洞已修复/复测通过"
    tools: List[str] = []
    evidence_lines: List[str] = []
    for item in result_data.get("retest_results") or []:
        if not isinstance(item, dict):
            continue
        for tool in item.get("context_checks") or []:
            if str(tool).strip():
                tools.append(str(tool))
        for vuln in item.get("vulnerabilities") or []:
            if not isinstance(vuln, dict):
                continue
            detail = str(vuln.get("detail") or vuln.get("evidence") or "").strip()
            if detail:
                evidence_lines.append(detail)
    for item in ai_judgement.get("evidence") or []:
        if str(item).strip():
            evidence_lines.append(str(item))
    reason = failure_reason or str(ai_judgement.get("reason") or result_data.get("reason") or "")
    if missing_model_verdict and not reason:
        reason = "模型未给出 reproduced/not_reproduced 判定，未由工具结果兜底。"
    return {
        "sourceFile": source_file,
        "sourceFileName": Path(source_file).name,
        "status": status,
        "statusLabel": status_label,
        "evidence": "\n".join(evidence_lines[:6]) or reason or "暂无证据摘要",
        "reason": reason,
        "reportPaths": list(report_paths or []),
        "tools": list(dict.fromkeys(tools))[:20],
        "riskCount": 1 if reproduced else 0,
        "manualCount": 0,
        "failedCount": 1 if failed or missing_model_verdict else int(result_data.get("failed_count") or 0),
    }


def _format_agent_result_message(file_path: str, result: Dict[str, Any], completion_item: Dict[str, Any], report_paths: List[str] | None = None) -> str:
    result_data = result.get("result_data") if isinstance(result.get("result_data"), dict) else {}
    ai_judgement = result_data.get("ai_judgement") if isinstance(result_data.get("ai_judgement"), dict) else {}
    final_verdict = _model_verdict_from_result_data(result_data)
    conclusion = str(ai_judgement.get("conclusion") or completion_item.get("statusLabel") or "").strip()
    judgement_label = "快速判定" if result_data.get("fast_mode") or ai_judgement.get("source") == "fast_rules" else "模型判定"
    missing_label = "快速规则未给出判定" if judgement_label == "快速判定" else "模型未给出判定"
    urls = [str(item) for item in (result_data.get("urls") or []) if str(item).strip()]
    lines = [
        f"文件: {Path(file_path).name}",
        f"复测结果: {completion_item.get('statusLabel') or missing_label}",
        f"{judgement_label}: {final_verdict or missing_label}{(' / ' + conclusion) if conclusion else ''}",
    ]
    reason = str(ai_judgement.get("reason") or completion_item.get("reason") or result_data.get("reason") or "").strip()
    if reason:
        lines.append(f"理由: {reason}")
    if urls:
        lines.append("目标: " + "；".join(urls[:4]))
    tools = completion_item.get("tools") or []
    if tools:
        lines.append("工具: " + ", ".join(str(item) for item in tools[:10]))
    evidence = str(completion_item.get("evidence") or "").strip()
    if evidence:
        lines.append("关键证据:\n" + "\n".join(evidence.splitlines()[:6]))
    reports = list(report_paths or completion_item.get("reportPaths") or [])
    if reports:
        lines.append("报告:\n" + "\n".join(f"{index + 1}. {path}" for index, path in enumerate(reports)))
    return "\n".join(lines)


def _format_agent_completion_overview(items: List[Dict[str, Any]]) -> str:
    if not items:
        return "复测结论总览\n暂无文件级结论。"
    labels = {
        "risk": "漏洞未修复/可复现",
        "clean": "漏洞已修复/复测通过",
        "failed": "执行失败",
    }
    lines = ["复测结论总览"]
    for status in ("risk", "clean", "failed"):
        group = [item for item in items if item.get("status") == status]
        lines.append("")
        lines.append(f"【{labels[status]}】{len(group)} 项")
        if not group:
            lines.append("- 无")
            continue
        for item in group:
            lines.append(f"- {item.get('sourceFileName') or Path(str(item.get('sourceFile') or '')).name}")
            evidence = str(item.get("evidence") or "").replace("\n", " / ")
            if evidence:
                lines.append(f"  证据: {evidence[:500]}")
            reason = str(item.get("reason") or "")
            if reason:
                lines.append(f"  原因: {reason[:500]}")
            reports = item.get("reportPaths") or []
            if reports:
                lines.append("  报告: " + "；".join(str(path) for path in reports))
    return "\n".join(lines)


def _format_report_artifact_content(message: str, reports: List[str]) -> str:
    lines = [str(message or ("报告生成完成" if reports else "未生成报告")).strip()]
    clean_reports = [str(path) for path in reports if str(path).strip()]
    if clean_reports:
        lines.append("报告路径:")
        lines.extend(f"{index + 1}. {path}" for index, path in enumerate(clean_reports))
    return "\n".join(line for line in lines if line)


def _existing_report_path(output_path: Any) -> str:
    if not output_path:
        return ""
    path = Path(str(output_path)).expanduser()
    if path.exists() and path.is_file():
        return str(path)
    return ""


def _generate_retest_reports_from_agent_summary(
    target_dir: Path,
    source_files: List[str],
    summary_text: str,
    logs: List[str],
    result_data: Dict[str, Any] | None = None,
    output_dir: Path | None = None,
    include_disposal_reports: bool = True,
) -> Dict[str, Any]:
    template_path = _retest_template_path()
    if not template_path.exists():
        return {"success": False, "message": f"未找到复测模板文件: {template_path}", "reports": [], "failures": [], "logs": logs}
    source_files = _source_notice_paths(source_files)
    if not source_files:
        return {"success": False, "message": "没有可生成报告的原始通报文件", "reports": [], "failures": [], "logs": logs}
    report_screenshot_path: Path | None = None
    report_screenshot_sections: List[Dict[str, str]] = []
    disposal_screenshot_path: Path | None = None
    reports: List[str] = []
    disposal_reports: List[Dict[str, Any]] = []
    failures: List[tuple[Path, str]] = []
    scratch_root = output_dir or target_dir
    try:
        detail_text = _format_report_text_explanation(summary_text, result_data or {})
        evidence_sections = _format_report_evidence_sections(summary_text, result_data or {})
        report_screenshot_sections = _save_retest_evidence_section_screenshots(scratch_root, evidence_sections)
        if report_screenshot_sections:
            report_screenshot_path = Path(report_screenshot_sections[0]["path"])
            logs.append(f"AI Agent 复测报告分段证据图已生成: {len(report_screenshot_sections)} 张")
        else:
            evidence_text = _format_report_evidence_screenshot_text(summary_text, result_data or {})
            report_screenshot_path = _save_retest_text_screenshot(scratch_root, evidence_text, "复测证据")
            report_screenshot_sections = [{"caption": detail_text, "path": str(report_screenshot_path)}]
            logs.append(f"AI Agent 复测报告证据图已生成: {report_screenshot_path}")
        disposal_text = _format_report_evidence_snapshot(summary_text, result_data or {})
        disposal_screenshot_path = _save_retest_text_screenshot(scratch_root, disposal_text, "复测证据总览")
        logs.append(f"处置文件复测证据总图已生成: {disposal_screenshot_path}")
        from modules.AI_Testing.retest.retest_report_generator import RetestReportGenerator

        for source_file in source_files:
            file_path = Path(source_file).expanduser()
            if not file_path.exists() or file_path.suffix.lower() not in WORD_SUFFIXES:
                failures.append((file_path, "通报文件不存在或不是 Word 文档"))
                continue
            if is_generated_retest_report_path(file_path):
                failures.append((file_path, "这是已生成的复测报告，不是原始通报文件"))
                continue
            generator = RetestReportGenerator(
                target_dir=str(file_path.parent),
                template_path=str(template_path),
                output_dir=str(output_dir) if output_dir is not None else None,
                screenshot_path=str(report_screenshot_path),
                screenshot_sections=report_screenshot_sections,
            )
            buffer = io.StringIO()
            with contextlib.redirect_stdout(buffer):
                generator_scan = generator.scan_document(file_path)
                generator_scan["report_detail"] = detail_text
                output_path = generator.generate_report(generator_scan)
            logs.extend(_captured_lines(buffer))
            report_path = _existing_report_path(output_path)
            if report_path:
                reports.append(report_path)
                logs.append(f"报告已生成: {report_path}")
                try:
                    if not include_disposal_reports:
                        continue
                    disposal_result = _prepare_retest_disposal_report(
                        file_path,
                        generator_scan,
                        disposal_screenshot_path,
                        logs,
                        output_dir=output_dir,
                    )
                    if disposal_result:
                        disposal_reports.append(disposal_result)
                        if disposal_result.get("pdf_error"):
                            failures.append((file_path, f"处置文件PDF转换失败: {disposal_result['pdf_error']}"))
                    else:
                        failures.append((file_path, "处置文件未生成"))
                except Exception as exc:
                    logs.append(traceback.format_exc())
                    failures.append((file_path, f"处置文件替换失败: {exc}"))
            else:
                failures.append((file_path, f"报告生成失败或文件未落盘: {output_path or '无输出路径'}"))
    except Exception as exc:
        logs.append(traceback.format_exc())
        return {"success": False, "message": f"报告生成失败: {exc}", "reports": reports, "failures": _failure_dicts(failures), "logs": logs}
    finally:
        _cleanup_retest_screenshot_dir(scratch_root, logs)
    success = bool(reports) and not failures
    artifacts = list(reports)
    for disposal_report in disposal_reports:
        for key in ("word", "pdf"):
            artifact = _existing_report_path(disposal_report.get(key))
            if artifact and artifact not in artifacts:
                artifacts.append(artifact)
    return {
        "success": success,
        "message": (
            f"报告生成完成: {len(reports)} 份，失败 {len(failures)} 份"
            if success
            else f"报告未生成或生成失败: 成功 {len(reports)} 份，失败 {len(failures)} 份"
        ),
        "reports": reports,
        "artifacts": artifacts,
        "disposal_reports": disposal_reports,
        "failures": _failure_dicts(failures),
        "logs": logs,
    }


class RetestAgentRunner:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.lock = threading.RLock()
        self.target_dir = ""
        self.workspace_root = str(_project_root())
        self.source_files: List[str] = []
        self.next_index = 0
        self.summaries: List[str] = []
        self.reports: List[str] = []
        self.report_evidence_summaries: Dict[str, str] = {}
        self.report_result_data: Dict[str, Dict[str, Any]] = {}
        self.completion_items: List[Dict[str, Any]] = []
        self.logs: List[str] = []
        self.latest_result_data: Dict[str, Any] | None = None
        self.current_file_resume: Dict[str, Any] | None = None
        self.generate_reports = False
        self.running = False
        self.blocked = False
        self.blocked_reason = ""
        self.blocked_stage = ""
        self.blocked_title = ""
        self.stopped = False
        self.thread: threading.Thread | None = None
        self.pending_messages: List[str] = []
        # 会话级 ReAct 完整消息历史（不含 system；含 user/assistant/tool 及 tool_calls 结构），
        # 跨轮持久化，让对话框记得"上一轮发过什么请求、拿到什么响应、调过什么工具"。
        self.conversation: List[Dict[str, Any]] = []
        self.frontend_context_fingerprint = ""
        self.frontend_context_payload: Dict[str, Any] = {}
        self.frontend_completed_file_names: List[str] = []
        self.frontend_completed_count_hint = 0
        self.frontend_next_index_hint = 0
        self.frontend_next_source_file_name = ""
        self.disk_completed_file_names: List[str] = []
        self.disk_completed_report_evidence: List[Dict[str, str]] = []
        self.turn_counter = 0
        self.current_turn_id = ""
        self.created_at = time.time()
        self.updated_at = time.time()
        self.hybrid_runtime = None
        # Ignore stale marker files left by previous turns. Stop markers are
        # retained so an older worker cannot be revived by a fast resume.
        self.cancel_epoch_ns = time.time_ns()

    def _agent_runtime(self):
        runtime = self.hybrid_runtime
        if runtime is None:
            runtime = _make_hybrid_agent_runtime(
                self.session_id,
                mode="retest",
                payload={"workspace_root": self.workspace_root, "target_dir": self.target_dir},
            )
            self.hybrid_runtime = runtime
        return runtime

    def _agent_runtime_snapshot(self) -> Dict[str, Any]:
        try:
            return self._agent_runtime().snapshot()
        except Exception:
            return {}

    def _reset_retest_state_locked(self) -> None:
        self.source_files = []
        self.next_index = 0
        self.summaries = []
        self.reports = []
        self.report_evidence_summaries = {}
        self.report_result_data = {}
        self.completion_items = []
        self.logs = []
        self.latest_result_data = None
        self.current_file_resume = None
        self.frontend_context_payload = {}
        self.frontend_completed_file_names = []
        self.frontend_completed_count_hint = 0
        self.frontend_next_index_hint = 0
        self.frontend_next_source_file_name = ""
        self.disk_completed_file_names = []
        self.disk_completed_report_evidence = []

    def _set_workspace_root_locked(self, value: Any) -> bool:
        text = str(value or "").strip()
        if not text:
            return False
        candidate = Path(text).expanduser()
        try:
            if not candidate.exists() or not candidate.is_dir():
                return False
            resolved = str(candidate.resolve())
        except Exception:
            return False
        changed = _retest_target_key(resolved) != _retest_target_key(self.workspace_root)
        if changed:
            self.workspace_root = resolved
            self.hybrid_runtime = None
        return changed

    def _set_target_dir_locked(self, target_dir: str, *, reset_on_change: bool = True) -> bool:
        cleaned = str(target_dir or "").strip()
        if not cleaned:
            return False
        changed = bool(self.target_dir) and _retest_target_key(cleaned) != _retest_target_key(self.target_dir)
        if changed and reset_on_change:
            self._reset_retest_state_locked()
            self.frontend_context_fingerprint = ""
            self.frontend_context_payload = {}
        self.target_dir = cleaned
        return changed

    def _is_turn_current_locked(self, turn_id: str) -> bool:
        return not turn_id or not self.current_turn_id or self.current_turn_id == turn_id

    def _turn_is_current(self, turn_id: str) -> bool:
        with self.lock:
            return self._is_turn_current_locked(turn_id)

    def _execution_cancelled_locked(self, turn_id: str = "") -> bool:
        # The desktop stop command may arrive while the sidecar request channel
        # is occupied by a long operation. Latch it at every safe checkpoint.
        self._cancel_signal_locked()
        return self.stopped or not self._is_turn_current_locked(turn_id)

    def _turn_is_cancelled(self, turn_id: str = "") -> bool:
        with self.lock:
            return self._execution_cancelled_locked(turn_id)

    def _needs_frontend_hydration_locked(self, context: Dict[str, Any]) -> bool:
        session = _as_record(context.get("session"))
        resume_state = _as_record(session.get("resumeState"))
        incoming_target = str(resume_state.get("targetDir") or session.get("targetDir") or "").strip()
        if incoming_target and self.target_dir and _retest_target_key(incoming_target) != _retest_target_key(self.target_dir):
            return True
        frontend_files = _as_text_list(resume_state.get("sourceFiles"))
        frontend_items = resume_state.get("completionItems") if isinstance(resume_state.get("completionItems"), list) else []
        progress_evidence = _as_record(context.get("progressEvidence"))
        completed_names = _as_text_list(progress_evidence.get("completedFileNames"))
        completed_count_hint = max(0, _as_int(progress_evidence.get("completedCountHint"), 0))
        next_index_hint = max(0, _as_int(progress_evidence.get("nextIndexHint"), 0))
        next_source_file_name = Path(str(progress_evidence.get("nextSourceFileName") or "")).name.strip()
        if not self.target_dir and (session.get("targetDir") or resume_state.get("targetDir")):
            return True
        if frontend_files and len(frontend_files) > len(self.source_files):
            return True
        try:
            frontend_next = int(resume_state.get("nextIndex"))
        except Exception:
            frontend_next = 0
        if frontend_next > self.next_index:
            return True
        if max(completed_count_hint, next_index_hint, len(completed_names), len(frontend_items)) > self.next_index:
            return True
        if next_source_file_name:
            return True
        if frontend_items and len(frontend_items) > len(self.completion_items):
            return True
        if completed_names and len(completed_names) > len(self.frontend_completed_file_names):
            return True
        return not self.conversation and self._has_frontend_recoverable_history(context)

    def _has_frontend_recoverable_history(self, context: Dict[str, Any]) -> bool:
        session = _as_record(context.get("session"))
        resume_state = _as_record(session.get("resumeState"))
        if str(session.get("memoryMarkdown") or "").strip():
            return True
        if resume_state.get("canContinue") or _as_text_list(resume_state.get("sourceFiles")):
            return True
        if resume_state.get("completionItems") or resume_state.get("summaries") or resume_state.get("reports"):
            return True
        if str(session.get("resultText") or "").strip() or str(session.get("latestResultDataText") or "").strip():
            return True
        if len(_as_text_list(session.get("logTail"))) > 3:
            return True

        progress_evidence = _as_record(context.get("progressEvidence"))
        if (
            _as_text_list(progress_evidence.get("completedFileNames"))
            or str(progress_evidence.get("latestSourceFileName") or "").strip()
            or _as_int(progress_evidence.get("completedCountHint"), 0) > 0
            or _as_int(progress_evidence.get("nextIndexHint"), 0) > 0
            or str(progress_evidence.get("nextSourceFileName") or "").strip()
        ):
            return True

        conversation = context.get("conversation") if isinstance(context.get("conversation"), list) else []
        if len(conversation) > 1:
            return True
        for item in conversation:
            if not isinstance(item, dict):
                continue
            if item.get("tools") or item.get("artifacts") or item.get("errors"):
                return True
            if str(item.get("role") or "") in {"agent", "system"}:
                return True

        recent_events = context.get("recentEvents") if isinstance(context.get("recentEvents"), list) else []
        if len(recent_events) > 2:
            return True
        for item in recent_events:
            if not isinstance(item, dict):
                continue
            if str(item.get("type") or "") != "chat" or str(item.get("role") or "") != "user":
                return True
        return False

    def hydrate_from_frontend_context(self, value: Any) -> bool:
        context = _as_record(value)
        if not context:
            return False
        session = _as_record(context.get("session"))
        resume_state = _as_record(session.get("resumeState"))
        with self.lock:
            self._set_workspace_root_locked(
                session.get("workspaceRoot")
                or context.get("workspaceRoot")
                or context.get("workspace_root")
                or ""
            )
            if not self._needs_frontend_hydration_locked(context):
                return False
            target_dir = str(resume_state.get("targetDir") or session.get("targetDir") or "").strip()
            target_changed = self._set_target_dir_locked(target_dir, reset_on_change=True) if target_dir else False

            source_files = _as_text_list(resume_state.get("sourceFiles"))
            if source_files and (target_changed or not self.source_files or len(source_files) >= len(self.source_files)):
                self.source_files = source_files

            try:
                next_index = int(resume_state.get("nextIndex"))
            except Exception:
                next_index = self.next_index
            if source_files:
                next_index = max(0, min(len(source_files), next_index))
            if target_changed or next_index > self.next_index or not self.completion_items:
                self.next_index = max(0, next_index)

            summaries = _as_text_list(resume_state.get("summaries"))
            if summaries and (target_changed or len(summaries) >= len(self.summaries)):
                self.summaries = summaries

            reports = _as_text_list(resume_state.get("reports"))
            if reports and (target_changed or len(reports) >= len(self.reports)):
                self.reports = reports

            completion_items = resume_state.get("completionItems")
            if isinstance(completion_items, list) and (target_changed or len(completion_items) >= len(self.completion_items)):
                self._set_completion_items_locked(completion_items)

            current_file = _as_record(resume_state.get("currentFile"))
            if current_file:
                self.current_file_resume = dict(current_file)

            progress_evidence = _as_record(context.get("progressEvidence"))
            completed_names = _as_text_list(progress_evidence.get("completedFileNames"), 500)
            if completed_names:
                if target_changed:
                    self.frontend_completed_file_names = []
                self._merge_completed_file_names_locked(completed_names, source="frontend")
            completed_count_hint = max(0, _as_int(progress_evidence.get("completedCountHint"), 0))
            next_index_hint = max(0, _as_int(progress_evidence.get("nextIndexHint"), 0))
            next_source_file_name = _source_notice_name(progress_evidence.get("nextSourceFileName"))
            completed_count_hint = max(
                completed_count_hint,
                len(self.frontend_completed_file_names),
                len(self.completion_items),
            )
            next_index_hint = max(
                next_index_hint,
                len(self.frontend_completed_file_names),
                len(self.completion_items),
            )
            self.frontend_completed_count_hint = max(self.frontend_completed_count_hint, completed_count_hint)
            self.frontend_next_index_hint = max(self.frontend_next_index_hint, next_index_hint)
            if next_source_file_name:
                self.frontend_next_source_file_name = next_source_file_name
            if (
                self.source_files
                and self.next_index > 0
                and not self._completed_file_display_names_locked()
                and not self.frontend_next_source_file_name
                and not self.current_file_resume
            ):
                self.next_index = 0
            if not self.source_files:
                if self._completed_file_display_names_locked():
                    self.next_index = max(self.next_index, len(self._completed_file_display_names_locked()))
            else:
                self._apply_frontend_progress_hints_locked()

            if self.source_files:
                self._apply_frontend_progress_hints_locked()

            logs = _as_text_list(resume_state.get("allLogs")) or _as_text_list(session.get("logTail"))
            if logs and len(logs) >= len(self.logs):
                self.logs = logs[-3000:]

            self.generate_reports = (
                _payload_bool(resume_state.get("generateReports"), self.generate_reports)
                or _payload_bool(session.get("generateReports"), False)
            )
            if not self.generate_reports and _frontend_context_requests_reports(context):
                self.generate_reports = True
            latest_text = str(session.get("latestResultDataText") or "").strip()
            if latest_text and self.latest_result_data is None:
                try:
                    parsed = json.loads(latest_text)
                    if isinstance(parsed, dict):
                        self.latest_result_data = parsed
                except Exception:
                    pass
            return True

    def _has_direct_resume_context_locked(self) -> bool:
        if self.source_files and self.next_index < len(self.source_files):
            return True
        if self.blocked and self.target_dir:
            return True
        if self.current_file_resume:
            return True
        context = _as_record(self.frontend_context_payload)
        session = _as_record(context.get("session"))
        resume_state = _as_record(session.get("resumeState"))
        if resume_state.get("canContinue") or _as_text_list(resume_state.get("sourceFiles")):
            return True
        if _as_record(resume_state.get("currentFile")):
            return True
        progress_evidence = _as_record(context.get("progressEvidence"))
        return bool(
            _as_text_list(progress_evidence.get("completedFileNames"))
            or _as_int(progress_evidence.get("completedCountHint"), 0) > 0
            or _as_int(progress_evidence.get("nextIndexHint"), 0) > 0
            or str(progress_evidence.get("nextSourceFileName") or "").strip()
        )

    def _current_file_resume_from_snapshot_locked(self, snapshot: Dict[str, Any] | None, source_file: str = "") -> Dict[str, Any] | None:
        if not isinstance(snapshot, dict) or not snapshot:
            return None
        source = str(snapshot.get("source_file") or source_file or "").strip()
        if not source:
            return None
        source_key = _retest_target_key(source)
        source_name = _source_notice_name(source)
        index = self.next_index
        matched = False
        for item_index, item in enumerate(self.source_files):
            if _retest_target_key(item) == source_key:
                index = item_index
                matched = True
                break
        if not matched and source_name:
            name_matches = [
                item_index
                for item_index, item in enumerate(self.source_files)
                if _source_notice_name(item).lower() == source_name.lower()
            ]
            if len(name_matches) == 1:
                index = name_matches[0]
        return {
            "index": max(0, min(index, max(0, len(self.source_files) - 1))),
            "sourceFile": source,
            "sourceFileName": source_name or Path(source).name,
            "stage": str(snapshot.get("stage") or snapshot.get("resume_stage") or ""),
            "resumeSnapshot": dict(snapshot),
        }

    def _save_current_file_checkpoint(self, snapshot: Dict[str, Any] | None, source_file: str, turn_id: str, round_id: str) -> None:
        if not isinstance(snapshot, dict) or not snapshot:
            return
        stage = str(snapshot.get("stage") or snapshot.get("resume_stage") or "checkpoint").strip().lower()
        result_data = snapshot.get("result_data") if isinstance(snapshot.get("result_data"), dict) else None
        with self.lock:
            # A stop may be observed while a Python repair call is unwinding.
            # Preserve that final checkpoint for the still-current turn, but
            # continue rejecting every checkpoint from a superseded turn.
            self._cancel_signal_locked()
            if not self._is_turn_current_locked(turn_id):
                return
            resume = self._current_file_resume_from_snapshot_locked(snapshot, source_file)
            if not resume:
                return
            self.current_file_resume = resume
            if result_data:
                self.latest_result_data = result_data
            completed_urls = _as_int(snapshot.get("completed_url_count"), _as_int(snapshot.get("next_url_index"), 0))
            total_urls = _as_int(snapshot.get("total_url_count"), 0)
            if stage == "execution":
                status_text = f"Checkpoint saved: verified {completed_urls}/{total_urls} URL(s)"
                content = "Completed URL results are stored; continue will resume from next_url_index without rerunning them."
            elif stage == "result":
                status_text = "Checkpoint saved: retest judgement complete"
                content = "Retest result is stored; continue will not rerun probes or judgement."
            elif stage in {"report", "report_generation"}:
                status_text = "Checkpoint saved: report generation"
                content = "Retest result is stored; continue can generate the report without rerunning probes."
            else:
                status_text = f"Checkpoint saved: {stage or 'unknown'}"
                content = "Current file checkpoint is stored for resume."
            patch = self._session_patch({
                "status": status_text,
                "latestResultData": result_data or self.latest_result_data,
                "resumeState": self._resume_state_locked(True),
            })
        self._publish(
            "status",
            status_text,
            content,
            "ok",
            metadata={
                "turnId": turn_id,
                "roundId": round_id,
                "phase": stage or "checkpoint",
                "sourceFileName": Path(source_file).name,
                "resumeSnapshot": dict(snapshot),
                "sessionPatch": patch,
            },
        )

    def _resume_snapshot_for_source_locked(self, index: int, source_file: str) -> Dict[str, Any]:
        candidates: List[Dict[str, Any]] = []
        if isinstance(self.current_file_resume, dict):
            candidates.append(self.current_file_resume)
        context = _as_record(self.frontend_context_payload)
        session = _as_record(context.get("session"))
        resume_state = _as_record(session.get("resumeState"))
        current_file = _as_record(resume_state.get("currentFile"))
        if current_file:
            candidates.append(current_file)
        source_key = _retest_target_key(source_file)
        source_name = _source_notice_name(source_file).lower()
        source_name_count = sum(
            1 for item in self.source_files
            if _source_notice_name(item).lower() == source_name
        )
        for candidate in candidates:
            snapshot = _as_record(candidate.get("resumeSnapshot") or candidate.get("resume_snapshot"))
            if not snapshot:
                continue
            stage = str(candidate.get("stage") or snapshot.get("stage") or snapshot.get("resume_stage") or "").strip().lower()
            if stage not in {"execution", "verification", "tool", "judgement", "result", "completed", "judgement_complete", "report", "report_generation"}:
                continue
            current_index = _as_int(candidate.get("index"), -1)
            current_source = str(candidate.get("sourceFile") or candidate.get("source_file") or snapshot.get("source_file") or "").strip()
            same_index = current_index == index
            same_path = bool(current_source and _retest_target_key(current_source) == source_key)
            same_name = bool(current_source and _source_notice_name(current_source).lower() == source_name)
            index_without_source = same_index and not current_source
            unique_name_match = same_name and source_name_count == 1
            if same_path or index_without_source or unique_name_match:
                return dict(snapshot)
        return {}

    def frontend_context_message(self, value: Any) -> Dict[str, Any] | None:
        context = _as_record(value)
        if not context:
            return None
        session = _as_record(context.get("session"))
        conversation = context.get("conversation") if isinstance(context.get("conversation"), list) else []
        recent_events = context.get("recentEvents") if isinstance(context.get("recentEvents"), list) else []
        lines: List[str] = ["[前端持久化会话上下文，用于恢复长时间间隔/后端重启后的同一对话记忆]"]
        if session:
            memory_markdown = _truncate_agent_context(session.get("memoryMarkdown"), 12000)
            if memory_markdown:
                lines.append("AI 语义压缩记忆（恢复/继续时必须优先使用，不能忽略）:\n" + memory_markdown)
            lines.append(f"会话: {session.get('title') or self.session_id}")
            lines.append(f"状态: {session.get('status') or ''}；进度: {session.get('progress') or 0}%")
            if session.get("targetDir"):
                lines.append(f"通报目录: {session.get('targetDir')}")
            resume_state = _as_record(session.get("resumeState"))
            if resume_state:
                files = _as_text_list(resume_state.get("sourceFiles"))
                lines.append(
                    "断点: "
                    f"canContinue={bool(resume_state.get('canContinue'))}, "
                    f"nextIndex={resume_state.get('nextIndex')}, total={len(files)}, "
                    f"reason={resume_state.get('blockedReason') or session.get('status') or ''}"
                )
            result_text = _truncate_agent_context(session.get("resultText"), 2500)
            if result_text:
                lines.append("最近复测结果:\n" + result_text)
            log_tail = _as_text_list(session.get("logTail"), 40)
            if log_tail:
                lines.append("最近日志:\n" + "\n".join(_truncate_agent_context(item, 500) for item in log_tail[-40:]))

        progress_evidence = _as_record(context.get("progressEvidence"))
        completed_names = _as_text_list(progress_evidence.get("completedFileNames"), 200)
        latest_source_name = str(progress_evidence.get("latestSourceFileName") or "").strip()
        completed_count_hint = max(0, _as_int(progress_evidence.get("completedCountHint"), 0))
        next_index_hint = max(0, _as_int(progress_evidence.get("nextIndexHint"), 0))
        next_source_file_name = Path(str(progress_evidence.get("nextSourceFileName") or "")).name
        if completed_count_hint or next_index_hint or next_source_file_name:
            lines.append(
                "Frontend recovery numeric hints: "
                f"completedCountHint={completed_count_hint}, "
                f"nextIndexHint={next_index_hint}, "
                f"nextSourceFileName={next_source_file_name or ''}. "
                "Treat completedCountHint/nextIndexHint as weak metadata only; do not skip notices from numeric hints alone. "
                "Skip only when completedFileNames, disk report evidence, or an exact nextSourceFileName confirms the queue position."
            )
        if completed_names or latest_source_name:
            lines.append(
                "前端恢复进度证据: "
                f"已完成文件 {len(completed_names)} 个"
                + (f"，最近处理 {latest_source_name}" if latest_source_name else "")
            )
            if completed_names:
                lines.append("已完成文件名:\n" + "\n".join(completed_names[-80:]))

        with self.lock:
            self._advance_next_index_past_completed_locked()
            queue_total = len(self.source_files)
            queue_next = self.next_index
            queue_next_name = Path(self.source_files[queue_next]).name if 0 <= queue_next < queue_total else ""
        if queue_total:
            if queue_next_name:
                lines.append(
                    f"后端确定性断点: nextIndex={queue_next}, total={queue_total}，"
                    f"下一份未完成是第 {queue_next + 1} 份: {queue_next_name}。"
                )
                lines.append("硬约束: 继续/恢复时必须从这份未完成通报开始，不要再从第 1 份或已完成文件开始。")
            else:
                lines.append(f"后端确定性断点: nextIndex={queue_next}, total={queue_total}，全部通报已完成。")

        if conversation:
            lines.append("最近对话轮次:")
            for item in conversation[-16:]:
                if not isinstance(item, dict):
                    continue
                role = str(item.get("role") or "agent")
                title = str(item.get("title") or role)
                content = _truncate_agent_context(item.get("content"), 1400)
                tool_bits: List[str] = []
                for tool_item in (item.get("tools") or [])[:6] if isinstance(item.get("tools"), list) else []:
                    if not isinstance(tool_item, dict):
                        continue
                    tool = _as_record(tool_item.get("tool"))
                    tool_bits.append(
                        f"{tool_item.get('title') or tool.get('label') or tool.get('toolId')}: "
                        f"{tool.get('status') or ''} {tool.get('target') or ''} {tool.get('resultPreview') or tool_item.get('content') or ''}"
                    )
                merged = f"- {role}/{title}: {content}".strip()
                if tool_bits:
                    merged += "\n  工具: " + "；".join(_truncate_agent_context(bit, 500) for bit in tool_bits)
                lines.append(merged)

        if not conversation and recent_events:
            lines.append("最近事件:")
            for item in recent_events[-30:]:
                if not isinstance(item, dict):
                    continue
                lines.append(
                    f"- {item.get('timestamp') or ''} {item.get('type') or ''}/{item.get('title') or ''}: "
                    f"{_truncate_agent_context(item.get('content'), 700)}"
                )
        return {
            "role": "user",
            "content": _truncate_agent_context("\n".join(lines), 24000),
            "metadata": {"frontendContext": True},
        }

    def snapshot(self) -> Dict[str, Any]:
        with self.lock:
            self._cancel_signal_locked()
            self._advance_next_index_past_completed_locked()
            return {
                "success": True,
                "session_id": self.session_id,
                "target_dir": self.target_dir,
            "source_files": list(self.source_files),
            "next_index": self.next_index,
            "completed_file_names": self._completed_file_display_names_locked(),
            "disk_completed_file_names": list(self.disk_completed_file_names),
            "disk_completed_report_evidence": list(self.disk_completed_report_evidence),
            "running": self.running,
                "stopped": self.stopped,
                "blocked": self.blocked,
                "blocked_reason": self.blocked_reason,
                "blocked_stage": self.blocked_stage,
                "blocked_title": self.blocked_title,
                "resume_state": self._resume_state_locked(True) if self.blocked or self.stopped else None,
                "generate_reports": self.generate_reports,
                "summaries": list(self.summaries),
                "reports": list(self.reports),
                "completion_items": list(self.completion_items),
                "logs": list(self.logs),
                "latest_result_data": self.latest_result_data,
                "progress": self._progress_locked(),
                "status": self._status_locked(),
                "agent_runtime": self._agent_runtime_snapshot(),
            }

    def _progress_locked(self) -> int:
        total = max(1, len(self.source_files))
        if not self.source_files:
            return 100 if self.completion_items else 0
        return max(0, min(100, int(round(self._effective_next_index_locked() / total * 100))))

    def _cancel_signal_locked(self) -> bool:
        if not self.stopped and _retest_cancel_requested(self.session_id, newer_than_ns=self.cancel_epoch_ns):
            self.stopped = True
            self.running = False
            self.blocked = False
            self.pending_messages = []
        return self.stopped

    def _status_locked(self) -> str:
        self._cancel_signal_locked()
        if self.running:
            return "Agent 正在处理..."
        if self.blocked:
            return self.blocked_title or "Agent 会话待继续"
        if self.stopped:
            return "已停止"
        if self.source_files and self._effective_next_index_locked() >= len(self.source_files):
            return "复测完成"
        if self.completion_items and not self.source_files:
            return "复测完成"
        return "等待 Agent 指令"

    def start(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        message = str(payload.get("message") or "一键复测并生成报告").strip()
        force_resume = _payload_bool(payload.get("force_resume"), False)
        one_click_queue = _payload_bool(payload.get("one_click_queue"), False)
        use_progress_evidence = _payload_bool(payload.get("use_progress_evidence"), True)
        hydrated = self.hydrate_from_frontend_context(payload.get("frontend_context"))
        self._merge_frontend_context(payload.get("frontend_context"), force=hydrated)
        with self.lock:
            target_dir = str(payload.get("target_dir") or self.target_dir or "").strip()
            if self.running and target_dir and self.target_dir and _retest_target_key(target_dir) != _retest_target_key(self.target_dir):
                return {**self.snapshot(), "success": False, "message": "当前 Agent 正在运行，不能切换通报目录。请先停止当前会话。"}
            if self.running and force_resume:
                self.stopped = True
                self.running = False
                self.blocked = False
                self.pending_messages = []
                self._publish(
                    "status", "接管旧会话",
                    "旧 Agent 运行态未正常退出，已废弃旧 turn 并从前端持久化上下文重新继续。",
                    "warn", metadata={"role": "agent", "turnId": self.current_turn_id, "phase": "frontend_context_restore"},
                )
            if target_dir:
                self._set_target_dir_locked(target_dir, reset_on_change=True)
            context_requests_reports = _frontend_context_requests_reports(payload.get("frontend_context"))
            self.generate_reports = (
                _payload_bool(payload.get("generate_reports"), self.generate_reports or _message_requests_report(message))
                or context_requests_reports
            )
            self.blocked = False
            self.blocked_reason = ""
            self.blocked_stage = ""
            self.blocked_title = ""
            direct_queue = bool(one_click_queue) or (force_resume and self._has_direct_resume_context_locked())
            reset_queue = not force_resume and not direct_queue and not bool(self.source_files)
        return self._launch(message, reset_queue=reset_queue, direct_queue=direct_queue, use_progress_evidence=use_progress_evidence)

    def message(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        message = str(payload.get("message") or "").strip()
        if not message:
            return {**self.snapshot(), "success": False, "message": "请输入 Agent 指令"}
        force_resume = _payload_bool(payload.get("force_resume"), False)
        use_progress_evidence = _payload_bool(payload.get("use_progress_evidence"), True)
        hydrated = self.hydrate_from_frontend_context(payload.get("frontend_context"))
        self._merge_frontend_context(payload.get("frontend_context"), force=hydrated)
        target_dir = str(payload.get("target_dir") or "").strip()
        with self.lock:
            if self.running and target_dir and self.target_dir and _retest_target_key(target_dir) != _retest_target_key(self.target_dir):
                return {**self.snapshot(), "success": False, "message": "当前 Agent 正在运行，不能切换通报目录。请先停止当前会话。"}
            if target_dir:
                self._set_target_dir_locked(target_dir, reset_on_change=True)
            context_requests_reports = _frontend_context_requests_reports(payload.get("frontend_context"))
            if "generate_reports" in payload:
                self.generate_reports = _payload_bool(payload.get("generate_reports"), self.generate_reports) or context_requests_reports
            elif _message_requests_report(message):
                self.generate_reports = True
            elif context_requests_reports:
                self.generate_reports = True
            if self.running:
                if force_resume:
                    self.stopped = True
                    self.running = False
                    self.blocked = False
                    self.pending_messages = []
                    self._publish(
                        "status", "接管旧会话",
                        "旧 Agent 运行态未正常退出，已废弃旧 turn 并从前端持久化上下文重新继续。",
                        "warn", metadata={"role": "agent", "turnId": self.current_turn_id, "phase": "frontend_context_restore"},
                    )
                else:
                    if len(self.pending_messages) >= _RETEST_AGENT_MAX_PENDING_MESSAGES:
                        self._publish("error", "指令队列已满", "当前 Agent 仍在执行，排队指令已达到上限。请等待当前任务结束，或先停止后再发送。", "warn", metadata={"role": "agent", "turnId": self.current_turn_id})
                        return {**self.snapshot(), "success": False, "message": "Agent 指令队列已满，请稍后再发送。"}
                    self.pending_messages.append(message)
                    self._publish("chat", "Agent", "我收到你的新指令了，当前工具执行结束后会继续处理。", "info", metadata={"role": "agent", "turnId": self.current_turn_id})
                    return {**self.snapshot(), "success": True, "message": "Agent 正在运行，已记录指令。"}
            direct_queue = force_resume and self._has_direct_resume_context_locked()
        return self._launch(message, reset_queue=False, direct_queue=direct_queue, use_progress_evidence=use_progress_evidence)

    def _merge_frontend_context(self, value: Any, force: bool = False) -> None:
        context = _as_record(value)
        session = _as_record(context.get("session"))
        has_semantic_memory = bool(str(session.get("memoryMarkdown") or "").strip())
        with self.lock:
            should_merge = (force or has_semantic_memory or self._needs_frontend_hydration_locked(context)) and self._has_frontend_recoverable_history(context)
        if not should_merge:
            return
        context_message = self.frontend_context_message(context)
        if not context_message:
            return
        fingerprint = str(hash(context_message.get("content") or ""))
        with self.lock:
            if fingerprint == self.frontend_context_fingerprint:
                self.frontend_context_payload = context
                return
            self.frontend_context_fingerprint = fingerprint
            self.frontend_context_payload = context
            self.conversation = [
                item for item in self.conversation
                if not (_as_record(item.get("metadata")).get("frontendContext"))
            ]
            self.conversation.append(context_message)
            if len(self.conversation) > 80:
                self.conversation = self.conversation[-80:]

    def stop(self) -> Dict[str, Any]:
        _write_retest_cancel_marker("session", self.session_id)
        with self.lock:
            self.stopped = True
            self.running = False
            self.blocked = False
            self.pending_messages = []
        self._publish("status", "Agent 已停止", "当前会话已收到停止指令。", "warn", metadata={"sessionPatch": self._session_patch()})
        return {"success": True, "message": "Agent 已停止", **self.snapshot()}

    def _launch(self, message: str, reset_queue: bool = False, direct_queue: bool = False, use_progress_evidence: bool = True) -> Dict[str, Any]:
        with self.lock:
            if self.running:
                return {"success": True, "message": "Agent 已在运行中", **self.snapshot()}
            # A stop marker belongs to the execution epoch that wrote it.  A
            # deliberate continue starts a fresh epoch; remove the old marker
            # first so filesystem timestamp granularity cannot immediately
            # cancel the new turn again.  Older workers remain cancelled by
            # their latched stopped flag / superseded turn id.
            _clear_retest_cancel_marker("session", self.session_id)
            self.cancel_epoch_ns = time.time_ns()
            self.running = True
            self.stopped = False
            self.blocked = False
            self.blocked_reason = ""
            self.blocked_stage = ""
            self.blocked_title = ""
            self.updated_at = time.time()
            self.turn_counter += 1
            turn_id = f"agent:{self.session_id}:turn:{self.turn_counter}:{uuid.uuid4().hex[:6]}"
            self.current_turn_id = turn_id
            if reset_queue:
                self._reset_retest_state_locked()
            self.pending_messages = []
            thread = threading.Thread(target=self._worker, args=(message, reset_queue, turn_id, direct_queue, use_progress_evidence), name=f"koi-retest-agent-{self.session_id[:8]}", daemon=True)
            self.thread = thread
            thread.start()
        return {**self.snapshot(), "success": True, "message": "Agent 已开始处理"}

    def _new_turn_id_locked(self) -> str:
        self.turn_counter += 1
        self.current_turn_id = f"agent:{self.session_id}:turn:{self.turn_counter}:{uuid.uuid4().hex[:6]}"
        return self.current_turn_id

    def _worker(self, message: str, reset_queue: bool, turn_id: str, direct_queue: bool = False, use_progress_evidence: bool = True) -> None:
        runtime_run = None
        runtime_status = "completed"
        try:
            runtime_run = self._agent_runtime().begin_run(message, mode="retest")
            self._publish("status", "Agent 正在思考", "正在理解你的消息，必要时会调用工具。", "info", metadata={"role": "agent", "turnId": turn_id, "sessionPatch": self._session_patch({"isRunning": True, "resumeState": None, "status": "Agent 正在处理..."})})
            self._handle_instruction(message, reset_queue, turn_id, direct_queue=direct_queue, use_progress_evidence=use_progress_evidence)
            while True:
                with self.lock:
                    if not self.pending_messages or self.stopped or self.blocked:
                        break
                    next_message = self.pending_messages.pop(0)
                    turn_id = self._new_turn_id_locked()
                if runtime_run:
                    self._agent_runtime().finish_run(runtime_run, runtime_status)
                runtime_status = "completed"
                runtime_run = self._agent_runtime().begin_run(next_message, mode="retest")
                self._publish("status", "Agent 正在处理下一条消息", "上一条消息处理完成，继续处理排队消息。", "info", metadata={"role": "agent", "turnId": turn_id, "sessionPatch": self._session_patch({"isRunning": True, "resumeState": None})})
                self._handle_instruction(next_message, False, turn_id)
        except RetestAIBlockedError as exc:
            runtime_status = "blocked"
            self._block(exc, turn_id)
        except Exception as exc:
            self.logs.append(traceback.format_exc())
            if _is_ai_runtime_block_message(exc):
                runtime_status = "blocked"
                self._block(RetestAIBlockedError(str(exc), "session_react"), turn_id)
            else:
                runtime_status = "failed"
                self._publish("error", "Agent 会话失败", str(exc), "error", metadata={"turnId": turn_id, "sessionPatch": self._session_patch({"isRunning": False, "status": f"Agent 会话失败: {exc}", "resumeState": None})})
        finally:
            with self.lock:
                if not self._is_turn_current_locked(turn_id):
                    if runtime_run:
                        self._agent_runtime().finish_run(runtime_run, "stopped")
                    return
                self.running = False
                self.updated_at = time.time()
                final_blocked = self.blocked
                final_stopped = self.stopped
            if final_blocked:
                runtime_status = "blocked"
                self._publish("status", self.blocked_title or "等待继续", self._status_locked(), "warn", metadata={"turnId": turn_id, "sessionPatch": self._session_patch({"isRunning": False, "resumeState": self._resume_state_locked(True)})})
            elif final_stopped:
                runtime_status = "stopped"
            else:
                self._publish("status", "Agent 空闲", self._status_locked(), "info", metadata={"turnId": turn_id, "sessionPatch": self._session_patch({"isRunning": False, "resumeState": None})})
            if runtime_run:
                self._agent_runtime().finish_run(runtime_run, runtime_status)

    def _handle_instruction(self, message: str, reset_queue: bool, turn_id: str, direct_queue: bool = False, use_progress_evidence: bool = True) -> None:
        """会话级 ReAct 入口：用户消息先经过模型理解，由模型自主调用会话工具。

        旧的关键词路由（if "报告"/"重测"/"工具" in message）已被完整 ReAct 循环取代；
        模型通过 list_reports / retest_report / retest_all_reports / retest_url /
        generate_reports / install_tools / tool_status 等工具完成动作。
        """
        ai_config = _ensure_retest_ai_ready("config")
        # 用户明确表达"重测/再跑一遍"时，先清空既有队列进度，让模型可以从头复测
        if reset_queue:
            with self.lock:
                self._reset_retest_state_locked()
                self.blocked = False
        if direct_queue and not reset_queue:
            with self.lock:
                if use_progress_evidence:
                    self._apply_frontend_progress_hints_locked()
                start = self.next_index
                total = len(self.source_files)
                next_name = Path(self.source_files[start]).name if 0 <= start < total else ""
                completed_count = len(self._completed_file_display_names_locked()) if use_progress_evidence else 0
                disk_completed_count = len(self.disk_completed_file_names) if use_progress_evidence else 0
            self._publish(
                "status",
                "断点续跑" if use_progress_evidence and (completed_count or disk_completed_count or start > 0) else "一键复测队列",
                (
                    f"已恢复旧会话进度，直接从第 {start + 1} 份未完成通报继续: {next_name}"
                    f"\n结构化已完成证据 {completed_count} 份，磁盘复测报告证据 {disk_completed_count} 份。"
                    if use_progress_evidence and next_name
                    else "已进入一键复测批量队列，将读取通报清单并按进度证据跳过已完成通报。"
                ),
                "ok",
                metadata={"role": "agent", "turnId": turn_id, "roundId": turn_id, "phase": "frontend_context_restore"},
            )
            result_text = self.tool_retest_all_reports(generate_reports=self.generate_reports, use_progress_evidence=use_progress_evidence, turn_id=turn_id)
            final_reply = f"已按断点继续复测。\n{result_text}"
            self._publish(
                "chat",
                "Agent",
                final_reply[:6000],
                "ok",
                metadata={"role": "agent", "turnId": turn_id, "roundId": turn_id, "phase": "frontend_context_restore"},
            )
            with self.lock:
                self.conversation = [
                    *self.conversation,
                    {"role": "user", "content": message},
                    {"role": "assistant", "content": final_reply[:6000]},
                ][-80:]
            return
        from modules.AI_Testing.retest.retest_session_agent import RetestSessionAgent

        with self.lock:
            prior_messages = list(self.conversation)

        agent = RetestSessionAgent(self, ai_config)
        _reply, persisted_messages = agent.run_turn(message, turn_id, prior_messages=prior_messages)

        # 回存本轮结束后的完整消息历史（不含 system），供下一轮继续对话。
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return
            self.conversation = persisted_messages

    # ============================ 会话级 ReAct 工具适配层 ============================
    # 下列 tool_* 方法是 RetestSessionAgent 的副作用出口：模型决定调用哪个工具，
    # 这里负责真正执行（跑流水线 / 改会话状态 / 推 WebSocket 事件），并返回一段
    # 文本结果回灌给模型。所有状态变更仍在 self.lock 内完成。

    def tool_session_state(self) -> Dict[str, Any]:
        with self.lock:
            if self.source_files:
                self._advance_next_index_past_completed_locked()
            completed = self._completed_file_display_names_locked()
            return {
                "workspace_root": self.workspace_root,
                "target_dir": self.target_dir,
                "has_target_dir": bool(self.target_dir),
                "source_files": [Path(item).name for item in self.source_files],
                "total_reports": len(self.source_files),
                "next_index": self.next_index,
                "completed_reports": completed,
                "completed_count": len(completed),
                "frontend_completed_count": len(self.frontend_completed_file_names),
                "disk_completed_count": len(self.disk_completed_file_names),
                "disk_completed_reports": list(self.disk_completed_file_names),
                "sandbox_workspace_root": str(self._agent_runtime().workspace_root),
                "generate_reports_default": self.generate_reports,
            }

    def _ensure_source_files_loaded(self, turn_id: str, use_progress_evidence: bool = True) -> None:
        with self.lock:
            loaded = bool(self.source_files)
        if not loaded:
            self._load_source_files(turn_id, use_progress_evidence=use_progress_evidence)
            return
        if not use_progress_evidence:
            with self.lock:
                self.next_index = max(0, min(self.next_index, len(self.source_files)))
            return
        with self.lock:
            previous_index = self.next_index
            recovered_next_index = self._apply_frontend_progress_hints_locked()
        if recovered_next_index > previous_index:
            self._publish(
                "status", "断点恢复",
                f"已根据旧会话事件恢复断点：跳过前 {recovered_next_index} 份已完成通报。",
                "ok", metadata={"turnId": turn_id, "roundId": turn_id, "phase": "frontend_context_restore"},
            )

    def _completed_file_display_names_locked(self) -> List[str]:
        seen: set[str] = set()
        names: List[str] = []
        for item in self.completion_items:
            if not isinstance(item, dict):
                continue
            if str(item.get("status") or "").lower() not in {"clean", "risk"}:
                continue
            raw_name = str(item.get("sourceFileName") or item.get("sourceFile") or "").strip()
            file_name = _source_notice_name(raw_name)
            key = file_name.lower()
            if key and key not in seen:
                seen.add(key)
                names.append(file_name)
        for raw_name in self.disk_completed_file_names:
            file_name = _source_notice_name(raw_name)
            key = file_name.lower()
            if key and key not in seen:
                seen.add(key)
                names.append(file_name)
        for raw_name in self.frontend_completed_file_names:
            file_name = _source_notice_name(raw_name)
            key = file_name.lower()
            if key and key not in seen:
                seen.add(key)
                names.append(file_name)
        return names

    def _completed_file_name_keys_locked(self) -> set[str]:
        return {name.lower() for name in self._completed_file_display_names_locked() if name}

    @staticmethod
    def _completion_item_source(item: Dict[str, Any]) -> str:
        return str(item.get("sourceFile") or item.get("source_file") or "").strip()

    def _completion_items_match_locked(self, left: Dict[str, Any], right: Dict[str, Any]) -> bool:
        left_source = self._completion_item_source(left)
        right_source = self._completion_item_source(right)
        if left_source and right_source and _retest_target_key(left_source) == _retest_target_key(right_source):
            return True
        left_name = _source_notice_name(left_source or left.get("sourceFileName")).lower()
        right_name = _source_notice_name(right_source or right.get("sourceFileName")).lower()
        if not left_name or left_name != right_name:
            return False
        same_name_count = sum(
            1 for source_file in self.source_files
            if _source_notice_name(source_file).lower() == left_name
        )
        return same_name_count <= 1

    def _upsert_completion_item_locked(self, item: Dict[str, Any]) -> None:
        if not isinstance(item, dict):
            return
        normalized = dict(item)
        self.completion_items = [
            existing for existing in self.completion_items
            if not isinstance(existing, dict) or not self._completion_items_match_locked(existing, normalized)
        ]
        self.completion_items.append(normalized)

    def _set_completion_items_locked(self, items: Iterable[Any]) -> None:
        self.completion_items = []
        for item in items:
            if isinstance(item, dict):
                self._upsert_completion_item_locked(item)

    def _completed_file_path_keys_locked(self) -> set[str]:
        paths: set[str] = set()
        for item in self.completion_items:
            if not isinstance(item, dict):
                continue
            if str(item.get("status") or "").lower() not in {"clean", "risk"}:
                continue
            source_file = str(item.get("sourceFile") or item.get("source_file") or "").strip()
            if source_file:
                paths.add(_retest_target_key(source_file))
        for item in self.disk_completed_report_evidence:
            if not isinstance(item, dict):
                continue
            source_file = str(item.get("source_file") or item.get("sourceFile") or "").strip()
            if source_file:
                paths.add(_retest_target_key(source_file))
        return {item for item in paths if item}

    def _source_file_has_completion_evidence_locked(self, index: int) -> bool:
        if index < 0 or index >= len(self.source_files):
            return False
        source_file = self.source_files[index]
        source_key = _retest_target_key(source_file)
        source_name = _source_notice_name(source_file).lower()
        same_name_count = sum(
            1 for item in self.source_files
            if _source_notice_name(item).lower() == source_name
        )
        for item in reversed(self.completion_items):
            if not isinstance(item, dict):
                continue
            item_source = self._completion_item_source(item)
            item_name = _source_notice_name(item_source or item.get("sourceFileName")).lower()
            same_path = bool(item_source and _retest_target_key(item_source) == source_key)
            unique_name_match = bool(source_name and item_name == source_name and same_name_count == 1)
            if same_path or unique_name_match:
                return str(item.get("status") or "").lower() in {"clean", "risk"}
        if _retest_target_key(source_file) in self._completed_file_path_keys_locked():
            return True
        return same_name_count == 1 and source_name in self._completed_file_name_keys_locked()

    def _current_file_resume_index_locked(self) -> int | None:
        current = self.current_file_resume if isinstance(self.current_file_resume, dict) else {}
        if not current or not self.source_files:
            return None
        snapshot = _as_record(current.get("resumeSnapshot") or current.get("resume_snapshot"))
        current_source = str(
            current.get("sourceFile")
            or current.get("source_file")
            or snapshot.get("source_file")
            or ""
        ).strip()
        if current_source:
            current_key = _retest_target_key(current_source)
            for index, source_file in enumerate(self.source_files):
                if _retest_target_key(source_file) == current_key:
                    return index
            current_name = _source_notice_name(current_source).lower()
            matches = [
                index for index, source_file in enumerate(self.source_files)
                if _source_notice_name(source_file).lower() == current_name
            ]
            if current_name and len(matches) == 1:
                return matches[0]
            return None
        current_index = _as_int(current.get("index"), -1)
        return current_index if 0 <= current_index < len(self.source_files) else None

    def _merge_completed_file_names_locked(self, names: Iterable[Any], *, source: str = "frontend") -> int:
        existing = list(self.disk_completed_file_names if source == "disk" else self.frontend_completed_file_names)
        seen = {name.lower() for name in existing if name}
        merged = existing
        for raw_name in names:
            file_name = _source_notice_name(raw_name)
            key = file_name.lower()
            if not key or key in seen:
                continue
            seen.add(key)
            merged.append(file_name)
        if source == "disk":
            self.disk_completed_file_names = merged
        else:
            self.frontend_completed_file_names = merged
        return len(merged)

    def _apply_frontend_progress_hints_locked(self) -> int:
        completed_names = self._completed_file_display_names_locked()
        completed_name_count = len(completed_names)
        numeric_hint = max(self.frontend_completed_count_hint, self.frontend_next_index_hint)
        hinted_index = max(0, self.next_index)
        if self.source_files:
            total = len(self.source_files)
            current_index = self._current_file_resume_index_locked()
            next_name = Path(str(self.frontend_next_source_file_name or "")).name.lower()
            if current_index is not None:
                hinted_index = current_index
            elif completed_names:
                for index, _source_file in enumerate(self.source_files):
                    if not self._source_file_has_completion_evidence_locked(index):
                        hinted_index = index
                        break
                else:
                    hinted_index = total
            elif next_name:
                name_matches = [
                    index for index, source_file in enumerate(self.source_files)
                    if Path(source_file).name.lower() == next_name
                ]
                matched_next_name = len(name_matches) == 1
                if matched_next_name and all(
                    self._source_file_has_completion_evidence_locked(index)
                    for index in range(name_matches[0])
                ):
                    hinted_index = name_matches[0]
                else:
                    hinted_index = 0
            else:
                if numeric_hint and self.next_index <= numeric_hint:
                    hinted_index = 0
                else:
                    hinted_index = max(0, self.next_index)
            self.next_index = max(0, min(total, hinted_index))
        else:
            if completed_names:
                hinted_index = max(hinted_index, completed_name_count)
            elif numeric_hint and self.next_index <= numeric_hint and not self.current_file_resume:
                hinted_index = 0
            self.next_index = max(0, hinted_index)
        return self._advance_next_index_past_completed_locked()

    def _effective_next_index_locked(self) -> int:
        if not self.source_files:
            return max(0, self.next_index)
        current_index = self._current_file_resume_index_locked()
        if current_index is not None:
            return current_index
        for index in range(len(self.source_files)):
            if not self._source_file_has_completion_evidence_locked(index):
                return index
        return len(self.source_files)

    def _advance_next_index_past_completed_locked(self) -> int:
        next_index = self._effective_next_index_locked()
        self.next_index = max(0, next_index)
        return self.next_index

    def _source_file_completed_locked(self, index: int) -> bool:
        if index < 0 or index >= len(self.source_files):
            return False
        return self._source_file_has_completion_evidence_locked(index)

    def _recover_next_index_from_frontend_completed_locked(self) -> int:
        return self._apply_frontend_progress_hints_locked()

    def tool_list_reports(self, turn_id: str, use_progress_evidence: bool = True) -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再扫描通报目录。"
            if not self.target_dir:
                return "当前会话还没有通报目录。请用户从一键复测入口选择目录，或在对话里提供 target_dir，再列通报。"
        self._ensure_source_files_loaded(turn_id, use_progress_evidence=use_progress_evidence)
        with self.lock:
            if use_progress_evidence:
                self._apply_frontend_progress_hints_locked()
            files = list(self.source_files)
            next_index = self.next_index
            completion_flags = [
                use_progress_evidence and self._source_file_completed_locked(index)
                for index in range(len(files))
            ]
        if not files:
            return "通报目录下没有发现可复测的通报文档（Word 报告）。"
        next_label = f"第 {next_index + 1} 份" if next_index < len(files) else "队列末尾"
        lines = [f"通报目录: {self.target_dir}", f"共 {len(files)} 份通报，断点在{next_label}："]
        for idx, item in enumerate(files, 1):
            done = "✓已复测" if idx - 1 < next_index or completion_flags[idx - 1] else "待复测"
            lines.append(f"{idx}. [{done}] {Path(item).name}")
        return "\n".join(lines)

    def _resolve_report_index(self, file_index: Any, file_name: str) -> int:
        with self.lock:
            self._advance_next_index_past_completed_locked()
            files = list(self.source_files)
            next_index = self.next_index
        if file_name:
            target = str(file_name).strip().lower()
            for idx, item in enumerate(files):
                if Path(item).name.lower() == target or target in Path(item).name.lower():
                    return idx
            return -1
        if file_index is not None:
            try:
                idx = int(file_index) - 1
            except Exception:
                return -1
            return idx if 0 <= idx < len(files) else -1
        # 未指定则取下一份未完成的通报
        return next_index if next_index < len(files) else -1

    def tool_retest_report(
        self,
        file_index: Any = None,
        file_name: str = "",
        generate_report: bool = False,
        use_progress_evidence: bool = True,
        turn_id: str = "",
    ) -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再执行复测。"
            if not self.target_dir:
                return "当前会话没有通报目录，无法复测通报文档。可改用 retest_url 对具体 URL 现场取证。"
        self._ensure_source_files_loaded(turn_id, use_progress_evidence=use_progress_evidence)
        index = self._resolve_report_index(file_index, file_name)
        if index < 0:
            return f"没找到要复测的通报（file_index={file_index}, file_name={file_name}）。可先调用 list_reports 查看清单。"
        with self.lock:
            file_label = Path(self.source_files[index]).name if 0 <= index < len(self.source_files) else ""
            resume_snapshot = self._resume_snapshot_for_source_locked(index, self.source_files[index]) if 0 <= index < len(self.source_files) else {}
            resume_stage = str(resume_snapshot.get("stage") or resume_snapshot.get("resume_stage") or "").strip().lower()
            pending_report_resume = resume_stage in {"report", "report_generation"}
            already_completed = use_progress_evidence and self._source_file_completed_locked(index) and not pending_report_resume
            if use_progress_evidence:
                self._apply_frontend_progress_hints_locked()
        if already_completed:
            return f"{file_label or '这份通报'} 已在当前/旧会话中完成复测，已跳过；如需重测，请先说『重新复测』。"
        outcome = self._retest_single_file(index, bool(generate_report), turn_id)
        return outcome.get("message") or "复测完成。"

    def tool_retest_all_reports(self, generate_reports: bool = False, use_progress_evidence: bool = True, turn_id: str = "") -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再执行批量复测。"
            if not self.target_dir:
                return "当前会话没有通报目录，无法批量复测。可改用 retest_url 对具体 URL 现场取证。"
        self._ensure_source_files_loaded(turn_id, use_progress_evidence=use_progress_evidence)
        with self.lock:
            total = len(self.source_files)
            start = self._apply_frontend_progress_hints_locked() if use_progress_evidence else self.next_index
            completed_names = self._completed_file_display_names_locked() if use_progress_evidence else []
            disk_completed_count = len(self.disk_completed_file_names) if use_progress_evidence else 0
            next_name = Path(self.source_files[start]).name if 0 <= start < total else ""
        if total <= 0:
            return "通报目录下没有可复测的通报文档。"
        if start >= total:
            return f"全部 {total} 份通报都已复测完成。如需重测，请告诉用户说『重新复测』。"
        self._publish(
            "thought_summary", "Agent 计划",
            (
                f"按队列从第 {start + 1} 份开始，依次复测剩余 {total - start} 份通报。"
                f"\n结构化已完成证据: {len(completed_names)} 份；磁盘复测报告证据: {disk_completed_count} 份。"
                f"\n下一份未完成通报: {next_name}"
            ),
            "info",
            metadata={
                "turnId": turn_id,
                "roundId": turn_id,
                "phase": "planning",
                "progressEvidence": {
                    "targetDir": self.target_dir,
                    "completedFileNames": completed_names,
                    "completedCountHint": len(completed_names),
                    "nextIndexHint": start,
                    "nextSourceFileName": next_name,
                },
            },
        )
        done = 0
        while True:
            with self.lock:
                if self._execution_cancelled_locked(turn_id):
                    break
                if use_progress_evidence:
                    self._apply_frontend_progress_hints_locked()
                index = self.next_index
                if index >= len(self.source_files):
                    break
            outcome = self._retest_single_file(index, bool(generate_reports), turn_id, use_progress_evidence=use_progress_evidence)
            if outcome.get("status") == "stopped":
                return "会话已停止，已保留断点；继续时会跳过已完成通报。"
            if outcome.get("status") != "skipped":
                done += 1
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，已保留断点；继续时会跳过已完成通报。"
        overview = _format_agent_completion_overview(self.completion_items)
        self._publish(
            "artifact", "复测结论总览", overview,
            "warn" if any(item.get("status") == "risk" for item in self.completion_items) else "ok",
            metadata={"turnId": turn_id, "roundId": turn_id, "phase": "completion_summary",
                      "completionItems": list(self.completion_items),
                      "sessionPatch": self._session_patch({"status": "复测完成", "progress": 100, "resumeState": None})},
        )
        return f"已复测 {done} 份通报。\n{overview}"

    def tool_retest_url(self, url: str, vuln_types: List[str], note: str = "", turn_id: str = "") -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再执行现场取证。"
        target = str(url or "").strip()
        if not _valid_http_target(target):
            return f"URL 无效或不是 http/https 目标: {url}"
        return self._retest_adhoc_url(target, vuln_types or [], note, turn_id)

    def tool_inspect_session_state(self, turn_id: str = "") -> str:
        state = self.tool_session_state()
        lines = [
            f"workspaceRoot: {state.get('workspace_root') or '-'}",
            f"sandboxWorkspaceRoot: {state.get('sandbox_workspace_root') or '-'}",
            f"targetDir: {state.get('target_dir') or '-'}",
            f"sourceFiles: {state.get('total_reports', 0)}",
            f"nextIndex: {state.get('next_index', 0)}",
            f"completed: {state.get('completed_count', 0)}",
            f"diskCompletedReports: {state.get('disk_completed_count', 0)}",
            f"generateReportsDefault: {bool(state.get('generate_reports_default'))}",
        ]
        disk_reports = state.get("disk_completed_reports") if isinstance(state.get("disk_completed_reports"), list) else []
        if disk_reports:
            lines.append("diskCompletedNames:\n" + "\n".join(str(item) for item in disk_reports[:50]))
        return "\n".join(lines)

    def tool_run_retest_queue(
        self,
        generate_reports: bool = False,
        use_progress_evidence: bool = True,
        turn_id: str = "",
    ) -> str:
        return self.tool_retest_all_reports(
            generate_reports=bool(generate_reports),
            use_progress_evidence=bool(use_progress_evidence),
            turn_id=turn_id,
        )

    def tool_delete_reports(self, turn_id: str = "") -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再删除旧报告。"
            target_dir = self.target_dir
        if not target_dir:
            return "当前会话没有通报目录，无法定位旧复测报告。"
        root = Path(target_dir).expanduser()
        if not root.exists() or not root.is_dir():
            return f"通报目录不存在，无法删除旧复测报告: {target_dir}"
        candidates: List[Path] = []
        for item in root.rglob("*"):
            try:
                if item.is_file() and item.suffix.lower() in WORD_SUFFIXES and is_generated_retest_report_path(item):
                    candidates.append(item)
            except Exception:
                continue
        deleted: List[str] = []
        failures: List[str] = []
        for path in candidates:
            try:
                path.unlink()
                deleted.append(str(path))
            except Exception as exc:
                failures.append(f"{path}: {exc}")
        with self.lock:
            self.disk_completed_file_names = []
            self.disk_completed_report_evidence = []
        if not candidates:
            return "没有找到可删除的旧复测报告。"
        lines = [f"已删除旧复测报告 {len(deleted)} 份，失败 {len(failures)} 份。"]
        if deleted:
            lines.append("deleted:\n" + "\n".join(deleted[:50]))
        if failures:
            lines.append("failures:\n" + "\n".join(failures[:20]))
        return "\n".join(lines)

    def tool_generate_reports(self, file_name: str = "", turn_id: str = "") -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再生成报告。"
            completed = [str(item.get("sourceFile") or "") for item in self.completion_items if item.get("sourceFile")]
        if not completed:
            return "当前会话还没有已完成的复测结果，无法生成报告。请先复测。"
        if file_name:
            target = str(file_name).strip().lower()
            match = next((item for item in completed if target in Path(item).name.lower()), "")
            if not match:
                return f"没找到已复测的通报 {file_name}，无法单独生成报告。"
            summary = self.report_evidence_summaries.get(match) or ""
            reports = self._generate_report_for_file(match, summary, turn_id, result_data=self.report_result_data.get(match))
            return f"已为 {Path(match).name} 生成 {len(reports)} 份报告。" if reports else "报告生成失败或无输出。"
        self._generate_reports_for_completed(turn_id)
        with self.lock:
            count = len(self.reports)
        return f"已为本会话已完成的通报生成报告，共 {count} 份。"

    def tool_install_tools(self, tools: List[str], turn_id: str = "") -> str:
        with self.lock:
            if self._execution_cancelled_locked(turn_id):
                return "会话已停止，不再下载外部工具。"
        selected = [t for t in (tools or []) if t in ("nmap", "sqlmap", "ffuf")] or ["nmap", "sqlmap", "ffuf"]
        message = "下载工具 " + " ".join(selected)
        self._install_external_tools(message, turn_id)
        result = _doc_retest_tools_status({})
        lines = []
        for item in result.get("tools") or []:
            lines.append(f"{item.get('name')}: {'已配置' if item.get('installed') else '未配置'}")
        return "外部工具安装流程已执行。当前状态:\n" + ("\n".join(lines) or "未知")

    def tool_tool_status(self, turn_id: str = "") -> str:
        self._show_external_tool_status(turn_id)
        result = _doc_retest_tools_status({})
        lines = []
        for item in result.get("tools") or []:
            command = " ".join(str(part) for part in item.get("command") or [])
            lines.append(f"{item.get('name')}: {'已配置' if item.get('installed') else '未配置'} {command}".rstrip())
        return "\n".join(lines) or (result.get("message") or "未获取到工具状态。")

    def _retest_single_file(self, index: int, generate_report: bool, turn_id: str, use_progress_evidence: bool = True) -> Dict[str, Any]:
        """复测单份通报（从队列循环抽取），事件/状态行为与批量队列保持一致。

        返回 {"status": ..., "message": ...}，message 会回灌给会话模型。
        """
        with self.lock:
            if index < 0 or index >= len(self.source_files):
                return {"status": "failed", "message": f"通报序号越界: {index + 1}"}
            source_file = self.source_files[index]
            resume_snapshot = self._resume_snapshot_for_source_locked(index, source_file)
            resume_stage = str(resume_snapshot.get("stage") or resume_snapshot.get("resume_stage") or "").strip().lower()
            pending_report_resume = resume_stage in {"report", "report_generation"}
            if use_progress_evidence and self._source_file_completed_locked(index) and not pending_report_resume:
                if index <= self.next_index:
                    self.next_index = max(self.next_index, index + 1)
                    self._advance_next_index_past_completed_locked()
                return {"status": "skipped", "message": f"{Path(source_file).name} 已完成复测，本轮跳过。"}
            total = len(self.source_files)
        file_path = Path(source_file)
        round_id = f"{turn_id}:file:{index + 1}"
        self._publish(
            "chat", f"通报 {index + 1}/{total}", f"开始复测: {file_path.name}", "info",
            metadata={"role": "agent", "turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name,
                      "sessionPatch": self._session_patch({"status": f"Agent 正在复测 ({index + 1}/{total}): {file_path.name}", "progress": int(index / max(1, total) * 100), "isRunning": True, "resumeState": None})},
        )
        file_logs: List[str] = []

        def on_event(event: Dict[str, Any]) -> None:
            try:
                with self.lock:
                    if self._execution_cancelled_locked(turn_id):
                        return
                from modules.backend_api.retest_event_stream import publish_retest_event

                # 内层复测产生的 tool_call/tool_result/status 默认只带 phase，缺 turnId/roundId。
                # 这里统一补全，保证前端能把这些事件归属到本轮、并让 tool_call↔tool_result 稳定配对。
                meta = event.get("metadata")
                if not isinstance(meta, dict):
                    meta = {}
                    event["metadata"] = meta
                meta.setdefault("turnId", turn_id)
                meta.setdefault("roundId", round_id)
                if not event.get("sourceFile"):
                    event["sourceFile"] = str(file_path)
                meta.setdefault("sourceFileName", file_path.name)

                publish_retest_event({"type": "retest_trace_event", "session_id": self.session_id, "task_id": "agent", "event": event})
            except Exception:
                pass

        def on_checkpoint(snapshot: Dict[str, Any]) -> None:
            self._save_current_file_checkpoint(snapshot, source_file, turn_id, round_id)

        try:
            with self.lock:
                frontend_context = dict(self.frontend_context_payload or {})
            run_payload = {
                "round_id": round_id,
                "turn_id": turn_id,
                "source_file_name": file_path.name,
                "session_id": self.session_id,
                "frontend_context": frontend_context,
            }
            if resume_snapshot:
                run_payload["resume_snapshot"] = resume_snapshot
                resume_stage = str(resume_snapshot.get("stage") or resume_snapshot.get("resume_stage") or "unknown").strip().lower()
                self.logs.append(f"{file_path.name} resume from {resume_stage or 'unknown'} snapshot")
            summary, result_data, _manual, _events = _run_retest_for_source_file(
                file_path,
                run_payload,
                file_logs,
                event_callback=on_event,
                stop_check=lambda: self._turn_is_cancelled(turn_id),
                checkpoint_callback=on_checkpoint,
            )
            with self.lock:
                if self._execution_cancelled_locked(turn_id):
                    return {"status": "stopped", "message": "会话已停止，不再记录本份通报结果。"}
            self.logs.extend(file_logs)
            result = {"success": True, "message": f"复测完成: {file_path.name}", "summary": summary, "result_data": result_data, "logs": file_logs}
            report_summary = _format_report_evidence_screenshot_text(summary, result_data)
            report_paths: List[str] = []
            if generate_report:
                report_snapshot = _report_resume_snapshot(source_file, summary, result_data, report_summary)
                report_resume = self._current_file_resume_from_snapshot_locked(report_snapshot, source_file)
                with self.lock:
                    if report_resume:
                        self.current_file_resume = report_resume
                    self.report_evidence_summaries[source_file] = report_summary
                    self.report_result_data[source_file] = result_data
                    self.latest_result_data = result_data
                    resume_patch = self._session_patch({
                        "status": f"Agent is generating report ({index + 1}/{total}): {file_path.name}",
                        "resultText": summary,
                        "latestResultData": result_data,
                        "resumeState": self._resume_state_locked(True),
                    })
                self._publish(
                    "status",
                    "Report checkpoint saved",
                    "Retest judgement is complete; report generation can resume without rerunning probes.",
                    "ok",
                    metadata={
                        "turnId": turn_id,
                        "roundId": round_id,
                        "sourceFileName": file_path.name,
                        "phase": "report",
                        "sessionPatch": resume_patch,
                    },
                )
                report_paths = self._generate_report_for_file(source_file, summary, turn_id, round_id, result_data)
                with self.lock:
                    if self._execution_cancelled_locked(turn_id):
                        return {"status": "stopped", "message": "会话已停止，不再记录本份通报结果。"}
            completion_item = _completion_item_from_result(source_file, result, report_paths)
            with self.lock:
                self.summaries.append(summary)
                self.reports.extend(report_paths)
                self.report_evidence_summaries[source_file] = report_summary
                self.report_result_data[source_file] = result_data
                self._upsert_completion_item_locked(completion_item)
                self.latest_result_data = result_data
                self.current_file_resume = None
                self.next_index = max(self.next_index, index + 1)
                self._advance_next_index_past_completed_locked()
            status = "risk" if completion_item.get("status") == "risk" else completion_item.get("status") or "clean"
            self._publish(
                "chat", "复测结果", _format_agent_result_message(source_file, result, completion_item, report_paths),
                "warn" if status == "risk" else "ok",
                metadata={"role": "agent", "turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name,
                          "fixStatus": "risk" if status == "risk" else "clean",
                          "sessionPatch": self._session_patch({"resultText": summary, "latestResultData": result_data, "progress": int((index + 1) / max(1, total) * 100), "resumeState": None})},
            )
            verdict = "漏洞可复现/未修复" if status == "risk" else "未见复现/复测通过"
            return {"status": status, "message": f"{file_path.name} 复测完成：{verdict}。" + (f" 已生成 {len(report_paths)} 份报告。" if report_paths else "")}
        except RetestAIBlockedError:
            raise
        except Exception as exc:
            with self.lock:
                if self._execution_cancelled_locked(turn_id):
                    return {"status": "stopped", "message": "会话已停止，不再记录本份通报结果。"}
            self.logs.append(traceback.format_exc())
            completion_item = _completion_item_from_result(source_file, None, [], str(exc))
            with self.lock:
                self._upsert_completion_item_locked(completion_item)
                self.next_index = min(self.next_index, index)
            self._publish(
                "error", "复测错误", f"{file_path.name} 处理失败: {exc}", "error",
                metadata={"turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name, "sessionPatch": self._session_patch()},
            )
            return {"status": "failed", "message": f"{file_path.name} 复测失败: {exc}"}

    def _retest_adhoc_url(self, url: str, vuln_types: List[str], note: str, turn_id: str) -> str:
        """对用户在对话里直接给出的 URL 现场取证（无需通报文档）。

        复用 VulnerabilityRetestScanner 的 ReAct 取证引擎 + judge 判定，事件推到当前会话流。
        """
        ai_config = _ensure_retest_ai_ready("config")
        from modules.AI_Testing.retest.vulnerability_batch_scanner import VulnerabilityRetestScanner

        round_id = f"{turn_id}:url:{uuid.uuid4().hex[:6]}"

        def on_event(event: Dict[str, Any]) -> None:
            with self.lock:
                if self._execution_cancelled_locked(turn_id):
                    return
            if isinstance(event, dict):
                metadata = event.get("metadata")
                if not isinstance(metadata, dict):
                    metadata = {}
                    event["metadata"] = metadata
                metadata.setdefault("turnId", turn_id)
                metadata.setdefault("roundId", round_id)
            try:
                from modules.backend_api.retest_event_stream import publish_retest_event

                publish_retest_event({"type": "retest_trace_event", "session_id": self.session_id, "task_id": "agent", "event": event})
            except Exception:
                pass

        self._publish(
            "chat", "现场取证", f"对用户指定 URL 启动黑盒取证: {url}", "info",
            metadata={"role": "agent", "turnId": turn_id, "roundId": round_id,
                      "sessionPatch": self._session_patch({"status": f"正在现场取证: {url}", "isRunning": True, "resumeState": None})},
        )

        context: Dict[str, Any] = {
            "target_urls": [url],
            "issue_tags": [],
            "raw_text": note or "",
        }
        scanner = VulnerabilityRetestScanner(
            timeout=15,
            max_workers=3,
            trace_callback=on_event,
            ai_config=ai_config,
            stop_check=lambda: self._turn_is_cancelled(turn_id),
        )
        try:
            result = scanner.scan_url_for_context(url, vuln_types or [], context)
        except RetestAIBlockedError:
            raise
        except Exception as exc:
            message = str(exc)
            if _is_ai_runtime_block_message(message):
                raise RetestAIBlockedError(f"AI Agent 现场取证阶段暂停: {message}", "execution") from exc
            self.logs.append(traceback.format_exc())
            self._publish(
                "error", "现场取证失败", f"{url} 取证失败: {exc}", "error",
                metadata={"turnId": turn_id, "roundId": round_id, "sessionPatch": self._session_patch()},
            )
            return f"对 {url} 现场取证失败: {exc}"

        if self._turn_is_cancelled(turn_id) or result.get("stopped"):
            return "现场取证已按用户指令停止，晚到结果未提交。"

        observations = [item for item in (result.get("vulnerabilities") or []) if isinstance(item, dict)]
        obs_count = len(observations)
        executed = result.get("context_checks") or []
        react_summary = ""
        advice = context.get("agent_advice") if isinstance(context.get("agent_advice"), dict) else {}
        if isinstance(advice, dict):
            react_summary = str(advice.get("react_summary") or "")
        lines = [f"对 {url} 的现场取证完成。", f"调用工具 {len(executed)} 次，记录 {obs_count} 条观察。"]
        if react_summary:
            lines.append(f"取证总结: {react_summary}")
        for item in observations[:8]:
            lines.append(f"- [{item.get('severity') or 'info'}] {item.get('type') or '观察'}: {str(item.get('detail') or item.get('evidence') or '')[:200]}")
        if result.get("target_unreachable"):
            lines.append("目标当前不可达。")
        summary_text = "\n".join(lines)
        self._publish(
            "chat", "现场取证结果", summary_text, "ok",
            metadata={"role": "agent", "turnId": turn_id, "roundId": round_id,
                      "sessionPatch": self._session_patch({"status": "现场取证完成", "resumeState": None})},
        )
        return (
            summary_text
            + "\n\n注意：这是现场取证的原始观察，不是对通报漏洞的最终二元判定。"
            "如需对某份通报给出 reproduced/not_reproduced 结论，请用 retest_report。"
        )

    def _load_source_files(self, turn_id: str, use_progress_evidence: bool = True) -> None:
        with self.lock:
            target_dir = self.target_dir
        if not target_dir:
            raise ValueError("当前会话没有通报目录。请先从一键复测入口启动，或在会话里提供 target_dir。")
        tool_call_id = f"list:{turn_id}"
        self._publish("tool_call", "列通报", f"扫描通报目录: {target_dir}", "info", tool={"toolId": "list_reports", "label": "列通报", "status": "running", "target": target_dir, "argsPreview": target_dir}, metadata={"toolCallId": tool_call_id, "turnId": turn_id, "roundId": turn_id, "phase": "tool"})
        result = _doc_retest_list_files({"target_dir": target_dir})
        self.logs.extend(result.get("logs") or [])
        source_files = result.get("source_files") or []
        disk_completed_names = result.get("completed_source_file_names") or []
        disk_report_evidence = result.get("existing_report_evidence") if isinstance(result.get("existing_report_evidence"), list) else []
        with self.lock:
            previous_index = self.next_index
            self.source_files = [str(item) for item in source_files]
            self.next_index = min(self.next_index, len(self.source_files))
            if use_progress_evidence:
                self._merge_completed_file_names_locked(disk_completed_names, source="disk")
            else:
                self.disk_completed_file_names = []
            self.disk_completed_report_evidence = [item for item in disk_report_evidence if isinstance(item, dict)]
            recovered_next_index = self._recover_next_index_from_frontend_completed_locked() if use_progress_evidence else self.next_index
        if use_progress_evidence and recovered_next_index > previous_index:
            self._publish(
                "status", "断点恢复",
                f"已根据结构化进度和磁盘复测报告恢复断点：跳过前 {recovered_next_index} 份已完成通报。",
                "ok", metadata={"turnId": turn_id, "roundId": turn_id, "phase": "frontend_context_restore"},
            )
        with self.lock:
            next_name = Path(self.source_files[self.next_index]).name if 0 <= self.next_index < len(self.source_files) else ""
            completed_preview = self._completed_file_display_names_locked()[:20] if use_progress_evidence else []
            progress_evidence = {
                "targetDir": target_dir,
                "completedFileNames": self._completed_file_display_names_locked() if use_progress_evidence else [],
                "completedCountHint": len(self._completed_file_display_names_locked()) if use_progress_evidence else 0,
                "nextIndexHint": self.next_index,
                "nextSourceFileName": next_name,
            }
        raw_output_lines = [str(item) for item in source_files]
        if disk_completed_names:
            raw_output_lines.append("")
            raw_output_lines.append("[已从磁盘复测报告识别完成]")
            raw_output_lines.extend(str(item) for item in disk_completed_names)
        result_preview = result.get("message") or f"发现 {len(source_files)} 份通报"
        if disk_completed_names:
            result_preview += f"；磁盘复测报告对应已完成 {len(disk_completed_names)} 份，下一份: {next_name or '队列末尾'}"
        self._publish(
            "tool_result",
            "列通报",
            result_preview,
            "ok" if source_files else "warn",
            tool={
                "toolId": "list_reports",
                "label": "列通报",
                "status": "completed",
                "target": target_dir,
                "resultPreview": result_preview,
                "rawCount": len(source_files),
                "rawOutput": "\n".join(raw_output_lines),
            },
            metadata={
                "toolCallId": tool_call_id,
                "turnId": turn_id,
                "roundId": turn_id,
                "phase": "tool",
                "progressEvidence": progress_evidence,
                "completedFileNames": completed_preview,
                "diskCompletedFileNames": list(disk_completed_names),
                "nextSourceFileName": next_name,
                "sessionPatch": self._session_patch({"targetDir": target_dir, "progress": self._progress_locked(), "resumeState": None}),
            },
        )
        if not source_files:
            raise ValueError(result.get("message") or "未找到通报文档。")

    def _run_retest_queue(self, turn_id: str) -> None:
        _ensure_retest_ai_ready("config")
        with self.lock:
            if not self.source_files:
                self._load_source_files(turn_id)
            self._apply_frontend_progress_hints_locked()
            total = len(self.source_files)
            start = self.next_index
        if total <= 0:
            return
        self._publish("thought_summary", "Agent 计划", f"我会按当前队列逐份读取通报、调用工具复测，再由 AI 给出二元结论。队列共 {total} 份，断点位置 {start + 1}。", "info", metadata={"turnId": turn_id, "roundId": turn_id, "phase": "planning", "sessionPatch": self._session_patch({"isRunning": True, "resumeState": None})})
        while True:
            with self.lock:
                if self._execution_cancelled_locked(turn_id):
                    return
                self._apply_frontend_progress_hints_locked()
                index = self.next_index
                if index >= len(self.source_files):
                    break
                if self._source_file_has_completion_evidence_locked(index):
                    # Defensive queue invariant: restored/non-contiguous
                    # completion evidence must never start another retest.
                    self.next_index = index + 1
                    self._advance_next_index_past_completed_locked()
                    continue
                source_file = self.source_files[index]
                generate_reports = self.generate_reports
            file_path = Path(source_file)
            round_id = f"{turn_id}:file:{index + 1}"
            self._publish("chat", f"通报 {index + 1}/{total}", f"开始复测: {file_path.name}", "info", metadata={"role": "agent", "turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name, "sessionPatch": self._session_patch({"status": f"Agent 正在复测 ({index + 1}/{total}): {file_path.name}", "progress": int(index / max(1, total) * 100), "isRunning": True, "resumeState": None})})
            file_logs: List[str] = []
            with self.lock:
                resume_snapshot = self._resume_snapshot_for_source_locked(index, source_file)

            def on_event(event: Dict[str, Any]) -> None:
                try:
                    with self.lock:
                        if self._execution_cancelled_locked(turn_id):
                            return
                    from modules.backend_api.retest_event_stream import publish_retest_event

                    publish_retest_event({"type": "retest_trace_event", "session_id": self.session_id, "task_id": "agent", "event": event})
                except Exception:
                    pass

            def on_checkpoint(snapshot: Dict[str, Any]) -> None:
                self._save_current_file_checkpoint(snapshot, source_file, turn_id, round_id)

            try:
                with self.lock:
                    frontend_context = dict(self.frontend_context_payload or {})
                run_payload = {
                    "round_id": round_id,
                    "turn_id": turn_id,
                    "source_file_name": file_path.name,
                    "session_id": self.session_id,
                    "frontend_context": frontend_context,
                }
                if resume_snapshot:
                    run_payload["resume_snapshot"] = resume_snapshot
                summary, result_data, _manual, _events = _run_retest_for_source_file(
                    file_path,
                    run_payload,
                    file_logs,
                    event_callback=on_event,
                    stop_check=lambda: self._turn_is_cancelled(turn_id),
                    checkpoint_callback=on_checkpoint,
                )
                # A forced resume invalidates the previous turn. Do not let an
                # in-flight old worker commit a result into the new queue.
                with self.lock:
                    if self._execution_cancelled_locked(turn_id):
                        return
                self.logs.extend(file_logs)
                result = {"success": True, "message": f"复测完成: {file_path.name}", "summary": summary, "result_data": result_data, "logs": file_logs}
                report_summary = _format_report_evidence_screenshot_text(summary, result_data)
                report_paths: List[str] = []
                if generate_reports:
                    self._save_current_file_checkpoint(
                        _report_resume_snapshot(source_file, summary, result_data, report_summary),
                        source_file,
                        turn_id,
                        round_id,
                    )
                    report_paths = self._generate_report_for_file(source_file, summary, turn_id, round_id, result_data)
                    with self.lock:
                        if self._execution_cancelled_locked(turn_id):
                            return
                completion_item = _completion_item_from_result(source_file, result, report_paths)
                with self.lock:
                    self.summaries.append(summary)
                    self.reports.extend(report_paths)
                    self.report_evidence_summaries[source_file] = report_summary
                    self.report_result_data[source_file] = result_data
                    self._upsert_completion_item_locked(completion_item)
                    self.latest_result_data = result_data
                    self.current_file_resume = None
                    self.next_index = index + 1
                    self._advance_next_index_past_completed_locked()
                self._publish("chat", "复测结果", _format_agent_result_message(source_file, result, completion_item, report_paths), "warn" if completion_item.get("status") == "risk" else "ok", metadata={"role": "agent", "turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name, "fixStatus": "risk" if completion_item.get("status") == "risk" else "clean", "sessionPatch": self._session_patch({"resultText": summary, "latestResultData": result_data, "progress": int((index + 1) / max(1, total) * 100), "resumeState": None})})
            except RetestAIBlockedError:
                raise
            except Exception as exc:
                with self.lock:
                    if self._execution_cancelled_locked(turn_id):
                        return
                self.logs.append(traceback.format_exc())
                completion_item = _completion_item_from_result(source_file, None, [], str(exc))
                with self.lock:
                    self._upsert_completion_item_locked(completion_item)
                    self.next_index = min(self.next_index, index)
                self._publish("error", "复测错误", f"{file_path.name} 处理失败: {exc}", "error", metadata={"turnId": turn_id, "roundId": round_id, "sourceFileName": file_path.name, "sessionPatch": self._session_patch()})
        overview = _format_agent_completion_overview(self.completion_items)
        final_result = "\n\n".join([overview] + self.summaries + (["生成报告:\n" + "\n".join(self.reports)] if self.reports else []))
        with self.lock:
            self.blocked = False
            self.blocked_reason = ""
            self.blocked_stage = ""
            self.blocked_title = ""
        self._publish("artifact", "复测结论总览", overview, "warn" if any(item.get("status") == "risk" for item in self.completion_items) else "ok", metadata={"turnId": turn_id, "roundId": turn_id, "phase": "completion_summary", "completionItems": list(self.completion_items), "sessionPatch": self._session_patch({"status": "复测完成", "progress": 100, "resultText": final_result, "lastReportPath": self.reports[0] if self.reports else self.target_dir, "resumeState": None})})

    def _generate_report_for_file(
        self,
        source_file: str,
        summary: str,
        turn_id: str = "",
        round_id: str = "",
        result_data: Dict[str, Any] | None = None,
    ) -> List[str]:
        if self._turn_is_cancelled(turn_id):
            return []
        tool_id = f"report:{self.session_id}:{Path(source_file).name}:{time.time()}"
        metadata = {"toolCallId": tool_id, "turnId": turn_id or self.current_turn_id, "roundId": round_id or turn_id or self.current_turn_id, "phase": "tool"}
        self._publish("tool_call", "生成报告", f"按用户明确要求为 {Path(source_file).name} 生成报告。", "info", tool={"toolId": "generate_report", "label": "生成报告", "status": "running", "target": source_file, "argsPreview": source_file}, metadata=metadata)
        logs: List[str] = []
        result: Dict[str, Any] = {}
        completed = threading.Event()
        source_path = Path(source_file).expanduser()
        staging_root = source_path.parent / ".koi_retest_staging"
        staging_dir = staging_root / f"{self.session_id[:16]}-{uuid.uuid4().hex}"

        def generate() -> None:
            try:
                staging_dir.mkdir(parents=True, exist_ok=True)
                result.update(_generate_retest_reports_from_agent_summary(
                    Path(self.target_dir),
                    [source_file],
                    summary,
                    logs,
                    result_data,
                    output_dir=staging_dir,
                    include_disposal_reports=True,
                ))
            except Exception as exc:
                logs.append(traceback.format_exc())
                result.update({
                    "success": False,
                    "message": f"报告生成异常: {exc}",
                    "reports": [],
                    "failures": [{"file": source_file, "name": Path(source_file).name, "reason": str(exc)}],
                    "logs": logs,
                })
            finally:
                if self._turn_is_cancelled(turn_id):
                    _cleanup_retest_staging_dir(staging_dir, logs)
                completed.set()

        # Document/Office conversion can block inside a native call.  Run it
        # behind a daemon worker so the cooperative stop signal can release the
        # Agent turn immediately; late output is deliberately discarded below.
        threading.Thread(
            target=generate,
            name=f"koi-retest-report-{self.session_id[:8]}",
            daemon=True,
        ).start()
        while not completed.wait(0.1):
            if self._turn_is_cancelled(turn_id):
                self.logs.append(f"{Path(source_file).name} 报告生成已收到停止信号，丢弃晚到结果")
                _cleanup_retest_staging_dir(staging_dir, self.logs)
                return []
        if self._turn_is_cancelled(turn_id):
            self.logs.append(f"{Path(source_file).name} 报告生成完成后检测到停止信号，丢弃晚到结果")
            shutil.rmtree(staging_dir, ignore_errors=True)
            return []
        self.logs.extend(result.get("logs") or logs)
        generated_success = bool(result.get("success"))
        staged_values = list(result.get("artifacts") or result.get("reports") or [])
        staged_paths: List[Path] = []
        resolved_root = staging_dir.resolve()
        invalid_manifest = False
        for staged_value in staged_values:
            staged_path = Path(str(staged_value)).expanduser()
            try:
                resolved_staged = staged_path.resolve()
            except Exception:
                resolved_staged = staged_path.absolute()
            if resolved_staged.parent != resolved_root or not resolved_staged.is_file():
                self.logs.append(f"跳过异常报告暂存路径: {staged_path}")
                invalid_manifest = True
                continue
            if resolved_staged not in staged_paths:
                staged_paths.append(resolved_staged)

        if not generated_success or invalid_manifest or not staged_paths:
            _cleanup_retest_staging_dir(staging_dir, self.logs)
            result["reports"] = []
            result["success"] = False
            if not result.get("message"):
                result["message"] = "报告暂存产物不完整，未提交到最终目录"
            report_summary = self.report_evidence_summaries.get(source_file) or _format_report_evidence_screenshot_text(
                summary,
                result_data or {},
            )
            raise RetestAIBlockedError(
                str(result.get("message") or "报告生成失败"),
                "report",
                _report_resume_snapshot(source_file, summary, result_data, report_summary),
            )

        if self._turn_is_cancelled(turn_id):
            self.logs.append(f"{source_path.name} 报告提交前检测到停止信号，丢弃暂存结果")
            _cleanup_retest_staging_dir(staging_dir, self.logs)
            return []

        commit_pairs = [(staged_path, source_path.parent / staged_path.name) for staged_path in staged_paths]
        backup_dir = staging_dir / ".commit_backup"
        committed: List[tuple[Path, Path]] = []
        backups: List[tuple[Path, Path]] = []
        try:
            for index, (staged_path, final_path) in enumerate(commit_pairs):
                if final_path.exists():
                    backup_dir.mkdir(parents=True, exist_ok=True)
                    backup_path = backup_dir / f"{index:04d}-{final_path.name}"
                    os.replace(str(final_path), str(backup_path))
                    backups.append((backup_path, final_path))
                staged_path.replace(final_path)
                committed.append((final_path, staged_path))
                if self._turn_is_cancelled(turn_id):
                    raise RuntimeError("报告提交过程中收到停止信号")
        except Exception as exc:
            self.logs.append(f"报告提交失败，正在回滚本次报告包: {exc}")
            for final_path, staged_path in reversed(committed):
                try:
                    if final_path.exists():
                        os.replace(str(final_path), str(staged_path))
                except Exception as rollback_exc:
                    self.logs.append(f"报告回滚失败 {final_path}: {rollback_exc}")
            for backup_path, final_path in reversed(backups):
                try:
                    if backup_path.exists():
                        os.replace(str(backup_path), str(final_path))
                except Exception as rollback_exc:
                    self.logs.append(f"原报告恢复失败 {final_path}: {rollback_exc}")
            _cleanup_retest_staging_dir(staging_dir, self.logs)
            result["reports"] = []
            result["success"] = False
            report_summary = self.report_evidence_summaries.get(source_file) or _format_report_evidence_screenshot_text(
                summary,
                result_data or {},
            )
            raise RetestAIBlockedError(
                f"报告未能提交，已回滚: {exc}",
                "report",
                _report_resume_snapshot(source_file, summary, result_data, report_summary),
            ) from exc

        reports = [str(final_path) for final_path, _ in committed]
        for final_path in reports:
            self.logs.append(f"报告已提交: {final_path}")
        _cleanup_retest_staging_dir(staging_dir, self.logs)
        result["reports"] = reports
        result["success"] = True
        artifact_content = _format_report_artifact_content(str(result.get("message") or ""), reports)
        self._publish("tool_result", "生成报告", artifact_content, "ok" if result.get("success") else "error", tool={"toolId": "generate_report", "label": "生成报告", "status": "completed" if result.get("success") else "failed", "target": source_file, "resultPreview": artifact_content, "rawOutput": "\n".join(result.get("logs") or []), "rawCount": len(reports), "failureReason": "" if result.get("success") else result.get("message")}, metadata={**metadata, "reports": reports, "sessionPatch": self._session_patch({"lastReportPath": reports[0] if reports else self.target_dir, "resumeState": None})})
        if result.get("success") or reports:
            self._publish(
                "artifact",
                "报告生成完成" if result.get("success") else "报告生成部分完成",
                artifact_content,
                "ok" if result.get("success") else "warn",
                metadata={
                    "turnId": metadata.get("turnId"),
                    "roundId": metadata.get("roundId"),
                    "phase": "artifact",
                    "reports": reports,
                    "sessionPatch": self._session_patch({"lastReportPath": reports[0] if reports else self.target_dir, "resumeState": None}),
                },
            )
        if not reports:
            report_summary = self.report_evidence_summaries.get(source_file) or _format_report_evidence_screenshot_text(
                summary,
                result_data or {},
            )
            raise RetestAIBlockedError(
                str(result.get("message") or "报告生成失败"),
                "report",
                _report_resume_snapshot(source_file, summary, result_data, report_summary),
            )
        return reports

    def _generate_reports_for_completed(self, turn_id: str) -> None:
        with self.lock:
            source_files = _source_notice_paths(str(item.get("sourceFile") or "") for item in self.completion_items if item.get("sourceFile"))
            summary = "\n\n".join(self.summaries) or _format_agent_completion_overview(self.completion_items)
        if not source_files:
            self._publish("chat", "Agent", "当前会话还没有已完成的复测结果，无法直接生成报告。", "warn", metadata={"role": "agent", "turnId": turn_id})
            return
        summary_parts: List[str] = []
        for source_file in source_files:
            evidence_summary = self.report_evidence_summaries.get(source_file)
            if evidence_summary:
                summary_parts.append(evidence_summary)
        if not summary_parts:
            summary_parts.append(summary)
        if self._turn_is_cancelled(turn_id):
            return
        reports: List[str] = []
        # Use the same cancellable per-file path as the main queue.  A single
        # all-files Office conversion used to make a stop request wait for the
        # entire batch and could also commit late reports after cancellation.
        for source_file in source_files:
            if self._turn_is_cancelled(turn_id):
                return
            reports.extend(self._generate_report_for_file(source_file, "\n\n".join(summary_parts), turn_id, turn_id, self.report_result_data.get(source_file)))
        if self._turn_is_cancelled(turn_id):
            return
        with self.lock:
            self.reports.extend(reports)
        artifact_content = _format_report_artifact_content("报告生成完成" if reports else "报告生成失败", reports)
        self._publish("artifact", "报告生成完成" if reports else "报告生成失败", artifact_content, "ok" if reports else "error", metadata={"turnId": turn_id, "roundId": turn_id, "phase": "artifact", "reports": reports, "sessionPatch": self._session_patch({"lastReportPath": reports[0] if reports else self.target_dir, "resumeState": None})})

    def _install_external_tools(self, message: str, turn_id: str) -> None:
        selected = []
        for name in ("nmap", "sqlmap", "ffuf"):
            if name in message.lower():
                selected.append(name)
        selected = selected or ["nmap", "sqlmap", "ffuf"]
        tool_call_id = f"install-tools:{self.session_id}:{uuid.uuid4().hex[:8]}"
        self._publish("tool_call", "一键下载外部工具", "下载并配置 nmap/sqlmap/ffuf 项目工具目录。", "info", tool={"toolId": "install_external_tools", "label": "一键下载外部工具", "status": "running", "target": ",".join(selected), "argsPreview": ",".join(selected)}, metadata={"toolCallId": tool_call_id, "turnId": turn_id, "roundId": turn_id, "phase": "tool", "sessionPatch": self._session_patch({"status": "正在下载外部工具", "isRunning": True, "resumeState": None})})
        result = _doc_retest_tools_install({"tools": selected})
        raw_output = "\n".join(result.get("logs") or []) or json.dumps(result, ensure_ascii=False, default=str, indent=2)
        self._publish("tool_result", "一键下载外部工具", result.get("message") or "", "ok" if result.get("success") else "error", tool={"toolId": "install_external_tools", "label": "一键下载外部工具", "status": "completed" if result.get("success") else "failed", "target": ",".join(selected), "resultPreview": result.get("message") or "", "rawOutput": raw_output, "failureReason": "" if result.get("success") else json.dumps(result.get("failures") or [], ensure_ascii=False)}, metadata={"toolCallId": tool_call_id, "turnId": turn_id, "roundId": turn_id, "phase": "tool", "sessionPatch": self._session_patch({"status": result.get("message") or "工具下载完成", "resumeState": None})})

    def _show_external_tool_status(self, turn_id: str) -> None:
        result = _doc_retest_tools_status({})
        lines = []
        for item in result.get("tools") or []:
            command = " ".join(str(part) for part in item.get("command") or [])
            lines.append(f"{item.get('name')}: {'已配置' if item.get('installed') else '未配置'} {command}")
        self._publish("chat", "外部工具状态", "\n".join(lines) or result.get("message") or "", "ok" if result.get("success") else "warn", metadata={"role": "agent", "turnId": turn_id})

    def _block(self, exc: RetestAIBlockedError, turn_id: str = "") -> None:
        with self.lock:
            self.running = False
            self.blocked = True
            self.blocked_reason = str(exc)
            self.blocked_stage = exc.stage
            self.blocked_title = _ai_blocked_title(exc)
            current_file_resume = self._current_file_resume_from_snapshot_locked(exc.resume_snapshot)
            if current_file_resume:
                self.current_file_resume = current_file_resume
        event_type = "status" if exc.stage in {"session_react", "session_compaction", "compact", "config"} else "error"
        self._publish(event_type, self.blocked_title, str(exc), "warn", metadata={"blockedByAiConfig": True, "turnId": turn_id or self.current_turn_id, "roundId": turn_id or self.current_turn_id, "phase": exc.stage, "sessionPatch": self._session_patch({"isRunning": False, "status": self.blocked_title, "resumeState": self._resume_state_locked(True)})})

    def _resume_state_locked(self, can_continue: bool) -> Dict[str, Any]:
        self._apply_frontend_progress_hints_locked()
        return {
            "canContinue": can_continue,
            "targetDir": self.target_dir,
            "sourceFiles": list(self.source_files),
            "nextIndex": self.next_index,
            "summaries": list(self.summaries),
            "reports": list(self.reports),
            "completionItems": list(self.completion_items),
            "diskCompletedFileNames": list(self.disk_completed_file_names),
            "diskCompletedReportEvidence": list(self.disk_completed_report_evidence),
            "allLogs": list(self.logs),
            "failedCount": len([item for item in self.completion_items if item.get("status") == "failed"]),
            "generateReports": self.generate_reports,
            "blockedReason": self.blocked_reason,
            "blockedStage": self.blocked_stage,
            "blockedTitle": self.blocked_title,
            "currentFile": dict(self.current_file_resume) if isinstance(self.current_file_resume, dict) else None,
        }

    def _session_patch(self, extra: Dict[str, Any] | None = None) -> Dict[str, Any]:
        with self.lock:
            patch = {
                "targetDir": self.target_dir,
                "status": self._status_locked(),
                "progress": self._progress_locked(),
                "isRunning": self.running,
                "log": "\n".join(self.logs[-3000:]),
                "lastReportPath": self.reports[0] if self.reports else self.target_dir,
                "latestResultData": self.latest_result_data,
                "resumeState": self._resume_state_locked(True) if self.blocked or self.stopped else None,
            }
        if extra:
            patch.update(extra)
        return patch

    def _publish(self, event_type: str, title: str, content: str = "", tone: str = "info", tool: Dict[str, Any] | None = None, metadata: Dict[str, Any] | None = None) -> None:
        event = _retest_trace_event(event_type, title, content, tone, tool=tool, metadata=metadata or {})
        try:
            self._agent_runtime().record_event(event)
        except Exception:
            pass
        try:
            from modules.backend_api.retest_event_stream import publish_retest_event

            publish_retest_event({"type": "retest_trace_event", "session_id": self.session_id, "task_id": "agent", "event": event})
        except Exception:
            pass


def _message_requests_report(message: str) -> bool:
    text = str(message or "")
    return any(word in text for word in ("生成报告", "写报告", "出报告", "导出报告", "报告"))


def _message_requests_rerun(message: str) -> bool:
    text = str(message or "")
    return any(word in text for word in ("重新", "再测", "重测", "重新测", "重新复测", "再跑", "重新跑", "测一遍"))


def _prune_retest_agent_runners_locked() -> None:
    if not _RETEST_AGENT_RUNNERS:
        return
    now = time.time()
    removable: List[str] = []
    for session_id, runner in _RETEST_AGENT_RUNNERS.items():
        with runner.lock:
            idle = not runner.running and not runner.pending_messages and (now - runner.updated_at) > _RETEST_AGENT_IDLE_TTL_SECONDS
        if idle:
            removable.append(session_id)

    if len(_RETEST_AGENT_RUNNERS) - len(removable) > _RETEST_AGENT_MAX_RUNNERS:
        candidates = []
        for session_id, runner in _RETEST_AGENT_RUNNERS.items():
            if session_id in removable:
                continue
            with runner.lock:
                if runner.running or runner.pending_messages:
                    continue
                candidates.append((runner.updated_at, session_id))
        candidates.sort()
        overflow = len(_RETEST_AGENT_RUNNERS) - len(removable) - _RETEST_AGENT_MAX_RUNNERS
        removable.extend(session_id for _updated_at, session_id in candidates[:max(0, overflow)])

    for session_id in removable:
        _RETEST_AGENT_RUNNERS.pop(session_id, None)


def _get_retest_agent_runner(session_id: str) -> RetestAgentRunner:
    if not session_id:
        session_id = f"agent-{uuid.uuid4().hex[:10]}"
    with _RETEST_AGENT_LOCK:
        _prune_retest_agent_runners_locked()
        runner = _RETEST_AGENT_RUNNERS.get(session_id)
        if runner is None:
            runner = RetestAgentRunner(session_id)
            _RETEST_AGENT_RUNNERS[session_id] = runner
        return runner


def _doc_retest_agent_start(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = str(payload.get("session_id") or "").strip() or f"session-{uuid.uuid4().hex[:10]}"
    runner = _get_retest_agent_runner(session_id)
    return runner.start(payload)


def _doc_retest_agent_message(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = str(payload.get("session_id") or "").strip()
    if not session_id:
        return {"success": False, "message": "缺少 session_id"}
    runner = _get_retest_agent_runner(session_id)
    return runner.message(payload)


def _doc_retest_agent_status(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = str(payload.get("session_id") or "").strip()
    if not session_id:
        return {"success": False, "message": "缺少 session_id"}
    with _RETEST_AGENT_LOCK:
        runner = _RETEST_AGENT_RUNNERS.get(session_id)
    if runner is None:
        return {
            "success": True,
            "active": False,
            "session_id": session_id,
            "running": False,
            "blocked": False,
            "message": "Agent 会话未在后端运行",
        }
    return {**runner.snapshot(), "success": True, "active": True}


def _doc_retest_agent_stop(payload: Dict[str, Any]) -> Dict[str, Any]:
    session_id = str(payload.get("session_id") or "").strip()
    if not session_id:
        return {"success": False, "message": "缺少 session_id"}
    _write_retest_cancel_marker("session", session_id)
    runner = _get_retest_agent_runner(session_id)
    return runner.stop()


def _doc_retest_run(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_dir = Path(_required_text(payload, "target_dir", "请选择通报目录")).expanduser()
    if not target_dir.exists() or not target_dir.is_dir():
        return {"success": False, "message": f"通报目录不存在: {target_dir}", "logs": []}

    logs: List[str] = []
    reports: List[str] = []
    source_files: List[str] = []
    summaries: List[str] = []
    risk_count = 0
    pass_count = 0
    failed_count = 0
    generate_reports = _payload_bool(payload.get("generate_reports"), True)

    try:
        _ensure_retest_ai_ready("config")
    except RetestAIBlockedError as exc:
        logs.append(str(exc))
        return _ai_blocked_payload(exc, "", logs, [])

    try:
        from modules.AI_Testing.retest.vulnerability_batch_scanner import VulnerabilityRetestScanner
        from modules.AI_Testing.retest.word_vulnerability_scanner import WordVulnerabilityScanner
        from modules.AI_Testing.retest.retest_report_generator import RetestReportGenerator
    except Exception as exc:
        return {"success": False, "message": f"导入复测模块失败: {exc}", "logs": [traceback.format_exc()]}

    template_path = _retest_template_path()
    if generate_reports and not template_path.exists():
        return {"success": False, "message": f"未找到复测模板文件: {template_path}", "logs": logs}

    scanner = WordVulnerabilityScanner(str(target_dir))
    buffer = io.StringIO()
    with contextlib.redirect_stdout(buffer):
        word_files = scanner.find_word_files()
    logs.extend(_captured_lines(buffer))
    logs.append(f"扫描完成，发现 {len(word_files)} 份通报文档")

    for index, file_path in enumerate(word_files, 1):
        source_files.append(str(file_path))
        logs.append(f"处理 ({index}/{len(word_files)}): {file_path.name}")
        try:
            summary, result_data, _manual_required, trace_events = _run_retest_for_source_file(file_path, payload, logs)
        except RetestAIBlockedError as exc:
            logs.append(str(exc))
            payload_blocked = _ai_blocked_payload(exc, str(file_path), logs, [])
            payload_blocked.update({
                "target_dir": str(target_dir),
                "source_files": source_files,
                "processed": max(0, index - 1),
                "reports": reports,
                "summary": "\n\n".join(summaries),
                "risk_count": risk_count,
                "pass_count": pass_count,
                "failed_count": failed_count,
                "manual_count": 0,
            })
            return payload_blocked
        final_verdict = _model_verdict_from_result_data(result_data)
        if final_verdict == "reproduced":
            risk_count += 1
        elif final_verdict == "not_reproduced":
            pass_count += 1
        else:
            failed_count += 1
            result_data["failed_count"] = max(1, int(result_data.get("failed_count") or 0))
            logs.append(f"{file_path.name} 模型未给出 reproduced/not_reproduced 判定，未按工具结果兜底。")
        summaries.append(summary)
        if trace_events:
            logs.append(f"{file_path.name} 记录 {len(trace_events)} 条 Agent 执行事件")

        if generate_reports:
            generator = RetestReportGenerator(
                target_dir=str(file_path.parent),
                template_path=str(template_path),
                output_dir=None,
                screenshot_path=None,
            )
            buffer = io.StringIO()
            with contextlib.redirect_stdout(buffer):
                generator_scan = generator.scan_document(file_path)
                output_path = generator.generate_report(generator_scan)
            logs.extend(_captured_lines(buffer))
            report_path = _existing_report_path(output_path)
            if report_path:
                reports.append(report_path)
                logs.append(f"报告已生成: {report_path}")
            else:
                failed_count += 1
                logs.append(f"报告生成失败或文件未落盘: {file_path.name} ({output_path or '无输出路径'})")

    message = (
        f"复测完成：处理 {len(word_files)} 份文档，生成 {len(reports)} 份报告，漏洞未修复/可复现 {risk_count} 份，复测通过 {pass_count} 份，执行失败 {failed_count} 份"
        if generate_reports
        else f"复测完成：处理 {len(word_files)} 份文档，等待截图写入报告，漏洞未修复/可复现 {risk_count} 份，复测通过 {pass_count} 份，执行失败 {failed_count} 份"
    )
    return {
        "success": True,
        "message": message,
        "target_dir": str(target_dir),
        "processed": len(word_files),
        "manual_count": 0,
        "risk_count": risk_count,
        "pass_count": pass_count,
        "failed_count": failed_count,
        "source_files": source_files,
        "reports": reports,
        "summary": "\n\n".join(summaries),
        "logs": logs,
    }


def _doc_retest_generate_reports_with_screenshot(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_dir = Path(_required_text(payload, "target_dir", "请选择通报目录")).expanduser()
    if not target_dir.exists() or not target_dir.is_dir():
        return {"success": False, "message": f"通报目录不存在: {target_dir}", "logs": []}

    source_files = _path_list(payload.get("source_files"))
    source_files = _source_notice_paths(source_files)
    if not source_files:
        return {"success": False, "message": "缺少待生成报告的原始通报文件列表", "logs": []}

    template_path = _retest_template_path()
    if not template_path.exists():
        return {"success": False, "message": f"未找到复测模板文件: {template_path}", "logs": []}

    logs: List[str] = []
    reports: List[str] = []
    disposal_reports: List[Dict[str, Any]] = []
    failures: List[tuple[Path, str]] = []
    report_screenshot_path: Path | None = None
    report_screenshot_sections: List[Dict[str, str]] = []
    disposal_screenshot_path: Path | None = None
    summary = repair_utf8_mojibake(payload.get("summary") or "")
    result_data = payload.get("result_data") if isinstance(payload.get("result_data"), dict) else {}
    detail_text = _format_report_text_explanation(summary, result_data) if (summary or result_data) else ""

    try:
        try:
            if summary or result_data:
                evidence_sections = _format_report_evidence_sections(summary, result_data)
                report_screenshot_sections = _save_retest_evidence_section_screenshots(target_dir, evidence_sections)
                if report_screenshot_sections:
                    report_screenshot_path = Path(report_screenshot_sections[0]["path"])
                    logs.append(f"复测报告分段证据图已生成: {len(report_screenshot_sections)} 张")
                else:
                    evidence_text = _format_report_evidence_screenshot_text(summary, result_data)
                    report_screenshot_path = _save_retest_text_screenshot(target_dir, evidence_text, "复测证据")
                    report_screenshot_sections = [{"caption": detail_text, "path": str(report_screenshot_path)}] if detail_text else []
                    logs.append(f"复测报告证据图已生成: {report_screenshot_path}")
                disposal_text = _format_report_evidence_snapshot(summary, result_data)
                disposal_screenshot_path = _save_retest_text_screenshot(target_dir, disposal_text, "复测证据总览")
                logs.append(f"处置文件复测证据总图已生成: {disposal_screenshot_path}")
            else:
                report_screenshot_path = _save_retest_screenshot_data(
                    target_dir,
                    _required_text(payload, "screenshot_data_url", "缺少复测截图数据"),
                )
                disposal_screenshot_path = report_screenshot_path
                report_screenshot_sections = [{"caption": "", "path": str(report_screenshot_path)}]
                logs.append(f"复测结果区域截图已保存: {report_screenshot_path}")
        except Exception as exc:
            logs.append(traceback.format_exc())
            return {"success": False, "message": f"保存复测截图失败: {exc}", "logs": logs}

        try:
            from modules.AI_Testing.retest.retest_report_generator import RetestReportGenerator
        except Exception as exc:
            logs.append(traceback.format_exc())
            return {"success": False, "message": f"导入复测报告生成器失败: {exc}", "logs": logs}

        for source_file in source_files:
            file_path = Path(source_file).expanduser()
            if not file_path.exists() or file_path.suffix.lower() not in WORD_SUFFIXES:
                failures.append((file_path, "通报文件不存在或不是 Word 文档"))
                continue
            if is_generated_retest_report_path(file_path):
                failures.append((file_path, "这是已生成的复测报告，不是原始通报文件"))
                continue

            generator = RetestReportGenerator(
                target_dir=str(file_path.parent),
                template_path=str(template_path),
                output_dir=None,
                screenshot_path=str(report_screenshot_path),
                screenshot_sections=report_screenshot_sections,
            )
            buffer = io.StringIO()
            with contextlib.redirect_stdout(buffer):
                generator_scan = generator.scan_document(file_path)
                if detail_text:
                    generator_scan["report_detail"] = detail_text
                output_path = generator.generate_report(generator_scan)
            logs.extend(_captured_lines(buffer))
            report_path = _existing_report_path(output_path)
            if report_path:
                reports.append(report_path)
                logs.append(f"报告已写入截图: {report_path}")
                try:
                    disposal_result = _prepare_retest_disposal_report(file_path, generator_scan, disposal_screenshot_path, logs)
                    if disposal_result:
                        disposal_reports.append(disposal_result)
                        if disposal_result.get("pdf_error"):
                            failures.append((file_path, f"处置文件PDF转换失败: {disposal_result['pdf_error']}"))
                except Exception as exc:
                    logs.append(traceback.format_exc())
                    failures.append((file_path, f"处置文件替换失败: {exc}"))
            else:
                failures.append((file_path, f"报告生成失败或文件未落盘: {output_path or '无输出路径'}"))

        success = bool(reports) and not failures
        message = (
            f"复测报告截图写入完成：生成 {len(reports)} 份，失败 {len(failures)} 份"
            if success
            else f"复测报告未生成或写入失败：生成 {len(reports)} 份，失败 {len(failures)} 份"
        )
        return {
            "success": success,
            "message": message,
            "target_dir": str(target_dir),
            "reports": reports,
            "disposal_reports": disposal_reports,
            "screenshot_path": str(report_screenshot_path),
            "failures": _failure_dicts(failures),
            "logs": logs,
        }
    finally:
        _cleanup_retest_screenshot_dir(target_dir, logs)


def _doc_retest_open_output(payload: Dict[str, Any]) -> Dict[str, Any]:
    target_dir = Path(_required_text(payload, "target_dir", "请选择通报目录")).expanduser()
    if not target_dir.exists() or not target_dir.is_dir():
        return {"success": False, "message": f"目录不存在: {target_dir}"}
    opened, error = _open_path_in_system(target_dir)
    return {
        "success": opened,
        "message": "已打开报告目录" if opened else f"无法打开报告目录: {error}",
        "path": str(target_dir),
    }

